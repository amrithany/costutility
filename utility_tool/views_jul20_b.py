# -*- coding: utf-8 -*-
from __future__ import unicode_literals

from django import forms
from django.template import Context, loader, RequestContext
from django.http import HttpResponse, HttpResponseRedirect,JsonResponse, FileResponse
from django.shortcuts import redirect,render, render_to_response
from django.template.loader import render_to_string
from django.forms.models import modelformset_factory,inlineformset_factory
from django.db.models import Q,F
from django.core.exceptions import ObjectDoesNotExist, MultipleObjectsReturned
from utility_tool.forms import DecisionForm, SolOptForm, SolOptForm2, ScrCriteriaForm, EvaCriteriaForm, LoginForm, RegisterForm, StakeholdersForm, SolOptView, VotesForm, ScoresForm, SolOptArchive, SetupForm, CostSetupForm, DecisionMadeForm, Solopt_Storage, License, SDForm_dec_file, SDForm_dec_link, SDForm_st_file, SDForm_st_link, SDForm_solopt_file, SDForm_solopt_link, SDForm_scr_file, SDForm_scr_link, SDForm_mapp_file, SDForm_mapp_link, SDForm_eva_file, SDForm_eva_link, SDForm_iw_file, SDForm_iw_link, SDForm_evam_file, SDForm_evam_link, SDForm_cost_file, SDForm_cost_link,SDForm_makedec_file, SDForm_makedec_link, ForgotForm, FileUploadForm   
from django.template import loader
from .models import Decisions, Solution_Options, Screening_Criteria, Evaluation_Criteria, Importance_Scores, Users, Stakeholders, Stakeholders_Decisions, MappingTable, SummaryTable, Evaluation_Measures, PA_Setup, EvaluationTable,  Cost_Setup, Cost_Utility, Decision_Made, Detailed_Costs, CBCSE_Screening_Criteria, Master_Screening_Criteria, CBCSE_Evaluation_Criteria, Master_Evaluation_Criteria, Scores_Setup, IdentifyTable, Solution_Options_Storage, Login, SD_dec_file, SD_dec_link, SD_st_file, SD_st_link, SD_solopt_file, SD_solopt_link, SD_scr_file, SD_scr_link, SD_mapp_file, SD_mapp_link, SD_eva_file, SD_eva_link, SD_iw_file, SD_iw_link, SD_evam_file, SD_evam_link, SD_cost_file, SD_cost_link,SD_makedec_file, SD_makedec_link, Duplicated_DecIds, SharedDec, Mapping_Data, Temp_Mapping, Cri_Temp_Mapping, Identify_Data  
import datetime
import json
import xlrd
import xlwt
import MySQLdb
import math
import types
import io
import itertools
from django.core import serializers
from django.http import Http404
from selectable.registry import registry
from utility_tool.functions import check_required, group_cal, individual_cal, further_cal, update_text_criteria, redistribution_func, IdforUser, StValue, PartValue, Participate, TotalST,LocVal,CleanupOpt,selectOption
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Flowable 
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from django.core.files.storage import FileSystemStorage
from reportlab.lib.colors import black, blue, lightblue
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
from reportlab.graphics.shapes import Drawing, Rect
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch, mm
import urllib
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
import inflect
from django.core.mail import EmailMessage,BadHeaderError, send_mail
import random
from datetime import timedelta
from django.utils import timezone
from dateutil.relativedelta import relativedelta
from datetime import date
from django.db.models import Count
from random import randint
from wsgiref.util import FileWrapper
import tempfile, zipfile
from zipfile import *
import os

def index(request):
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    #return HttpResponse("Hello, world. You're at the costutility index.")
    return render(request, 'index.html', {'loggedinuser':loggedinuser})

def Steps(request):
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    return render(request, 'Steps.html', {'loggedinuser':loggedinuser})

def Home(request):
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    return render(request, 'Home.html',{'loggedinuser':loggedinuser}) 

def ContactUs(request):
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    return render(request, 'Contact-Us.html', {'loggedinuser':loggedinuser})

def RG(request):
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    
    if loggedinuser == 'not found':
       return HttpResponseRedirect('/utility_tool/decisions/nouser_message.html')     
    else:    
       return render(request, 'Resources-Guidance.html',{'loggedinuser':loggedinuser})

def WhoWeAre(request):
    '''if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if loggedinuser == 'not found':
       return HttpResponseRedirect('/utility_tool/decisions/nouser_message.html')     
    else:'''
    return render(request, 'Who-We-Are.html') 

def OurTeam(request):                                                                                                                      
    '''if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if loggedinuser == 'not found':
       return HttpResponseRedirect('/utility_tool/decisions/nouser_message.html')     
    else:'''
    return render(request, 'Our-Team.html') 

def Online(request):                                                                                                                     
    return render(request, 'Online.html') 

def tutorials(request):      
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if loggedinuser == 'not found':
       return HttpResponseRedirect('/utility_tool/decisions/nouser_message.html')     
    else:
       return render(request, 'tutorials.html',{'loggedinuser':loggedinuser})        

def tutorials_dm(request):                                                                                                                               
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    '''if loggedinuser == 'not found':
       return HttpResponseRedirect('/utility_tool/decisions/nouser_message.html')     
    else:'''
    return render(request, 'tutorials_dm.html',{'loggedinuser':loggedinuser})      

def tutorials_ca(request): 
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    return render(request, 'tutorials_ca.html',{'loggedinuser':loggedinuser})

def tutorials_rc(request): 
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else: 
       loggedinuser = 'not found'
    return render(request, 'tutorials_rc.html',{'loggedinuser':loggedinuser})

def options(request):     
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'    
    return render(request, 'admin/options.html',{'loggedinuser':loggedinuser}) 

def genadmin(request):          
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    return render(request, 'admin/genadmin.html',{'loggedinuser':loggedinuser}) 

def upload(request):
    context = RequestContext(request) 
    DocList = []

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if request.method == 'POST':
        sform = FileUploadForm(request.POST, request.FILES)
        if sform.is_valid():
           id = sform.save(commit=False)
           id.docName = request.FILES['docfile']  
           id.docfile = '/home/amritha/costutility/costutility/static/' + request.FILES['docfile'].name
           id.save()
           try:  
              getfile = request.POST.get('docfile', False)
              loc = '/home/amritha/costutility/costutility/static/' + request.FILES['docfile'].name
              f = request.FILES['docfile']
              with open(loc, 'wb+') as destination:
                   for chunk in f.chunks():
                       destination.write(chunk)
           except:  
              print 'Please upload a file.'
    else:
        sform = FileUploadForm()
    # Load documents for the list page
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")
    cursor = database.cursor ()
    mysql = """SELECT docName, MAX(CONVERT_TZ(docDate,'GMT','EST')) from utility_tool_fileupload group by docName"""
    try:
      cursor.execute(mysql)
      results = cursor.fetchall()
      for row in results:
        ret = {}
        ret['Name'] = row[0]
        ret['Date'] = row[1]
        DocList.append(ret)
    except:
      print "Error: unable to fetch data"

    return render(request, 'admin/upload.html',{'DocList':DocList,'sform':sform,'loggedinuser':loggedinuser})

def updatedemo(request):          
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if request.method == 'POST':
       print request.POST.get('decidfordemoadmin')
       returned_decId = request.POST.get('decidfordemoadmin')
       try:
          val = int(returned_decId)
       except ValueError:
          return render(request,'admin/updatedemo.html',{'loggedinuser':loggedinuser, 'err': 'The Decision Id you have entered is not an integer.'})     
       try:
          #aug14
          dec = Decisions.objects.get(id = returned_decId)
          dec.demoDec = 'Y'
          dec.save(update_fields=['demoDec','updated_by', 'updated_date'])
          dupl_dec = Duplicated_DecIds(dec_id_for_dupl = dec.id, created_by = loggedinuser,created_date = datetime.datetime.now())
          dupl_dec.save() 
       except ObjectDoesNotExist:
          return render(request,'admin/updatedemo.html',{'loggedinuser':loggedinuser, 'err': 'The Decision Id you have entered does not exist.'})     

    all_dupldec = Duplicated_DecIds.objects.all() 
    return render(request, 'admin/updatedemo.html',{'loggedinuser':loggedinuser,'all_dupldec':all_dupldec}) 

def remove_dec(request, dec_id):                                                                                                                                                                                
    try:
       dupl_dec = Duplicated_DecIds.objects.get(id = dec_id)
       dec = Decisions.objects.get(id = dupl_dec.dec_id_for_dupl)
       dec.demoDec = 'N'
       dec.save(update_fields=['demoDec','updated_by', 'updated_date'])
       dupl_dec = Duplicated_DecIds.objects.get(id = dec_id).delete()
    except ObjectDoesNotExist:
       print 'nothing to do'

    return HttpResponseRedirect('/utility_tool/admin/updatedemo.html')

def userlist(request):   
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    users = Users.objects.all()
    users_count = users.count()
    UsersList = []
    for u in users:
        ret = {} 
        ret['id'] = u.id 
        ret['user'] = u.user
        ret['email'] = u.email
        ret['firstName'] = u.firstName
        ret['lastName'] = u.lastName
        ret['addressline1'] = u.addressline1
        ret['addressline2'] = u.addressline2
        ret['city'] = u.city
        ret['state'] = u.state
        ret['zip'] = u.zip
        ret['country'] = u.country
        ret['phone'] = u.phone
        ret['organisation'] = u.organisation
        ret['type_of_org'] = u.type_of_org
        ret['other_org'] = u.other_org
        ret['position'] = u.position
        ret['other_pos'] = u.other_pos
        ret['hearaboutus'] = u.hearaboutus  
        ret['other_hear'] = u.other_hear
        ret['updates'] = u.updates
        ret['education'] = u.education
        ret['age'] = u.age
        ret['gender'] = u.gender
        ret['race'] = u.race
        ret['other_race'] = u.other_race
        ret['publicOrPrivate'] = u.publicOrPrivate
        ret['startDate'] = str(u.startDate)
        ret['endDate'] = str(u.endDate)
        ret['lastLogin'] = str(u.lastLogin)
        ret['timesLoggedin'] = u.timesLoggedin
        UsersList.append(ret)
    if request.method == 'POST':
       print 'in user list' 
       print request.POST
       for val in request.POST.getlist('deleted'):
           print val
           val = val.strip()
           y = val.replace('[','')
           z = y.replace(']','')     
           temp_list = [] 
           # adding each id to a temporary list
           for l2 in z.split(','):                                                                                                                 
               l3 = l2.replace('"', '')
               temp_list.append(l3) 
           print temp_list
           for uid in temp_list:
               print uid
               try:
                  u = Users.objects.get(id = uid)
                  for d in Decisions.objects.filter(created_by = u.user):
                      Solution_Options.objects.filter(dec_id=d.id).delete()
                      Solution_Options_Storage.objects.filter(dec_id=d.id).delete()    
                      Screening_Criteria.objects.filter(dec_id=d.id).delete()
                      Evaluation_Criteria.objects.filter(dec_id=d.id).delete()
                      Stakeholders_Decisions.objects.filter(dec_id=d.id).delete()
                      Cost_Utility.objects.filter(dec_id=d.id).delete()
                      Cost_Setup.objects.filter(dec_id=d.id).delete()
                      Decision_Made.objects.filter(dec_id=d.id).delete()
                      Evaluation_Measures.objects.filter(dec_id=d.id).delete()
                      EvaluationTable.objects.filter(dec_id=d.id).delete()
                      Importance_Scores.objects.filter(dec_id=d.id).delete()
                      MappingTable.objects.filter(dec_id=d.id).delete()
                      Scores_Setup.objects.filter(dec_id=d.id).delete()
                      SummaryTable.objects.filter(dec_id=d.id).delete()
                      IdentifyTable.objects.filter(dec_id=d.id).delete()
                      Master_Screening_Criteria.objects.filter(dec_id=d.id).delete()
                      Master_Evaluation_Criteria.objects.filter(dec_id=d.id).delete()
                      Detailed_Costs.objects.filter(dec_id=d.id).delete()
                  Decisions.objects.filter(created_by = u.user).delete()
               except ObjectDoesNotExist:                                                                                                                   
                 print 'user does not exist'
               Users.objects.get(id = uid).delete() 
    return render(request, 'admin/userlist.html', {'UsersList': UsersList, 'users_count':users_count, 'loggedinuser':loggedinuser})   

#def reglist(request):     
    #return render(request, 'admin/reglist.html')   

def usageinfo(request):     
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    users = Users.objects.all()
    users_count = users.count()
    current_users = users.filter(startDate__lt=F('endDate'))
    current_users_count = current_users.count()

    real_dec_list = []
    dec1 = Decisions.objects.all()
    #aug7
    print dec1.count()
    dec2 = dec1.exclude(real_dec_yn = 'X')
    dec = dec2.exclude(demoDec = 'Y')
    print dec2.count()
    dec_count = dec.count()
    print dec_count
    real_dec1 = dec1.filter(real_dec_yn = 'R')
    real_dec = real_dec1.exclude(demoDec = 'Y')
    real_dec_count = real_dec.count()
    print 'Nov24'
    print real_dec_count
    for r in real_dec:
        real_dec_list.append(r.id) 
   
    tr_dec1 = dec1.filter(real_dec_yn = 'T')
    tr_dec = tr_dec1.exclude(demoDec = 'Y')                                                                                           
    tr_dec_count = tr_dec.count()
    test_dec1 = dec1.filter(real_dec_yn = 'X')
    test_dec = test_dec1.exclude(demoDec = 'Y')                                                                                           
    test_dec_count = test_dec.count()
 
    #aug14
    demo_dec_list = []
    demo_dec = dec1.filter(demoDec = 'Y')
    demo_dec_count = demo_dec.count()
    #print 'demo_dec_count'
    #print demo_dec_count
    for d in demo_dec:
        demo_dec_list.append(d.id) 
 
    # add test decisions to demo_dec_list
    for d in test_dec:
        demo_dec_list.append(d.id)  
    #Nov25 this is incorrect
    dist_users= demo_dec.values_list('created_by',flat = True).distinct()    
    dist_users_count = dist_users.count()
    print 'dist_users'
    print dist_users_count
    # select distinct(created_by) from utility_tool_decisions where demoDec = 'Y';

    iw2  = Importance_Scores.objects.all()
    iw1 = iw2.exclude(deleted = 'Y') 
    iw_count = iw1.count()                                                                                                                           
    #print 'all imp sc count'
    #print iw_count
    iw_nodemo = iw1.exclude(dec_id__in=demo_dec_list)
    #print iw_nodemo.count()
    final_iw = iw_nodemo.values_list('dec_id',flat = True).distinct()    
    iw1_count = final_iw.count()
    #print 'final number iw'
    #print iw1_count
    perc_iw = round(float(iw1_count) / float(dec_count),2) * 100

    mapp1  = MappingTable.objects.all()
    mapp1_count = mapp1.count()                                                                                                                           
    #print 'all mapp count'
    #print mapp1_count
    mapp_nodemo = mapp1.exclude(dec_id__in=demo_dec_list)
    #print mapp_nodemo.count()
    mapp = mapp_nodemo.values_list('dec_id',flat = True).distinct()    
    mapp_count = mapp.count()
    #print mapp_count 
    perc_mapp = round(float(mapp_count) / float(dec_count),2) * 100

    stdec2 = Stakeholders_Decisions.objects.all()
    stdec1 = stdec2.exclude(deleted = 'Y') 
    stdec = stdec1.exclude(dec_id__in=demo_dec_list)
    stdec_count = stdec.count()
    stdec_participate = stdec.filter(evacr_type = 'Y') | stdec.filter(scrcr_type = 'Y') | stdec.filter(iw_type = 'Y') | stdec.filter(solopt_type = 'Y')
    #we must remove the stakeholder who is the decision owner 
    stdec_participate_count = stdec_participate.count() - dec_count
    ave_stdec = round(float(stdec_count) / float(dec_count),2)

    stdec_real = stdec.filter(dec_id__in=real_dec_list)
    stdec_real_participate = stdec_real.filter(evacr_type = 'Y') | stdec_real.filter(scrcr_type = 'Y') | stdec_real.filter(iw_type = 'Y') | stdec_real.filter(solopt_type = 'Y')
    stdec_real_participate_count = stdec_real_participate.count() - real_dec_count

    solopt3 = Solution_Options.objects.all()
    solopt2 = solopt3.exclude(deleted = 'Y') 
    solopt1 = solopt2.exclude(archived = 'Y')  
    solopt = solopt1.exclude(dec_id__in=demo_dec_list) 
    solopt_real = solopt.filter(dec_id__in=real_dec_list)
    solopt_count = solopt.count()
    solopt_real_count = solopt_real.count()
    #print 'solopt_count'       
    #print solopt_count
    ave_solopt = round(float(solopt_count) / float(dec_count),2)      
    ave_real_solopt = round(float(solopt_real_count) / float(real_dec_count),2)

    scrcr1 = Screening_Criteria.objects.all()
    scrcr = scrcr1.exclude(dec_id__in=demo_dec_list)
    scrcr_real = scrcr.filter(dec_id__in=real_dec_list)
    scrcr_count = scrcr.count()
    scrcr_real_count = scrcr_real.count()
    print 'scrcr_count'
    print scrcr_count
    ave_scrcr = round(float(scrcr_count) / float(dec_count),2)     
    ave_real_scrcr = round(float(scrcr_real_count) / float(real_dec_count),2)
    scrcr1_count = scrcr.filter(criterion = 'Fits within available budget').count() 
    scrcr2_count = scrcr.filter(criterion = 'Can be implemented by date required').count() 
    scrcr3_count = scrcr.filter(criterion = 'Meets privacy standards').count() 
    scrcr4_count = scrcr.filter(criterion = 'Evidence of effectiveness exists').count() 
    scrcr5_count = scrcr.filter(criterion = 'Fits within school schedule').count() 
    scrcr6_count = scrcr.filter(criterion = 'Meets content requirements/ learning objectives').count() 
    scrcr7_count = scrcr.filter(criterion = 'Meets state code and/or other regulations').count() 
    scrcr8_count = scrcr.filter(criterion = 'Serves target population (grade level, ESL etc.)').count() 
    scrcr1_real_count = scrcr_real.filter(criterion = 'Fits within available budget').count() 
    scrcr2_real_count = scrcr_real.filter(criterion = 'Can be implemented by date required').count() 
    scrcr3_real_count = scrcr_real.filter(criterion = 'Meets privacy standards').count() 
    scrcr4_real_count = scrcr_real.filter(criterion = 'Evidence of effectiveness exists').count() 
    scrcr5_real_count = scrcr_real.filter(criterion = 'Fits within school schedule').count() 
    scrcr6_real_count = scrcr_real.filter(criterion = 'Meets content requirements/ learning objectives').count() 
    scrcr7_real_count = scrcr_real.filter(criterion = 'Meets state code and/or other regulations').count() 
    scrcr8_real_count = scrcr_real.filter(criterion = 'Serves target population (grade level, ESL etc.)').count()  
 
    evacr2 = Evaluation_Criteria.objects.all()
    evacr1 = evacr2.exclude(deleted = 'Y')  
    evacr = evacr1.exclude(dec_id__in=demo_dec_list)
    evacr_real = evacr.filter(dec_id__in=real_dec_list)
    evacr_count = evacr.count()
    evacr_real_count = evacr_real.count()
    print evacr_count
    print 'nov24_2'
    print evacr_real_count
    dec_evacr_real = evacr_real.values_list('dec_id',flat = True).distinct()
    print 'nov24_4'
    print dec_evacr_real.count()
    #evacr_real_count = dec_evacr_real.count()
    ave_evacr = round(float(evacr_count) / float(dec_count),2)      
    ave_real_evacr = round(float(evacr_real_count) / float(real_dec_count),2)    
    evacr1_countx = evacr.filter(criterion = 'Content meets learning objectives') 
    evacr1_count = evacr1_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr2_countx = evacr.filter(criterion = 'Number of students in need who can be served') 
    evacr2_count = evacr2_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr3_countx = evacr.filter(criterion = 'Accessible to target population') 
    evacr3_count = evacr3_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr4_countx = evacr.filter(criterion = 'Distribution of resources across population to be served') 
    evacr4_count = evacr4_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr5_countx = evacr.filter(criterion = 'Number of students participating') 
    evacr5_count = evacr5_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr6_countx = evacr.filter(criterion = 'Serves historically underserved groups') 
    evacr6_count = evacr6_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr7_countx = evacr.filter(criterion = 'Recommendations from external experts') 
    evacr7_count = evacr7_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr8_countx = evacr.filter(criterion = 'Recommendations from external peers') 
    evacr8_count = evacr8_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr9_countx = evacr.filter(criterion = 'Accessibility of physical location') 
    evacr9_count = evacr9_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr10_countx = evacr.filter(criterion = 'Amount of change in personnel resource requirements (e.g., time and number of staff) from the current status') 
    evacr10_count = evacr10_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr11_countx = evacr.filter(criterion = 'Amount of technical support needed') 
    evacr11_count = evacr11_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr12_countx = evacr.filter(criterion = 'Amount of training/PD needed') 
    evacr12_count = evacr12_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr13_countx = evacr.filter(criterion = 'Availability of necessary personnel, facilities, materials and equipment') 
    evacr13_count = evacr13_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr14_countx = evacr.filter(criterion = 'Availability of technical support to support implementers') 
    evacr14_count = evacr14_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr15_countx = evacr.filter(criterion = 'Availability of training/PD to support implementers')
    evacr15_count = evacr15_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr16_countx = evacr.filter(criterion = 'Can be implemented in desired timeline') 
    evacr16_count = evacr16_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr17_countx = evacr.filter(criterion = 'Capacity/skill level of current teachers/staff to implement option with fidelity') 
    evacr17_count = evacr17_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr18_countx = evacr.filter(criterion = 'Compatibility with existing systems') 
    evacr18_count = evacr18_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr19_countx = evacr.filter(criterion = 'Ease of use/ User friendliness') 
    evacr19_count = evacr19_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr20_countx = evacr.filter(criterion = 'Effect on teacher/staff workload') 
    evacr20_count = evacr20_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr21_countx = evacr.filter(criterion = 'Evidence of successful implementation in similar schools/districts/states') 
    evacr21_count = evacr21_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr22_countx = evacr.filter(criterion = 'Financial sustainability over time') 
    evacr22_count = evacr22_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr23_countx = evacr.filter(criterion = 'Fit with school calendar/schedule') 
    evacr23_count = evacr23_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr24_countx = evacr.filter(criterion = 'Likelihood this option will continue to be implemented with fidelity over time') 
    evacr24_count = evacr24_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr25_countx = evacr.filter(criterion = 'Scalability') 
    evacr25_count = evacr25_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr26_countx = evacr.filter(criterion = 'Solid plan proposed for financing') 
    evacr26_count = evacr26_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr27_countx = evacr.filter(criterion = 'Solid plan proposed for implementation') 
    evacr27_count = evacr27_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr28_countx = evacr.filter(criterion = 'Alignment to state/district/school mission and/or vision') 
    evacr28_count = evacr28_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr29_countx = evacr.filter(criterion = 'Alignment with current school/district/state curriculum') 
    evacr29_count = evacr29_countx.values_list('dec_id',flat = True).distinct().count()
    evacr30_countx = evacr.filter(criterion = 'Alignment with current state/district/school priorities') 
    evacr30_count = evacr30_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr31_countx = evacr.filter(criterion = 'Appropriate for student/staff demographics') 
    evacr31_count = evacr31_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr32_countx = evacr.filter(criterion = 'Customizability of solution to local needs') 
    evacr32_count = evacr32_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr33_countx = evacr.filter(criterion = 'Fit with local cultural values') 
    evacr33_count = evacr33_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr34_countx = evacr.filter(criterion = 'Impact on central control') 
    evacr34_count = evacr34_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr35_countx = evacr.filter(criterion = 'Impact on local autonomy') 
    evacr35_count = evacr35_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr36_countx = evacr.filter(criterion = 'Political value') 
    evacr36_count = evacr36_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr37_countx = evacr.filter(criterion = 'Viable in current political context') 
    evacr37_count = evacr37_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr38_countx = evacr.filter(criterion = 'Number of times parents call school') 
    evacr38_count = evacr38_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr39_countx = evacr.filter(criterion = 'Number of hours parents help children with homework') 
    evacr39_count = evacr39_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr40_countx = evacr.filter(criterion = 'Encourages parents to show up to PTA meetings') 
    evacr40_count = evacr40_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr41_countx = evacr.filter(criterion = 'Number of hours parents read to their child') 
    evacr41_count = evacr41_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr42_countx = evacr.filter(criterion = 'Impact on standardized test scores') 
    evacr42_count = evacr42_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr43_countx = evacr.filter(criterion = 'Impact on student grades')
    evacr43_count = evacr43_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr44_countx = evacr.filter(criterion = 'Impact on college admission') 
    evacr44_count = evacr44_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr45_countx = evacr.filter(criterion = 'Impact on course completion') 
    evacr45_count = evacr45_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr46_countx = evacr.filter(criterion = 'Impact on graduation') 
    evacr46_count = evacr46_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr47_countx = evacr.filter(criterion = 'Impact on progression to higher grade') 
    evacr47_count = evacr47_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr48_countx = evacr.filter(criterion = 'Impact on progression towards graduation') 
    evacr48_count = evacr48_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr49_countx = evacr.filter(criterion = 'Impact on closing the achievement gap') 
    evacr49_count = evacr49_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr50_countx = evacr.filter(criterion = 'Impact on student sense of belongingness') 
    evacr50_count = evacr50_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr51_countx = evacr.filter(criterion = 'Impact on incidence of misbehavior') 
    evacr51_count = evacr51_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr52_countx = evacr.filter(criterion = 'Impact on school climate') 
    evacr52_count = evacr52_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr53_countx = evacr.filter(criterion = 'Impact on suspensions') 
    evacr53_count = evacr53_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr54_countx = evacr.filter(criterion = 'Impact on attendance') 
    evacr54_count = evacr54_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr55_countx = evacr.filter(criterion = 'Impact on staff absenteeism') 
    evacr55_count = evacr55_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr56_countx = evacr.filter(criterion = 'Impact on staff/teacher/student effort') 
    evacr56_count = evacr56_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr57_countx = evacr.filter(criterion = 'Impact on student-teacher interaction') 
    evacr57_count = evacr57_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr58_countx = evacr.filter(criterion = 'Improves teacher time-on-task') 
    evacr58_count = evacr58_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr59_countx = evacr.filter(criterion = 'Improves student time-on-task') 
    evacr59_count = evacr59_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr60_countx = evacr.filter(criterion = 'Improves teacher value-added') 
    evacr60_count = evacr60_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr61_countx = evacr.filter(criterion = 'Improves teacher pedogogical skills') 
    evacr61_count = evacr61_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr62_countx = evacr.filter(criterion = 'Improves teacher content knowledge') 
    evacr62_count = evacr62_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr63_countx = evacr.filter(criterion = 'Addresses safety concerns') 
    evacr63_count = evacr63_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr64_countx = evacr.filter(criterion = 'Alignment with state code and other regulations') 
    evacr64_count = evacr64_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr65_countx = evacr.filter(criterion = 'Compliance with regulations') 
    evacr65_count = evacr65_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr66_countx = evacr.filter(criterion = 'Meets privacy standards') 
    evacr66_count = evacr66_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr67_countx = evacr.filter(criterion = 'Content maintains level of rigor/quality of learning') 
    evacr67_count = evacr67_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr68_countx = evacr.filter(criterion = 'Meets state/district/school standards') 
    evacr68_count = evacr68_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr69_countx = evacr.filter(criterion = 'Consistency of implementation across sites') 
    evacr69_count = evacr69_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr70_countx = evacr.filter(criterion = 'Fidelity of implementation')
    evacr70_count = evacr70_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr71_countx = evacr.filter(criterion = 'Level of monitoring conducted') 
    evacr71_count = evacr71_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr72_countx = evacr.filter(criterion = 'Quality of technical support provided to implementers') 
    evacr72_count = evacr72_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr73_countx = evacr.filter(criterion = 'Quality of training delivered') 
    evacr73_count = evacr73_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr74_countx = evacr.filter(criterion = 'Change in teacher pedagogy') 
    evacr74_count = evacr74_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr75_countx = evacr.filter(criterion = 'Board preference/buy-in/support') 
    evacr75_count = evacr75_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr76_countx = evacr.filter(criterion = 'Community preference/buy-in/support') 
    evacr76_count = evacr76_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr77_countx = evacr.filter(criterion = 'Parent preference/buy-in/support') 
    evacr77_count = evacr77_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr78_countx = evacr.filter(criterion = 'Student preference/buy-in/support') 
    evacr78_count = evacr78_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr79_countx = evacr.filter(criterion = 'Teacher preference/buy-in/support') 
    evacr79_count = evacr79_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr80_countx = evacr.filter(criterion = 'Union preference/buy-in/support') 
    evacr80_count = evacr80_countx.values_list('dec_id',flat = True).distinct().count()

    evacr_totalcount1 = evacr1_count + evacr2_count
    evacr_totalcount2 = evacr3_count + evacr4_count + evacr5_count + evacr6_count
    evacr_totalcount3 = evacr7_count + evacr8_count
    evacr_totalcount4 = evacr9_count + evacr10_count + evacr11_count + evacr12_count + evacr13_count + evacr14_count + evacr15_count + evacr16_count + evacr17_count + evacr18_count + evacr19_count + evacr20_count + evacr21_count + evacr22_count + evacr23_count + evacr24_count + evacr25_count + evacr26_count + evacr27_count 
    evacr_totalcount5 = evacr28_count + evacr29_count + evacr30_count + evacr31_count + evacr32_count + evacr33_count + evacr34_count + evacr35_count + evacr36_count + evacr37_count
    evacr_totalcount6 = evacr38_count + evacr39_count + evacr40_count + evacr41_count
    evacr_totalcount7 = evacr42_count + evacr43_count + evacr44_count + evacr45_count + evacr46_count + evacr47_count + evacr48_count + evacr49_count
    evacr_totalcount8 = evacr50_count + evacr51_count + evacr52_count + evacr53_count
    evacr_totalcount9 = evacr54_count + evacr55_count + evacr56_count + evacr57_count
    evacr_totalcount10 = evacr58_count + evacr59_count + evacr60_count + evacr61_count + evacr62_count 
    evacr_totalcount11 = evacr63_count + evacr64_count + evacr65_count + evacr66_count + evacr67_count + evacr68_count
    evacr_totalcount12 = evacr69_count + evacr70_count + evacr71_count + evacr72_count + evacr73_count + evacr74_count
    evacr_totalcount13 = evacr75_count + evacr76_count + evacr77_count + evacr78_count + evacr79_count + evacr80_count

    or1_countx = evacr.filter(or_criterion = 'Addresses the identified need') 
    or1_count = or1_countx.values_list('dec_id',flat = True).distinct().count() 
    or2_countx = evacr.filter(or_criterion = 'Equity') 
    or2_count = or2_countx.values_list('dec_id',flat = True).distinct().count() 
    or3_countx = evacr.filter(or_criterion = 'External recommendations') 
    or3_count = or3_countx.values_list('dec_id',flat = True).distinct().count() 
    or4_countx = evacr.filter(or_criterion = 'Feasibility of implementation') 
    or4_count = or4_countx.values_list('dec_id',flat = True).distinct().count() 
    or5_countx = evacr.filter(or_criterion = 'Fit with local context') 
    or5_count = or5_countx.values_list('dec_id',flat = True).distinct().count() 
    or6_countx = evacr.filter(or_criterion = 'Impact on parent engagement') 
    or6_count = or6_countx.values_list('dec_id',flat = True).distinct().count() 
    or7_countx = evacr.filter(or_criterion = 'Impact on student academic performance') 
    or7_count = or7_countx.values_list('dec_id',flat = True).distinct().count() 
    or8_countx = evacr.filter(or_criterion = 'Impact on student socio-emotional development') 
    or8_count = or8_countx.values_list('dec_id',flat = True).distinct().count() 
    or9_countx = evacr.filter(or_criterion = 'Impact on student/staff engagement') 
    or9_count = or9_countx.values_list('dec_id',flat = True).distinct().count() 
    or10_countx = evacr.filter(or_criterion = 'Improves teacher performance') 
    or10_count = or10_countx.values_list('dec_id',flat = True).distinct().count() 
    or11_countx = evacr.filter(or_criterion = 'Meets required standards and regulations') 
    or11_count = or11_countx.values_list('dec_id',flat = True).distinct().count() 
    or12_countx = evacr.filter(or_criterion = 'Quality of implementation (for programs/strategies/tools already in place)') 
    or12_count = or12_countx.values_list('dec_id',flat = True).distinct().count() 
    or13_countx = evacr.filter(or_criterion = 'Support from stakeholders') 
    or13_count = or13_countx.values_list('dec_id',flat = True).distinct().count()

    or1_real_countx = evacr_real.filter(or_criterion = 'Addresses the identified need')
    or1_real_count = or1_real_countx.values_list('dec_id',flat = True).distinct().count()
    or2_real_countx = evacr_real.filter(or_criterion = 'Equity')
    or2_real_count = or2_real_countx.values_list('dec_id',flat = True).distinct().count()
    or3_real_countx = evacr_real.filter(or_criterion = 'External recommendations')
    or3_real_count = or3_real_countx.values_list('dec_id',flat = True).distinct().count()
    or4_real_countx = evacr_real.filter(or_criterion = 'Feasibility of implementation')
    or4_real_count = or4_real_countx.values_list('dec_id',flat = True).distinct().count()
    or5_real_countx = evacr_real.filter(or_criterion = 'Fit with local context')
    or5_real_count = or5_real_countx.values_list('dec_id',flat = True).distinct().count()
    or6_real_countx = evacr_real.filter(or_criterion = 'Impact on parent engagement')
    or6_real_count = or6_real_countx.values_list('dec_id',flat = True).distinct().count()
    or7_real_countx = evacr_real.filter(or_criterion = 'Impact on student academic performance')
    or7_real_count = or7_real_countx.values_list('dec_id',flat = True).distinct().count()
    or8_real_countx = evacr_real.filter(or_criterion = 'Impact on student socio-emotional development')
    or8_real_count = or8_real_countx.values_list('dec_id',flat = True).distinct().count()
    or9_real_countx = evacr_real.filter(or_criterion = 'Impact on student/staff engagement')
    or9_real_count = or9_real_countx.values_list('dec_id',flat = True).distinct().count()
    or10_real_countx = evacr_real.filter(or_criterion = 'Improves teacher performance')
    or10_real_count = or10_real_countx.values_list('dec_id',flat = True).distinct().count()
    or11_real_countx = evacr_real.filter(or_criterion = 'Meets required standards and regulations')
    or11_real_count = or11_real_countx.values_list('dec_id',flat = True).distinct().count()
    or12_real_countx = evacr_real.filter(or_criterion = 'Quality of implementation (for programs/strategies/tools already in place)')
    or12_real_count = or12_real_countx.values_list('dec_id',flat = True).distinct().count()
    or13_real_countx = evacr_real.filter(or_criterion = 'Support from stakeholders')
    or13_real_count = or13_real_countx.values_list('dec_id',flat = True).distinct().count()

    evacr1_real_countx = evacr_real.filter(criterion = 'Content meets learning objectives')
    evacr1_real_count = evacr1_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr2_real_countx = evacr_real.filter(criterion = 'Number of students in need who can be served')
    evacr2_real_count = evacr2_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr3_real_countx = evacr_real.filter(criterion = 'Accessible to target population')
    evacr3_real_count = evacr3_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr4_real_countx = evacr_real.filter(criterion = 'Distribution of resources across population to be served')
    evacr4_real_count = evacr4_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr5_real_countx = evacr_real.filter(criterion = 'Number of students participating')
    evacr5_real_count = evacr5_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr6_real_countx = evacr_real.filter(criterion = 'Serves historically underserved groups')
    evacr6_real_count = evacr6_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr7_real_countx = evacr_real.filter(criterion = 'Recommendations from external experts')
    evacr7_real_count = evacr7_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr8_real_countx = evacr_real.filter(criterion = 'Recommendations from external peers')
    evacr8_real_count = evacr8_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr9_real_countx = evacr_real.filter(criterion = 'Accessibility of physical location')
    evacr9_real_count = evacr9_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr10_real_countx = evacr_real.filter(criterion = 'Amount of change in personnel resource requirements (e.g., time and number of staff) from the current status')
    evacr10_real_count = evacr10_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr11_real_countx = evacr_real.filter(criterion = 'Amount of technical support needed')
    evacr11_real_count = evacr11_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr12_real_countx = evacr_real.filter(criterion = 'Amount of training/PD needed')
    evacr12_real_count = evacr12_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr13_real_countx = evacr_real.filter(criterion = 'Availability of necessary personnel, facilities, materials and equipment')
    evacr13_real_count = evacr13_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr14_real_countx = evacr_real.filter(criterion = 'Availability of technical support to support implementers')
    evacr14_real_count = evacr14_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr15_real_countx = evacr_real.filter(criterion = 'Availability of training/PD to support implementers')
    evacr15_real_count = evacr15_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr16_real_countx = evacr_real.filter(criterion = 'Can be implemented in desired timeline')
    evacr16_real_count = evacr16_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr17_real_countx = evacr_real.filter(criterion = 'Capacity/skill level of current teachers/staff to implement option with fidelity')
    evacr17_real_count = evacr17_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr18_real_countx = evacr_real.filter(criterion = 'Compatibility with existing systems')
    evacr18_real_count = evacr18_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr19_real_countx = evacr_real.filter(criterion = 'Ease of use/ User friendliness')
    evacr19_real_count = evacr19_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr20_real_countx = evacr_real.filter(criterion = 'Effect on teacher/staff workload')
    evacr20_real_count = evacr20_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr21_real_countx = evacr_real.filter(criterion = 'Evidence of successful implementation in similar schools/districts/states')
    evacr21_real_count = evacr21_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr22_real_countx = evacr_real.filter(criterion = 'Financial sustainability over time')
    evacr22_real_count = evacr22_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr23_real_countx = evacr_real.filter(criterion = 'Fit with school calendar/schedule')
    evacr23_real_count = evacr23_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr24_real_countx = evacr_real.filter(criterion = 'Likelihood this option will continue to be implemented with fidelity over time')
    evacr24_real_count = evacr24_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr25_real_countx = evacr_real.filter(criterion = 'Scalability')
    evacr25_real_count = evacr25_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr26_real_countx = evacr_real.filter(criterion = 'Solid plan proposed for financing')
    evacr26_real_count = evacr26_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr27_real_countx = evacr_real.filter(criterion = 'Solid plan proposed for implementation')
    evacr27_real_count = evacr27_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr28_real_countx = evacr_real.filter(criterion = 'Alignment to state/district/school mission and/or vision')
    evacr28_real_count = evacr28_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr29_real_countx = evacr_real.filter(criterion = 'Alignment with current school/district/state curriculum')
    evacr29_real_count = evacr29_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr30_real_countx = evacr_real.filter(criterion = 'Alignment with current state/district/school priorities')
    evacr30_real_count = evacr30_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr31_real_countx = evacr_real.filter(criterion = 'Appropriate for student/staff demographics')
    evacr31_real_count = evacr31_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr32_real_countx = evacr_real.filter(criterion = 'Customizability of solution to local needs')
    evacr32_real_count = evacr32_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr33_real_countx = evacr_real.filter(criterion = 'Fit with local cultural values')
    evacr33_real_count = evacr33_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr34_real_countx = evacr_real.filter(criterion = 'Impact on central control')
    evacr34_real_count = evacr34_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr35_real_countx = evacr_real.filter(criterion = 'Impact on local autonomy')
    evacr35_real_count = evacr35_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr36_real_countx = evacr_real.filter(criterion = 'Political value')
    evacr36_real_count = evacr36_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr37_real_countx = evacr_real.filter(criterion = 'Viable in current political context')
    evacr37_real_count = evacr37_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr38_real_countx = evacr_real.filter(criterion = 'Number of times parents call school')
    evacr38_real_count = evacr38_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr39_real_countx = evacr_real.filter(criterion = 'Number of hours parents help children with homework')
    evacr39_real_count = evacr39_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr40_real_countx = evacr_real.filter(criterion = 'Encourages parents to show up to PTA meetings')
    evacr40_real_count = evacr40_real_countx.values_list('dec_id',flat = True).distinct().count() 
    evacr41_real_countx = evacr_real.filter(criterion = 'Number of hours parents read to their child')
    evacr41_real_count = evacr41_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr42_real_countx = evacr_real.filter(criterion = 'Impact on standardized test scores')
    evacr42_real_count = evacr42_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr43_real_countx = evacr_real.filter(criterion = 'Impact on student grades')
    evacr43_real_count = evacr43_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr44_real_countx = evacr_real.filter(criterion = 'Impact on college admission')
    evacr44_real_count = evacr44_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr45_real_countx = evacr_real.filter(criterion = 'Impact on course completion')
    evacr45_real_count = evacr45_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr46_real_countx = evacr_real.filter(criterion = 'Impact on graduation') 
    evacr46_real_count = evacr46_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr47_real_countx = evacr_real.filter(criterion = 'Impact on progression to higher grade')
    evacr47_real_count = evacr47_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr48_real_countx = evacr_real.filter(criterion = 'Impact on progression towards graduation')
    evacr48_real_count = evacr48_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr49_real_countx = evacr_real.filter(criterion = 'Impact on closing the achievement gap')
    evacr49_real_count = evacr49_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr50_real_countx = evacr_real.filter(criterion = 'Impact on student sense of belongingness')
    evacr50_real_count = evacr50_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr51_real_countx = evacr_real.filter(criterion = 'Impact on incidence of misbehavior')
    evacr51_real_count = evacr51_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr52_real_countx = evacr_real.filter(criterion = 'Impact on school climate')
    evacr52_real_count = evacr52_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr53_real_countx = evacr_real.filter(criterion = 'Impact on suspensions')
    evacr53_real_count = evacr53_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr54_real_countx = evacr_real.filter(criterion = 'Impact on attendance')
    evacr54_real_count = evacr54_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr55_real_countx = evacr_real.filter(criterion = 'Impact on staff absenteeism')
    evacr55_real_count = evacr55_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr56_real_countx = evacr_real.filter(criterion = 'Impact on staff/teacher/student effort')
    evacr56_real_count = evacr56_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr57_real_countx = evacr_real.filter(criterion = 'Impact on student-teacher interaction')
    evacr57_real_count = evacr57_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr58_real_countx = evacr_real.filter(criterion = 'Improves teacher time-on-task') 
    evacr58_real_count = evacr58_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr59_real_countx = evacr_real.filter(criterion = 'Improves student time-on-task')
    evacr59_real_count = evacr59_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr60_real_countx = evacr_real.filter(criterion = 'Improves teacher value-added')
    evacr60_real_count = evacr60_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr61_real_countx = evacr_real.filter(criterion = 'Improves teacher pedogogical skills')
    evacr61_real_count = evacr61_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr62_real_countx = evacr_real.filter(criterion = 'Improves teacher content knowledge') 
    evacr62_real_count = evacr62_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr63_real_countx = evacr_real.filter(criterion = 'Addresses safety concerns')
    evacr63_real_count = evacr63_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr64_real_countx = evacr_real.filter(criterion = 'Alignment with state code and other regulations')
    evacr64_real_count = evacr64_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr65_real_countx = evacr_real.filter(criterion = 'Compliance with regulations')
    evacr65_real_count = evacr65_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr66_real_countx = evacr_real.filter(criterion = 'Meets privacy standards')
    evacr66_real_count = evacr66_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr67_real_countx = evacr_real.filter(criterion = 'Content maintains level of rigor/quality of learning')
    evacr67_real_count = evacr67_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr68_real_countx = evacr_real.filter(criterion = 'Meets state/district/school standards')
    evacr68_real_count = evacr68_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr69_real_countx = evacr_real.filter(criterion = 'Consistency of implementation across sites')
    evacr69_real_count = evacr69_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr70_real_countx = evacr_real.filter(criterion = 'Fidelity of implementation')
    evacr70_real_count = evacr70_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr71_real_countx = evacr_real.filter(criterion = 'Level of monitoring conducted')
    evacr71_real_count = evacr71_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr72_real_countx = evacr_real.filter(criterion = 'Quality of technical support provided to implementers')
    evacr72_real_count = evacr72_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr73_real_countx = evacr_real.filter(criterion = 'Quality of training delivered')
    evacr73_real_count = evacr73_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr74_real_countx = evacr_real.filter(criterion = 'Change in teacher pedagogy')
    evacr74_real_count = evacr74_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr75_real_countx = evacr_real.filter(criterion = 'Board preference/buy-in/support') 
    evacr75_real_count = evacr75_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr76_real_countx = evacr_real.filter(criterion = 'Community preference/buy-in/support')
    evacr76_real_count = evacr76_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr77_real_countx = evacr_real.filter(criterion = 'Parent preference/buy-in/support')
    evacr77_real_count = evacr77_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr78_real_countx = evacr_real.filter(criterion = 'Student preference/buy-in/support')
    evacr78_real_count = evacr78_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr79_real_countx = evacr_real.filter(criterion = 'Teacher preference/buy-in/support')
    evacr79_real_count = evacr79_real_countx.values_list('dec_id',flat = True).distinct().count()
    evacr80_real_countx = evacr_real.filter(criterion = 'Union preference/buy-in/support')
    evacr80_real_count = evacr80_real_countx.values_list('dec_id',flat = True).distinct().count()

    r_evacr_totalcount1 = evacr1_real_count + evacr2_real_count
    r_evacr_totalcount2 = evacr3_real_count + evacr4_real_count + evacr5_real_count + evacr6_real_count
    r_evacr_totalcount3 = evacr7_real_count + evacr8_real_count
    r_evacr_totalcount4 = evacr9_real_count + evacr10_real_count + evacr11_real_count + evacr12_real_count + evacr13_real_count + evacr14_real_count + evacr15_real_count + evacr16_real_count + evacr17_real_count + evacr18_real_count + evacr19_real_count + evacr20_real_count + evacr21_real_count + evacr22_real_count + evacr23_real_count + evacr24_real_count + evacr25_real_count + evacr26_real_count + evacr27_real_count 
    r_evacr_totalcount5 = evacr28_real_count + evacr29_real_count + evacr30_real_count + evacr31_real_count + evacr32_real_count + evacr33_real_count + evacr34_real_count + evacr35_real_count + evacr36_real_count + evacr37_real_count
    r_evacr_totalcount6 = evacr38_real_count + evacr39_real_count + evacr40_real_count + evacr41_real_count
    r_evacr_totalcount7 = evacr42_real_count + evacr43_real_count + evacr44_real_count + evacr45_real_count + evacr46_real_count + evacr47_real_count + evacr48_real_count + evacr49_real_count
    r_evacr_totalcount8 = evacr50_real_count + evacr51_real_count + evacr52_real_count + evacr53_real_count
    r_evacr_totalcount9 = evacr54_real_count + evacr55_real_count + evacr56_real_count + evacr57_real_count
    r_evacr_totalcount10 = evacr58_real_count + evacr59_real_count + evacr60_real_count + evacr61_real_count + evacr62_real_count 
    r_evacr_totalcount11 = evacr63_real_count + evacr64_real_count + evacr65_real_count + evacr66_real_count + evacr67_real_count + evacr68_real_count
    r_evacr_totalcount12 = evacr69_real_count + evacr70_real_count + evacr71_real_count + evacr72_real_count + evacr73_real_count + evacr74_real_count
    r_evacr_totalcount13 = evacr75_real_count + evacr76_real_count + evacr77_real_count + evacr78_real_count + evacr79_real_count + evacr80_real_count

    sol  = solopt.values_list('dec_id',flat = True).distinct()
    sol_count = sol.count()
    perc_sol = round(float(sol_count) / float(dec_count),2) * 100

    scr  = scrcr.values_list('dec_id',flat = True).distinct()
    scr_count = scr.count()
    perc_scr = round(float(scr_count) / float(dec_count),2) * 100

    eva  = evacr.values_list('dec_id',flat = True).distinct()
    eva_count = eva.count()
    perc_eva = round(float(eva_count) / float(dec_count),2) * 100

    std  = stdec.values_list('dec_id',flat = True).distinct()
    std_count = std.count()
    print 'nov25_3'
    print std_count
    print dec_count
    perc_std = round(float(std_count) / float(dec_count),2) * 100
    print perc_std

    costs3 = Cost_Utility.objects.all()
    costs1 = costs3.exclude(archived = 'Y')
    costs2 = costs1.exclude(dec_id__in=demo_dec_list)
    costs = costs2.values_list('dec_id',flat = True).distinct()
    costs_real = costs.filter(dec_id__in=real_dec_list).values_list('dec_id',flat = True).distinct()
    costs_count = costs.count()
    costs_real_count = costs_real.count()
    print 'costs count'
    print costs_count
    ave_costs = round(float(costs_count) / float(dec_count),2) * 100
    ave_real_costs = round(float(costs_real_count) / float(real_dec_count),2) * 100

    # when you try to exclude votes not equal to 10 results are weird so am using greater than 10 
    iw = stdec.filter(iw_type = 'Y',votes__gt=10).values_list('dec_id',flat = True).distinct()
    iw_count =iw.count()
    print 'iw_count'
    print iw_count
    ave_iw = round(float(iw_count) / float(dec_count),2) * 100   
    iw_real = stdec.filter(iw_type = 'Y', dec_id__in=real_dec_list, votes__gt=10).values_list('dec_id',flat = True).distinct()    
    iw_real_count =iw_real.count()
    print iw_real_count
    ave_real_iw = round(float(iw_real_count) / float(real_dec_count),2) * 100  

    evm3 = Evaluation_Measures.objects.all()
    evm1 = evm3.exclude(deleted = 'Y')
    evm2 = evm1.exclude(dec_id__in=demo_dec_list)
    evm = evm2.values_list('dec_id',flat = True).distinct()
    evm_count = evm.count()
    print 'evm count'
    print evm_count
    ave_evm = round(float(evm_count) / float(dec_count),2) * 100   
    evm_real = evm.filter(dec_id__in=real_dec_list).values_list('dec_id',flat = True).distinct()
    evm_real_count = evm_real.count()
    ave_real_evm = round(float(evm_real_count) / float(real_dec_count),2) * 100  

    dm1 = Decision_Made.objects.all()
    dm2 = dm1.exclude(dec_id__in=demo_dec_list)
    dm = dm2.values_list('dec_id',flat = True).distinct()
    print 'dm count'
    dm_count = dm.count()
    print dm_count
    perc_dm = round(float(dm_count) / float(dec_count),2) * 100

    decmade = dm2.filter(sol_option__isnull=False)
    decmade_real = dm2.filter(dec_id__in=real_dec_list, sol_option__isnull=False)
    decmade_1 = decmade.exclude(sol_option = 'NULL')
    decmade_2 = decmade_1.exclude(sol_option = '[]')
    decmade_2_count = decmade_2.count()
    decmade_r1 = decmade_real.exclude(sol_option = 'NULL')
    decmade_r2 = decmade_r1.exclude(sol_option = '[]')
    decmade_r2_count = decmade_r2.count()
    print decmade_2_count
    ave_decmade = round(float(decmade_2_count) / float(dec_count),2) * 100 
    ave_real_decmade = round(float(decmade_r2_count) / float(real_dec_count),2) * 100 

    start = datetime.date(2019, 8, 31) 
    today = datetime.date.today()
    end = today - start
    login_sinceAug31 = Users.objects.filter(lastLogin__gte = start).values_list('lastLogin',flat = True).distinct() 
    login31_count = login_sinceAug31.count()
    print 'login_sinceAug31_count'
    print login31_count

    one_week_back = timezone.now().date() - timedelta(days=7)
    #monday_of_last_week = some_day_last_week - timedelta(days=(some_day_last_week.isocalendar()[2] - 1))
    #monday_of_this_week = monday_of_last_week + timedelta(days=7)
    users_lastweek = Users.objects.filter(lastLogin__gte=one_week_back)
    users_lastweek_count = users_lastweek.count()
    login_lastweek = Login.objects.filter(loggedindate__gte=one_week_back)
    login_lastweek_count = login_lastweek.count()

    one_month_back = timezone.now().date() - timedelta(days=30)
    print one_month_back
    users_lastmonth = Users.objects.filter(lastLogin__gte=one_month_back)
    users_lastmonth_count = users_lastmonth.count()
    login_lastmonth = Login.objects.filter(loggedindate__gte=one_month_back)
    login_lastmonth_count = login_lastmonth.count()

    request.session['r_evacr_totalcount1'] = r_evacr_totalcount1
    request.session['r_evacr_totalcount2'] = r_evacr_totalcount2
    request.session['r_evacr_totalcount3'] = r_evacr_totalcount3
    request.session['r_evacr_totalcount4'] = r_evacr_totalcount4
    request.session['r_evacr_totalcount5'] = r_evacr_totalcount5
    request.session['r_evacr_totalcount6'] = r_evacr_totalcount6
    request.session['r_evacr_totalcount7'] = r_evacr_totalcount7
    request.session['r_evacr_totalcount8'] = r_evacr_totalcount8
    request.session['r_evacr_totalcount9'] = r_evacr_totalcount9
    request.session['r_evacr_totalcount10'] = r_evacr_totalcount10
    request.session['r_evacr_totalcount11'] = r_evacr_totalcount11
    request.session['r_evacr_totalcount12'] = r_evacr_totalcount12
    request.session['r_evacr_totalcount13'] = r_evacr_totalcount13
    request.session['evacr_totalcount1'] = evacr_totalcount1
    request.session['evacr_totalcount2'] = evacr_totalcount2
    request.session['evacr_totalcount3'] = evacr_totalcount3
    request.session['evacr_totalcount4'] = evacr_totalcount4
    request.session['evacr_totalcount5'] = evacr_totalcount5
    request.session['evacr_totalcount6'] = evacr_totalcount6
    request.session['evacr_totalcount7'] = evacr_totalcount7
    request.session['evacr_totalcount8'] = evacr_totalcount8
    request.session['evacr_totalcount9'] = evacr_totalcount9
    request.session['evacr_totalcount10'] = evacr_totalcount10
    request.session['evacr_totalcount11'] = evacr_totalcount11
    request.session['evacr_totalcount12'] = evacr_totalcount12
    request.session['evacr_totalcount13'] = evacr_totalcount13
    request.session['end'] = end.days
    request.session['login31_count'] = login31_count
    request.session['loggedinuser'] = loggedinuser
    request.session['users_count'] = users_count
    request.session['real_dec_count'] = real_dec_count
    request.session['stdec_count'] = stdec_count
    request.session['stdec_participate_count'] = stdec_participate_count
    request.session['ave_solopt'] = ave_solopt
    request.session['ave_scrcr'] = ave_scrcr
    request.session['ave_evacr'] = ave_evacr
    request.session['ave_costs'] = ave_costs
    request.session['ave_decmade'] = ave_decmade
    request.session['users_lastweek_count'] = users_lastweek_count
    request.session['users_lastmonth_count'] = users_lastmonth_count
    request.session['scrcr1_count'] = scrcr1_count
    request.session['scrcr2_count'] = scrcr2_count
    request.session['scrcr3_count'] = scrcr3_count
    request.session['scrcr4_count'] = scrcr4_count
    request.session['scrcr5_count'] = scrcr5_count
    request.session['scrcr6_count'] = scrcr6_count
    request.session['scrcr7_count'] = scrcr7_count
    request.session['scrcr8_count'] = scrcr8_count
    request.session['dec_count'] = dec_count
    request.session['ave_stdec'] = ave_stdec
    request.session['stdec_real_participate_count'] = stdec_real_participate_count
    request.session['ave_real_solopt'] = ave_real_solopt
    request.session['ave_real_scrcr'] = ave_real_scrcr
    request.session['scrcr1_real_count'] = scrcr1_real_count
    request.session['scrcr2_real_count'] = scrcr2_real_count
    request.session['scrcr3_real_count'] = scrcr3_real_count
    request.session['scrcr4_real_count'] = scrcr4_real_count
    request.session['scrcr5_real_count'] = scrcr5_real_count
    request.session['scrcr6_real_count'] = scrcr6_real_count
    request.session['scrcr7_real_count'] = scrcr7_real_count
    request.session['scrcr8_real_count'] = scrcr8_real_count
    request.session['ave_real_evacr'] = ave_real_evacr
    request.session['evacr1_count'] = evacr1_count
    request.session['evacr2_count'] = evacr2_count
    request.session['evacr3_count'] = evacr3_count
    request.session['evacr4_count'] = evacr4_count
    request.session['evacr5_count'] = evacr5_count
    request.session['evacr6_count'] = evacr6_count
    request.session['evacr7_count'] = evacr7_count
    request.session['evacr8_count'] = evacr8_count
    request.session['evacr9_count'] = evacr9_count
    request.session['evacr10_count'] = evacr10_count
    request.session['evacr11_count'] = evacr11_count
    request.session['evacr12_count'] = evacr12_count
    request.session['evacr13_count'] = evacr13_count
    request.session['evacr14_count'] = evacr14_count
    request.session['evacr15_count'] = evacr15_count
    request.session['evacr16_count'] = evacr16_count
    request.session['evacr17_count'] = evacr17_count
    request.session['evacr18_count'] = evacr18_count
    request.session['evacr19_count'] = evacr19_count
    request.session['evacr20_count'] = evacr20_count
    request.session['evacr21_count'] = evacr21_count
    request.session['evacr22_count'] = evacr22_count
    request.session['evacr23_count'] = evacr23_count
    request.session['evacr24_count'] = evacr24_count
    request.session['evacr25_count'] = evacr25_count
    request.session['evacr26_count'] = evacr26_count
    request.session['evacr27_count'] = evacr27_count
    request.session['evacr28_count'] = evacr28_count
    request.session['evacr29_count'] = evacr29_count
    request.session['evacr30_count'] = evacr30_count
    request.session['evacr31_count'] = evacr31_count
    request.session['evacr32_count'] = evacr32_count
    request.session['evacr33_count'] = evacr33_count
    request.session['evacr34_count'] = evacr34_count
    request.session['evacr35_count'] = evacr35_count
    request.session['evacr36_count'] = evacr36_count
    request.session['evacr37_count'] = evacr37_count
    request.session['evacr38_count'] = evacr38_count
    request.session['evacr39_count'] = evacr39_count
    request.session['evacr40_count'] = evacr40_count
    request.session['evacr41_count'] = evacr41_count
    request.session['evacr42_count'] = evacr42_count
    request.session['evacr43_count'] = evacr43_count
    request.session['evacr44_count'] = evacr44_count
    request.session['evacr45_count'] = evacr45_count
    request.session['evacr46_count'] = evacr46_count
    request.session['evacr47_count'] = evacr47_count
    request.session['evacr48_count'] = evacr48_count
    request.session['evacr49_count'] = evacr49_count
    request.session['evacr50_count'] = evacr50_count
    request.session['evacr51_count'] = evacr51_count
    request.session['evacr52_count'] = evacr52_count
    request.session['evacr53_count'] = evacr53_count
    request.session['evacr54_count'] = evacr54_count
    request.session['evacr55_count'] = evacr55_count
    request.session['evacr56_count'] = evacr56_count
    request.session['evacr57_count'] = evacr57_count
    request.session['evacr58_count'] = evacr58_count
    request.session['evacr59_count'] = evacr59_count
    request.session['evacr60_count'] = evacr60_count
    request.session['evacr61_count'] = evacr61_count
    request.session['evacr62_count'] = evacr62_count
    request.session['evacr63_count'] = evacr63_count
    request.session['evacr64_count'] = evacr64_count
    request.session['evacr65_count'] = evacr65_count
    request.session['evacr66_count'] = evacr66_count
    request.session['evacr67_count'] = evacr67_count
    request.session['evacr68_count'] = evacr68_count
    request.session['evacr69_count'] = evacr69_count
    request.session['evacr70_count'] = evacr70_count
    request.session['evacr71_count'] = evacr71_count
    request.session['evacr72_count'] = evacr72_count
    request.session['evacr73_count'] = evacr73_count
    request.session['evacr74_count'] = evacr74_count
    request.session['evacr75_count'] = evacr75_count
    request.session['evacr76_count'] = evacr76_count
    request.session['evacr77_count'] = evacr77_count
    request.session['evacr78_count'] = evacr78_count
    request.session['evacr79_count'] = evacr79_count
    request.session['evacr80_count'] = evacr80_count
    request.session['evacr1_real_count'] = evacr1_real_count
    request.session['evacr2_real_count'] = evacr2_real_count
    request.session['evacr3_real_count'] = evacr3_real_count
    request.session['evacr4_real_count'] = evacr4_real_count
    request.session['evacr5_real_count'] = evacr5_real_count
    request.session['evacr6_real_count'] = evacr6_real_count
    request.session['evacr7_real_count'] = evacr7_real_count
    request.session['evacr8_real_count'] = evacr8_real_count
    request.session['evacr9_real_count'] = evacr9_real_count
    request.session['evacr10_real_count'] = evacr10_real_count
    request.session['evacr11_real_count'] = evacr11_real_count
    request.session['evacr12_real_count'] = evacr12_real_count
    request.session['evacr13_real_count'] = evacr13_real_count
    request.session['evacr14_real_count'] = evacr14_real_count
    request.session['evacr15_real_count'] = evacr15_real_count
    request.session['evacr16_real_count'] = evacr16_real_count
    request.session['evacr17_real_count'] = evacr17_real_count
    request.session['evacr18_real_count'] = evacr18_real_count
    request.session['evacr19_real_count'] = evacr19_real_count
    request.session['evacr20_real_count'] = evacr20_real_count
    request.session['evacr21_real_count'] = evacr21_real_count
    request.session['evacr22_real_count'] = evacr22_real_count
    request.session['evacr23_real_count'] = evacr23_real_count
    request.session['evacr24_real_count'] = evacr24_real_count
    request.session['evacr25_real_count'] = evacr25_real_count
    request.session['evacr26_real_count'] = evacr26_real_count
    request.session['evacr27_real_count'] = evacr27_real_count
    request.session['evacr28_real_count'] = evacr28_real_count
    request.session['evacr29_real_count'] = evacr29_real_count
    request.session['evacr30_real_count'] = evacr30_real_count
    request.session['evacr31_real_count'] = evacr31_real_count
    request.session['evacr32_real_count'] = evacr32_real_count
    request.session['evacr33_real_count'] = evacr33_real_count
    request.session['evacr34_real_count'] = evacr34_real_count
    request.session['evacr35_real_count'] = evacr35_real_count
    request.session['evacr36_real_count'] = evacr36_real_count
    request.session['evacr37_real_count'] = evacr37_real_count
    request.session['evacr38_real_count'] = evacr38_real_count
    request.session['evacr39_real_count'] = evacr39_real_count
    request.session['evacr40_real_count'] = evacr40_real_count
    request.session['evacr41_real_count'] = evacr41_real_count
    request.session['evacr42_real_count'] = evacr42_real_count
    request.session['evacr43_real_count'] = evacr43_real_count
    request.session['evacr44_real_count'] = evacr44_real_count
    request.session['evacr45_real_count'] = evacr45_real_count
    request.session['evacr46_real_count'] = evacr46_real_count
    request.session['evacr47_real_count'] = evacr47_real_count
    request.session['evacr48_real_count'] = evacr48_real_count
    request.session['evacr49_real_count'] = evacr49_real_count
    request.session['evacr50_real_count'] = evacr50_real_count
    request.session['evacr51_real_count'] = evacr51_real_count
    request.session['evacr52_real_count'] = evacr52_real_count
    request.session['evacr53_real_count'] = evacr53_real_count
    request.session['evacr54_real_count'] = evacr54_real_count
    request.session['evacr55_real_count'] = evacr55_real_count
    request.session['evacr56_real_count'] = evacr56_real_count
    request.session['evacr57_real_count'] = evacr57_real_count
    request.session['evacr58_real_count'] = evacr58_real_count
    request.session['evacr59_real_count'] = evacr59_real_count
    request.session['evacr60_real_count'] = evacr60_real_count
    request.session['evacr61_real_count'] = evacr61_real_count
    request.session['evacr62_real_count'] = evacr62_real_count
    request.session['evacr63_real_count'] = evacr63_real_count
    request.session['evacr64_real_count'] = evacr64_real_count
    request.session['evacr65_real_count'] = evacr65_real_count
    request.session['evacr66_real_count'] = evacr66_real_count
    request.session['evacr67_real_count'] = evacr67_real_count
    request.session['evacr68_real_count'] = evacr68_real_count
    request.session['evacr69_real_count'] = evacr69_real_count
    request.session['evacr70_real_count'] = evacr70_real_count
    request.session['evacr71_real_count'] = evacr71_real_count
    request.session['evacr72_real_count'] = evacr72_real_count
    request.session['evacr73_real_count'] = evacr73_real_count
    request.session['evacr74_real_count'] = evacr74_real_count
    request.session['evacr75_real_count'] = evacr75_real_count
    request.session['evacr76_real_count'] = evacr76_real_count
    request.session['evacr77_real_count'] = evacr77_real_count
    request.session['evacr78_real_count'] = evacr78_real_count
    request.session['evacr79_real_count'] = evacr79_real_count
    request.session['evacr80_real_count'] = evacr80_real_count
    request.session['ave_real_costs'] = ave_real_costs
    request.session['ave_real_decmade'] = ave_real_decmade
    request.session['ave_evm'] = ave_evm
    request.session['ave_real_evm'] = ave_real_evm
    request.session['ave_iw'] = ave_iw
    request.session['ave_real_iw'] = ave_real_iw
    request.session['perc_sol'] = perc_sol
    request.session['perc_scr'] = perc_scr
    request.session['perc_eva'] = perc_eva
    request.session['perc_iw'] = perc_iw
    request.session['perc_mapp'] = perc_mapp
    request.session['perc_dm'] = perc_dm
    request.session['perc_std'] = perc_std
    request.session['login_lastmonth_count'] = login_lastmonth_count
    request.session['login_lastweek_count'] = login_lastweek_count
    request.session['current_users_count'] = current_users_count
    request.session['dist_users_count'] = dist_users_count

    if loggedinuser == 'not found':
       return HttpResponseRedirect('/utility_tool/decisions/nouser_message.html')     
    else:
        return render(request, 'admin/usage-info.html',{'or1_count':or1_count,'or2_count':or2_count,'or3_count':or3_count,'or4_count':or4_count,'or5_count':or5_count,'or6_count':or6_count,'or7_count':or7_count,'or8_count':or8_count,'or9_count':or9_count,'or10_count':or10_count,'or11_count':or11_count,'or12_count':or12_count,'or13_count':or13_count,'or1_real_count':or1_real_count,'or2_real_count':or2_real_count,'or3_real_count':or3_real_count,'or4_real_count':or4_real_count,'or5_real_count':or5_real_count,'or6_real_count':or6_real_count,'or7_real_count':or7_real_count,'or8_real_count':or8_real_count,'or9_real_count':or9_real_count,'or10_real_count':or10_real_count,'or11_real_count':or11_real_count,'or12_real_count':or12_real_count,'or13_real_count':or13_real_count, 'tr_dec_count':tr_dec_count,'test_dec_count':test_dec_count,'dist_users_count':dist_users_count,'r_evacr_totalcount1':r_evacr_totalcount1, 'r_evacr_totalcount2':r_evacr_totalcount2, 'r_evacr_totalcount3':r_evacr_totalcount3, 'r_evacr_totalcount4':r_evacr_totalcount4, 'r_evacr_totalcount5':r_evacr_totalcount5, 'r_evacr_totalcount6':r_evacr_totalcount6, 'r_evacr_totalcount7':r_evacr_totalcount7, 'r_evacr_totalcount8':r_evacr_totalcount8, 'r_evacr_totalcount9':r_evacr_totalcount9, 'r_evacr_totalcount10':r_evacr_totalcount10, 'r_evacr_totalcount11':r_evacr_totalcount11, 'r_evacr_totalcount12':r_evacr_totalcount12, 'r_evacr_totalcount13':r_evacr_totalcount13,'evacr_totalcount1':evacr_totalcount1, 'evacr_totalcount2':evacr_totalcount2, 'evacr_totalcount3':evacr_totalcount3, 'evacr_totalcount4':evacr_totalcount4, 'evacr_totalcount5':evacr_totalcount5, 'evacr_totalcount6':evacr_totalcount6, 'evacr_totalcount7':evacr_totalcount7, 'evacr_totalcount8':evacr_totalcount8, 'evacr_totalcount9':evacr_totalcount9, 'evacr_totalcount10':evacr_totalcount10, 'evacr_totalcount11':evacr_totalcount11, 'evacr_totalcount12':evacr_totalcount12, 'evacr_totalcount13':evacr_totalcount13, 'end':end.days,'login31_count':login31_count, 'loggedinuser':loggedinuser, 'users_count':users_count,'real_dec_count':real_dec_count, 'stdec_count':stdec_count, 'stdec_participate_count':stdec_participate_count, 'ave_solopt':ave_solopt, 'ave_scrcr':ave_scrcr, 'ave_evacr':ave_evacr, 'ave_costs':ave_costs, 'ave_decmade':ave_decmade, 'users_lastweek_count':users_lastweek_count, 'users_lastmonth_count':users_lastmonth_count, 'scrcr1_count':scrcr1_count, 'scrcr2_count':scrcr2_count,'scrcr3_count':scrcr3_count,'scrcr4_count':scrcr4_count,'scrcr5_count':scrcr5_count,'scrcr6_count':scrcr6_count,'scrcr7_count':scrcr7_count,'scrcr8_count':scrcr8_count,'dec_count':dec_count, 'ave_stdec':ave_stdec, 'stdec_real_participate_count':stdec_real_participate_count, 'ave_real_solopt':ave_real_solopt, 'ave_real_scrcr':ave_real_scrcr,'scrcr1_real_count':scrcr1_real_count, 'scrcr2_real_count':scrcr2_real_count,'scrcr3_real_count':scrcr3_real_count,'scrcr4_real_count':scrcr4_real_count,'scrcr5_real_count':scrcr5_real_count,'scrcr6_real_count':scrcr6_real_count,'scrcr7_real_count':scrcr7_real_count,'scrcr8_real_count':scrcr8_real_count, 'ave_real_evacr':ave_real_evacr, 'evacr1_count':evacr1_count,'evacr2_count':evacr2_count,'evacr3_count':evacr3_count,'evacr4_count':evacr4_count,'evacr5_count':evacr5_count,'evacr6_count':evacr6_count,'evacr7_count':evacr7_count,'evacr8_count':evacr8_count,'evacr9_count':evacr9_count,'evacr10_count':evacr10_count,'evacr11_count':evacr11_count,'evacr12_count':evacr12_count,'evacr13_count':evacr13_count,'evacr14_count':evacr14_count,'evacr15_count':evacr15_count,'evacr16_count':evacr16_count,'evacr17_count':evacr17_count,'evacr18_count':evacr18_count,'evacr19_count':evacr19_count,'evacr20_count':evacr20_count,'evacr21_count':evacr21_count,'evacr22_count':evacr22_count,'evacr23_count':evacr23_count,'evacr24_count':evacr24_count,'evacr25_count':evacr25_count,'evacr26_count':evacr26_count,'evacr27_count':evacr27_count,'evacr28_count':evacr28_count,'evacr29_count':evacr29_count,'evacr30_count':evacr30_count,'evacr31_count':evacr31_count,'evacr32_count':evacr32_count,'evacr33_count':evacr33_count,'evacr34_count':evacr34_count,'evacr35_count':evacr35_count,'evacr36_count':evacr36_count,'evacr37_count':evacr37_count,'evacr38_count':evacr38_count,'evacr39_count':evacr39_count,'evacr40_count':evacr40_count,'evacr41_count':evacr41_count,'evacr42_count':evacr42_count,'evacr43_count':evacr43_count,'evacr44_count':evacr44_count,'evacr45_count':evacr45_count,'evacr46_count':evacr46_count,'evacr47_count':evacr47_count,'evacr48_count':evacr48_count,'evacr49_count':evacr49_count,'evacr50_count':evacr50_count,'evacr51_count':evacr51_count,'evacr52_count':evacr52_count,'evacr53_count':evacr53_count,'evacr54_count':evacr54_count,'evacr55_count':evacr55_count,'evacr56_count':evacr56_count, 'evacr57_count':evacr57_count,'evacr58_count':evacr58_count,'evacr59_count':evacr59_count,'evacr60_count':evacr60_count,'evacr61_count':evacr61_count,'evacr62_count':evacr62_count,'evacr63_count':evacr63_count,'evacr64_count':evacr64_count,'evacr65_count':evacr65_count,'evacr66_count':evacr66_count,'evacr67_count':evacr67_count,'evacr68_count':evacr68_count,'evacr69_count':evacr69_count,'evacr70_count':evacr70_count,'evacr71_count':evacr71_count,'evacr72_count':evacr72_count,'evacr73_count':evacr73_count,'evacr74_count':evacr74_count,'evacr75_count':evacr75_count,'evacr76_count':evacr76_count,'evacr77_count':evacr77_count,'evacr78_count':evacr78_count,'evacr79_count':evacr79_count,'evacr80_count':evacr80_count, 'evacr1_real_count':evacr1_real_count,'evacr2_real_count':evacr2_real_count,'evacr3_real_count':evacr3_real_count,'evacr4_real_count':evacr4_real_count,'evacr5_real_count':evacr5_real_count,'evacr6_real_count':evacr6_real_count,'evacr7_real_count':evacr7_real_count,'evacr8_real_count':evacr8_real_count,'evacr9_real_count':evacr9_real_count,'evacr10_real_count':evacr10_real_count,'evacr11_real_count':evacr11_real_count,'evacr12_real_count':evacr12_real_count,'evacr13_real_count':evacr13_real_count,'evacr14_real_count':evacr14_real_count,'evacr15_real_count':evacr15_real_count,'evacr16_real_count':evacr16_real_count,'evacr17_real_count':evacr17_real_count,'evacr18_real_count':evacr18_real_count,'evacr19_real_count':evacr19_real_count,'evacr20_real_count':evacr20_real_count,'evacr21_real_count':evacr21_real_count,'evacr22_real_count':evacr22_real_count,'evacr23_real_count':evacr23_real_count,'evacr24_real_count':evacr24_real_count,'evacr25_real_count':evacr25_real_count,'evacr26_real_count':evacr26_real_count,'evacr27_real_count':evacr27_real_count,'evacr28_real_count':evacr28_real_count,'evacr29_real_count':evacr29_real_count,'evacr30_real_count':evacr30_real_count,'evacr31_real_count':evacr31_real_count,'evacr32_real_count':evacr32_real_count,'evacr33_real_count':evacr33_real_count,'evacr34_real_count':evacr34_real_count,'evacr35_real_count':evacr35_real_count,'evacr36_real_count':evacr36_real_count,'evacr37_real_count':evacr37_real_count,'evacr38_real_count':evacr38_real_count,'evacr39_real_count':evacr39_real_count,'evacr40_real_count':evacr40_real_count,'evacr41_real_count':evacr41_real_count,'evacr42_real_count':evacr42_real_count,'evacr43_real_count':evacr43_real_count,'evacr44_real_count':evacr44_real_count,'evacr45_real_count':evacr45_real_count,'evacr46_real_count':evacr46_real_count,'evacr47_real_count':evacr47_real_count,'evacr48_real_count':evacr48_real_count,'evacr49_real_count':evacr49_real_count,'evacr50_real_count':evacr50_real_count,'evacr51_real_count':evacr51_real_count,'evacr52_real_count':evacr52_real_count,'evacr53_real_count':evacr53_real_count,'evacr54_real_count':evacr54_real_count,'evacr55_real_count':evacr55_real_count,'evacr56_real_count':evacr56_real_count, 'evacr57_real_count':evacr57_real_count,'evacr58_real_count':evacr58_real_count,'evacr59_real_count':evacr59_real_count,'evacr60_real_count':evacr60_real_count,'evacr61_real_count':evacr61_real_count,'evacr62_real_count':evacr62_real_count,'evacr63_real_count':evacr63_real_count,'evacr64_real_count':evacr64_real_count,'evacr65_real_count':evacr65_real_count,'evacr66_real_count':evacr66_real_count,'evacr67_real_count':evacr67_real_count,'evacr68_real_count':evacr68_real_count,'evacr69_real_count':evacr69_real_count,'evacr70_real_count':evacr70_real_count,'evacr71_real_count':evacr71_real_count,'evacr72_real_count':evacr72_real_count,'evacr73_real_count':evacr73_real_count,'evacr74_real_count':evacr74_real_count,'evacr75_real_count':evacr75_real_count,'evacr76_real_count':evacr76_real_count,'evacr77_real_count':evacr77_real_count,'evacr78_real_count':evacr78_real_count,'evacr79_real_count':evacr79_real_count,'evacr80_real_count':evacr80_real_count, 'ave_real_costs':ave_real_costs, 'ave_real_decmade':ave_real_decmade, 'ave_evm':ave_evm,'ave_real_evm':ave_real_evm, 'ave_iw':ave_iw,'ave_real_iw':ave_real_iw, 'perc_sol':perc_sol, 'perc_scr':perc_scr, 'perc_eva':perc_eva, 'perc_iw':perc_iw, 'perc_mapp':perc_mapp, 'perc_dm':perc_dm, 'perc_std':perc_std, 'login_lastmonth_count':login_lastmonth_count, 'login_lastweek_count':login_lastweek_count, 'current_users_count':current_users_count })

def export_info(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename=usage_info.xls'      
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet("Usage Information")
    row_num = 0  
    col_width = 256 * 75 
    ws.col(0).width = col_width
    #Heading of tables
    a = xlwt.Alignment()
    a.wrap = True 
    a.vert = a.VERT_CENTER
    a.horz = a.HORZ_CENTER
    font_style = xlwt.XFStyle()
    font_style.font.bold = True 
    font_style.alignment = a
    pattern = xlwt.Pattern()
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = 22 
    pattern2 = xlwt.Pattern()
    pattern2.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern2.pattern_fore_colour = 22 
    font_style.pattern = pattern2
    pattern3 = xlwt.Pattern()
    pattern3.pattern_fore_colour = 1
    aR = xlwt.Alignment()     
    aR.horz = a.HORZ_RIGHT
    aR.wrap = True
    aL = xlwt.Alignment()     
    aL.horz = a.HORZ_LEFT
    aL.wrap = True 
    font_style2 = xlwt.XFStyle()                                                                                                                                                                                 
    font_style2.pattern = pattern3                                                                                                                                                                               
    font_style2.alignment = aL
    font_style4 = xlwt.XFStyle()                                                                                                                                                                                 
    font_style4.pattern = pattern3
    font_style4.alignment = aR 
    num_style = xlwt.XFStyle()
    num_style.num_format_str = '0.0\\%'
    row_num = 0
    col_num = 0
    ws.write(row_num, col_num, "Usage data", font_style)
    row_num += 1
    ws.write(row_num, col_num, "Number of days since official launch (Aug 31, 2019)", font_style2)
    col_num += 1   
    ws.write(row_num, col_num, request.session['end'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Number of days from Aug 31 2019 on which there is any user activity at all in DM", font_style2)
    col_num += 1   
    ws.write(row_num, col_num, request.session['login31_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Total number of people ever registered", font_style2) 
    col_num += 1                                                                                                                                                                                                 
    ws.write(row_num, col_num, request.session['users_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Total number of people currently registered", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['current_users_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Number of unique users who have logged into the tool over the last week", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['users_lastweek_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Number of unique users who have logged into the tool over the last month", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['users_lastmonth_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Number of logins over the last week", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['login_lastweek_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Number of logins over the last month", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['login_lastmonth_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Total number of decisions created based on ID creation", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['dec_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Total number of real decisions created based on ID creation", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['real_dec_count'] , font_style4)
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "", font_style4)
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "Completion of decision steps", font_style)
    row_num += 1
    ws.write(row_num, col_num, "For Decision Problem step, what % of decision IDs has been touched?", font_style2)
    col_num += 1 
    ws.write(row_num, col_num, "100.0%" , font_style4)
    row_num += 1 
    col_num = 0 
    ws.write(row_num, col_num, "For Identify Stakeholders step, what % of decision IDs has been touched?", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['perc_std'] , num_style)
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "For Solution Options step, what % of decision IDs has been touched?", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['perc_sol'] , num_style)  
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "For Screening Criteria step, what % of decision IDs has been touched?", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['perc_scr'] , num_style)  
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "For Screen Solution Options step, what % of decision IDs has been touched?", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['perc_mapp'] , num_style)  
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "For Evaluation Criteria step, what % of decision IDs has been touched?", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['perc_eva'] , num_style)  
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "For Importance Scores step, what % of decision IDs has been touched?", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['perc_iw'] , num_style)  
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "For Evaluation Table step, what % of decision IDs has been touched?", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_evm'] , num_style)  
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "For Costs step, what % of decision IDs has been touched?", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_costs'] , num_style)  
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "For Make a Decision step, what % of decision IDs has been touched?", font_style2) 
    col_num += 1 
    ws.write(row_num, col_num, request.session['perc_dm'] , num_style)  
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "", font_style4)
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "Stakeholders", font_style)
    row_num += 1  
    ws.write(row_num, col_num, "Total number of stakeholders across all directories", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['stdec_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Average number of stakeholders across all directories", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_stdec'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Total # of decisions that invite s/h to participate in at least one of the four activities in assign tasks", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['stdec_participate_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Total # of real decisions that invite s/h to participate in at least one of the four activities in assign tasks", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['stdec_real_participate_count'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "", font_style4)
    row_num += 1 
    col_num = 0  
    ws.write(row_num, col_num, "Solution Options", font_style)
    row_num += 1
    ws.write(row_num, col_num, "Average number of Solution Options per decision", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_solopt'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "Average number of Solution Options per real decision", font_style2)   
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_real_solopt'] , font_style4)
    row_num += 1 
    col_num = 0
    ws.write(row_num, col_num, "", font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Screening Criteria", font_style)
    row_num += 1
    ws.write(row_num, col_num, "Average number of Screening Criteria per decision", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_scrcr'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Average number of Screening Criteria per real decision", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_real_scrcr'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this screening criterion - Fits within available budget", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr1_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this screening criterion - Can be implemented by date required", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr2_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this screening criterion - Meets privacy standards", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr3_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this screening criterion - Evidence of effectiveness exists", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr4_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this screening criterion - Fits within school schedule", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr5_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this screening criterion - Meets content requirements/ learning objectives", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr6_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this screening criterion - Meets state code and/or other regulations", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr7_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this screening criterion - Serves target population (grade level, ESL etc.)", font_style2)   
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr8_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this screening criterion - Fits within available budget", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr1_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this screening criterion - Can be implemented by date required", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr2_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this screening criterion - Meets privacy standards", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr3_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this screening criterion - Evidence of effectiveness exists", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr4_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this screening criterion - Fits within school schedule", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr5_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this screening criterion - Meets content requirements/ learning objectives", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr6_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this screening criterion - Meets state code and/or other regulations", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr7_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this screening criterion - Serves target population (grade level, ESL etc.)", font_style2)   
    col_num += 1 
    ws.write(row_num, col_num, request.session['scrcr8_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "", font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Evaluation Criteria", font_style)
    row_num += 1
    ws.write(row_num, col_num, "Average number of Evaluation Criteria per decision", font_style2)   
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_evacr'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Average number of Evaluation Criteria per real decision", font_style2)   
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_real_evacr'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Addresses the identified need", font_style)   
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr_totalcount1'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Content meets learning objectives", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr1_count'] , font_style4) 
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of students in need who can be served", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr2_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Equity", font_style)      
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr_totalcount2'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Accessible to target population", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr3_count'] , font_style4) 
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Distribution of resources across population to be served", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr4_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of students participating", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr5_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Serves historically underserved groups", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr6_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - External recommendations", font_style)      
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr_totalcount3'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Recommendations from external experts", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr7_count'] , font_style4)  
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Recommendations from external peers", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr8_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Feasibility of implementation", font_style)  
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr_totalcount4'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Accessibility of physical location", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr9_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Amount of change in personnel resource requirements (e.g., time and number of staff) from the current status", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr10_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Amount of technical support needed", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr11_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Amount of training/PD needed", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr12_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Availability of necessary personnel, facilities, materials and equipment", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr13_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Availability of technical support to support implementers", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr14_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Availability of training/PD to support implementers", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr15_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Can be implemented in desired timeline", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr16_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Capacity/skill level of current teachers/staff to implement option with fidelity", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr17_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Compatibility with existing systems", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr18_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Ease of use/ User friendliness", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr19_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Effect on teacher/staff workload", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr20_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Evidence of successful implementation in similar schools/districts/states", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr21_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Financial sustainability over time", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr22_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Fit with school calendar/schedule", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr23_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Likelihood this option will continue to be implemented with fidelity over time", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr24_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Scalability", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr25_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Solid plan proposed for financing", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr26_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Solid plan proposed for implementation", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr27_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Fit with local context", font_style)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr_totalcount5'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Alignment to state/district/school mission and/or vision", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr28_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Alignment with current school/district/state curriculum", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr29_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Alignment with current state/district/school priorities", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr30_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Appropriate for student/staff demographics", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr31_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Customizability of solution to local needs", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr32_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Fit with local cultural values", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr33_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on central control", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr34_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on local autonomy", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr35_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Political value", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr36_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Viable in current political context", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr37_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Impact on parent engagement", font_style)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr_totalcount6'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of times parents call school", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr38_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of hours parents help children with homework", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr39_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Encourages parents to show up to PTA meetings", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr40_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of hours parents read to their child", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr41_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Impact on student academic performance", font_style)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr_totalcount7'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on standardized test scores", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr42_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on student grades", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr43_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on college admission", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr44_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on course completion", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr45_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on graduation", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr46_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on progression to higher grade", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr47_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on progression towards graduation", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr48_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on closing the achievement gap", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr49_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Impact on student socio-emotional development", font_style)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr_totalcount8'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on student sense of belongingness", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr50_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on incidence of misbehavior", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr51_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on school climate", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr52_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on suspensions", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr53_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Impact on student/staff engagement", font_style)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr_totalcount9'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on attendance", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr54_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on staff absenteeism", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr55_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on staff/teacher/student effort", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr56_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on student-teacher interaction", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr57_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Improves teacher performance", font_style)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr_totalcount10'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves teacher time-on-task", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr58_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves student time-on-task", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr59_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves teacher value-added", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr60_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves teacher pedogogical skills", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr61_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves teacher content knowledge", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr62_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Meets required standards and regulations", font_style)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr_totalcount11'] , font_style4)
    row_num += 1
    col_num = 0  
    ws.write(row_num, col_num, "Addresses safety concerns", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr63_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Alignment with state code and other regulations", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr64_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Compliance with regulations", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr65_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Meets privacy standards", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr66_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Content maintains level of rigor/quality of learning", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr67_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Meets state/district/school standards", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr68_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Quality of implementation (for programs/strategies/tools already in place)", font_style)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr_totalcount12'] , font_style4)
    row_num += 1
    col_num = 0 
    ws.write(row_num, col_num, "Consistency of implementation across sites", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr69_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Fidelity of implementation", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr70_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Level of monitoring conducted", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr71_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Quality of technical support provided to implementers", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr72_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Quality of training delivered", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr73_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Change in teacher pedagogy", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr74_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of decisions which selects this evaluation criterion - Support from stakeholders", font_style)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr_totalcount13'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Board preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr75_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Community preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr76_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Parent preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr77_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Student preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr78_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Teacher preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr79_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Union preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr80_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Average number of Evaluation Criteria per real decision", font_style2)   
    col_num += 1 
    ws.write(row_num, col_num, request.session['ave_real_evacr'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Addresses the identified need", font_style)   
    col_num += 1 
    ws.write(row_num, col_num, request.session['r_evacr_totalcount1'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Content meets learning objectives", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr1_real_count'] , font_style4) 
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of students in need who can be served", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr2_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Equity", font_style)      
    col_num += 1 
    ws.write(row_num, col_num, request.session['r_evacr_totalcount2'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Accessible to target population", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr3_real_count'] , font_style4) 
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Distribution of resources across population to be served", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr4_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of students participating", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr5_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Serves historically underserved groups", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr6_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - External recommendations", font_style)      
    col_num += 1 
    ws.write(row_num, col_num, request.session['r_evacr_totalcount3'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Recommendations from external experts", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr7_real_count'] , font_style4)  
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Recommendations from external peers", font_style2)  
    col_num += 1 
    ws.write(row_num, col_num, request.session['evacr8_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Feasibility of implementation", font_style)  
    col_num += 1  
    ws.write(row_num, col_num, request.session['r_evacr_totalcount4'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Accessibility of physical location", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr9_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Amount of change in personnel resource requirements (e.g., time and number of staff) from the current status", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr10_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Amount of technical support needed", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr11_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Amount of training/PD needed", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr12_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Availability of necessary personnel, facilities, materials and equipment", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr13_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Availability of technical support to support implementers", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr14_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Availability of training/PD to support implementers", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr15_real_count'] , font_style4)                                                                                                                              
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Can be implemented in desired timeline", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr16_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Capacity/skill level of current teachers/staff to implement option with fidelity", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr17_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Compatibility with existing systems", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr18_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Ease of use/ User friendliness", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr19_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Effect on teacher/staff workload", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr20_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Evidence of successful implementation in similar schools/districts/states", font_style2)   
    col_num += 1  
    ws.write(row_num, col_num, request.session['evacr21_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Financial sustainability over time", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr22_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Fit with school calendar/schedule", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr23_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Likelihood this option will continue to be implemented with fidelity over time", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr24_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Scalability", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr25_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Solid plan proposed for financing", font_style2)                                                                                                                                 
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr26_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Solid plan proposed for implementation", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr27_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Fit with local context", font_style)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['r_evacr_totalcount5'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Alignment to state/district/school mission and/or vision", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr28_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Alignment with current school/district/state curriculum", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr29_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Alignment with current state/district/school priorities", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr30_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Appropriate for student/staff demographics", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr31_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Customizability of solution to local needs", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr32_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Fit with local cultural values", font_style2)   
    col_num += 1   
    ws.write(row_num, col_num, request.session['evacr33_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on central control", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr34_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on local autonomy", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr35_real_count'] , font_style4)
    row_num += 1                                                                                                                                                                                                 
    col_num = 0
    ws.write(row_num, col_num, "Political value", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr36_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Viable in current political context", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr37_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Impact on parent engagement", font_style)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['r_evacr_totalcount6'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of times parents call school", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr38_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of hours parents help children with homework", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr39_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Encourages parents to show up to PTA meetings", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr40_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of hours parents read to their child", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr41_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Impact on student academic performance", font_style)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['r_evacr_totalcount7'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on standardized test scores", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr42_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on student grades", font_style2)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['evacr43_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on college admission", font_style2)   
    col_num += 1                                                                                                                                                                                                 
    ws.write(row_num, col_num, request.session['evacr44_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on course completion", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr45_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on graduation", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr46_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on progression to higher grade", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr47_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on progression towards graduation", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr48_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on closing the achievement gap", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr49_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Impact on student socio-emotional development", font_style)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['r_evacr_totalcount8'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on student sense of belongingness", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr50_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on incidence of misbehavior", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr51_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on school climate", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr52_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on suspensions", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr53_real_count'] , font_style4)
    row_num += 1
    col_num = 0                                                                                                                                                                                                  
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Impact on student/staff engagement", font_style)   
    col_num += 1    
    ws.write(row_num, col_num, request.session['r_evacr_totalcount9'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on attendance", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr54_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on staff absenteeism", font_style2)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['evacr55_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on staff/teacher/student effort", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr56_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Impact on student-teacher interaction", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr57_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Improves teacher performance", font_style)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['r_evacr_totalcount10'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves teacher time-on-task", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr58_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves student time-on-task", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr59_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves teacher value-added", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr60_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves teacher pedogogical skills", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr61_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Improves teacher content knowledge", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr62_real_count'] , font_style4)                                                                                                                              
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Meets required standards and regulations", font_style) 
    col_num += 1      
    ws.write(row_num, col_num, request.session['r_evacr_totalcount11'] , font_style4)
    row_num += 1
    col_num = 0  
    ws.write(row_num, col_num, "Addresses safety concerns", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr63_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Alignment with state code and other regulations", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr64_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Compliance with regulations", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr65_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Meets privacy standards", font_style2)   
    col_num += 1      
    ws.write(row_num, col_num, request.session['evacr66_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Content maintains level of rigor/quality of learning", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr67_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Meets state/district/school standards", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr68_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Quality of implementation (for programs/strategies/tools already in place)", font_style)   
    col_num += 1     
    ws.write(row_num, col_num, request.session['r_evacr_totalcount12'] , font_style4)
    row_num += 1
    col_num = 0 
    ws.write(row_num, col_num, "Consistency of implementation across sites", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr69_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Fidelity of implementation", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr70_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Level of monitoring conducted", font_style2)                                                                                                                                     
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr71_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Quality of technical support provided to implementers", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr72_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Quality of training delivered", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr73_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Change in teacher pedagogy", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr74_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Number of real decisions which selects this evaluation criterion - Support from stakeholders", font_style)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['r_evacr_totalcount13'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Board preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr75_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Community preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr76_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Parent preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr77_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Student preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr78_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Teacher preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr79_real_count'] , font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Union preference/buy-in/support", font_style2)   
    col_num += 1       
    ws.write(row_num, col_num, request.session['evacr80_real_count'] , font_style4)
    row_num += 1                                                                                                                                                                                                 
    col_num = 0
    ws.write(row_num, col_num, "", font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Importance Scores", font_style)   
    row_num += 1        
    ws.write(row_num, col_num, "In what % of decisions does the PA change the votes from the default of 10 per person?", font_style2)   
    col_num += 1        
    ws.write(row_num, col_num, request.session['ave_iw'] , num_style)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "In what % of real decisions does the PA change the votes from the default of 10 per person?", font_style2)   
    col_num += 1        
    ws.write(row_num, col_num, request.session['ave_real_iw'] , num_style)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "", font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Evaluation Measures Table and Evaluation Data Tables", font_style)   
    row_num += 1         
    ws.write(row_num, col_num, "In what % of decisions created has the PA filled out the Evaluation Data Table?", font_style2)   
    col_num += 1         
    ws.write(row_num, col_num, request.session['ave_evm'] , num_style)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "In what % of real decisions created has the PA filled out the Evaluation Data Table?", font_style2)   
    col_num += 1         
    ws.write(row_num, col_num, request.session['ave_real_evm'] , num_style)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "", font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Costs", font_style)   
    row_num += 1        
    ws.write(row_num, col_num, "In what % of decisions created have the cost table been filled out?", font_style2)   
    col_num += 1         
    ws.write(row_num, col_num, request.session['ave_costs'] , num_style)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "In what % of real decisions created have the cost table been filled out?", font_style2)   
    col_num += 1         
    ws.write(row_num, col_num, request.session['ave_real_costs'] , num_style)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "", font_style4)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "Make a Decision", font_style)   
    row_num += 1
    ws.write(row_num, col_num, "In what % of decisions does the PA choose at least one solution option for the question Which option(s) did you choose?", font_style2)   
    col_num += 1        
    ws.write(row_num, col_num, request.session['ave_decmade'] , num_style)
    row_num += 1
    col_num = 0
    ws.write(row_num, col_num, "In what % of real decisions does the PA choose at least one solution option for the question Which option(s) did you choose?", font_style2)   
    col_num += 1        
    ws.write(row_num, col_num, request.session['ave_real_decmade'] , num_style)
    row_num += 1
    col_num = 0

    wb.save(response)
    return response

def summary_report(request, dec_id):
    '''if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0''' 
    dec_id2 = dec_id
    dec_id = int(dec_id)
    if 'user_email' in request.session: 
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    '''if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'''
    # get all the data required for the report
    try: 
       dec = Decisions.objects.get(id=dec_id)
       name_decisionmaker = dec.name_decisionmaker
       title = dec.title
       decision_prob = dec.decision_prob
       evidence = dec.evidence
       goal = dec.goal
       target_audience = dec.target_audience
       stakeholders = dec.stakeholders
       participating_stakeholders = dec.participating_stakeholders
       by_when = dec.by_when
       created_by = dec.created_by
       real_dec_yn = dec.real_dec_yn
    except:
        name_decisionmaker = ''  
        title = ''
        decision_prob = ''
        evidence = ''
        goal = ''
        target_audience = ''
        stakeholders = ''
        participating_stakeholders = ''
        by_when = ''
        real_dec_yn = ''

    try:
       std = Stakeholders_Decisions.objects.filter(dec_id=dec_id)
       stdec_count = std.exclude(email = user_email).count()
    except:
       stdec_count = 0
    
    text = ""
    sugg_solopt = 0
    sugg_scr = 0
    sugg_evacr = 0
    sugg_iw = 0 
    st_solopt_contrib_count = 0
    st_scr_contrib_count = 0
    st_evacr_contrib_count = 0
    st_iw_contrib_count = 0
    try:
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id, solopt_type = 'Y') 
       std_count = std.exclude(email = user_email).count()    
       sugg_solopt = std_count 
       if std_count > 0: 
          text = "\n suggesting Solution Options"
       for ss in std:
           try:
              user_solopt = Users.objects.get(email = ss.email)
              try:
                 solopt_user = Solution_Options.objects.filter(dec_id=dec_id, created_by = user_solopt.user)
                 print 'I AM HERE'
                 solopt_user_count = solopt_user.count()   
                 print solopt_user_count
                 if solopt_user_count > 0:
                    st_solopt_contrib_count = st_solopt_contrib_count + 1
              except:
                 print 'something wrong with the solopt user count'
           except ObjectDoesNotExist:
              print 'user does not exist'
    except ObjectDoesNotExist:
       text= text + ""
    except MultipleObjectsReturned:                                                                                                           
       text = "\n suggesting Solution Options"   
       #sugg_solopt = "Y"

    try:
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id, scrcr_type = 'Y') 
       std_count = std.exclude(email = user_email).count()     
       sugg_scr = std_count
       if std_count > 0:
          if text <> "": 
             text = text + ",\n providing Screening Criteria"
          else:
             text = text + "\n providing Screening Criteria"
       for ss in std:
           try:
              user_scr = Users.objects.get(email = ss.email)
              try:
                 scr_user = Screening_Criteria.objects.filter(dec_id=dec_id, created_by = user_scr.user)
                 scr_user_count = scr_user.count()   
                 if scr_user_count > 0:
                    st_scr_contrib_count = st_scr_contrib_count + 1
              except:
                 print 'something wrong with the scr user count'
           except ObjectDoesNotExist:
              print 'user does not exist'
    except ObjectDoesNotExist:
       text= text + ""
    except MultipleObjectsReturned:                                                                                                           
          if text <> "": 
             text = text + ",\n providing Screening Criteria"
          else:
             text = text + "\n providing Screening Criteria"

    try:
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id, evacr_type = 'Y') 
       std_count = std.exclude(email = user_email).count()     
       sugg_evacr = std_count
       if std_count > 0: 
          if text <> "": 
             text = text + ",\n developing a list of Evaluation Criteria"
          else:
             text = text + "\n developing a list of Evaluation Criteria"
       for ss in std:
           try:
              user_eva = Users.objects.get(email = ss.email)
              try:
                 eva_user = Evaluation_Criteria.objects.filter(dec_id=dec_id, created_by = user_eva.user)
                 eva_user_count = eva_user.count()   
                 if eva_user_count > 0:
                    st_evacr_contrib_count = st_evacr_contrib_count + 1
              except:
                 print 'something wrong with the evacr user count'
           except ObjectDoesNotExist:
              print 'user does not exist'
    except ObjectDoesNotExist:
       text= text + ""
    except MultipleObjectsReturned:                       
          if text <> "": 
             text = text + ",\n developing a list of Evaluation Criteria"
          else:
             text = text + "\n developing a list of Evaluation Criteria"   

    stakeholdersNow = '' 
    try:
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id, iw_type = 'Y') 
       std1 = std.exclude(email = user_email)
       std_count = std1.count()     
       
       sugg_iw = std_count
       if std_count > 0: 
          stakeholdersNow = 'Y'      
          if text <> "": 
             text = text + ",\n contributing Importance Scores"
          else:
             text = text + "\n contributing Importance Scores"
       else:
          stakeholdersNow = 'N'     
       for ss in std1:
           print 'ss.email'
           print ss.email
           try:
              user_iw = Users.objects.get(email = ss.email)
              try:
                 iw_user = Importance_Scores.objects.filter(dec_id=dec_id, created_by = user_iw.user)
                 iw_user_count = iw_user.count()   
                 if iw_user_count > 0:
                    st_iw_contrib_count = st_iw_contrib_count + 1
              except:
                 print 'something wrong with the iw user count'
           except ObjectDoesNotExist:
              print 'user does not exist' 
    except ObjectDoesNotExist:
       text= text + ""
       stakeholdersNow = 'N'       
    except MultipleObjectsReturned:                                                                                                          
          stakeholdersNow = 'Y'      
          if text <> "": 
             text = text + ",\n contributing Importance Scores"
          else:
             text = text + "\n contributing Importance Scores"  

    try:
       solopt1 = Solution_Options.objects.filter(dec_id=dec_id)
       #total_solopt_count = solopt1.count()
       solopt_notdeleted = solopt1.filter(deleted='N')
       total_solopt_count = solopt_notdeleted.count()
       solopt = solopt1.filter(archived='N', deleted='N')
       solopt_count = solopt.count()
       solopt_archived = solopt1.filter(archived='Y')
       total_archived = solopt_archived.count()
    except:
       print 'no solution options'

    scrcr_count = 0
    try:
       scrcr = Screening_Criteria.objects.filter(dec_id=dec_id)
       scrcr_count = scrcr.count()
    except:
       print 'no screening criteria'
    
    mapp_count = 0
    try:
       mapp = MappingTable.objects.filter(dec_id=dec_id)
       mapp_count = mapp.count()
    except:
       print 'no mapping table'

    evacr_count = 0
    try:
       evacr = Evaluation_Criteria.objects.filter(dec_id=dec_id)
       evacr_notdeleted = evacr.exclude(deleted = 'Y')  
       evacr_count = evacr_notdeleted.count()
    except:
       print 'no evaluation criteria'

    iw_count = 0
    try:
       iw = Importance_Scores.objects.filter(dec_id=dec_id).order_by('criterion', '-score') 
       iw_count = iw.count()
    except:
       print 'no iw'

    try:
       adj_w = evacr_notdeleted.order_by('criterion', '-adjusted_weight') 
    except:
       print 'no adj_w'

    evm_count = 0
    try:
       evm = Evaluation_Measures.objects.filter(dec_id=dec_id)
       evm_forever = evm.exclude(archived = 'Y').exclude(deleted = 'Y')
       evm2 = evm.exclude(archived = 'Y').exclude(deleted = 'Y').order_by('sol_option') 
       onerec_evm = evm2.first() 
       onerec_cri = onerec_evm.criterion
       onerec_opt = onerec_evm.sol_option
       evm3 = evm.exclude(archived = 'Y').exclude(deleted = 'Y').order_by('criterion')
       evm_distinct = evm3.values_list('criterion','measure','unit','lowest_value','highest_value','higher_better').distinct()
       # Had to do a lot of tinkering to make it work - like not using order by if you need a values_list distinct afterwards - some things don't really make sense
       evm_distinct_cri = evm_forever.values_list('criterion').distinct()
       evm_distinct_opt = evm2.values_list('sol_option').distinct()
       evm_opt_count = evm_distinct_opt.count()
       print 'evm_opt_count'
       print evm_opt_count
       evm_count = evm.count()
    except:
       print 'no evm'

    for e in evm2:
        print 'TTTY'
        print e.criterion
        print e.sol_option
    det_count = 0
    try:
       det = Detailed_Costs.objects.filter(dec_id=dec_id)
       for d in det:
           if d.personnel_cost <> 0 or d.facilities_cost <> 0 or d.materials_cost <> 0 or d.training_cost <> 0 or d.other_cost <> 0:
              det_count = det_count + 1
    except:
       print 'no det'

    type_of_cost = ''
    cost_text = ''  
    costzero = ''
    so1_one = ''
    wu_w = '' 
    try: 
       c = Cost_Setup.objects.get(dec_id = dec_id) 
       type_of_cost = c.type_of_cost
       if type_of_cost == "Total":
          cost_text = "Total Cost"                                                                                                                                                                               
          type_of_cost = 'Total'
       elif type_of_cost == "Avg":
          cost_text = "Average Cost"
          type_of_cost = 'Average'
       else:
          cost_text = "Marginal Cost"
          type_of_cost = 'Marginal'
    except ObjectDoesNotExist: 
       print 'no cost setup'

    wu = ''
    so1 = ''
    co = ''
    so2 = ''
    cur = ''
    so3 = ''
    cu_rec_exists = ''
    cu_rec_one = ''
    lowest_cost = ''
    highest_cost = ''
    cost_utility_exists = 'N'
    so2_one = ''                                                                                                                                      
    so3_one = ''
    co_w = ''
    cur_w = ''
    try:
       cost_utility1 = Cost_Utility.objects.filter(dec_id = dec_id)
       cost_utility = cost_utility1.exclude(archived = 'Y')
       cu_count = cost_utility.count()
       cc = cost_utility.order_by('-weighted_utility')
       if cu_count > 1:
          cu_rec_one = 'N'
       else:
          cu_rec_one = 'Y' 
       onerec = cc.first()
       if not onerec:
           costzero = 'Y'
           cu_rec_exists = 'N'
       else:    
          cu_rec_exists = 'Y' 
          wu = onerec.weighted_utility 
          if onerec.cost <= 0:                                                                                                               
             costzero = 'Y'
          else:   
             costzero = 'N' 
       print 'aug3408'
       print costzero
       if wu is not None and  wu <> '' and wu <> ' ':
          wu_w = "%.1f" % onerec.weighted_utility
       so1 = ""
       so1_one = "Y"
       if wu == 0:
          wuzero = 'Y'
       else:                                                                                                                                                                                                        
          wuzero = 'N'
       for c in cc:
           print c.cost
           if c.cost <> 0:                                                                                                                                                                                          
              costzero = 'N'
           if c.weighted_utility == 0:                                                                                                                                                                              
              wuzero = 'Y'
           if wu == c.weighted_utility and (c.weighted_utility is not None and c.weighted_utility <> '' and c.weighted_utility <> ' '):
              if so1 <> "":
                 so1 = so1 + ", " + c.sol_option
                 so1_one = "N"
              else:
                 so1 = so1 + " " + c.sol_option 
       if costzero == 'N': 
          cd = cost_utility.exclude(cost = 0).order_by('cost')
          onerec2 = cd.first()                                                                                                                                                                                     
          co_w = "%.1f" % onerec2.cost
          co = onerec2.cost
          lowest_cost = onerec2.sol_option
          onerec_last = cd.last()
          highest_cost = onerec_last.sol_option
          so2 = ""
          so2_one = 'Y'
          for c in cd:
             if co == c.cost:
                if so2 <> "":
                   so2 = so2 + ", " + c.sol_option
                   so2_one = 'N'
                else:
                   so2 = so2 + " " + c.sol_option

          ce1 = cost_utility.exclude(weighted_utility = 0) 
          ce = ce1.order_by('cost_utility_ratio') 
          onerec3 = ce.first()                                                                                                                                                                                    
          if cur_w is not None and  cur_w <> '' and cur_w <> ' ':
             cur_w = "%.1f" % onerec3.cost_utility_ratio
          cur = onerec3.cost_utility_ratio
          so3 = ""
          so3_one = 'Y'
          for c in ce:
             if cur == c.cost_utility_ratio and (c.cost_utility_ratio is not None and c.cost_utility_ratio <> '' and c.cost_utility_ratio <> ' '):
                if so3 <> "":
                   so3 = so3 + ", " + c.sol_option
                   so3_one = "Y" 
                else:
                   so3 = so3 + " " + c.sol_option 

       # atleast one sol option with cost_utility_ratio not equal to zero
       for counter in cost_utility:
           if counter.cost_utility_ratio <> 0:
              cost_utility_exists = 'Y' 
    except ObjectDoesNotExist:
       try:
          c1 = Cost_Utility.objects.get(dec_id = dec_id)             
          wu = c1.weighted_utility
          so1 = c1.sol_option
          co = c1.cost
          so2 = c1.sol_option
          cur = c1.cost_utility_ratio
          so3 = c1.sol_option
          cu_rec_exists = 'Y'
          cu_rec_one = 'Y'
          print 'cu_rec two'
          print cu_rec_exists
          co_w = "%.1f" % co
          if wu is not None:
             wu_w = "%.1f" % wu
          if cur is not None:
             cur_w = "%.1f" % cur
          if co <= 0:                                                                                                               
             costzero = 'Y'
          else:   
             costzero = 'N'              
       except ObjectDoesNotExist:
          cu_rec_exists = 'N'
          print 'cu_rec three'
          print cu_rec_exists  
    dm_rec_exists = ''
    chosen = ''
    x = ''
    try:
       dec_made = Decision_Made.objects.get(dec_id = dec_id)
       opt = dec_made.sol_option  
       reason = dec_made.reason
       primary_factor = dec_made.primary_factor
       other_cons = dec_made.other_cons
       print opt
       z = opt.replace('[', '')
       print z
       yy = z.replace("u", "")
       y = yy.replace("'", "") 
       print y
       m = y.replace("L", "") 
       print 'm'
       print m
       x = m.replace("]","")
       print x
       for sc in x.split(','): 
           print sc
           s = Solution_Options.objects.get(id=sc)
           print 'chosen'
           print s.sol_option
           chosen = s.sol_option + ", " + chosen
       dm_rec_exists = 'Y'
    except:
       dm_rec_exists = 'N'
    
    chosen = chosen[:-2] 
    print chosen

    # using a function here
    retval = check_required(request, dec_id)
    dec_mesg = ''
    check_again = 'Y'
    print 'retval'
    print retval
    if 'solopt' in retval:
        dec_mesg = dec_mesg + 'Solution Options'
    if 'eva' in retval:   
        if dec_mesg == '':
           dec_mesg = 'Evaluation Criteria'     
        else: 
           dec_mesg = dec_mesg + ', Evaluation Criteria' 
    if 'iw' in retval or 'listerr3' in retval:
        if dec_mesg == '':
           dec_mesg = 'Importance Scores'     
        else: 
           dec_mesg =  dec_mesg + ', Importance Scores' 
    if 'mea' in retval or 'listerr1' in retval or 'listerr2' in retval:
        check_again = 'N'    
        if dec_mesg == '':
           dec_mesg = 'Evaluation Measures'     
        else: 
           dec_mesg =  dec_mesg + ', Evaluation Measures' 

    if check_again == 'Y':
       if stakeholdersNow == 'Y':
          individual_cal(dec_id, created_by, request)
       else:                                                                                                                                                                                                        
          group_cal(dec_id, created_by, request)
       retval = further_cal(dec_id, created_by, request)   
       if retval == 'em':
          if dec_mesg == '':
             dec_mesg = 'Evaluation Measures'     
          else:  
             dec_mesg =  dec_mesg + ', Evaluation Measures'

    print dec_mesg
    document = Document()
    p4 = inflect.engine()
    p2 = document.add_heading(' ',0)
    p2.add_run('DecisionMaker').italic = True
    p2.add_run(' Summary Report')

    #p = document.add_paragraph('A plain paragraph having some ')
    #p.add_run('bold').bold = True
    #p.add_run(' and some ')
    #p.add_run('italic.').italic = True
    #Executive Summary page
    document.add_heading('Executive Summary', level=1)
    p = document.add_paragraph('Using ')
    p.style = document.styles['Normal'] 
    #p.font = p.style.font
    #p.font.name = 'Calibri'
    p.add_run('DecisionMaker').italic = True
    print 'cu_rec four'
    print cu_rec_exists
    p.add_run('’s cost-utility decision-making framework, ' +  name_decisionmaker + " engaged in a")
    if real_dec_yn == 'R':
       p.add_run(" real ")
    elif real_dec_yn == 'T':
       p.add_run(" training ")
    else:
       p.add_run(" test ")
    p.add_run("decision about: ")
    p.add_run(decision_prob).bold = True
    print 'last char'
    if decision_prob[-1:] <> '.' and decision_prob[-1:] <> '?' and decision_prob[-1:] <> ';' and decision_prob[-1:] <> ':' and decision_prob[-1:] <> '!' and decision_prob[-1:] <> ',':
       p.add_run(".") 
    print stdec_count
    print text
    p.add_run(" The decision ID is " + str(dec_id) + ".") 
    run = p.add_run()
    run.add_break() 
    run.add_break()   
    if stdec_count > 0 and text <> "":
       if stdec_count == 1:
          p.add_run(p4.number_to_words(stdec_count).capitalize() + " stakeholder was ")
       else:
          p.add_run(p4.number_to_words(stdec_count).capitalize() + " stakeholders were ")  
       p.add_run("invited to contribute via ")
       p.add_run('DecisionMaker').italic = True
       p.add_run(" to the following stages of the decision-making process: ")
       p.add_run(text).bold = True
       p.add_run("." )
    run = p.add_run()
    run.add_break() 
    run.add_break()
    if cu_rec_exists == 'Y':
       print 'aug25 cu_rec_one'
       print cu_rec_one
       print so3
       if cu_rec_one == 'Y': 
          p.add_run('Among ' +  p4.number_to_words(solopt_count) + ' solution option that was evaluated, it appears that, before considering costs, ') 
       else:
          p.add_run('Among ' +  p4.number_to_words(solopt_count) + ' solution options that were evaluated, it appears that, before considering costs, ')
       print 'tuesday'
       print wu_w  
       print cur_w
       #if wu_w is not None and wu_w <> '' and wu_w <> ' ':
       p.add_run(so1).bold = True   
       if so1_one == 'Y':
          p.add_run(' yields the highest stakeholder satisfaction, i.e., it best meets your stakeholders’ criteria, earning an overall utility value of ' + str(wu_w) + ' out of 10.')
       else:
          p.add_run(' yields the highest stakeholder satisfaction, i.e., they best meet your stakeholders’ criteria, earning an overall utility value of ' + str(wu_w) + ' out of 10.')
       run = p.add_run()
       run.add_break()
       run.add_break()
       if costzero == 'N':
          if so2_one == 'Y':
             p.add_run(so2).bold = True  
             #if type_of_cost is not None:
             p.add_run(' is the least costly: ' + type_of_cost + ' cost is ' + str(co_w) + '.')
          else:
             p.add_run(so2).bold = True  
             #if type_of_cost is not None:
             p.add_run(' are the least costly: ' + type_of_cost + ' cost is ' + str(co_w) + '.')
          run = p.add_run()
          run.add_break()
          run.add_break()
       if so3 is not None and so3 <> '' and so3 <> ' ':
          if cur_w is not None and cur_w <> '' and cur_w <> ' ': 
             if so3_one == 'Y':
                p.add_run(so3).bold = True 
                p.add_run(' provides the best return on investment because it costs the least per unit of utility or stakeholder satisfaction: ' + str(cur_w) + '.') 
             else:
                p.add_run(so3).bold = True 
                p.add_run(' provide the best return on investment because they cost the least per unit of utility or stakeholder satisfaction: ' + str(cur_w) + '.') 
             run = p.add_run()
             run.add_break()
             run.add_break()

    if dm_rec_exists == 'Y': 
       if chosen <> '' and chosen <> ' ' and chosen is not None:  
          p.add_run(name_decisionmaker + ' chose ')
          p.add_run(chosen).bold = True 
       if reason <> '' and reason <> ' ' and reason is not None: 
          reason = reason.replace('#', '')
          p.add_run(' based on the following rationale or consideration: ' + reason)
          if reason[-1:] <> '.' and reason[-1:] <> '?' and reason[-1:] <> ';' and reason[-1:] <> ':' and reason[-1:] <> '!' and reason[-1:] <> ',':
             p.add_run(".") 
          run = p.add_run()
          run.add_break()
          run.add_break()
       if primary_factor <> '' and primary_factor <> ' ' and primary_factor is not None:   
          p.add_run('The primary factor in making this decision was: ' + primary_factor)
          if primary_factor[-1:] <> '.' and primary_factor[-1:] <> '?' and primary_factor[-1:] <> ';' and primary_factor[-1:] <> ':' and primary_factor[-1:] <> '!' and primary_factor[-1:] <> ',':                        
             p.add_run(".") 
          run = p.add_run()
          run.add_break()
          run.add_break()
       if other_cons <> '' and other_cons <> ' ' and other_cons is not None:   
          p.add_run('Considerations outside of the cost-utility analysis that were important included: ' + other_cons)
          if other_cons[-1:] <> '.' and other_cons[-1:] <> '?' and other_cons[-1:] <> ';' and other_cons[-1:] <> ':' and other_cons[-1:] <> '!' and other_cons[-1:] <> ',': 
             p.add_run(".") 
    document.add_page_break()
 
    # Decision Problem page 
    document.add_heading('Decision Problem', level=1)
    d = document.add_paragraph('')
    d.add_run(name_decisionmaker).bold = True
    d.add_run(' needs to make a decision about: ')
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    runner = d.add_run(decision_prob)
    runner.bold = True
    runner.italic = True
    print 'DP' 
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    run.add_break()
    d.add_run('Decision needs to be made by:').bold = True 
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    d.add_run(str(by_when))
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    run.add_break()
    d.add_run('Description of the problem being addressed: ').bold = True
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    d.add_run(title) 
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    run.add_break()
    d.add_run('Evidence to show that this problem exists: ').bold = True
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    d.add_run(evidence) 
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    run.add_break()
    d.add_run('The goal of this decision: ').bold = True
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    d.add_run(goal) 
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    run.add_break()
    d.add_run('The target audience being served is: ').bold = True
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    d.add_run(target_audience) 
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    run.add_break()
    d.add_run('The stakeholders in this decision were identified as: ').bold = True
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    d.add_run(stakeholders) 
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    run.add_break()
    d.add_run('The stakeholders who will be invited to participate in this decision are: ').bold = True
    run = d.add_run()                                                                                                                                                                                            
    run.add_break()
    d.add_run(participating_stakeholders) 
    document.add_page_break()

    # Solution Options
    document.add_heading('Solution Options', level=1)
    d = document.add_paragraph('')
    print 'SO'
    if sugg_solopt > 0:
       if sugg_solopt == 1:
          d.add_run(p4.number_to_words(sugg_solopt).capitalize() + ' stakeholder was invited to contribute to the list of possible Solution Options, out of which ' + p4.number_to_words(st_solopt_contrib_count) + ' did so (see Appendix A for details).')  
       else:  
          d.add_run(p4.number_to_words(sugg_solopt).capitalize() + ' stakeholders were invited to contribute to the list of possible Solution Options, out of which ' + p4.number_to_words(st_solopt_contrib_count) + ' did so (see Appendix A for details).')  
    else:
       d.add_run('No stakeholders have been invited to suggest Solution Options.')

    if total_solopt_count > 0: 
       d = document.add_paragraph('')
       run = d.add_run()     
       run.add_break() 
       if total_solopt_count > 1:
          d.add_run(p4.number_to_words(total_solopt_count).capitalize() + ' possible Solution Options were initially considered. ' + p4.number_to_words(solopt_count).capitalize() + ' made it through the screening process.')
       else:    
          d.add_run(p4.number_to_words(total_solopt_count).capitalize() + ' possible Solution Option was initially considered:')
       run = d.add_run()     
       run.add_break()    
       '''d.add_run('Table 1. Solution Options Considered').italic = True
       table = document.add_table(rows=1, cols=4)
       table.style = 'LightShading-Accent1'
       hdr_cells = table.rows[0].cells
       hdr_cells[0].text = '#'
       hdr_cells[1].text = 'Solution Option'
       hdr_cells[2].text = 'Source of data'
       hdr_cells[3].text = 'Description'''
       counter = 1
       d = document.add_paragraph('')
       for s in solopt:
          '''row_cells = table.add_row().cells
          row_cells[0].text = str(counter)
          row_cells[1].text = s.sol_option
          row_cells[2].text = s.source
          row_cells[3].text = s.option_details'''
          runner = d.add_run('Solution Option ' + str(counter) + ': '+ s.sol_option)
          runner.bold = True
          runner.underline = True
          run = d.add_run()                                                                                                                                                                                    
          run.add_break() 
          d.add_run('Source of idea: ').bold = True 
          d.add_run(s.source)
          run = d.add_run()                                                                                                                                                                                      
          run.add_break() 
          d.add_run('Description: ').bold = True 
          d.add_run(s.option_details.strip())
          run = d.add_run()                                                                                                                                                                                      
          run.add_break() 
          counter = counter + 1
          run = d.add_run()
          run.add_break()
       run = d.add_run()     
       run.add_break() 
       d.add_run('Information about the options can be accessed at the following link: ')
       run = d.add_run()     
       run.add_break()
       d.add_run('http://amritha.pythonanywhere.com/utility_tool/decisions/' + dec_id2 + '/solution_options/view_solopt_det.html')
       run = d.add_run()     
       run.add_break()
       if total_archived > 0:
          d = document.add_paragraph('')
          if total_archived > 1:
             d.add_run('The following additional Solution Options were suggested but were not moved forward in the decision process: ') 
          else: 
             d.add_run('The following additional Solution Option was suggested but was not moved forward in the decision process: ')
          for s in solopt_archived:
            document.add_paragraph(s.sol_option, style='List Bullet')
    else:
        run = d.add_run()  
        run.add_break()
        run.add_break()      
        d.add_run('No Solution Options have been listed.')
    document.add_page_break()
    
    # Screening Criteria
    document.add_heading('Screening Criteria', level=1)
    d = document.add_paragraph('')
    print 'SC'
    if sugg_scr > 0:
       if sugg_scr == 1:
          if st_scr_contrib_count > 0:
             d.add_run(p4.number_to_words(sugg_scr).capitalize() + ' stakeholder was invited to contribute to the list of Screening Criteria (non-negotiable requirements or “deal-breakers”) to help narrow down the list of possible Solution Options to a number that would be feasible to evaluate fully. ' + p4.number_to_words(st_scr_contrib_count).capitalize() + ' of the invited stakeholders contributed one or more Screening Criteria (see Appendix A for details).')
          else:
             d.add_run(p4.number_to_words(sugg_scr).capitalize() + ' stakeholder was invited to contribute to the list of Screening Criteria (non-negotiable requirements or “deal-breakers”) to help narrow down the list of possible Solution Options to a number that would be feasible to evaluate fully. The invited stakeholder did not contribute to the Screening Criteria.') 
       else:  
          if st_scr_contrib_count > 0:
             d.add_run(p4.number_to_words(sugg_scr).capitalize() + ' stakeholders were invited to contribute to the list of Screening Criteria (non-negotiable requirements or “deal-breakers”) to help narrow down the list of possible Solution Options to a number that would be feasible to evaluate fully. ' + p4.number_to_words(st_scr_contrib_count).capitalize() + ' did so (see Appendix A for details).')  
          else:
             d.add_run(p4.number_to_words(sugg_scr).capitalize() + ' stakeholders were invited to contribute to the list of Screening Criteria (non-negotiable requirements or “deal-breakers”) to help narrow down the list of possible Solution Options to a number that would be feasible to evaluate fully. None of the invited stakeholders contributed to the Screening Criteria.') 
    else:
       d.add_run('No stakeholders have been invited to suggest Screening Criteria.')

    if scrcr_count > 0: 
       d = document.add_paragraph('')
       run = d.add_run()     
       run.add_break() 
       d.add_run('The following Screening Criteria were proposed:')
       for s in scrcr:
           document.add_paragraph(s.criterion, style='List Bullet')
    else:
        run = d.add_run()  
        run.add_break()
        run.add_break()      
        d.add_run('No Screening Criteria have been listed.')

    if mapp_count > 0:
       d = document.add_paragraph('')
       run = d.add_run()
       run.add_break() 
       d.add_run('The Solution Options that met the Screening Criteria and remained on the list for further evaluation were: ')                                                                                     
       for s in solopt:
          document.add_paragraph(s.sol_option, style='List Bullet')

       d = document.add_paragraph('')
       run = d.add_run()
       run.add_break()
       d.add_run('Solution Options that were eliminated were: ')
       for s in solopt_archived:
          document.add_paragraph(s.sol_option, style='List Bullet')

       '''d = document.add_paragraph('')
       run = d.add_run()
       run.add_break()
       d.add_run('The Solution Options that met the Screening Criteria were: ')
       for s in solopt:
          document.add_paragraph(s.sol_option, style='List Bullet')'''

       d = document.add_paragraph('')
       run = d.add_run()
       run.add_break()
       d.add_run('Table B1 in Appendix B shows how each Solution Option performed against the Screening Criteria.')
    else:
        d = document.add_paragraph('')   
        run = d.add_run()
        run.add_break()
        d.add_run('The decision-maker has not used the table in the “Screen Solution Options” step to assess whether each Solution Option meets each Screening Criterion.')

    document.add_page_break()
    # Evaluation Criteria
    document.add_heading('Evaluation Criteria', level=1)
    d = document.add_paragraph('')
    if sugg_evacr > 0:
       if sugg_evacr == 1:
          if st_evacr_contrib_count > 0:
             d.add_run(p4.number_to_words(sugg_evacr).capitalize() + ' stakeholder was invited to contribute to the list of Evaluation Criteria, i.e., factors to consider in assessing each of the Solution Options to determine how well each one would meet their needs. ' + p4.number_to_words(st_evacr_contrib_count).capitalize() + ' of these invited stakeholders provided one or more Evaluation Criteria (see Appendix A for details).')
          else:
             d.add_run(p4.number_to_words(sugg_evacr).capitalize() + ' stakeholder was invited to contribute to the list of Evaluation Criteria, i.e., factors to consider in assessing each of the Solution Options to determine how well each one would meet their needs. The invited stakeholder did not provide Evaluation Criteria.')  
       else:  
          if st_evacr_contrib_count > 0:
             d.add_run(p4.number_to_words(sugg_evacr).capitalize() + ' stakeholders were invited to contribute to the list of Evaluation Criteria, i.e., factors to consider in assessing each of the Solution Options to determine how well each one would meet their needs. ' + p4.number_to_words(st_evacr_contrib_count).capitalize() + ' of these invited stakeholders provided one or more Evaluation Criteria (see Appendix A for details).')  
          else:
             d.add_run(p4.number_to_words(sugg_evacr).capitalize() + ' stakeholders were invited to contribute to the list of Evaluation Criteria, i.e., factors to consider in assessing each of the Solution Options to determine how well each one would meet their needs. None of these invited stakeholders provided Evaluation Criteria.')  
    else:
       d.add_run('No stakeholders have been invited to suggest Evaluation Criteria.')
    if evacr_count > 0:
       d = document.add_paragraph('')
       run = d.add_run()
       run.add_break() 
       d.add_run('The Evaluation Criteria listed were:')
       for e in evacr_notdeleted:
           document.add_paragraph(e.criterion, style='List Bullet')
    else:
        run = d.add_run()
        run.add_break()
        run.add_break() 
        d.add_run('No Evaluation Criteria have been listed.')

    document.add_page_break()
    # Importance Scores
    document.add_heading('Importance Scores and Weights', level=1)
    d = document.add_paragraph('')
    if sugg_iw > 0: 
       if sugg_iw == 1:
          if st_iw_contrib_count > 0: 
             d.add_run(p4.number_to_words(sugg_iw).capitalize() + ' stakeholder was invited to contribute Importance Scores, out of which ' + p4.number_to_words(st_iw_contrib_count) + ' did so (see Appendix A for details).')
          else:
             d.add_run(p4.number_to_words(sugg_iw).capitalize() + ' stakeholder was invited to contribute Importance Scores but they did not.')
       else:  
          if st_iw_contrib_count > 0: 
             d.add_run(p4.number_to_words(sugg_iw).capitalize() + ' stakeholders were invited to contribute Importance Scores, out of which ' + p4.number_to_words(st_iw_contrib_count) + ' did so (see Appendix A for details).')
          else:
             d.add_run(p4.number_to_words(sugg_iw).capitalize() + ' stakeholders were invited to contribute Importance Scores but none did so.')
    if iw_count > 0: 
       d = document.add_paragraph('')
       d.add_run('Importance Scores were used to assign a weight to each Evaluation Criterion indicating its relative importance to the stakeholder(s). The weights (which are between 0 and 1 for each criterion and sum to 1 across all Evaluation Criteria) are shown in Table 1. Refer to Appendix C for a summary of Importance Scores and how the importance weights are calculated.')
       run = d.add_run()
       run.add_break()
       run.add_break()
       d.add_run('Table 1. Stakeholder Weights for Each Evaluation Criterion').italic = True
       table = document.add_table(rows=1, cols=2)
       table.style = 'LightShading-Accent1'
       hdr_cells = table.rows[0].cells
       hdr_cells[0].text = 'Evaluation Criteria'
       hdr_cells[1].text = 'Importance Weight (0 - 1)'
       for i in adj_w:
          row_cells = table.add_row().cells 
          row_cells[0].text = i.criterion
          if i.adjusted_weight is not None:
             adjw = "%.2f" % i.adjusted_weight
             row_cells[1].text = str(adjw.rjust(30))
    else:
        run = d.add_run()
        run.add_break()
        run.add_break() 
        d.add_run('No Importance Scores have been provided.')

    document.add_page_break()
    # Evaluation Measures
    # https://stackoverflow.com/questions/31893557/python-docx-sections-page-orientation
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    print new_section.orientation
    document.add_heading('Evaluation Measures and Data to Collect', level=1)
    d = document.add_paragraph('')
    if evm_count > 0: 
       d = document.add_paragraph('')
       d.add_run('The evaluation measures used to assess the Solution Options are shown in Table 2.')
       run = d.add_run()
       run.add_break()
       run.add_break()
       d.add_run('Table 2. Evaluation Measures').italic = True
       table = document.add_table(rows=1, cols=6)
       table.style = 'LightShading-Accent1'
       hdr_cells = table.rows[0].cells
       hdr_cells[0].text = 'Evaluation Criterion'
       hdr_cells[1].text = 'How will you measure this?'      
       hdr_cells[2].text = 'Data to collect' 
       hdr_cells[3].text = 'Likely lowest score'     
       hdr_cells[4].text = 'Likely highest score'
       hdr_cells[5].text = 'Higher scores are better? (Yes/No)'

       for e in evm_distinct:
          row_cells = table.add_row().cells
          if e[0] is not None:                                                                                                                        
             row_cells[0].text = e[0]
          if e[1] is not None:
             row_cells[1].text = e[1]
          if e[2] is not None:
             row_cells[2].text = e[2]
          if e[3] is not None:
             row_cells[3].text = str(e[3])
          if e[4] is not None:
             row_cells[4].text = str(e[4])
          if e[5] is not None:
             row_cells[5].text = e[5]
    else:
        run = d.add_run()
        run.add_break()
        run.add_break() 
        d.add_run('No evaluation measures have been entered.')
    
    d = document.add_paragraph('')
    run = d.add_run()
    run.add_break()
    d.add_run('See Appendix D for details on what information was used to assess each Solution Option against each Evaluation Criterion.')
    run = d.add_run()
    run.add_break()
    run.add_break()
    d.add_run('Table 3 shows:')
    document.add_paragraph("The rating or score that each Solution Option earned for each evaluation measure;", style='List Bullet')
    document.add_paragraph("How well, on average, each Solution Option performed against each Evaluation Criterion on a scale of 0-10 (i.e., the “criterion-level unweighted utility values”)", style='List Bullet')
    document.add_paragraph("The overall utility value for each Solution Option.", style='List Bullet')
    d = document.add_paragraph('')
    run = d.add_run()
    run.add_break()
    d.add_run('Note: ').bold = True
    d.add_run(': If you have completed the Evaluation Data Table but Criterion-Level Utility Values are not appearing in Table 3, click on the “Make a Decision” icon in the DecisionMaker flowchart and re-export the Summary Report. (This will automatically validate the data you have entered).')   
    d = document.add_paragraph('')
    run = d.add_run()
    run.add_break()
    d.add_run('See “How Utility Values are Calculated” section of this document to learn more about how Utility Values are calculated.')
    if evm_count > 0:                                                                                                                                                                                            
       d = document.add_paragraph('')
       run = d.add_run()
       run.add_break()
       d.add_run('Table 3. Scores on Evaluation Measures and the Criterion-level Utility Values.').italic = True
       no_of_cols = 2 + evm_opt_count + evm_opt_count
       counter = 1
       print 'no of cols'
       print no_of_cols 
       table = document.add_table(rows=1, cols=no_of_cols)
       table.style = 'LightShading-Accent1'
       hdr_cells = table.rows[0].cells
       hdr_cells[0].text = 'Evaluation Criterion'
       #hdr_cells[1].text = 'Importance Weight'
       for e in evm_distinct_opt:     
           if counter > (no_of_cols - 2):
              counter = 1
           hdr_cells[counter].text = e[0] + ' Average Rating or Score'
           counter = counter + 1
           hdr_cells[counter].text = e[0] + ' Criterion-level Utility Value(0-10)'     
           counter = counter + 1
           onerec_opt = e[0] 
       counter = 1 
       for e in evm_distinct_cri:
          counter = 1
          row_cells = table.add_row().cells 
          row_cells[0].text = e[0]
          #row_cells[1].text = ''
          onerec_cri = e[0]  
          for ee in evm2: 
             if ee.criterion == onerec_cri:
                if ee.option_value is not None: 
                   row_cells[counter].text = "%.1f" % float(ee.option_value)
                   counter = counter + 1
                if ee.utility_value is not None: 
                   row_cells[counter].text = "%.1f" % float(ee.utility_value)
                   counter = counter + 1


       '''total_utility_value = 0
       counter = 1  
       for ee in evm2:
          if ee.sol_option == onerec_opt:
             if ee.utility_value is not None: 
                total_utility_value = total_utility_value + float(ee.utility_value)
          onerec_opt = ee.sol_option
       run = d.add_run()  
       run.add_break()
       d.add_run('The overall utility value of ' + onerec_opt + ' is ' + "%.1f" % float(total_utility_value) + '.')'''


    document.add_page_break()
    # Costs
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = new_width
    new_section.page_height = new_height
    print new_section.orientation  
    document.add_heading('Costs and Numbers of Participants Served', level=1)
    d = document.add_paragraph('')
    if costzero == 'N':
    #if cu_count > 0:
       if cu_count > 1: 
          d.add_run('"' + lowest_cost + '" is expected to be the least costly to implement while "' + highest_cost + '" is expected to be the most costly to implement. ')
       else:
          d.add_run('The ' + cost_text +  ' of ' + so1 + ' is ' + str(co_w) + '.') 
       run = d.add_run()  
       run.add_break()
       run.add_break()
       d.add_run('Table 4. Expected Costs of Implementing each Solution Option and the Number of Participants Served.').italic = True
       table = document.add_table(rows=1, cols=3)
       table.style = 'LightShading-Accent1'
       hdr_cells = table.rows[0].cells
       hdr_cells[0].text = 'Solution Options'
       hdr_cells[1].text = 'Number of participants served'      
       hdr_cells[2].text = cost_text
       for c in cost_utility:
          row_cells = table.add_row().cells 
          row_cells[0].text = c.sol_option
          row_cells[1].text = str(c.no_of_participants).rjust(15)
          row_cells[2].text = str(c.cost).rjust(15)
    else:
        run = d.add_run()
        run.add_break()
        run.add_break() 
        d.add_run('No cost information has been provided.')

    d = document.add_paragraph('')
    if det_count > 0: 
       run = d.add_run()
       run.add_break()                                                                                                                                                                                          
       d.add_run('For more details on the costs, refer to Appendix E.')
    else:
        run = d.add_run()
        run.add_break()
        d.add_run('No further details on the cost data have been entered.')

    document.add_page_break()
    # Decision Made
    document.add_heading('Results: Cost, Utility and Cost-Utility Ratios', level=1)
    d = document.add_paragraph('')
    if dec_mesg <> '':
       d.add_run('The utility values cannot be calculated because information on ' +  dec_mesg + ' is missing.')
    else:
       d.add_run('Utility is a measure of stakeholder satisfaction or “usefulness” reported in ')
       d.add_run('DecisionMaker').italic = True
       d.add_run(' on a scale of 0 to 10. ')
       if so1_one == 'Y':
          d.add_run('It appears that, before considering costs, ')
          d.add_run(so1).bold = True
          d.add_run(' yields the highest stakeholder satisfaction, i.e., it best meets your stakeholders’ criteria, earning an overall utility of ' + str(wu_w) + ' out of 10. Refer to Appendix E to see how the utility value for each Solution Option is calculated.')
       else:
          d.add_run('It appears that, before considering costs, ')
          d.add_run(so1).bold = True
          d.add_run(' yields the highest stakeholder satisfaction, i.e., they best meet your stakeholders’ criteria, earning an overall utility of ' + str(wu_w) + ' out of 10. Refer to Appendix E to see how the utility value for each Solution Option is calculated.')
    '''run = d.add_run()                                                                                                                                                                                         
    run.add_break()
    run.add_break()'''

    d = document.add_paragraph('')
    if cu_count > 0: 
       d.add_run('Costs are the resource requirements for each Solution Option, which may include personnel time, training, facilities, materials and equipment, and other resources.')
       if so2_one == 'Y':
          d.add_run(so2).bold = True  
          #if type_of_cost is not None:
          if costzero == 'N':     
             d.add_run(' is the least costly: ' + cost_text + ' is ' + str(co_w) + '.')
       else:
          d.add_run(so2).bold = True  
          #if type_of_cost is not None:
          if costzero == 'N':     
             d.add_run(' are the least costly: ' + cost_text + ' is ' + str(co_w) + '.')
    else:
        d.add_run('No cost information has been provided.')
    run = d.add_run()                                                                                                                              
    run.add_break()
    run.add_break()
    if cost_utility_exists == 'Y':
        d.add_run('The cost-utility ratio is the cost per unit of stakeholder satisfaction (costs divided by utility value). Lower cost-utility ratios indicate greater return on investment.')
        if so3 is not None:
           if so3_one == 'Y':
              d.add_run(so3).bold = True 
              d.add_run(' provides the best return on investment because it costs the least per unit of utility (or unit of stakeholder satisfaction): ' + str(cur_w))  
           else:  
              d.add_run(so3).bold = True  
              d.add_run(' provide the best return on investment because it costs the least per unit of utility (or unit of stakeholder satisfaction): '  + str(cur_w))

           run = d.add_run()
           run.add_break()
    run.add_break()
    d.add_run('Decision-makers should consider the costs and utility of each Solution Option as well as the cost-utility ratio before making decisions. Options with higher utility values are likely to be better received by stakeholders but high-cost options may strain the budget and divert resources from other initiatives or programs. Cost-utility ratios can be used to rank Solution Options in terms of overall value for money.  The table below shows the utility, costs and cost-utility ratio for each Solution Option. The Solution Option that provides the most stakeholder satisfaction is at the top.')
    run = d.add_run()
    run.add_break()
    run.add_break() 
    d.add_run('Table 5. Overall Utility, Costs and Cost-utility Ratios for Each Solution Option').italic = True
    table = document.add_table(rows=1, cols=4)
    table.style = 'LightShading-Accent1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Solution Option'
    hdr_cells[1].text = 'Overall Utility Value or Stakeholder Satisfaction (0-10)'      
    hdr_cells[2].text = cost_text
    hdr_cells[3].text = 'Cost per Unit of Utility'
    for c in cc:
        row_cells = table.add_row().cells 
        row_cells[0].text = c.sol_option
        if c.weighted_utility is not None:
           wut = "%.1f" % c.weighted_utility
        else:
           wut = "" 
        row_cells[1].text = str(wut.rjust(15))
        if c.cost is not None: 
           ct = "%.1f" % c.cost
        else:
           ct = ""
        row_cells[2].text = str(ct.rjust(15))
        if c.cost_utility_ratio is not None:
           curatio = "%.1f" % c.cost_utility_ratio 
        else:
           curatio = ""
        row_cells[3].text = str(curatio.rjust(15))


    d = document.add_paragraph('')
    if dm_rec_exists == 'Y':
       #document.add_heading('Final decision and rationale', level=1)          
       d = document.add_paragraph('')
       if chosen <> '' and chosen <> ' ' and chosen is not None:  
          d.add_run(name_decisionmaker + ' chose ')
          d.add_run(chosen).bold = True
       if reason <> '' and reason <> ' ' and reason is not None: 
          reason = reason.replace('#', '')
          d.add_run(' based on the following rationale or consideration: ' + reason)
          if reason[-1:] <> '.' and reason[-1:] <> '?' and reason[-1:] <> ';' and reason[-1:] <> ':' and reason[-1:] <> '!' and reason[-1:] <> ',':
             d.add_run(".") 
          run = d.add_run()
          run.add_break()
          run.add_break()
       if primary_factor <> '' and primary_factor <> ' ' and primary_factor is not None:   
          d.add_run('The primary factor in making this decision was: ' + primary_factor)
          if primary_factor[-1:] <> '.' and primary_factor[-1:] <> '?' and primary_factor[-1:] <> ';' and primary_factor[-1:] <> ':' and primary_factor[-1:] <> '!' and primary_factor[-1:] <> ',':              
             d.add_run(".") 
          run = d.add_run()
          run.add_break()
          run.add_break()
       if other_cons <> '' and other_cons <> ' ' and other_cons is not None:   
          d.add_run('Considerations outside of the cost-utility analysis that were important included: ' + other_cons)
          if other_cons[-1:] <> '.' and other_cons[-1:] <> '?' and other_cons[-1:] <> ';' and other_cons[-1:] <> ':' and other_cons[-1:] <> '!' and other_cons[-1:] <> ',': 
             d.add_run(".") 
    else:
       d = document.add_paragraph('') 
       run = d.add_run()
       run.add_break()
       d.add_run("Solution Options to address the decision problem have not yet been selected.")

    document.add_page_break()
    # Utility Values
    document.add_heading('How Utility Values are Calculated', level=1)
    d = document.add_paragraph('')
    d.add_run('DecisionMaker').italic = True
    d.add_run(' uses the data you entered in the Evaluation Data Table to calculate utility. The overall utility value earned by a Solution Option is the sum of the utility scores it earns on each of the Evaluation Criteria (Criterion-level unweighted utility values) multiplied by the importance weights assigned by stakeholders to the criteria.')
    run = d.add_run()
    run.add_break()
    document.add_heading('Criterion-level unweighted utility value', level=1)
    d = document.add_paragraph('')
    d.add_run('Each measure you used in your evaluation is rescaled to convert your results to a common utility scale with a minimum of 0 and maximum of 10, where 0 means that you and your stakeholders are not at all satisfied and 10 means you are completely satisfied. The likely lowest score and likely highest score you entered for each measure are used to set the extremes of the scale and a straight line connects the two points. This assumes that utility changes in direct proportion to the changes in the evaluation measure.')
    d = document.add_paragraph('')
    d.add_run('When the rating or score on an evaluation measure is ')
    d.add_run('positively').bold = True
    d.add_run(' associated with the utility values (i.e., higher scores are better), the likely lowest score you entered is assumed to provide zero utility and the likely highest score you entered is assumed to provide a utility value of 10. The criterion-level unweighted utility value for a Solution Option is:')
    document.add_picture('/home/amritha/decisionmaker/utility_tool/positiveformula2.PNG',width=Inches(6))
    d = document.add_paragraph('')
    d.add_run('When the rating or score on an evaluation measure is ')
    d.add_run('negatively').bold = True
    d.add_run(' associated with the utility values (i.e., lower scores are better), the likely lowest score you entered is assumed to provide a utility value of 10 while the likely highest score you entered is now assumed to provide zero utility. The criterion-level unweighted utility value for a Solution Option is:')
    document.add_picture('/home/amritha/decisionmaker/utility_tool/negativeformula2.PNG',width=Inches(6))                                              
    d = document.add_paragraph('Visit ')
    add_hyperlink(d, 'https://amritha.pythonanywhere.com/Steps.html#Utility', "https://amritha.pythonanywhere.com/Steps.html#Utility")
    d.add_run(' to see an example of this calculation.')
    document.add_heading('Overall Utility Value', level=1)
    d = document.add_paragraph('')
    d.add_run('The overall utility value is the sum of the criterion-level utility values multiplied by the importance weights. Visit ')
    add_hyperlink(d, 'https://amritha.pythonanywhere.com/Steps.html#Utility', "https://amritha.pythonanywhere.com/Steps.html#Utility")
    d.add_run(' to see an example of this calculation.')
    
    if real_dec_yn == 'R':
       filename = 'SummaryReport-' + str(dec_id) + '-R.docx'
    elif real_dec_yn == 'T':
       filename = 'SummaryReport-' + str(dec_id) + '-T.docx'         
    else:
       filename = 'SummaryReport-' + str(dec_id) + '-X.docx'      
    document.save(filename)
    #fs = FileSystemStorage("/tmp")
    with open(filename) as docx:
        response = HttpResponse(docx.read(), content_type='application/docx')
        response['Content-Disposition'] = 'inline; filename="%s"' % filename
        return response

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    return hyperlink              

def temp(request):
    return render(request, 'temp.html')

def message(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser})

def nouser_message(request):
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'decisions/nouser_message.html',  {'loggedinuser':loggedinuser})

def dec_info(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0 

    if 'whereamI' in request.session:
       whereamI = request.session['whereamI']
    else:
       whereamI  = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
 
    return render(request, 'resources/dec_info.html', {'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser})  

def st_info(request):                                                                                                                                                                                           
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'whereamI' in request.session:
       whereamI = request.session['whereamI']
    else:
       whereamI  = 0  
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    return render(request, 'resources/st_info.html', {'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser})  

def solopt_info(request):     
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'whereamI' in request.session:
       whereamI = request.session['whereamI']
    else:
       whereamI  = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'resources/solopt_info.html', {'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser})    

def scr_info(request):                                                                                                                            
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'whereamI' in request.session:
       whereamI = request.session['whereamI']
    else:
       whereamI  = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'resources/scr_info.html', {'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser})    

def eva_info(request):                                                                                                                            
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'whereamI' in request.session:
       whereamI = request.session['whereamI']
    else:
       whereamI  = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'resources/eva_info.html', {'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser})  

def score_info(request):                                                                                                                            
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'whereamI' in request.session:
       whereamI = request.session['whereamI']
    else:
       whereamI  = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'resources/score_info.html', {'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser}) 


def evamea_info(request):     
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'whereamI' in request.session:
       whereamI = request.session['whereamI']
    else:
       whereamI  = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'resources/evamea_info.html', {'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser}) 

def utility_info(request):     
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    return render(request, 'resources/utility_info.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser})  

def costs_info(request):     
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0 

    if 'whereamI' in request.session:
       whereamI = request.session['whereamI']
    else:
       whereamI  = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
 
    return render(request, 'resources/costs_info.html', {'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser})     

def makedec_info(request):     
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'whereamI' in request.session:
       whereamI = request.session['whereamI']
    else:
       whereamI  = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'resources/makedec_info.html', {'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser})   

def gen_info(request):     
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'resources/gen_info.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser})  

def about(request):
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request, 'about.html', {'loggedinuser':loggedinuser})

def about_maker(request):                                                                                                                                  
    request.session['whereamI'] = 'aboutmaker'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    return render(request, 'resources/about_maker.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser})


def menu3(request, dec_id):
    request.session['dec_id'] = dec_id

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    try: 
       dec = Decisions.objects.get(id=dec_id)
       request.session['dec_title'] = dec.title
       created_by = dec.created_by
    except:
       request.session['dec_title'] = 'not found' 
       created_by = 'not found'

    return render(request, 'decisions/menu3.html',{'dec_title':request.session['dec_title'], 'loggedinuser':loggedinuser})

def menu4(request, dec_id):
    request.session['dec_id'] = dec_id
                                                                                                                                                            
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    try: 
       dec = Decisions.objects.get(id=dec_id)
       request.session['dec_title'] = dec.title
       created_by = dec.created_by
    except:
       request.session['dec_title'] = 'not found' 
       created_by = 'not found'
    return render(request, 'decisions/menu4.html',{'dec_title':request.session['dec_title'], 'loggedinuser':loggedinuser})

def stakeholders(request):
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
   
    if 'whereamI' in request.session:                                                                                                                                                                            
        whereamI = request.session['whereamI']
    else:
        whereamI = 'stakeholders'

    if 'idList' in request.session:                                                                                                                                                                            
        idList = request.session['idList']
    else:
        idList = 'no idList'

    if request.method == 'POST':
       print request.POST.getlist('id') 
       if whereamI == 'dec_list':
          return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html')
       else: 
          if 'id' in request.POST:                                                                                                                      
             if 'submit' in request.POST:
                 for value in request.POST.getlist('id'):                                                                                               
                     try: 
                        old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value)
                     except ObjectDoesNotExist:
                        st = Stakeholders.objects.get(pk=value)
                        st_name = st.firstName + ' ' + st.lastName
                        st_dec = Stakeholders_Decisions(st_id = value, name = st_name, email=st.email, dec_id = request.session['dec_id'],created_by = request.session['user'],created_date = datetime.datetime.now())  
                        st_dec.save() 
             return HttpResponseRedirect('/utility_tool/decisions/solution_options/assign_tasks.html')
          else:
             return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id) 
    if loggedinuser == 'not found':
       return HttpResponseRedirect('/utility_tool/decisions/nouser_message.html')           
    else:    
       stakeholders = Stakeholders.objects.filter(created_by=loggedinuser).order_by('firstName')
       return render(request,'stakeholders/stakeholders.html',{'stakeholders':stakeholders, 'loggedinuser':loggedinuser, 'dec_id':dec_id, 'whereamI':whereamI, 'idList':idList})

def add_stakeholder(request):
    context = RequestContext(request)
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
   
    if 'whereamI' in request.session:                                                                                                                                                                            
        whereamI = request.session['whereamI']
    else:
        whereamI = 'stakeholders'

    idList = []
    request.session['idList'] = []
    MFormSet = modelformset_factory(Stakeholders, form=StakeholdersForm, extra=6) 
    if request.method == 'POST':
        stform = MFormSet(request.POST,request.FILES, prefix="stform" )

        if stform.is_valid():
            id = stform.save(commit=False)
            for recs in id:
                if recs.email == '':
                   sterr = "Please enter an email address for the stakeholder/s you have added." 
                   return render(request,'stakeholders/add_stakeholder.html',{'stform':stform, 'sterr':sterr,'dec_id':dec_id, 'whereamI':whereamI}) 
                else:
                    try:
                       Stakeholders.objects.get(firstName = recs.firstName, lastName = recs.lastName, created_by=loggedinuser) 
                       sterr = "You have already added a stakeholder with the same name." 
                       return render(request,'stakeholders/add_stakeholder.html',{'stform':stform, 'sterr':sterr,'dec_id':dec_id, 'whereamI':whereamI}) 
                    except ObjectDoesNotExist:
                       try:
                           Stakeholders.objects.get(email = recs.email, created_by=loggedinuser) 
                           sterr = "You have already added a stakeholder with the same email address." 
                           return render(request,'stakeholders/add_stakeholder.html',{'stform':stform, 'sterr':sterr,'dec_id':dec_id, 'whereamI':whereamI}) 
                       except ObjectDoesNotExist:
                           recs.created_by = request.session['user']
                           recs.save()
                           idList.append(recs.id)
            request.session['idList'] = idList               
            if 'submit' in request.POST:
               return HttpResponseRedirect('/utility_tool/stakeholders/stakeholders.html')
            elif 'menu' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
        else:
            print stform.errors

    else:
        qset = Stakeholders.objects.none()
        stform = MFormSet(queryset=qset,prefix="stform" )
        #stform = StakeholdersForm()

    return render(request,'stakeholders/add_stakeholder.html',{'stform':stform,'dec_id':dec_id, 'whereamI':whereamI, 'loggedinuser':loggedinuser})

def edit_stakeholder(request, st_id):
    context = RequestContext(request)

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    st = Stakeholders.objects.get(pk=st_id)
    if request.method == 'POST':
        stform = StakeholdersForm(data=request.POST,instance=st)
        if stform.is_valid():
            id = stform.save(commit=False)
            if id.email == '':
                sterr = "Please enter an email address." 
                return render(request,'stakeholders/edit_stakeholder.html',{'stform':stform, 'sterr':sterr}) 
            else:
                st1 = Stakeholders.objects.filter(firstName = id.firstName, lastName = id.lastName, created_by=loggedinuser).count() 
                if st1 > 1:
                   sterr = "You have already added a stakeholder with the same name." 
                   return render(request,'stakeholders/edit_stakeholder.html',{'stform':stform, 'sterr':sterr}) 
                st1 = Stakeholders.objects.filter(email = id.email, created_by=loggedinuser).count()
                if st1 > 1:
                   sterr = "You have already added a stakeholder with the same email address." 
                   return render(request,'stakeholders/edit_stakeholder.html',{'stform':stform, 'sterr':sterr}) 
                id.updated_by = request.session['user']   
                id.updated_date = datetime.datetime.now()
                id.save(update_fields=['firstName', 'lastName', 'title', 'email', 'phone', 'organisation', 'notes','updated_by', 'updated_date']) 
                st_name = id.firstName + ' ' + id.lastName
                stdec = Stakeholders_Decisions.objects.filter(st_id=id.id)
                for s in stdec:
                    s.updated_by = request.session['user']   
                    s.updated_date = datetime.datetime.now()
                    s.name = st_name
                    s.email = id.email 
                    s.save(update_fields=['name','email','updated_by', 'updated_date']) 
            return HttpResponseRedirect('/utility_tool/stakeholders/stakeholders.html')
        else:
            print stform.errors
    else:    
        stform = StakeholdersForm(instance=st)
    t = loader.get_template('stakeholders/edit_stakeholder.html')
    c = Context({'stform' :stform})
    return render(request,'stakeholders/edit_stakeholder.html',{'st_id':st_id, 'stform':stform, 'loggedinuser':loggedinuser})

def delete_stakeholder(request, st_id):
    context = RequestContext(request)
    stdec = Stakeholders_Decisions.objects.filter(st_id=st_id)
    for s in stdec:
        s.updated_by = request.session['user']   
        s.updated_date = datetime.datetime.now()
        s.deleted = 'Y'
        s.save(update_fields=['deleted','updated_by', 'updated_date']) 
    Stakeholders.objects.get(pk=st_id).delete()                                                                                                                                                               
    return HttpResponseRedirect('/utility_tool/stakeholders/stakeholders.html')

def send_email(request):
    context = RequestContext(request)
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']
         
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'user_email' in request.session:                                                                                                                                                                          
       user_email = request.session['user_email']
    else:   
       user_email = 'not found'

    dec = Decisions.objects.get(pk=dec_id) 
    stdec_solopt = Stakeholders_Decisions.objects.raw("SELECT id, st_id, name, email, solopt_date from utility_tool_stakeholders_decisions where created_by=%s and dec_id = %s and deleted is null and solopt_type = 'Y'", [loggedinuser, dec_id])          
    stdec_scrcr = Stakeholders_Decisions.objects.raw("SELECT id, st_id, name, email, scrcr_date from utility_tool_stakeholders_decisions where created_by=%s and dec_id = %s and deleted is null and scrcr_type = 'Y'", [loggedinuser, dec_id])  
    stdec_evacr = Stakeholders_Decisions.objects.raw("SELECT id, st_id, name, email, evacr_date from utility_tool_stakeholders_decisions where created_by=%s and dec_id = %s and deleted is null and evacr_type = 'Y'", [loggedinuser, dec_id])  
    stdec_iw = Stakeholders_Decisions.objects.raw("SELECT id, st_id, name, email, iw_date from utility_tool_stakeholders_decisions where created_by=%s and dec_id = %s and deleted is null and iw_type = 'Y' and email != %s", [loggedinuser, dec_id, user_email])  
 

    if 'selected_solopt' in request.POST:
       for val in request.POST.getlist('selected_solopt'):
           if val != "[]":
              val = val.strip()
              y = val.replace('[','')
              z = y.replace(']','')          
              temp_list = [] 
              # adding each id to a temporary list
              for l2 in z.split(','):                                                                                                                                                                          
                  l3 = l2.replace('"', '')
                  temp_list.append(l3) 
              print temp_list
              for l in temp_list:
                   try: 
                      solopt = Stakeholders_Decisions.objects.get(dec_id=dec_id, id=l)
                      st = Stakeholders.objects.get(id=solopt.st_id)
                      print st.firstName
                      user = Users.objects.get(email = user_email)
                      username = user.firstName + ' '+ user.lastName
                      solopt.solopt_date = datetime.datetime.strptime(request.POST.get('sol_date'),'%m/%d/%Y')
                      if user.organisation <> '' and user.organisation <> ' ' and user.organisation is not None:
                         message = 'Dear ' + st.firstName + ',\n' + username + ' from ' +  user.organisation + ' has invited you to participate in the following decision: \n\n' + dec.short_title + '\n\nPlease click on decisionmaker (https://amritha.pythonanywhere.com)  and log in to contribute:\nIdeas for Solution Options by ' + request.POST.get('sol_date') + '\n\nIf this is the first time you will be using DecisionMaker, please register using this email address: ' + solopt.email + '\n\nYou may want to visit the “About” page in DecisionMaker and some of the Resources & Guidance pages to learn more about the decision-making framework in which you will be participating.\n\nThank you for contributing to this decision process.\n\n'+ username
                      else:
                          message = 'Dear ' + st.firstName + ',\n' + username +  ' has invited you to participate in the following decision: \n\n' + dec.short_title + '\n\nPlease click on DecisionMaker (https://amritha.pythonanywhere.com) and log in to contribute:\nIdeas for Solution Options by ' + request.POST.get('sol_date') + '\n\nIf this is the first time you will be using DecisionMaker, please register using this email address: ' + solopt.email + '\n\nYou may want to visit the “About” page in DecisionMaker and some of the Resources & Guidance pages to learn more about the decision-making framework in which you will be participating.\n\nThank you for contributing to this decision process.\n\n'+ username
                      subject = 'Invitation to Contribute Solution Options'
                      from_email = user_email
                      to_email = solopt.email
                      if subject and message and from_email and to_email:
                         try:
                            send_mail(subject, message, from_email,[to_email])
                         except BadHeaderError:
                            return HttpResponse('Invalid header found.')
                      else:
                         return HttpResponse('Make sure all fields are entered and valid.') 
                      #print request.POST.get('sol_date')
                      solopt.updated_by = request.session['user'] 
                      solopt.updated_date = datetime.datetime.now()
                      solopt.save(update_fields=['solopt_date','updated_by', 'updated_date'])  
                   except ObjectDoesNotExist:
                      print 'id does not exist'
    #Suggestions for Screening Criteria by [date entered by PA]
    if 'selected_scrcr' in request.POST:
       for val in request.POST.getlist('selected_scrcr'):
           if val != "[]":
              val = val.strip()
              y = val.replace('[','')
              z = y.replace(']','')          
              temp_list = [] 
              # adding each id to a temporary list
              for l2 in z.split(','):                                                                                                                                                                          
                  l3 = l2.replace('"', '')
                  temp_listetappend(l3) 
              print temp_list
              for l in temp_list:
                   try: 
                      scrcr = Stakeholders_Decisions.objects.get(dec_id=dec_id, id=l)
                      st = Stakeholders.objects.get(id=scrcr.st_id)
                      print st.firstName
                      user = Users.objects.get(email = user_email)
                      username = user.firstName + ' '+ user.lastName
                      scrcr.scrcr_date = datetime.datetime.strptime(request.POST.get('scr_date'),'%m/%d/%Y')
                      if user.organisation <> '' and user.organisation <> ' ' and user.organisation is not None:
                         message = 'Dear ' + st.firstName + ',\n' + username + ' from ' +  user.organisation + ' has invited you to participate in the following decision: \n\n' + dec.short_title + '\n\nPlease click on DecisionMaker (https://amritha.pythonanywhere.com)  and log in to contribute:\nSuggestions for Screening Criteria by ' + request.POST.get('scr_date') + '\n\nIf this is the first time you will be using DecisionMaker, please register using this email address: ' + scrcr.email + '\n\nYou may want to visit the “About” page in DecisionMaker and some of the Resources & Guidance pages to learn more about the decision-making framework in which you will be participating.\n\nThank you for contributing to this decision process.\n\n'+ username                                                          
                      else:
                          message = 'Dear ' + st.firstName + ',\n' + username +  ' has invited you to participate in the following decision: \n\n' + dec.short_title + '\n\nPlease click on DecisionMaker (https://amritha.pythonanywhere.com) and log in to contribute:\nSuggestions for Screening Criteria by ' + request.POST.get('scr_date') + '\n\nIf this is the first time you will be using DecisionMaker, please register using this email address: ' + scrcr.email + '\n\nYou may want to visit the “About” page in DecisionMaker and some of the Resources & Guidance pages to learn more about the decision-making framework in which you will be participating.\n\nThank you for contributing to this decision process.\n\n'+ username
                      subject = 'Invitation to provide Suggestions for Screening Criteria'
                      from_email = user_email
                      to_email = scrcr.email
                      if subject and message and from_email and to_email:
                         try:
                            send_mail(subject, message, from_email,[to_email])
                         except BadHeaderError:
                            return HttpResponse('Invalid header found.')
                      else:
                         return HttpResponse('Make sure all fields are entered and valid.') 
                      #print request.POST.get('sol_date')
                      scrcr.updated_by = request.session['user'] 
                      scrcr.updated_date = datetime.datetime.now()
                      scrcr.save(update_fields=['scrcr_date','updated_by', 'updated_date'])  
                   except ObjectDoesNotExist:
                      print 'id does not exist'
    #Suggestions for Evaluation Criteria by [date entered by PA]
    if 'selected_evacr' in request.POST:
       for val in request.POST.getlist('selected_evacr'):
           if val != "[]":
              val = val.strip()
              y = val.replace('[','')
              z = y.replace(']','')          
              temp_list = [] 
              # adding each id to a temporary list
              for l2 in z.split(','):                                                                                                                                                                          
                  l3 = l2.replace('"', '')
                  temp_list.append(l3) 
              print temp_list
              for l in temp_list:
                   try: 
                      evacr = Stakeholders_Decisions.objects.get(dec_id=dec_id, id=l)
                      st = Stakeholders.objects.get(id=evacr.st_id)
                      print st.firstName
                      user = Users.objects.get(email = user_email)
                      username = user.firstName + ' '+ user.lastName
                      evacr.evacr_date = datetime.datetime.strptime(request.POST.get('eva_date'),'%m/%d/%Y')
                      if user.organisation <> '' and user.organisation <> ' ' and user.organisation is not None:
                         message = 'Dear ' + st.firstName + ',\n' + username + ' from ' +  user.organisation + ' has invited you to participate in the following decision: \n\n' + dec.short_title + '\n\nPlease click on DecisionMaker (https://amritha.pythonanywhere.com)  and log in to contribute:\nSuggestions for Evaluation Criteria by ' + request.POST.get('eva_date') + '\n\nIf this is the first time you will be using DecisionMaker, please register using this email address: ' + evacr.email + '\n\nYou may want to visit the “About” page in DecisionMaker and some of the Resources & Guidance pages to learn more about the decision making framework in which you will be participating.\n\nThank you for contributing to this decision process.\n\n'+ username                                                    
                      else:
                          message = 'Dear ' + st.firstName + ',\n' + username +  ' has invited you to participate in the following decision: \n\n' + dec.short_title + '\n\nPlease click on DecisionMaker (https://amritha.pythonanywhere.com) and log in to contribute:\nSuggestions for Evaluation Criteria by ' + request.POST.get('eva_date') + '\n\nIf this is the first time you will be using DecisionMaker, please register using this email address: ' + evacr.email + '\n\nYou may want to visit the “About” page in DecisionMaker and some of the Resources & Guidance pages to learn more about the decision-making framework in which you will be participating.\n\nThank you or contributing to this decision process.\n\n'+ username
                      subject = 'Invitation to provide Suggestions for Evaluation Criteria'
                      from_email = user_email
                      to_email = evacr.email
                      if subject and message and from_email and to_email:
                         try:
                            send_mail(subject, message, from_email,[to_email])
                         except BadHeaderError:
                            return HttpResponse('Invalid header found.')
                      else:
                         return HttpResponse('Make sure all fields are entered and valid.') 
                      #print request.POST.get('sol_date')
                      evacr.updated_by = request.session['user'] 
                      evacr.updated_date = datetime.datetime.now()
                      evacr.save(update_fields=['evacr_date','updated_by', 'updated_date'])  
                   except ObjectDoesNotExist:
                      print 'id does not exist'
    #Importance Scores for the evaluation criteria by [date entered by PA]].  
    if 'selected_iw' in request.POST:
       for val in request.POST.getlist('selected_iw'):
           if val != "[]":
              val = val.strip()
              y = val.replace('[','')
              z = y.replace(']','')          
              temp_list = [] 
              # adding each id to a temporary list
              for l2 in z.split(','):                                                                                                                                                                          
                  l3 = l2.replace('"', '')
                  temp_list.append(l3) 
              print temp_list
              for l in temp_list:
                   try: 
                      iw = Stakeholders_Decisions.objects.get(dec_id=dec_id, id=l)
                      st = Stakeholders.objects.get(id=iw.st_id)
                      print st.firstName
                      user = Users.objects.get(email = user_email)
                      username = user.firstName + ' '+ user.lastName
                      iw.iw_date = datetime.datetime.strptime(request.POST.get('iw_date'),'%m/%d/%Y')
                      if user.organisation <> '' and user.organisation <> ' ' and user.organisation is not None:
                         message = 'Dear ' + st.firstName + ',\n' + username + ' from ' +  user.organisation + ' has invited you to participate in the following decision: \n\n' + dec.short_title + '\n\nPlease click on DecisionMaker (https://amritha.pythonanywhere.com)  and log in to contribute:\nImportance Scores by ' + request.POST.get('iw_date') + '\n\nIf this is the first time you will be using DecisionMaker, please register using this email address: ' + iw.email + '\n\nYou may want to visit the “About” page in DecisionMaker and some of the Resources & Guidance pages to learn more about the decision-making framework in which you will be participating.\n\nThank you for contributing to this decision process.\n\n'+ username
                      else:
                          message = 'Dear ' + st.firstName + ',\n' + username +  ' has invited you to participate in the following decision: \n\n' + dec.short_title + '\n\nPlease click on DecisionMaker (https://amritha.pythonanywhere.com) and log in to contribute:\nImportance Scores by ' + request.POST.get('iw_date') + '\n\nIf this is the first time you will be using DecisionMaker, please register using this email address: ' + iw.email + '\n\nYou may want to visit the “About” page in DecisionMaker and some of the Resources & Guidance pages to learn more about the decision-making framework in which you will be participating.\n\nThank you for contributing to this decision process.\n\n'+ username
                      subject = 'Invitation to Contribute Importance Scores'
                      from_email = user_email
                      to_email = iw.email
                      if subject and message and from_email and to_email:
                         try:
                            send_mail(subject, message, from_email,[to_email])
                         except BadHeaderError:
                            return HttpResponse('Invalid header found.')
                      else:
                         return HttpResponse('Make sure all fields are entered and valid.') 
                      #print request.POST.get('sol_date')
                      iw.updated_by = request.session['user'] 
                      iw.updated_date = datetime.datetime.now()
                      iw.save(update_fields=['iw_date','updated_by', 'updated_date'])  
                   except ObjectDoesNotExist:
                      print 'id does not exist'
    return render(request,'decisions/solution_options/send_email.html',{'dec_id':dec_id, 'dec_title':dec_title, 'loggedinuser':loggedinuser,'stdec_solopt':stdec_solopt, 'stdec_scrcr':stdec_scrcr, 'stdec_evacr':stdec_evacr,'stdec_iw':stdec_iw })

def add_st_privs(request):
    return render(request, 'stakeholders/add_st_privs.html') 

'''
def handsontable(request):
    return render(request, 'decisions/handsontable.html')
'''

def add_decision(request):
    context = RequestContext(request)

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'user_email' in request.session:                                                                                                                                                                          
       user_email = request.session['user_email']
    else:
       user_email = 'not found'
    print request.session['user']
    print user_email    
    if request.method == 'POST':
        decform = DecisionForm(data=request.POST)

        if decform.is_valid():
            id = decform.save(commit=False)
            id.created_by = request.session['user']

            try: 
               d = Decisions.objects.filter(short_title = id.short_title, created_by = id.created_by).count()
               if d > 0: 
                  return render(request, 'decisions/add_decision.html',{'decform':decform,'err':'This title is already taken. Please enter a unique name.'})
            except ObjectDoesNotExist:
                print 'something wrong in add decision unique check'

            id.save()
            try:
               st = Stakeholders.objects.get(created_by = request.session['user'], email = user_email) 
               name = st.firstName + ' ' + st.lastName
               st_dec = Stakeholders_Decisions(st_id = st.id, name = name, email=st.email, dec_id = id.id,solopt_type = 'Y',scrcr_type = 'Y',evacr_type = 'Y',iw_type = 'Y',PA = 'Y', created_by = request.session['user'],created_date = datetime.datetime.now())      
               st_dec.save()
            except ObjectDoesNotExist:
                print 'stakeholder does not exist'
            #return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html') 
            return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % id.id)
        else:
            print decform.errors

    else:
        decform = DecisionForm()

    t = loader.get_template('decisions/add_decision.html')
    c = Context({'decform' :decform})
    #html = t.render({'decform': decform})
    #return HttpResponse(html)

    return render(request,'decisions/add_decision.html',{'decform':decform, 'loggedinuser':loggedinuser})
    #return render_to_string(
            #'decisions/add_decision.html',
            #{'decform': decform}, context)

def decisions_list(request):
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if loggedinuser == 'not found':
       return HttpResponseRedirect('/utility_tool/decisions/nouser_message.html')           
    else:   
       if 'user_email' in request.session: 
          user_email = request.session['user_email']
       else:
          user_email = 'not found'
       if 'idList' in request.session:   
          del request.session['idList']
       if 'dec_id' in request.session:   
          del request.session['dec_id']    
       request.session['whereamI'] = 'dec_list'   

       #f = open( '/home/amritha/costtool/documents/f.txt', 'w+' )
       #f.write('\n') 
       #f.close()

       declist = [] 
       st = Stakeholders_Decisions.objects.filter(email = user_email)
       qset = st.exclude(deleted = 'Y')

       for d in qset:
          declist.append(d.dec_id)

       for s in SharedDec.objects.filter(shared_user = loggedinuser, shared = "Y"):
          declist.append(s.dec_id)

       #get unique dec ids   
       myset = set(declist)
       #print myset
       alldecisions = Decisions.objects.filter(created_by=loggedinuser) | Decisions.objects.filter(id__in=myset)
       alldecisions = alldecisions.order_by('-id')

       return render(request,'decisions/decisions_list.html',{'alldecisions':alldecisions, 'loggedinuser':loggedinuser})  

def view_decision(request, dec_id):
    decision = Decisions.objects.get(pk=dec_id)
    return render(request,'decisions/view_decision.html',{'decision':decision})

def edit_decision(request, dec_id):
    context = RequestContext(request)
    '''if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0'''
    decision = Decisions.objects.get(pk=dec_id)

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    #if 'created_by' in request.session:
       #created_by = request.session['created_by']
    #else:
    created_by = decision.created_by

    if 'whereamI' in request.session:
        whereamI = request.session['whereamI']
    else:
        whereamI = 'dec_list'

    if request.method == 'POST':
        decform = DecisionForm(data=request.POST,instance=decision)
        if decform.is_valid():
            id = decform.save(commit=False)
            id.updated_date = datetime.datetime.now()
            id.updated_by = request.session['user']
            try: 
               d = Decisions.objects.filter(short_title = id.short_title, created_by = id.created_by).count()
               if d > 1: 
                  return render(request, 'decisions/edit_decision.html',{'decform':decform,'err':'This title is already taken. Please enter a unique name.'})
            except ObjectDoesNotExist:
               print 'something wrong in edit decision unique check'
            print id.by_when   
            id.save(update_fields=['short_title','title','name_decisionmaker', 'evidence', 'real_dec_yn', 'decision_prob','goal','target_audience','stakeholders', 'participating_stakeholders', 'potential_sources','by_when', 'updated_date','updated_by']) 
            if whereamI == 'dec_list':
               return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html')
            else:
               return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % id.id)
        else:
            print decform.errors
    else:                                                                                                                                                                                                        
        decform = DecisionForm(instance=decision)
        if loggedinuser != created_by:
           decform.fields['title'].widget.attrs['disabled'] = True
           decform.fields['name_decisionmaker'].widget.attrs['disabled'] = True
           decform.fields['short_title'].widget.attrs['disabled'] = True
           decform.fields['decision_prob'].widget.attrs['disabled'] = True
           decform.fields['evidence'].widget.attrs['disabled'] = True
           decform.fields['goal'].widget.attrs['disabled'] = True
           decform.fields['target_audience'].widget.attrs['disabled'] = True
           decform.fields['by_when'].widget.attrs['disabled'] = True
           decform.fields['stakeholders'].widget.attrs['disabled'] = True  
           decform.fields['participating_stakeholders'].widget.attrs['disabled'] = True 
           decform.fields['real_dec_yn'].widget.attrs['disabled'] = True  
           decform.fields['potential_sources'].widget.attrs['disabled'] = True
    t = loader.get_template('decisions/edit_decision.html')
    c = Context({'decform' :decform})
    return render(request,'decisions/edit_decision.html',{'dec_id':dec_id, 'decform':decform, 'whereamI':whereamI, 'loggedinuser':loggedinuser, 'created_by':created_by})

def delete_decision(request, dec_id):
    context = RequestContext(request)

    Solution_Options.objects.filter(dec_id=dec_id).delete()
    Solution_Options_Storage.objects.filter(dec_id=dec_id).delete()    
    Screening_Criteria.objects.filter(dec_id=dec_id).delete()
    Evaluation_Criteria.objects.filter(dec_id=dec_id).delete()
    Stakeholders_Decisions.objects.filter(dec_id=dec_id).delete()
    Cost_Utility.objects.filter(dec_id=dec_id).delete()
    Cost_Setup.objects.filter(dec_id=dec_id).delete()
    Decision_Made.objects.filter(dec_id=dec_id).delete()
    Evaluation_Measures.objects.filter(dec_id=dec_id).delete()
    EvaluationTable.objects.filter(dec_id=dec_id).delete()
    Importance_Scores.objects.filter(dec_id=dec_id).delete()
    MappingTable.objects.filter(dec_id=dec_id).delete()
    Scores_Setup.objects.filter(dec_id=dec_id).delete()
    SummaryTable.objects.filter(dec_id=dec_id).delete()
    IdentifyTable.objects.filter(dec_id=dec_id).delete()
    Master_Screening_Criteria.objects.filter(dec_id=dec_id).delete()
    Master_Evaluation_Criteria.objects.filter(dec_id=dec_id).delete()
    Detailed_Costs.objects.filter(dec_id=dec_id).delete()
    Decisions.objects.get(pk=dec_id).delete()                                                                                                                                                                  
    return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html') 

def share_dec(request, dec_id):
    try:
       sh = SharedDec.objects.get(dec_id=dec_id)
       sh.updated_by = request.session['user']
       sh.updated_date = datetime.datetime.now()
       sh.shared = "Y"
       sh.shared_user = "DM_Admin"
       sh.save(update_fields=['shared','shared_user','updated_by', 'updated_date'])
    except ObjectDoesNotExist:
       sh = SharedDec(dec_id = dec_id, shared_user = "DM_Admin", shared = "Y", created_by = request.session['user'], created_date = datetime.datetime.now())
       sh.save()

    dec = Decisions.objects.get(pk=dec_id)
    dec.updated_by = request.session['user']
    dec.updated_date = datetime.datetime.now()
    dec.shared = "Y"
    dec.save(update_fields=['shared','updated_by', 'updated_date'])
    return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html') 

def unshare_dec(request, dec_id):
    try:
       sh = SharedDec.objects.get(dec_id=dec_id)
       sh.updated_by = request.session['user']
       sh.updated_date = datetime.datetime.now()
       sh.shared = "N"
       sh.save(update_fields=['shared','updated_by', 'updated_date'])
       dec = Decisions.objects.get(pk=dec_id)
       dec.updated_by = request.session['user']
       dec.updated_date = datetime.datetime.now()
       dec.shared = "N"
       dec.save(update_fields=['shared','updated_by', 'updated_date'])
    except ObjectDoesNotExist:
       print 'nothing to do' 
    return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html') 

def dupl_decision(request, dec_id):
    dupl(request, dec_id, 'dupl')
    return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html') 

def download_demo(request):
    print '1'
    dupl_dec = Duplicated_DecIds.objects.all()
    dupl_dec_count = dupl_dec.count()
    print dupl_dec_count 
    if dupl_dec_count > 0: 
       for d in dupl_dec:
          dupl(request, d.dec_id_for_dupl, 'demo') 
    else: 
       print 'Duplicated Decision Ids do not exist.'
       dupl(request, 767, 'demo')
    print '4'
    return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html') 

def dupl(request, dec_id, dtype):
    if 'user_email' in request.session:                                                                                                                                                                          
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    context = RequestContext(request)
    dec = Decisions.objects.get(pk=dec_id)
    dec_orig_user = dec.created_by
    dec.created_by = request.session['user']
    dec.created_date = datetime.datetime.now()
    #if dec.short_title[:5] == "DEMO:":
    #dec.short_title = dec.short_title[6:len(dec.short_title)] + ' COPY' 
    #else:    
    dec.short_title = dec.short_title + ' COPY' 
    #str(randint(0,100))
    #aug14
    print 'dtype'
    print dtype
    if dtype == 'demo':
       dec.demoDec = 'Y'
    else:   
       dec.demoDec = None 
    dec.updated_by = None
    dec.updated_date = None
    dec.pk = None 
    dec.save()

    if dtype == 'demo':
       try:
          s2 = Stakeholders_Decisions.objects.filter(dec_id=dec_id)
          for s in s2:
              try:
                 st = Stakeholders.objects.get(created_by = request.session['user'], email = s.email)
              except ObjectDoesNotExist:
                 st_dm = Stakeholders.objects.get(created_by = 'DM_Admin', email = s.email)
                 if st_dm.email <> 'dm@admin.edu':
                    st = Stakeholders(firstName = st_dm.firstName, lastName = st_dm.lastName, email=s.email, title = st_dm.title, organisation = st_dm.organisation, created_by = request.session['user'], created_date = datetime.datetime.now())
                    st.save()

              try: 
                 st_dec = Stakeholders_Decisions.objects.get(dec_id = dec.id, email = s.email)
              except ObjectDoesNotExist:
                 if s.email == 'dm@admin.edu':
                    st_myself = Stakeholders.objects.get(created_by = request.session['user'], email = user_email)
                    name = st_myself.firstName + ' ' + st_myself.lastName
                    st_dec = Stakeholders_Decisions(st_id = st_myself.id, name = name, email=st_myself.email, dec_id = dec.id, votes = s.votes, solopt_type = s.solopt_type,scrcr_type = s.scrcr_type,evacr_type = s.evacr_type,iw_type = s.iw_type, created_by = request.session['user'],created_date = datetime.datetime.now())    
                 else:
                    name = st.firstName + ' ' + st.lastName 
                    st_dec = Stakeholders_Decisions(st_id = st.id, name = name, email=st.email, dec_id = dec.id,votes = s.votes, solopt_type = s.solopt_type,scrcr_type = s.scrcr_type,evacr_type = s.evacr_type,iw_type = s.iw_type,created_by = request.session['user'],created_date = datetime.datetime.now())      
                 st_dec.save()
       except ObjectDoesNotExist:
          print 'no stakeholders for decisions in DM_Admin'

    else:
       try:
          for s in Stakeholders_Decisions.objects.filter(dec_id=dec_id):
             s = Stakeholders_Decisions.objects.get(pk = s.id)
             s.dec_id = dec.id
             s.pk = None
             s.created_by = request.session['user']
             s.created_date = datetime.datetime.now()
             s.updated_by = None 
             s.updated_date = None 
             s.save()
       except ObjectDoesNotExist:
          print 'Stakeholders Decisions does not exist'

    try: 
       for s in Master_Screening_Criteria.objects.filter(dec_id=dec_id):
          s = Master_Screening_Criteria.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None     
          s.save()
    except ObjectDoesNotExist:
          print 'Master Screening Criteria do not exist'

    try: 
       for s in Master_Evaluation_Criteria.objects.filter(dec_id=dec_id):
          s = Master_Evaluation_Criteria.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None 
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Master Evaluation Criteria do not exist'                                                                                  

    try: 
       for s in Screening_Criteria.objects.filter(dec_id=dec_id):
          s = Screening_Criteria.objects.get(pk = s.id)
          if s.orig_scr_id is not None:
             s2 = Master_Screening_Criteria.objects.get(criterion = s.criterion, dec_id = dec.id)
             s.orig_scr_id = s2.id
          else:
             s.orig_scr_id = None
          s.dec_id = dec.id
          s.pk = None 
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Screening Criteria do not exist'                                                                                  

    try: 
       for s in Solution_Options.objects.filter(dec_id=dec_id):
          s = Solution_Options.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None 
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Solution Options do not exist' 

    try: 
       for s in Evaluation_Criteria.objects.filter(dec_id=dec_id):
          s = Evaluation_Criteria.objects.get(pk = s.id)
          if s.orig_eva_id is not None:
             s2 = Master_Evaluation_Criteria.objects.get(granular_ec = s.criterion, dec_id = dec.id)
             s.orig_eva_id = s2.id
          else:
             s.orig_eva_id = None
          s.dec_id = dec.id
          s.pk = None 
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Evaluation Criteria do not exist' 

    try: 
       for s in Importance_Scores.objects.filter(dec_id=dec_id):
          s = Importance_Scores.objects.get(pk = s.id)
          e = Evaluation_Criteria.objects.get(combined = s.criterion, dec_id = dec.id)
          s.dec_id = dec.id
          s.eva_id = e.id
          s.pk = None 
          if s.created_by == dec_orig_user:
             s.created_by = request.session['user']
          if s.email == 'dm@admin.edu':
             s.email = user_email
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Importance Scores do not exist' 

    try:
       for s in MappingTable.objects.filter(dec_id=dec_id):
          s = MappingTable.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Mapping Table does not exist'


    try:
       for s in Cost_Utility.objects.filter(dec_id=dec_id):
          s = Cost_Utility.objects.get(pk = s.id)
          o = Solution_Options.objects.get(sol_option = s.sol_option, dec_id = dec.id)
          s.opt_id = o.id 
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Cost Utility does not exist'

    try:
       for s in Cost_Setup.objects.filter(dec_id=dec_id):
          s = Cost_Setup.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Cost Setup does not exist'

    try:
       for s in IdentifyTable.objects.filter(dec_id=dec_id):
          s = IdentifyTable.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Identify Table does not exist'

    try:
       for s in EvaluationTable.objects.filter(dec_id=dec_id):
          s = EvaluationTable.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Evaluation Table does not exist'

    try:
       for s in Evaluation_Measures.objects.filter(dec_id=dec_id):
          s = Evaluation_Measures.objects.get(pk = s.id)
          e = Evaluation_Criteria.objects.get(combined = s.criterion, dec_id = dec.id)
          s.eva_id = e.id 
          o = Solution_Options.objects.get(sol_option = s.sol_option, dec_id = dec.id)
          s.opt_id = o.id 
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Evaluation Measures does not exist'

    try:
       for s in SummaryTable.objects.filter(dec_id=dec_id):
          s = SummaryTable.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Summary Table does not exist'

    try:
       for s in Detailed_Costs.objects.filter(dec_id=dec_id):
          s = Detailed_Costs.objects.get(pk = s.id)
          o = Solution_Options.objects.get(sol_option = s.sol_option, dec_id = dec.id)
          s.opt_id = o.id 
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Detailed Costs do not exist'

    justalist = []      
    try:
       for s in Decision_Made.objects.filter(dec_id=dec_id):
          s = Decision_Made.objects.get(pk = s.id)
          
          i = s.sol_option
          print 'i'
          print i
          if i == '[]':
             s.sol_option = '[]' 
          else: 
             a = i.replace('[','')
             a1 = a.replace(']','')
             print a1
             a2 = a1.replace('u','')
             b = a2.replace('L','')
             print 'b'
             print b
             for l in b.split(','):
                c = l.replace("'","")
                print c.strip()
                oldrec = Solution_Options.objects.get(id = c.strip(), dec_id = dec_id)
                newrec = Solution_Options.objects.get(sol_option = oldrec.sol_option, dec_id = dec.id)
                justalist.append(newrec.id)
             s.sol_option = justalist 
             print justalist
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()
    except ObjectDoesNotExist:
          print 'Decision Made does not exist'
    try:
       for s in Scores_Setup.objects.filter(dec_id=dec_id):
          s = Scores_Setup.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'Scores Setup does not exist'

    try:
       for s in SD_dec_file.objects.filter(dec_id=dec_id):
          s = SD_dec_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_dec_file does not exist'

    try:
       for s in SD_dec_link.objects.filter(dec_id=dec_id):
          s = SD_dec_link.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_dec_link does not exist'     
    try:
       for s in SD_st_file.objects.filter(dec_id=dec_id):
          s = SD_st_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_st_file does not exist'     
    try:
       for s in SD_st_link.objects.filter(dec_id=dec_id):
          s = SD_st_link.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_st_link does not exist'     
    try:
       for s in SD_solopt_file.objects.filter(dec_id=dec_id):
          s = SD_solopt_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_solopt_file does not exist'     
    try:
       for s in SD_solopt_link.objects.filter(dec_id=dec_id):
          s = SD_solopt_link.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_solopt_link does not exist'     
    try:
       for s in SD_scr_file.objects.filter(dec_id=dec_id):
          s = SD_scr_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_scr_file does not exist'     
    try:
       for s in SD_scr_link.objects.filter(dec_id=dec_id):
          s = SD_scr_link.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_scr_link does not exist'     
    try:
       for s in SD_mapp_file.objects.filter(dec_id=dec_id):
          s = SD_mapp_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_mapp_file does not exist'     
    try:
       for s in SD_mapp_link.objects.filter(dec_id=dec_id):
          s = SD_mapp_link.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_mapp_link does not exist'     
    try:
       for s in SD_eva_file.objects.filter(dec_id=dec_id):
          s = SD_eva_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_eva_file does not exist'     
    try:
       for s in SD_eva_link.objects.filter(dec_id=dec_id):
          s = SD_eva_link.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_eva_link does not exist'     
    try:
       for s in SD_iw_file.objects.filter(dec_id=dec_id):
          s = SD_iw_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_iw_file does not exist'     
    try:
       for s in SD_evam_file.objects.filter(dec_id=dec_id):
          s = SD_evam_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_evam_file does not exist'     
    try:
       for s in SD_evam_link.objects.filter(dec_id=dec_id):
          s = SD_evam_link.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_evam_link does not exist'     
    try:
       for s in SD_cost_file.objects.filter(dec_id=dec_id):
          s = SD_cost_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_cost_file does not exist'     
    try:
       for s in SD_cost_link.objects.filter(dec_id=dec_id):
          s = SD_cost_link.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_cost_link does not exist'     
    try:
       for s in SD_makedec_file.objects.filter(dec_id=dec_id):
          s = SD_makedec_file.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_makedec_file does not exist'     
    try:
       for s in SD_makedec_link.objects.filter(dec_id=dec_id):
          s = SD_makedec_link.objects.get(pk = s.id)
          s.dec_id = dec.id
          s.pk = None
          s.created_by = request.session['user']
          s.created_date = datetime.datetime.now()
          s.updated_by = None 
          s.updated_date = None 
          s.save()                                                                                                                          
    except ObjectDoesNotExist:
          print 'SD_makedec_file does not exist'     

    return 1      
    #return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html') 

def solutions_options_menu(request):
    return render(request,'decisions/solution_options/menu.html')

def question1(request, dec_id):
    request.session['dec_id'] = dec_id
    return render(request,'decisions/solution_options/question1.html')

def pa_setup(request):
    context = RequestContext(request)
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0 

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                       
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']                                                                                                                                                                    
    else:
       loggedinuser = 'not found'

    if 'user_email' in request.session:                                                                                                               
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    try:
        SharedDec.objects.get(dec_id = dec_id,shared_user = loggedinuser) 
        shared = 'Y'
    except ObjectDoesNotExist:
        shared = 'N'

    try:
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id, iw_type = 'Y') 
       std_count = std.exclude(email = user_email).count()     
       if std_count > 0: 
          stakeholdersNow = 'Y'   
       else:
          stakeholdersNow = 'N'  
    except ObjectDoesNotExist:
       stakeholdersNow = 'N'
    except MultipleObjectsReturned:                                                                                                           
       stakeholdersNow = 'Y'                                                                                                                          

    #return render(request,'decisions/pa_setup.html',{'dec_id':dec_id,'dec_title':dec_title, 'stakeholdersNow':stakeholdersNow})

    qset = Stakeholders_Decisions.objects.filter(dec_id=dec_id,iw_type = 'Y') 
    #qset_count = Stakeholders_Decisions.objects.filter(dec_id=dec_id, created_by=loggedinuser, iw_type = 'Y').count()
    qset_count = qset.count()
    total_votes = 10 * qset_count
    allowed_votes = 0
    MFormSet = modelformset_factory(Stakeholders_Decisions, form=VotesForm, extra=0)
    dec = Decisions.objects.get(pk=dec_id)
    something_saved = 'no' 
    if request.method == 'POST':
        votesform = MFormSet(request.POST,request.FILES,prefix="votesform" )
        if votesform.is_valid():
           id = votesform.save(commit=False)
           for recs in id:
               if recs.votes is None:
                  errtext = 'Please enter the number of votes'
                  return render(request,'decisions/pa_setup.html',{'votesform':votesform,'dec_id':dec_id, 'dec_title':dec_title, 'errtext':errtext, 'total_votes':total_votes,'total_voters':qset_count})  
               allowed_votes = recs.votes + allowed_votes
           print allowed_votes
           if allowed_votes > total_votes:
              errtext = 'The total number of votes cannot exceed '  + str(total_votes) + ', i.e., ten times the number of Stakeholders.' 
              return render(request,'decisions/pa_setup.html',{'votesform':votesform,'dec_id':dec_id, 'dec_title':dec_title, 'errtext':errtext, 'total_votes':total_votes, 'total_voters':qset_count})  
           elif allowed_votes <> total_votes:
              errtext = 'The total number of votes must be equal to ' + str(total_votes) + ', i.e., ten times the number of Stakeholders.'
              return render(request,'decisions/pa_setup.html',{'votesform':votesform,'dec_id':dec_id, 'dec_title':dec_title, 'errtext':errtext, 'total_votes':total_votes, 'total_voters':qset_count})  
           else:
               for recs in id:                                                                                                                                                                                  
                  #votes_zero = 'N' 
                  if recs.votes <> '':
                     recs.updated_by = request.session['user'] 
                     recs.updated_date = datetime.datetime.now()
                     recs.save(update_fields=['votes','updated_by', 'updated_date'])
                     something_saved ='yes'
                     #if recs.votes == 0:
                        #votes_zero = 'Y' 
               if something_saved == 'yes':  
                  dec.updated_by = request.session['user'] 
                  dec.updated_date = datetime.datetime.now()
                  dec.save(update_fields=['updated_by','updated_date']) 
               if 'submit' in request.POST:   
                  return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
               elif 'scores' in request.POST or 'next' in request.POST:
                  return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_scores.html')                    
               elif 'summary' in request.POST:
                  return HttpResponseRedirect('/utility_tool/decisions/solution_options/is_summary.html')       
               elif 'st' in request.POST:
                  return HttpResponseRedirect('/utility_tool/stakeholders/stakeholders.html')       
    else:
        votesform = MFormSet(queryset = qset,prefix="votesform")
        for form in votesform:                                                                                                                                                                                   
           form.fields['name'].widget.attrs['readonly'] = True 
           form.fields['updated_by'].widget.attrs['readonly'] = True
           if shared == "Y":
              form.fields['votes'].widget.attrs['readonly'] = True
           instance = getattr(form, 'instance', None)
           if not instance.votes:
              if instance.votes <> 0:
                 form.initial['votes'] = 10 
    return render(request,'decisions/pa_setup.html',{'dec_id':dec_id,'dec_title':dec_title, 'loggedinuser':loggedinuser, 'stakeholdersNow':stakeholdersNow, 'votesform':votesform,'total_votes':total_votes, 'total_voters':qset_count, 'shared':shared})


'''
def pa_setup(request):                                                                                                                                                                                 
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0 

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                  
    else:
       dec_title = 'not found'

    if 'user_email' in request.session:                                                                                                               
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    try: 
        setup = PA_Setup.objects.get(dec_id=dec_id)
        group_yn = setup.scores_group_yn
        votes_yn = setup.votes_yn 
    except ObjectDoesNotExist:
        #print 'error in setup'
        setup = PA_Setup(dec_id = dec_id, scores_group_yn = 'Y', votes_yn = 'N', created_date = datetime.datetime.now(), created_by = request.session['user'])
        group_yn = setup.scores_group_yn                                                                                                               
        votes_yn = setup.votes_yn
        setup.save()
    dec = Decisions.objects.get(pk=dec_id) 
    if request.method == 'POST':
       print request.POST.get('group_yn')
       print request.POST.get('votes_yn')
       if request.POST.get('group_yn') or request.POST.get('votes_yn'):
          setup.scores_group_yn = request.POST.get('group_yn')
          setup.votes_yn = request.POST.get('votes_yn') 
          setup.updated_date = datetime.datetime.now()
          setup.save(update_fields=['scores_group_yn','votes_yn','updated_date'])   
          dec.updated_by = request.session['user'] 
          dec.updated_date = datetime.datetime.now()
          dec.save(update_fields=['updated_by','updated_date'])
            if 'submit' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
            elif 'votes' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_iw_votes.html')
            elif 'scores' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_scores.html')
    return render(request,'decisions/pa_setup.html',{'dec_id':dec_id,'dec_title':dec_title, 'group_yn':group_yn, 'votes_yn':votes_yn})
'''
def flowchart(request):                                                                                                                          
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    return render(request, 'decisions/Flowchart.html',{'loggedinuser':loggedinuser}) 

def menu(request, dec_id):
    request.session['dec_id'] = dec_id

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'user_email' in request.session: 
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    if 'error_message' in request.session:
       del request.session['error_message']

    if 'whereamI2' in request.session:
       del request.session['whereamI2']

    try:
       dec = Decisions.objects.get(id=dec_id)
       request.session['dec_title'] = dec.short_title
       created_by = dec.created_by
       request.session['created_by'] = created_by
    except:
       request.session['dec_title'] = 'not found' 
       created_by = 'not found'
       request.session['created_by'] = created_by 
                                                                                                                                                                                                                 
    request.session['whereamI'] = 'menu' 

    try:
       SharedDec.objects.get(dec_id = dec_id, shared_user = loggedinuser)                                                                                                           
       shared = 'Y'
    except ObjectDoesNotExist:
       shared = 'N'

    try: 
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id)
       std_count = std.exclude(email = user_email).count()                                                                                                                                                       
       if std_count > 0:
          stakeholdersNow = 'Y'   
       else:
          stakeholdersNow = 'N' 
    except ObjectDoesNotExist:
       stakeholdersNow = 'N'
    except MultipleObjectsReturned:                                                                                                           
       stakeholdersNow = 'Y' 

    # if the logged in user has created the decision, he has access to everything
    # if the user is a stakeholder for a decision type, he has access to that type 
    try: 
       Stakeholders_Decisions.objects.get(email = user_email, dec_id = dec_id, solopt_type = 'Y')
       solopt_allowed = 'Y'
    except:
       if loggedinuser == created_by:
          solopt_allowed = 'Y'
       else: 
          solopt_allowed = 'N'

    try: 
       st1 = Stakeholders_Decisions.objects.filter(dec_id = dec_id, solopt_type = 'Y')
       st2_count = st1.exclude(email = user_email).count()
       if st2_count > 0:
          stsolopt_created = 'Y'   
       else:
          stsolopt_created = 'N'
    except ObjectDoesNotExist:
       stsolopt_created = 'N'

    try: 
       Stakeholders_Decisions.objects.get(email = user_email, dec_id = dec_id, scrcr_type = 'Y')
       scrcr_allowed = 'Y' 
    except:
       if loggedinuser == created_by:
          scrcr_allowed = 'Y'
       else: 
          scrcr_allowed = 'N'

    try: 
       st1 = Stakeholders_Decisions.objects.filter(dec_id = dec_id, scrcr_type = 'Y')
       st2_count = st1.exclude(email = user_email).count()
       if st2_count > 0:
          stscr_created = 'Y'   
       else:
          stscr_created = 'N'
    except ObjectDoesNotExist:
       stscr_created = 'N'

    try:                                                                                                                                                                                                         
       Stakeholders_Decisions.objects.get(email = user_email, dec_id = dec_id, evacr_type = 'Y')
       evacr_allowed = 'Y' 
       steva_created = 'Y'
    except:
       steva_created = 'N'
       if loggedinuser == created_by:
          evacr_allowed = 'Y'
       else: 
          evacr_allowed = 'N'  

    try: 
       st1 = Stakeholders_Decisions.objects.filter(dec_id = dec_id, evacr_type = 'Y')
       st2_count = st1.exclude(email = user_email).count()
       if st2_count > 0:
          steva_created = 'Y'   
       else:
          steva_created = 'N'
    except ObjectDoesNotExist:
       steva_created = 'N'

    try:                                                                                                                                                                                                         
       Stakeholders_Decisions.objects.get(email = user_email, dec_id = dec_id, iw_type = 'Y')
       iw_allowed = 'Y' 
       stiw_created = 'Y' 
    except:
       stiw_created = 'N'
       if loggedinuser == created_by:
          iw_allowed = 'Y'
       else: 
          iw_allowed = 'N'       
     
    try: 
       st1 = Stakeholders_Decisions.objects.filter(dec_id = dec_id, iw_type = 'Y')
       st2_count = st1.exclude(email = user_email).count()
       if st2_count > 0:
          stiw_created = 'Y'   
       else:
          stiw_created = 'N'
    except ObjectDoesNotExist:
       stiw_created = 'N'

    # if PA decides scores are decided by the group and not individually then scores screen should NOT be shown to other stakeholders
    #if group_yn == 'Y' and loggedinuser <> created_by:
    #if stiw_created == 'Y': 
       #iw_allowed = 'Y' 

    request.session['solopt_allowed'] = solopt_allowed
    request.session['scrcr_allowed'] = scrcr_allowed 
    request.session['evacr_allowed'] = evacr_allowed 
    request.session['iw_allowed'] = iw_allowed 

    try:
       solopt = Solution_Options.objects.get(dec_id=dec_id, created_by = loggedinuser, archived='N', deleted='N')
       solopt_created = 'Y'   
    except ObjectDoesNotExist:
       solopt_created = 'N'
    except MultipleObjectsReturned:                                                                                                           
       solopt_created = 'Y'     

    try: 
       soloptP = Solution_Options.objects.get(dec_id=dec_id, archived='N', deleted='N')
       soloptP_created = 'Y'   
    except ObjectDoesNotExist:
       soloptP_created = 'N'
    except MultipleObjectsReturned:                                                                                                                
       soloptP_created = 'Y'  
 
    try: 
       scr = Screening_Criteria.objects.get(dec_id=dec_id, created_by = loggedinuser)
       scr_created = 'Y' 
    except ObjectDoesNotExist:
       scr_created = 'N' 
    except MultipleObjectsReturned:                                                                                                                
       scr_created = 'Y' 

    if scr_created == 'N':
       scrrec = Screening_Criteria.objects.raw("SELECT id FROM utility_tool_screening_criteria WHERE dec_id = %s AND fieldname IS NOT NULL AND criterion != '' and updated_by = %s", [dec_id, loggedinuser])    
       if len(list(scrrec)) > 0:
          scr_created = 'Y'  
       else:
          scr_created = 'N'  

    try:
       scrP = Screening_Criteria.objects.get(dec_id=dec_id)
       scrP_created = 'Y'
    except ObjectDoesNotExist:
       scrP_created = 'N'
    except MultipleObjectsReturned:
       scrP_created = 'Y' 

    try: 
       mapp = MappingTable.objects.get(dec_id=dec_id) 
       mapp_created = 'Y'  
    except ObjectDoesNotExist:
       mapp_created = 'N'  

    try: 
       eva = Evaluation_Criteria.objects.get(dec_id=dec_id,created_by = loggedinuser) 
       if eva.deleted == 'Y':
          eva_created = 'N'
       else:
          eva_created = 'Y' 
    except ObjectDoesNotExist:
       eva_created = 'N' 
    except MultipleObjectsReturned:                                                                                                                
       evaC = Evaluation_Criteria.objects.filter(dec_id=dec_id,created_by = loggedinuser).exclude(deleted = 'Y').count() 
       if evaC > 0:
          eva_created = 'Y'
       else:
          eva_created = 'N' 

    if eva_created == 'N':
       evarec = Evaluation_Criteria.objects.raw("SELECT id FROM utility_tool_evaluation_criteria WHERE dec_id = %s AND fieldname IS NOT NULL AND criterion != ''  and updated_by = %s", [dec_id, loggedinuser])
       if len(list(evarec)) > 0:
          eva_created = 'Y'  
       else:
          eva_created = 'N'  

    try: 
       evaP = Evaluation_Criteria.objects.get(dec_id=dec_id)
       if evaP.deleted == 'Y':
          evaP_created = 'N'
       else:
          evaP_created = 'Y' 
    except ObjectDoesNotExist:
       evaP_created = 'N' 
    except MultipleObjectsReturned:                                                                                                                
       evaC = Evaluation_Criteria.objects.filter(dec_id=dec_id).exclude(deleted = 'Y').count() 
       if evaC > 0:
          evaP_created = 'Y'
       else:
          evaP_created = 'N'   
 
    evam_created = 'N'   
    try:
       qset = Evaluation_Measures.objects.get(dec_id=dec_id) 
       if qset.measure is None or qset.unit is None or qset.lowest_value is None or qset.highest_value is None or qset.higher_better is None or qset.option_value is None:
          evam_created = 'N'
       else:   
          evam_created = 'Y' 
    except ObjectDoesNotExist:
       evam_created = 'N' 
    except MultipleObjectsReturned:                                                                                                                
       q1 = Evaluation_Measures.objects.filter(dec_id=dec_id)
       q2 = q1.exclude(archived = 'Y')
       q = q2.exclude(deleted = 'Y')   
       for qset in q: 
           if qset.measure is None or qset.unit is None or qset.lowest_value is None or qset.highest_value is None or qset.higher_better is None or qset.option_value is None:
              evam_created = 'N'
           else:   
              evam_created = 'Y'  
              break
    
    try: 
       iw = Importance_Scores.objects.get(dec_id = dec_id,created_by = loggedinuser)
       if iw.deleted == 'Y':
          iw_created = 'N'
       else: 
          iw_created = 'Y' 
    except ObjectDoesNotExist:
       iw_created = 'N' 
    except MultipleObjectsReturned:                                                                                                                
       iwc = Importance_Scores.objects.filter(dec_id=dec_id,created_by = loggedinuser).exclude(deleted = 'Y').count()                                                                                                                   
       if iwc > 0: 
          iw_created = 'Y'
       else:
          iw_created = 'N'  

    try: 
       iwP = Importance_Scores.objects.get(dec_id = dec_id)
       iwP_created = 'Y' 
    except ObjectDoesNotExist:
       iwP_created = 'N' 
    except MultipleObjectsReturned:                                                                                                                
       iwc = Importance_Scores.objects.filter(dec_id=dec_id).exclude(deleted = 'Y').count()                                                                                            
       if iwc > 0: 
          iwP_created = 'Y'
       else:
          iwP_created = 'N'  
    try:
       cu = Cost_Utility.objects.get(dec_id = dec_id)
       cu_created = 'Y'
    except ObjectDoesNotExist:
       cu_created = 'N'
    except MultipleObjectsReturned:
       cu_created = 'Y'

    try: 
       cost = Cost_Setup.objects.get(dec_id = dec_id)                                                                                                                                                            
       cost_created = 'Y'
    except ObjectDoesNotExist:
       cost_created = 'N'
    except MultipleObjectsReturned:
       cost_created = 'Y'

    try:
       decmade = Decision_Made.objects.get(dec_id = dec_id)
       decmade_created = 'Y'
    except ObjectDoesNotExist:
       decmade_created = 'N'
    except MultipleObjectsReturned:
       decmade_created = 'Y'
    #print iw_allowed 
    #print iw_created
    #print setup_created
    #print created_by
    #print loggedinuser
    return render(request,'decisions/menu.html',{'dec_title':request.session['dec_title'], 'dec_id':dec_id, 'loggedinuser':loggedinuser, 'created_by' :created_by, 'solopt_allowed':solopt_allowed, 'scrcr_allowed':scrcr_allowed, 'evacr_allowed': evacr_allowed, 'iw_allowed':iw_allowed, 'solopt_created':solopt_created, 'scr_created': scr_created, 'eva_created':eva_created, 'evam_created':evam_created, 'iw_created':iw_created, 'decmade_created':decmade_created, 'cu_created':cu_created, 'cost_created':cost_created, 'stsolopt_created':stsolopt_created, 'stscr_created':stscr_created, 'steva_created':steva_created, 'stiw_created':stiw_created, 'stakeholdersNow':stakeholdersNow, 'mapp_created':mapp_created, 'soloptP_created':soloptP_created, 'scrP_created':scrP_created, 'evaP_created':evaP_created, 'iwP_created':iwP_created, 'shared':shared})

def question2(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if request.method == 'POST':
        print request.POST 
        if 'whoEntersSO' in request.POST:
            request.session['whoEntersSO'] = request.POST['whoEntersSO']
            request.session['listType'] = request.POST['listType']
            '''try:
               import smtplib
               smtp = smtplib.SMTP('smtp.gmail.com',587)
               smtp.ehlo()
               smtp.starttls()
               smtp.login('amrithany@gmail.com', 'Daff1911')
               smtp.sendmail('amrithany@gmail.com', 'amritha_mm@yahoo.com', 'test message from cost utility')
               smtp.quit()
            except smtplib.SMTPException, error:
               #return render_to_response('login/forgot.html',{'registerform':registerform,'err':str(error)}, context) 
               return HttpResponse('failure')
            '''
            return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_st_solopt.html')       
        else: 
            return HttpResponse('FAIL!!!!!')
     
    return render(request,'decisions/solution_options/question2.html', {'dec_id':dec_id})

def question3(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    return render(request,'decisions/solution_options/question3.html', {'dec_id':dec_id})

def identify_st(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                  
    else:
       dec_title = 'not found'

    return render(request,'decisions/solution_options/identify_st.html', {'dec_id':dec_id, 'dec_title':dec_title})

def guidance(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    request.session['whereamI']  = 0
    if loggedinuser == 'not found':
       return HttpResponseRedirect('/utility_tool/decisions/nouser_message.html')           
    else:    
       return render(request,'decisions/solution_options/guidance.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser})

def add_supp_doc(request):
    context = RequestContext(request)
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0 
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'

    try:
       SharedDec.objects.get(dec_id = dec_id, shared_user = loggedinuser)                                                                                                               
       shared = 'Y'
    except ObjectDoesNotExist:
       shared = 'N' 

    MFormSet = modelformset_factory(SD_dec_file, form=SDForm_dec_file, extra=1)
    MFormSet2 = modelformset_factory(SD_dec_link, form=SDForm_dec_link, extra=1)
    MFormSet3 = modelformset_factory(SD_st_file, form=SDForm_st_file, extra=1)
    MFormSet4 = modelformset_factory(SD_st_link, form=SDForm_st_link, extra=1)
    MFormSet5 = modelformset_factory(SD_solopt_file, form=SDForm_solopt_file, extra=1)
    MFormSet6 = modelformset_factory(SD_solopt_link, form=SDForm_solopt_link, extra=1)
    MFormSet7 = modelformset_factory(SD_scr_file, form=SDForm_scr_file, extra=1)
    MFormSet8 = modelformset_factory(SD_scr_link, form=SDForm_scr_link, extra=1)
    MFormSet9 = modelformset_factory(SD_mapp_file, form=SDForm_mapp_file, extra=1)
    MFormSet10 = modelformset_factory(SD_mapp_link, form=SDForm_mapp_link, extra=1)
    MFormSet11 = modelformset_factory(SD_eva_file, form=SDForm_eva_file, extra=1)
    MFormSet12 = modelformset_factory(SD_eva_link, form=SDForm_eva_link, extra=1)
    MFormSet13 = modelformset_factory(SD_iw_file, form=SDForm_iw_file, extra=1)
    MFormSet14 = modelformset_factory(SD_iw_link, form=SDForm_iw_link, extra=1)
    MFormSet15 = modelformset_factory(SD_evam_file, form=SDForm_evam_file, extra=1)
    MFormSet16 = modelformset_factory(SD_evam_link, form=SDForm_evam_link, extra=1)
    MFormSet17 = modelformset_factory(SD_cost_file, form=SDForm_cost_file, extra=1)
    MFormSet18 = modelformset_factory(SD_cost_link, form=SDForm_cost_link, extra=1)
    MFormSet19 = modelformset_factory(SD_makedec_file, form=SDForm_makedec_file, extra=1)
    MFormSet20 = modelformset_factory(SD_makedec_link, form=SDForm_makedec_link, extra=1)

    dec = Decisions.objects.get(pk=dec_id)   
    something_saved = 'no'
    if request.method == 'POST':
        sdform1 = MFormSet(request.POST,request.FILES, prefix="sdform1" )
        sdform2 = MFormSet2(request.POST,request.FILES, prefix="sdform2" )
        sdform3 = MFormSet3(request.POST,request.FILES, prefix="sdform3" )
        sdform4 = MFormSet4(request.POST,request.FILES, prefix="sdform4" )
        sdform5 = MFormSet5(request.POST,request.FILES, prefix="sdform5" )
        sdform6 = MFormSet6(request.POST,request.FILES, prefix="sdform6" )
        sdform7 = MFormSet7(request.POST,request.FILES, prefix="sdform7" )
        sdform8 = MFormSet8(request.POST,request.FILES, prefix="sdform8" )
        sdform9 = MFormSet9(request.POST,request.FILES, prefix="sdform9" )
        sdform10 = MFormSet10(request.POST,request.FILES, prefix="sdform10" )
        sdform11 = MFormSet11(request.POST,request.FILES, prefix="sdform11" )
        sdform12 = MFormSet12(request.POST,request.FILES, prefix="sdform12" )
        sdform13 = MFormSet13(request.POST,request.FILES, prefix="sdform13" )
        sdform14 = MFormSet14(request.POST,request.FILES, prefix="sdform14" )
        sdform15 = MFormSet15(request.POST,request.FILES, prefix="sdform15" )
        sdform16 = MFormSet16(request.POST,request.FILES, prefix="sdform16" )
        sdform17 = MFormSet17(request.POST,request.FILES, prefix="sdform17" )
        sdform18 = MFormSet18(request.POST,request.FILES, prefix="sdform18" )
        sdform19 = MFormSet19(request.POST,request.FILES, prefix="sdform19" )
        sdform20 = MFormSet20(request.POST,request.FILES, prefix="sdform20" )

        if sdform1.is_valid() and sdform2.is_valid() and sdform3.is_valid() and sdform4.is_valid() and sdform5.is_valid() and sdform6.is_valid() and sdform7.is_valid() and sdform8.is_valid() and sdform9.is_valid() and sdform10.is_valid() and sdform11.is_valid() and sdform12.is_valid() and sdform13.is_valid() and sdform14.is_valid() and sdform15.is_valid() and sdform16.is_valid() and sdform17.is_valid() and sdform18.is_valid() and sdform19.is_valid() and sdform20.is_valid():
           id = sdform1.save(commit=False)
           for recs in id:
               print 'add supp doc'
               print recs.filename                                                                                                                      
               print recs.file_attachment  
               if recs.filename == '' and recs.file_attachment == '':
                  recs.delete()
               else:
                  recs.dec_id = dec_id
                  recs.created_by = request.session['user']
                  recs.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs.save()
           id2 = sdform2.save(commit=False)
           for recs2 in id2:
               if recs2.linkname == '' and recs2.link == '':
                  recs2.delete()
               else:
                  recs2.dec_id = dec_id     
                  recs2.created_by = request.session['user']
                  recs2.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs2.save()
           id3 = sdform3.save(commit=False)
           for recs3 in id3:
               if recs3.filename == '' and recs3.file_attachment == '':
                  recs3.delete()
               else:
                  recs3.dec_id = dec_id
                  recs3.created_by = request.session['user']
                  recs3.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs3.save()
           id4 = sdform4.save(commit=False)
           for recs4 in id4:
               if recs4.linkname == '' and recs4.link == '':
                  recs4.delete()
               else:
                  recs4.dec_id = dec_id     
                  recs4.created_by = request.session['user']
                  recs4.created_date = datetime.datetime.now()
                  something_saved ='yes'                                                                                                                                                                            
                  recs4.save()
           id5 = sdform5.save(commit=False)
           for recs5 in id5:
               if recs5.filename == '' and recs5.file_attachment == '':
                  recs5.delete()
               else:
                  recs5.dec_id = dec_id
                  recs5.created_by = request.session['user']
                  recs5.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs5.save()
           id6 = sdform6.save(commit=False)
           for recs6 in id6:
               if recs6.linkname == '' and recs6.link == '':
                  recs6.delete()
               else:
                  recs6.dec_id = dec_id     
                  recs6.created_by = request.session['user']
                  recs6.created_date = datetime.datetime.now()
                  something_saved ='yes'                                                                                                                                                                            
                  recs6.save()
           id7 = sdform7.save(commit=False)
           for recs7 in id7:
               if recs7.filename == '' and recs7.file_attachment == '':
                  recs7.delete()
               else:
                  recs7.dec_id = dec_id
                  recs7.created_by = request.session['user']
                  recs7.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs7.save()
           id8 = sdform8.save(commit=False)
           for recs8 in id8:
               if recs8.linkname == '' and recs8.link == '':
                  recs8.delete()
               else:
                  recs8.dec_id = dec_id     
                  recs8.created_by = request.session['user']
                  recs8.created_date = datetime.datetime.now()
                  something_saved ='yes'                                                                                                                                                                            
                  recs8.save()
           id9 = sdform9.save(commit=False)
           for recs9 in id9:
               if recs9.filename == '' and recs9.file_attachment == '':
                  recs9.delete()
               else:
                  recs9.dec_id = dec_id
                  recs9.created_by = request.session['user']
                  recs9.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs9.save()
           id10 = sdform10.save(commit=False)
           for recs10 in id10:
               if recs10.linkname == '' and recs10.link == '':
                  recs10.delete()
               else:
                  recs10.dec_id = dec_id     
                  recs10.created_by = request.session['user']
                  recs10.created_date = datetime.datetime.now()
                  something_saved ='yes'                                                                                                                                                                            
                  recs10.save()
           id11 = sdform11.save(commit=False)
           for recs11 in id11:
               if recs11.filename == '' and recs11.file_attachment == '':
                  recs11.delete()
               else:
                  recs11.dec_id = dec_id
                  recs11.created_by = request.session['user']
                  recs11.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs11.save()
           id12 = sdform12.save(commit=False)
           for recs12 in id12:
               if recs12.linkname == '' and recs12.link == '':
                  recs12.delete()
               else:
                  recs12.dec_id = dec_id     
                  recs12.created_by = request.session['user']
                  recs12.created_date = datetime.datetime.now()
                  something_saved ='yes'                                                                                                                                                                            
                  recs12.save()
           id13 = sdform13.save(commit=False)
           for recs13 in id13:
               if recs13.filename == '' and recs13.file_attachment == '':
                  recs13.delete()
               else:
                  recs13.dec_id = dec_id
                  recs13.created_by = request.session['user']
                  recs13.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs13.save()
           id14 = sdform14.save(commit=False)
           for recs14 in id14:
               if recs14.linkname == '' and recs14.link == '':
                  recs14.delete()
               else:
                  recs14.dec_id = dec_id     
                  recs14.created_by = request.session['user']
                  recs14.created_date = datetime.datetime.now()
                  something_saved ='yes'                                                                                                                                                                            
                  recs14.save()
           id15 = sdform15.save(commit=False)
           for recs15 in id15:
               if recs15.filename == '' and recs15.file_attachment == '':
                  recs15.delete()
               else:
                  recs15.dec_id = dec_id
                  recs15.created_by = request.session['user']
                  recs15.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs15.save()
           id16 = sdform16.save(commit=False)
           for recs16 in id16:
               if recs16.linkname == '' and recs16.link == '':
                  recs16.delete()
               else:
                  recs16.dec_id = dec_id     
                  recs16.created_by = request.session['user']
                  recs16.created_date = datetime.datetime.now()
                  something_saved ='yes'                                                                                                                                                                            
                  recs16.save()
           id17 = sdform17.save(commit=False)
           for recs17 in id17:
               if recs17.filename == '' and recs17.file_attachment == '':
                  recs17.delete()
               else:
                  recs17.dec_id = dec_id
                  recs17.created_by = request.session['user']
                  recs17.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs17.save()
           id18 = sdform18.save(commit=False)
           for recs18 in id18:
               if recs18.linkname == '' and recs18.link == '':
                  recs18.delete()
               else:
                  recs18.dec_id = dec_id     
                  recs18.created_by = request.session['user']
                  recs18.created_date = datetime.datetime.now()
                  something_saved ='yes'                                                                                                                                                                            
                  recs18.save()
           id19 = sdform19.save(commit=False)
           for recs19 in id19:
               if recs19.filename == '' and recs19.file_attachment == '':
                  recs19.delete()
               else:
                  recs19.dec_id = dec_id
                  recs19.created_by = request.session['user']
                  recs19.created_date = datetime.datetime.now()
                  something_saved ='yes'
                  recs19.save()
           id20 = sdform20.save(commit=False)
           for recs20 in id20:
               if recs20.linkname == '' and recs20.link == '':
                  recs20.delete()
               else:
                  recs20.dec_id = dec_id     
                  recs20.created_by = request.session['user']
                  recs20.created_date = datetime.datetime.now()
                  something_saved ='yes'                                                                                                                                                                            
                  recs20.save()
           if something_saved == 'yes':  
              dec.updated_by = request.session['user'] 
              dec.updated_date = datetime.datetime.now()
              dec.save(update_fields=['updated_by','updated_date'])
           if 'submit' in request.POST:   
              return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)  
           else:
              return HttpResponseRedirect('/utility_tool/decisions/add_supp_doc.html') 
    else:
        qset = SD_dec_file.objects.filter(dec_id=dec_id)
        sdform1 = MFormSet(queryset=qset,prefix="sdform1" )
        qset2 = SD_dec_link.objects.filter(dec_id=dec_id)                                                                                                                                                         
        sdform2 = MFormSet2(queryset=qset2,prefix="sdform2" )
        
        qset3 = SD_st_file.objects.filter(dec_id=dec_id)
        sdform3 = MFormSet3(queryset=qset3,prefix="sdform3" )
        qset4 = SD_st_link.objects.filter(dec_id=dec_id)
        sdform4 = MFormSet4(queryset=qset4,prefix="sdform4" )

        qset5 = SD_solopt_file.objects.filter(dec_id=dec_id)
        sdform5 = MFormSet5(queryset=qset5,prefix="sdform5" )
        qset6 = SD_solopt_link.objects.filter(dec_id=dec_id)
        sdform6 = MFormSet6(queryset=qset6,prefix="sdform6" )
        qset7 = SD_scr_file.objects.filter(dec_id=dec_id)
        sdform7 = MFormSet7(queryset=qset7,prefix="sdform7" )
        qset8 = SD_scr_link.objects.filter(dec_id=dec_id)
        sdform8 = MFormSet8(queryset=qset8,prefix="sdform8" )
        qset9 = SD_mapp_file.objects.filter(dec_id=dec_id)
        sdform9 = MFormSet9(queryset=qset9,prefix="sdform9" )
        qset10 = SD_mapp_link.objects.filter(dec_id=dec_id)
        sdform10 = MFormSet10(queryset=qset10,prefix="sdform10" )
        qset11 = SD_eva_file.objects.filter(dec_id=dec_id)
        sdform11 = MFormSet11(queryset=qset11,prefix="sdform11" )
        qset12 = SD_eva_link.objects.filter(dec_id=dec_id)
        sdform12 = MFormSet12(queryset=qset12,prefix="sdform12" )
        qset13 = SD_iw_file.objects.filter(dec_id=dec_id)
        sdform13 = MFormSet13(queryset=qset13,prefix="sdform13" )
        qset14 = SD_iw_link.objects.filter(dec_id=dec_id)
        sdform14 = MFormSet14(queryset=qset14,prefix="sdform14" )

        qset15 = SD_evam_file.objects.filter(dec_id=dec_id)
        sdform15 = MFormSet15(queryset=qset15,prefix="sdform15" )
        qset16 = SD_evam_link.objects.filter(dec_id=dec_id)     
        sdform16 = MFormSet16(queryset=qset16,prefix="sdform16" )
        qset17 = SD_cost_file.objects.filter(dec_id=dec_id)
        sdform17 = MFormSet17(queryset=qset17,prefix="sdform17" )
        qset18 = SD_cost_link.objects.filter(dec_id=dec_id)     
        sdform18 = MFormSet18(queryset=qset18,prefix="sdform18" )
        qset19 = SD_makedec_file.objects.filter(dec_id=dec_id)
        sdform19 = MFormSet19(queryset=qset19,prefix="sdform19" )
        qset20 = SD_makedec_link.objects.filter(dec_id=dec_id)     
        sdform20 = MFormSet20(queryset=qset20,prefix="sdform20" )

    return render(request,'decisions/add_supp_doc.html',{'sdform1':sdform1,'sdform2':sdform2,'sdform3':sdform3,'sdform4':sdform4,'sdform5':sdform5,'sdform6':sdform6,'sdform7':sdform7,'sdform8':sdform8,'sdform9':sdform9,'sdform10':sdform10,'sdform11':sdform11,'sdform12':sdform12,'sdform13':sdform13,'sdform14':sdform14,'sdform15':sdform15,'sdform16':sdform16,'sdform17':sdform17,'sdform18':sdform18, 'sdform19':sdform19,'sdform20':sdform20,'dec_id':dec_id, 'dec_title':dec_title, 'loggedinuser':loggedinuser, 'shared':shared})



def add_solopt_det(request):
    context = RequestContext(request)
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0 

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                  
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'

    #if solopt_allowed == 'Y' and scrcr_allowed == 'Y': 
    if loggedinuser == created_by:
       mapping_allowed = 'Y'
    else:
       mapping_allowed = 'N'

    if 'user_email' in request.session: 
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    try: 
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id)
       std_count = std.exclude(email = user_email).count()                                                                                                                                                       
       if std_count > 0: 
          stakeholdersNow = 'Y'   
       else:
          stakeholdersNow = 'N'  
    except ObjectDoesNotExist:
       stakeholdersNow = 'N'
    except MultipleObjectsReturned:                                                                                                                
       stakeholdersNow = 'Y'  
    
    try:
       SharedDec.objects.get(dec_id = dec_id, shared_user = loggedinuser)
       shared = 'Y'
    except ObjectDoesNotExist:
       shared = 'N'

    print 'SHARED'
    print shared 
    sol_perm = Stakeholders_Decisions.objects.filter(dec_id=dec_id, email = user_email)
    sol_permC = sol_perm.exclude(solopt_type = 'Y').count()
    print 'SOL_PERMC'
    print sol_permC
    scr_count = Screening_Criteria.objects.filter(dec_id=dec_id).count()
    MFormSet = modelformset_factory(Solution_Options, form=SolOptForm2, extra=6)
    dec = Decisions.objects.get(pk=dec_id)   
    something_saved = 'no'
    if request.method == 'POST':
        soloptform = MFormSet(request.POST,request.FILES, prefix="soloptform" )
        if soloptform.is_valid():
           id = soloptform.save(commit=False)
           for recs in id:
              if recs.sol_option <> '':   
                 recs.dec_id = request.session['dec_id']  
                 if recs.deleted == 'Y': 
                    try:
                       solopt_check = Solution_Options.objects.get(sol_option = recs.sol_option + 'has been deleted', dec_id=dec_id)
                       recs.sol_option = recs.sol_option + 'has been deleted' + str(randint(0,100))  
                    except ObjectDoesNotExist:
                       recs.sol_option = recs.sol_option + 'has been deleted'
                    recs.updated_by = request.session['user'] 
                    recs.updated_date = datetime.datetime.now()
                    recs.save()
                    try:
                       Evaluation_Measures.objects.filter(opt_id = recs.id, dec_id = dec_id).delete()
                    except ObjectDoesNotExist:
                       print 'evaluation measures do not exist'
                    try:
                       Cost_Utility.objects.get(opt_id = recs.id, dec_id = dec_id).delete()
                    except ObjectDoesNotExist:
                       print 'cost utility does not exist'  
                    try:
                       Detailed_Costs.objects.get(opt_id = recs.id, dec_id = dec_id).delete()
                    except ObjectDoesNotExist:
                       print 'detailed costs do not exist'
                 else:
                    recs.deleted = 'N'
                    if recs.source == '':
                       return render(request,'decisions/solution_options/add_solopt_det.html',{'soloptform':soloptform,'dec_id':dec_id, 'dec_title':dec_title, 'mapping_allowed': mapping_allowed, 'scr_count':scr_count, 'err':'Please enter the Source of the Option entered.'})  
                    try:
                       ss = Solution_Options.objects.get(id = recs.id)
                       screated_by = ss.created_by
                       supdated_by = request.session['user'] 
                    except:
                       screated_by = request.session['user']  
                       supdated_by = ''

                    if recs.archived == 'Y':
                       recs.archived_by = request.session['user'] 
                       recs.archived_date = datetime.datetime.now()
                       recs.created_by = screated_by 
                       recs.updated_by = supdated_by 
                       print 'why am i in here'
                       for e in Evaluation_Measures.objects.filter(opt_id = recs.id):
                           e.archived = 'Y'
                           e.sol_option = recs.sol_option
                           e.updated_by = request.session['user'] 
                           e.updated_date = datetime.datetime.now()
                           e.save(update_fields=['archived','sol_option','updated_by','updated_date'])
                       try:    
                          c =  Cost_Utility.objects.get(opt_id = recs.id)
                          c.archived = 'Y'
                          c.sol_option = recs.sol_option
                          c.updated_by = request.session['user'] 
                          c.updated_date = datetime.datetime.now()
                          c.save(update_fields=['archived','sol_option','updated_by','updated_date'])      
                       except:
                          print 'does not exist yet' 
                    else:
                       recs.archived = 'N'
                       if screated_by == '':
                          recs.created_by = request.session['user'] 
                          recs.created_date = datetime.datetime.now()
                       else:   
                          recs.created_by = screated_by
                          recs.updated_by = request.session['user'] 
                          recs.updated_date = datetime.datetime.now()
                       for e in Evaluation_Measures.objects.filter(opt_id = recs.id):
                          e.archived = 'N'
                          e.sol_option = recs.sol_option
                          e.updated_by = request.session['user'] 
                          e.updated_date = datetime.datetime.now()
                          e.save(update_fields=['archived','sol_option','updated_by','updated_date'])
                       try:    
                          c =  Cost_Utility.objects.get(opt_id = recs.id)
                          c.archived = 'N'
                          c.sol_option = recs.sol_option
                          c.updated_by = request.session['user'] 
                          c.updated_date = datetime.datetime.now()
                          c.save(update_fields=['archived','sol_option','updated_by','updated_date'])           
                       except:
                          print 'XXX does not exist yet' 
                    something_saved ='yes'
                    request.session['error_message'] = ''
                    recs.save()
                    #print recs.id
           
           try: 
              dupes = Solution_Options.objects.values('sol_option').annotate(Count('id')).order_by().filter(dec_id = dec_id, id__count__gt=1)
              for item in dupes:
                  if item['id__count'] > 1:
                     request.session['error_message'] = '"' + item['sol_option'] + '" has already been suggested as a Solution Option. Please remove it from the list below.'   
                     return redirect('/utility_tool/decisions/solution_options/add_solopt_det.html',{'soloptform':soloptform,'dec_id':dec_id,'dec_title':dec_title, 'mapping_allowed': mapping_allowed, 'scr_count':scr_count,'err':request.session['error_message']})            
           except ObjectDoesNotExist:                                                                                                                                                                        
                 print 'solution option does not exist' 
           if something_saved == 'yes':  
              dec.updated_by = request.session['user'] 
              dec.updated_date = datetime.datetime.now()
              dec.save(update_fields=['updated_by','updated_date'])
           if 'submit' in request.POST:
              return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
           elif 'map' in request.POST:
              '''
              if something_saved == 'no':
                 err = 'You cannot Map Solution Options to Screening Criteria until you have entered at least one option.'
                 return render(request,'decisions/solution_options/add_solopt_det.html',{'err':err, 'soloptform':soloptform,'dec_id':dec_id, 'dec_title':dec_title, 'mapping_allowed': mapping_allowed, 'scr_count':scr_count })
              else:
              '''
              return HttpResponseRedirect('/utility_tool/decisions/solution_options/handsontable.html')
           elif 'scrcr' in request.POST:
              return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_scr_criteria.html')
           elif 'archived' in request.POST:
              return HttpResponseRedirect('/utility_tool/decisions/solution_options/solopt_archive.html')
           elif 'st' in request.POST:
              if stakeholdersNow == 'Y':
                 return HttpResponseRedirect('/utility_tool/decisions/solution_options/assign_tasks.html') 
              else:
                 return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_st_all.html') 
           elif 'link' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/solution_options/link.html')
           elif 'import' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/solution_options/solopt_storage.html')    
        else:
            print soloptform.errors
    else:
        qset = Solution_Options.objects.filter(dec_id=dec_id, archived='N', deleted='N')
        soloptform = MFormSet(queryset=qset,prefix="soloptform" )
        for form in soloptform:
            form.fields['created_by'].widget.attrs['readonly'] = True
            form.fields['updated_by'].widget.attrs['readonly'] = True
            if loggedinuser != created_by:
               form.fields['archived'].widget.attrs['disabled'] = True
               form.fields['deleted'].widget.attrs['disabled'] = True
               if sol_permC > 0 or shared == 'Y':
                  form.fields['sol_option'].widget.attrs['disabled'] = True 
                  form.fields['source'].widget.attrs['disabled'] = True 
                  form.fields['option_details'].widget.attrs['disabled'] = True 
                  form.fields['filename1'].widget.attrs['disabled'] = True 
                  form.fields['file_attachment1'].widget.attrs['disabled'] = True 
                  form.fields['filename2'].widget.attrs['disabled'] = True 
                  form.fields['file_attachment2'].widget.attrs['disabled'] = True 
                  form.fields['filename3'].widget.attrs['disabled'] = True 
                  form.fields['file_attachment3'].widget.attrs['disabled'] = True 
                  form.fields['filename4'].widget.attrs['disabled'] = True 
                  form.fields['file_attachment4'].widget.attrs['disabled'] = True 
                  form.fields['linkname1'].widget.attrs['disabled'] = True 
                  form.fields['link1'].widget.attrs['disabled'] = True 
                  form.fields['linkname2'].widget.attrs['disabled'] = True 
                  form.fields['link2'].widget.attrs['disabled'] = True 
                  form.fields['linkname3'].widget.attrs['disabled'] = True 
                  form.fields['link3'].widget.attrs['disabled'] = True 
                  form.fields['linkname4'].widget.attrs['disabled'] = True 
                  form.fields['link4'].widget.attrs['disabled'] = True 

    if 'error_message' in request.session: 
       return render(request,'decisions/solution_options/add_solopt_det.html',{'soloptform':soloptform,'dec_id':dec_id, 'dec_title':dec_title, 'mapping_allowed': mapping_allowed, 'scr_count':scr_count,'created_by':created_by, 'loggedinuser':loggedinuser, 'err':request.session['error_message'], 'sol_permC':sol_permC, 'shared': shared})
    else:
       return render(request,'decisions/solution_options/add_solopt_det.html',{'soloptform':soloptform,'dec_id':dec_id, 'dec_title':dec_title, 'mapping_allowed': mapping_allowed, 'scr_count':scr_count,'created_by':created_by, 'loggedinuser':loggedinuser, 'sol_permC':sol_permC, 'shared':shared})

 
def solopt_archive(request):
    context = RequestContext(request)

    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                  
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'

    MFormSet = modelformset_factory(Solution_Options, form=SolOptArchive, extra=0)
    dec = Decisions.objects.get(pk=dec_id) 
    something_saved = 'no'
    qset_count = 0
    button_shown = 'no'

    if request.method == 'POST':
        soloptform = MFormSet(request.POST,request.FILES, prefix="soloptform" )
        if soloptform.is_valid():
           id = soloptform.save(commit=False)
           for recs in id:
               if recs.unarchived == 'Y':
                  recs.archived = 'N'
                  recs.unarchived = ''
                  recs.unarchived_by = request.session['user'] 
                  recs.unarchived_date = datetime.datetime.now()
                  something_saved = 'yes'
                  recs.save(update_fields=['archived','unarchived','unarchived_by','unarchived_date'])
           if something_saved == 'yes':  
              dec.updated_by = request.session['user'] 
              dec.updated_date = datetime.datetime.now()
              dec.save(update_fields=['updated_by','updated_date'])   
           #if 'submit' in request.POST:
               #return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
           #elif 'solopt' in request.POST:
           return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_solopt_det.html')
        else:
            print soloptform.errors
    else:
        qset = Solution_Options.objects.filter(dec_id=dec_id, archived='Y')
        qset_count = Solution_Options.objects.filter(dec_id=dec_id, archived='Y').count()
        if qset_count == 0:
           button_shown = 'no'
        else:   
           button_shown = 'yes' 
        soloptform = MFormSet(queryset=qset,prefix="soloptform" )
        for form in soloptform:
            if loggedinuser != created_by:
               form.fields['unarchived'].widget.attrs['disabled'] = True

    return render(request,'decisions/solution_options/solopt_archive.html',{'soloptform':soloptform,'dec_id':dec_id, 'dec_title':dec_title, 'button_shown':button_shown, 'loggedinuser':loggedinuser })


def view_solopt_det(request, dec_id):
    context = RequestContext(request)                                                                                                                                                                            
    MFormSet = modelformset_factory(Solution_Options, form=SolOptView)
    try:
       dec = Decisions.objects.get(id=dec_id)
       dec_title = dec.title
    except:
       dec_title = 'not found' 
        
    qset = Solution_Options.objects.filter(dec_id=dec_id, archived='N', deleted='N')
    soloptform = MFormSet(queryset=qset,prefix="soloptform" )
    return render(request,'decisions/solution_options/view_solopt_det.html',{'soloptform':soloptform, 'dec_title':dec_title})

def xlsx1(request):
    with open('/home/amritha/costutility/documents/Solution Options template.xlsx', 'r') as xlsx:                                           
       response = HttpResponse(xlsx.read(), content_type='application/xlsx')
       response['Content-Disposition'] = 'inline;filename=Solution Options template.xlsx'                                                   
       return response
    xlsx.closed

def solopt_storage(request):
    context = RequestContext(request)     

    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0  

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']     
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if request.method == 'POST':
        sform = Solopt_Storage(request.POST, request.FILES)
        if sform.is_valid():
           id = sform.save(commit=False)
           id.dec_id = dec_id
           print 'what is in here'
           print id.solopt_file
           id.save()
           try:
              getfile = request.POST.get('solopt_file', False)
              loc = '/home/amritha/costutility/documents/' + request.FILES['solopt_file'].name
              f = request.FILES['solopt_file']
              with open(loc, 'wb+') as destination:
                   for chunk in f.chunks():
                       destination.write(chunk)
              try:    
                 # Open the workbook and define the worksheet                                                                                            
                 book = xlrd.open_workbook(loc)
                 sheet = book.sheet_by_name("Sheet1")
                 # Establish a MySQL connection
                 database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")    
                 # Get the cursor, which is used to traverse the database, line by line
                 cursor = database.cursor()
                 # Create the INSERT INTO sql query
                 query = """INSERT INTO utility_tool_solution_options (dec_id, sol_option, source, option_details, linkname1, link1, linkname2, link2,linkname3, link3,linkname4, link4, archived, created_by, created_date) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"""
                 # Create a For loop to iterate through each row in the XLS file, starting at row 2 to skip the headers
                 for r in range(1, sheet.nrows):
                    option = sheet.cell(r,0).value
                    source = sheet.cell(r,1).value
                    details = sheet.cell(r,2).value
                    linkname1 = sheet.cell(r,3).value
                    link1 = sheet.cell(r,4).value
                    linkname2 = sheet.cell(r,5).value
                    link2 = sheet.cell(r,6).value   
                    linkname3 = sheet.cell(r,7).value
                    link3 = sheet.cell(r,8).value   
                    linkname4 = sheet.cell(r,9).value
                    link4 = sheet.cell(r,10).value    
                    values  = (dec_id, option, source, details, linkname1, link1, linkname2, link2, linkname3, link3, linkname4, link4, 'N', loggedinuser, datetime.datetime.now())
                    if (link1 <> '' and 'http://' not in link1 and 'https://' not in link1) or (link2 <> '' and 'http://' not in link2 and 'https://' not in link2) or (link3 <> '' and 'http://' not in link3 and 'https://' not in link3) or (link4 <> '' and 'http://' not in link4 and 'https://' not in link4):
                        err = 'Enter a valid URL that starts with http:// or https:// for the link fields. Copying and pasting from your browser may be helpful.'
                        return render(request,'decisions/solution_options/solopt_storage.html',{'dec_title':dec_title, 'sform':sform, 'err':err, 'loggedinuser':loggedinuser})
                    try:
                       s = Solution_Options.objects.get(sol_option = option, dec_id = dec_id)
                    except ObjectDoesNotExist: 
                       # Execute sql Query
                       cursor.execute(query, values)
                 # Close the cursor
                 cursor.close()
                 # Commit the transaction
                 database.commit()
                 # Close the database connection
                 database.close()
                 #columns = str(sheet.ncols)     
                 #rows = str(sheet.nrows)
                 return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_solopt_det.html') 
              except Exception as e:
                 print e 
                 if e == "argument of type 'float' is not iterable": 
                    err = 'Enter a valid URL that starts with http:// or https:// for the link fields. Copying and pasting from your browser may be helpful.'
                    print 'err'
                    print err
                 else:    
                    err = e 
                 #err = 'Please check the Excel sheet you have imported. It does not match the template we have provided.'
                 return render(request,'decisions/solution_options/solopt_storage.html',{'dec_title':dec_title, 'sform':sform, 'err':err, 'loggedinuser':loggedinuser})
           except Exception as e:  
               err = 'Please upload an Excel sheet.'
               return render(request,'decisions/solution_options/solopt_storage.html',{'dec_title':dec_title, 'sform':sform, 'err':err, 'loggedinuser':loggedinuser})
        else:
           print sform.errors
    else:
        sform = Solopt_Storage()

    return render(request,'decisions/solution_options/solopt_storage.html',{'dec_title':dec_title, 'sform':sform, 'loggedinuser':loggedinuser})

def link(request):
    context = RequestContext(request)   
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0 
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']          
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'     
    return render(request,'decisions/solution_options/link.html',{'dec_id':dec_id, 'dec_title':dec_title, 'loggedinuser':loggedinuser})  

def add_st_all(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                  
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    stakeholders_decisions = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and deleted is null)", [loggedinuser, dec_id])
    stakeholders = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id not in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s) order by firstName", [loggedinuser, dec_id])
   
    if request.method == 'POST':
        #print request.POST.getlist('id') 
        if 'id' in request.POST:
            if 'submit' in request.POST: 
               for value in request.POST.getlist('id'):
                  try: 
                     old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value)
                  except ObjectDoesNotExist:
                     st = Stakeholders.objects.get(pk=value)
                     st_name = st.firstName + ' ' + st.lastName
                     st_dec = Stakeholders_Decisions(st_id = value, name = st_name, email=st.email, dec_id = request.session['dec_id'],created_by = request.session['user'],created_date = datetime.datetime.now())
                     st_dec.save() 
        return HttpResponseRedirect('/utility_tool/decisions/solution_options/assign_tasks.html')       
    return render(request,'decisions/solution_options/add_st_all.html',{'stakeholders':stakeholders,'st_dec': stakeholders_decisions, 'dec_id':dec_id, 'dec_title':dec_title})

def assign_tasks(request):                                                                                                                                                                                         
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                  
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'user_email' in request.session:                                                                                                                                                                          
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    try: 
       SharedDec.objects.get(dec_id = dec_id, shared_user = loggedinuser)
       shared = 'Y'
    except ObjectDoesNotExist:
       shared = 'N'

    try: 
      Importance_Scores.objects.get(dec_id = dec_id, created_by = loggedinuser)                                                                     
      impexists = 'Y'
    except ObjectDoesNotExist:
       print 'is' 
       impexists = 'N'
    except MultipleObjectsReturned:                                                                                                                 
       impexists = 'Y'

    dec = Decisions.objects.get(pk=dec_id)
    something_saved = 'no' 
    redistribute_yn = 'N'  
    stakeholders_decisions = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and deleted is null) order by firstName", [loggedinuser, dec_id]) 
    #stakeholders = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id not in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and deleted = 'Y')", [loggedinuser, dec_id])
    stdec_type = Stakeholders_Decisions.objects.raw("SELECT id, st_id, solopt_type, scrcr_type, evacr_type, iw_type from utility_tool_stakeholders_decisions where dec_id=%s and deleted is null", [dec_id]) 
    if request.method == 'POST':
        print request.POST
        id_to_use = request.POST.get('id')
        solopt_type = ''
        scrcr_type = ''
        evacr_type = ''
        iw_type = ''
        if 'selected' in request.POST:
            for value in request.POST.getlist('selected'):
                if value != "[]":
                   #print value
                   if "solopt" in value:
                       if "soloptY" in value: 
                          solopt_type = 'Y'
                       else:   
                          solopt_type = ''
                   if "scrcr" in value:
                       print value 
                       if "scrcrY" in value: 
                          scrcr_type = 'Y'
                       else:   
                          scrcr_type = ''   
                   if "evacr" in value: 
                       if "evacrY" in value: 
                          evacr_type = 'Y'
                       else:   
                          evacr_type = ''   
                   if "iw" in value: 
                       if "iwY" in value: 
                          iw_type = 'Y'
                       else:   
                          iw_type = ''   

                   stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=id_to_use) 
                   stdec.solopt_type = solopt_type
                   stdec.scrcr_type = scrcr_type 
                   stdec.evacr_type = evacr_type 
                   stdec.iw_type = iw_type
                   if iw_type == 'Y':
                      if stdec.votes is None or stdec.votes == 0: 
                         stdec.votes = 10
                   else:
                      stdec.votes = 0 
                   redistribute_yn = 'Y'
                   stdec.updated_by = request.session['user']
                   stdec.updated_date = datetime.datetime.now()
                   stdec.save(update_fields=['solopt_type','scrcr_type','evacr_type','iw_type','votes','updated_by','updated_date'])        
                   something_saved ='yes' 
        else:
            print 'in DELETED'
            for val in request.POST.getlist('deleted'):
                print val
                val = val.strip()
                y = val.replace('[','')
                z = y.replace(']','')     
                temp_list = [] 
                # adding each id to a temporary list
                for l2 in z.split(','):                                                                                                                                                                             
                    l3 = l2.replace('"', '')
                    temp_list.append(l3) 
                print temp_list
                for l in temp_list:
                   try:
                      print dec_id
                      print l
                      delsolopt = Stakeholders_Decisions.objects.get(dec_id=dec_id, st_id=l)
                      print delsolopt.st_id
                      delsolopt.delete()
                      redistribute_yn = 'Y'
                      something_saved ='yes' 
                   except ObjectDoesNotExist:
                      print 'id does not exist'
                      return HttpResponse('Selected Id does not exist in database. Please contact your Administrator.')
                
        if redistribute_yn == 'Y':
           redistribution_func(dec_id, loggedinuser, request)
        if 'idList' in request.session: 
           del request.session['idList'] 
        if something_saved == 'yes':
              dec.updated_by = request.session['user']
              dec.updated_date = datetime.datetime.now()
              dec.save(update_fields=['updated_by','updated_date'])
    return render(request,'decisions/solution_options/assign_tasks.html',{'st_dec': stakeholders_decisions, 'stdec_type':stdec_type, 'dec_id':dec_id, 'dec_title':dec_title, 'loggedinuser':loggedinuser, 'user_email':user_email, 'impexists':impexists, 'shared':shared})

def add_st_solopt(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                        
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    stakeholders_decisions = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and solopt_type = 'Y' and deleted is null) order by firstName", [loggedinuser, dec_id])

    stakeholders = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id not in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and (solopt_type = 'Y' or deleted = 'Y')) order by firstName", [loggedinuser, dec_id])
   
    if request.method == 'POST':
        if 'id' in request.POST:
            if 'submit' in request.POST: 
               for value in request.POST.getlist('id'):
                  try:
                     old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value) 
                     old_stdec.solopt_type = 'Y' 
                     old_stdec.updated_by = request.session['user']
                     old_stdec.updated_date = datetime.datetime.now()
                     old_stdec.save(update_fields=['solopt_type','updated_by','updated_date'])
                  except ObjectDoesNotExist:
                     st = Stakeholders.objects.get(pk=value)
                     st_dec = Stakeholders_Decisions(st_id = value, email=st.email, dec_id = request.session['dec_id'],solopt_type = 'Y' ,created_by = request.session['user'],created_date = datetime.datetime.now())
                     st_dec.save()
               #return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)      
            else:    
                for value in request.POST.getlist('id'):
                    try: 
                       old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value) 
                       old_stdec.solopt_type = 'N'  
                       old_stdec.updated_by = request.session['user']
                       old_stdec.updated_date = datetime.datetime.now()
                       old_stdec.save(update_fields=['solopt_type','updated_by','updated_date'])
                    except ObjectDoesNotExist:
                       print 'id does not exist'
                       return HttpResponse('Selected Id does not exist in database. Please contact your Administrator.')

        else:
            return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)

    return render(request,'decisions/solution_options/add_st_solopt.html',{'stakeholders':stakeholders,'st_dec': stakeholders_decisions, 'dec_id':dec_id, 'dec_title':dec_title})

def add_scr_criteria(request):
    context = RequestContext(request)
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'
    #if solopt_allowed == 'Y' and scrcr_allowed == 'Y': 
    if loggedinuser == created_by: 
       mapping_allowed = 'Y'
    else:
       mapping_allowed = 'N'
    if 'user_email' in request.session: 
       user_email = request.session['user_email']
    else:                                                                                                                                                
       user_email = 'not found'

    try:
       SharedDec.objects.get(dec_id = dec_id, shared_user = loggedinuser)                                                                                                           
       shared = 'Y'
    except ObjectDoesNotExist:
       shared = 'N'

    try: 
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id)
       std_count = std.exclude(email = user_email).count()                                                                                               
       if std_count > 0:                                                                                                                                 
          stakeholdersNow = 'Y'
       else:
          stakeholdersNow = 'N'
    except ObjectDoesNotExist:
       stakeholdersNow = 'N'
    except MultipleObjectsReturned:                                                                                                                
       stakeholdersNow = 'Y'  
    solopt_count = Solution_Options.objects.filter(dec_id=dec_id, archived='N', deleted='N').count()
    dec = Decisions.objects.get(pk=dec_id) 
    something_saved = 'no'
    
    try: 
       firstrec = Master_Screening_Criteria.objects.get(dec_id=dec_id)
    except ObjectDoesNotExist:
       orig_qset = CBCSE_Screening_Criteria.objects.all()                                                                                                   
       for orig in orig_qset:
           orig_scr_save = Master_Screening_Criteria(criterion = orig.criterion, dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())  
           orig_scr_save.save()
    except MultipleObjectsReturned:                                                                                                                    
       print 'multiple rows in master screening criteria'
 
    qset = Master_Screening_Criteria.objects.filter(dec_id=dec_id)
    qset2 = Screening_Criteria.objects.filter(dec_id=dec_id)
    qset3 = qset2.exclude(orig_scr_id__isnull=False)
    request.session['map_list'] = []  
    # First part is a set of checkboxes. Second part is a set of text boxes.
    # Checkbox values come from Master Screening Criteria for each decision (which comes from CBCSE screening criteria excel)
    # Delete is based on a hiddenfield text box in the H
    # There are 15 text boxes - you can add and edit the criterion, you cannot delete but you can clear the field so it looks as if it has been deleted
    if request.method == 'POST':
       #print request.POST 
       for val in request.POST.getlist('hiddenfield'):   
           if val.endswith('U'):
              print val
              print val[:-1]
              try: 
                 scrdel = Screening_Criteria.objects.get(orig_scr_id = val[:-1], dec_id=dec_id)
                 request.session['map_list'].append(scrdel.id)
                 scrdel.delete()
              except ObjectDoesNotExist:
                 print 'cannot delete something that does not exist'  
       for value in request.POST.getlist('scrcr'):
           #print value
            
           try:
              cbcse_scr = Master_Screening_Criteria.objects.get(id = value, dec_id=dec_id)
              crit =  cbcse_scr.criterion
           except ObjectDoesNotExist:
              crit = ''

           try: 
              old_scr = Screening_Criteria.objects.get(criterion = crit, dec_id=dec_id)
              old_scr_exists = 'Y'   
           except ObjectDoesNotExist:
              old_scr_exists = 'N'

           fieldname = value + 'two'
           if old_scr_exists == 'N':  
              scr_save = Screening_Criteria(criterion = crit, criterion2 = request.POST.get(fieldname), orig_scr_id = value, dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now()) 
              scr_save.save() 
           else:
               old_scr.criterion2 = request.POST.get(fieldname)
               old_scr.orig_scr_id = value
               old_scr.updated_by = request.session['user']
               old_scr.updated_date = datetime.datetime.now()
               old_scr.save(update_fields=['criterion2','orig_scr_id','updated_by','updated_date'])

       try:
          old_scr1 = Screening_Criteria.objects.get(fieldname = 'cri1', dec_id=dec_id)
          old_scr1.criterion = request.POST.get('cri1')
          old_scr1.updated_by = request.session['user']
          old_scr1.updated_date = datetime.datetime.now()
          old_scr1.save(update_fields=['criterion','updated_by','updated_date'])
          print old_scr1.criterion
          if old_scr1.criterion == '' or old_scr1.criterion is None:
             request.session['map_list'].append(old_scr1.id)
             old_scr1.delete()   
       except ObjectDoesNotExist:   
          if request.POST.get('cri1') <> '':                                                                                                                                                                     
             scr_save1 = Screening_Criteria(criterion = request.POST.get('cri1'), fieldname = 'cri1', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save1.save()
       try:
          old_scr2 = Screening_Criteria.objects.get(fieldname = 'cri2', dec_id=dec_id)
          old_scr2.criterion = request.POST.get('cri2')
          old_scr2.updated_by = request.session['user']
          old_scr2.updated_date = datetime.datetime.now()
          old_scr2.save(update_fields=['criterion','updated_by','updated_date']) 
          if old_scr2.criterion == '' or old_scr2.criterion is None:
             request.session['map_list'].append(old_scr2.id)
             old_scr2.delete()  
       except ObjectDoesNotExist:         
          if request.POST.get('cri2') <> '':                                                                                                                                                                    
             scr_save2 = Screening_Criteria(criterion = request.POST.get('cri2'), fieldname = 'cri2', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save2.save()  
       try:
          old_scr3 = Screening_Criteria.objects.get(fieldname = 'cri3', dec_id=dec_id)
          old_scr3.criterion = request.POST.get('cri3')
          old_scr3.updated_by = request.session['user']
          old_scr3.updated_date = datetime.datetime.now()
          old_scr3.save(update_fields=['criterion','updated_by','updated_date']) 
          if old_scr3.criterion == '' or old_scr3.criterion is None:
             request.session['map_list'].append(old_scr3.id)
             old_scr3.delete()  
       except ObjectDoesNotExist:              
          if request.POST.get('cri3') <> '':                                                                                                                                                               
             scr_save3 = Screening_Criteria(criterion = request.POST.get('cri3'), fieldname = 'cri3', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save3.save()

       try:
          old_scr4 = Screening_Criteria.objects.get(fieldname = 'cri4', dec_id=dec_id)
          old_scr4.criterion = request.POST.get('cri4')
          old_scr4.updated_by = request.session['user']
          old_scr4.updated_date = datetime.datetime.now()
          old_scr4.save(update_fields=['criterion','updated_by','updated_date']) 
          if old_scr4.criterion == '' or old_scr4.criterion is None:
             request.session['map_list'].append(old_scr4.id)
             old_scr4.delete()  
       except ObjectDoesNotExist:              
          if request.POST.get('cri4') <> '':                                                                                                                                                               
             scr_save4 = Screening_Criteria(criterion = request.POST.get('cri4'), fieldname = 'cri4', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save4.save() 
       try:
          old_scr5 = Screening_Criteria.objects.get(fieldname = 'cri5', dec_id=dec_id)
          old_scr5.criterion = request.POST.get('cri5')
          old_scr5.updated_by = request.session['user']
          old_scr5.updated_date = datetime.datetime.now()
          old_scr5.save(update_fields=['criterion','updated_by','updated_date']) 
          if old_scr5.criterion == '' or old_scr5.criterion is None:
             request.session['map_list'].append(old_scr5.id)
             old_scr5.delete()  
       except ObjectDoesNotExist:              
          if request.POST.get('cri5') <> '':                                                                                                                                                               
             scr_save5 = Screening_Criteria(criterion = request.POST.get('cri5'), fieldname = 'cri5', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save5.save() 
       try:
          old_scr6 = Screening_Criteria.objects.get(fieldname = 'cri6', dec_id=dec_id)
          old_scr6.criterion = request.POST.get('cri6')
          old_scr6.updated_by = request.session['user']
          old_scr6.updated_date = datetime.datetime.now()
          old_scr6.save(update_fields=['criterion','updated_by','updated_date']) 
          if old_scr6.criterion == '' or old_scr6.criterion is None:
             request.session['map_list'].append(old_scr6.id)
             old_scr6.delete()  
       except ObjectDoesNotExist:              
          if request.POST.get('cri6') <> '':                                                                                                                                                               
             scr_save6 = Screening_Criteria(criterion = request.POST.get('cri6'), fieldname = 'cri6', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save6.save()  
       try:
          old_scr7 = Screening_Criteria.objects.get(fieldname = 'cri7', dec_id=dec_id)
          old_scr7.criterion = request.POST.get('cri7')
          old_scr7.updated_by = request.session['user']
          old_scr7.updated_date = datetime.datetime.now()
          old_scr7.save(update_fields=['criterion','updated_by','updated_date']) 
          if old_scr7.criterion == '' or old_scr7.criterion is None:
             request.session['map_list'].append(old_scr7.id)
             old_scr7.delete()  
       except ObjectDoesNotExist:  
          if request.POST.get('cri7') <> '':                                                                                                                                                                     
             scr_save7 = Screening_Criteria(criterion = request.POST.get('cri7'), fieldname = 'cri7', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save7.save()  
       try:
          old_scr8 = Screening_Criteria.objects.get(fieldname = 'cri8', dec_id=dec_id)
          old_scr8.criterion = request.POST.get('cri8')
          old_scr8.updated_by = request.session['user']
          old_scr8.updated_date = datetime.datetime.now()
          old_scr8.save(update_fields=['criterion','updated_by','updated_date']) 
          if old_scr8.criterion == '' or old_scr8.criterion is None:
             request.session['map_list'].append(old_scr8.id)
             old_scr8.delete()  
       except ObjectDoesNotExist:              
          if request.POST.get('cri8') <> '':                                                                                                                                                               
             scr_save8 = Screening_Criteria(criterion = request.POST.get('cri8'), fieldname = 'cri8', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save8.save()  
       try:
          old_scr9 = Screening_Criteria.objects.get(fieldname = 'cri9', dec_id=dec_id)
          old_scr9.criterion = request.POST.get('cri9')
          old_scr9.updated_by = request.session['user']
          old_scr9.updated_date = datetime.datetime.now()
          old_scr9.save(update_fields=['criterion','updated_by','updated_date']) 
          if old_scr9.criterion == '' or old_scr9.criterion is None:
             request.session['map_list'].append(old_scr9.id)
             old_scr9.delete()  
       except ObjectDoesNotExist:              
          if request.POST.get('cri9') <> '':                                                                                                                                                               
             scr_save9 = Screening_Criteria(criterion = request.POST.get('cri9'), fieldname = 'cri9', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save9.save()  
       try:
          old_scr10 = Screening_Criteria.objects.get(fieldname = 'cri10', dec_id=dec_id)
          old_scr10.criterion = request.POST.get('cri10')
          old_scr10.updated_by = request.session['user']
          old_scr10.updated_date = datetime.datetime.now()
          old_scr10.save(update_fields=['criterion','updated_by','updated_date']) 
          if old_scr10.criterion == '' or old_scr10.criterion is None:
             request.session['map_list'].append(old_scr10.id)
             old_scr10.delete()   
       except ObjectDoesNotExist:              
          if request.POST.get('cri10') <> '':                                                                                                                                                               
             scr_save10 = Screening_Criteria(criterion = request.POST.get('cri10'), fieldname = 'cri10', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save10.save() 
       try:
          old_scr11 = Screening_Criteria.objects.get(fieldname = 'cri11', dec_id=dec_id)
          old_scr11.criterion = request.POST.get('cri11')
          old_scr11.updated_by = request.session['user']
          old_scr11.updated_date = datetime.datetime.now()
          old_scr11.save(update_fields=['criterion','updated_by','updated_date'])
          if old_scr11.criterion == '' or old_scr11.criterion is None:
             request.session['map_list'].append(old_scr11.id)
             old_scr11.delete()  
       except ObjectDoesNotExist:
          if request.POST.get('cri11') <> '':
             scr_save11 = Screening_Criteria(criterion = request.POST.get('cri11'), fieldname = 'cri11', dec_id = request.session['dec_id'], created_by =request.session['user'],created_date = datetime.datetime.now())         
             scr_save11.save()
       try:
          old_scr12 = Screening_Criteria.objects.get(fieldname = 'cri12', dec_id=dec_id)
          old_scr12.criterion = request.POST.get('cri12')
          old_scr12.updated_by = request.session['user']
          old_scr12.updated_date = datetime.datetime.now()
          old_scr12.save(update_fields=['criterion','updated_by','updated_date'])
          if old_scr12.criterion == '' or old_scr12.criterion is None:
             request.session['map_list'].append(old_scr12.id)
             old_scr12.delete()  
       except ObjectDoesNotExist:
          if request.POST.get('cri12') <> '':
             scr_save12 = Screening_Criteria(criterion = request.POST.get('cri12'), fieldname = 'cri12', dec_id = request.session['dec_id'], created_by =request.session['user'],created_date = datetime.datetime.now())         
             scr_save12.save()
       try:
          old_scr13 = Screening_Criteria.objects.get(fieldname = 'cri13', dec_id=dec_id)
          old_scr13.criterion = request.POST.get('cri13')
          old_scr13.updated_by = request.session['user']
          old_scr13.updated_date = datetime.datetime.now()
          old_scr13.save(update_fields=['criterion','updated_by','updated_date'])
          if old_scr13.criterion == '' or old_scr13.criterion is None:
             request.session['map_list'].append(old_scr13.id)
             old_scr13.delete()   
       except ObjectDoesNotExist:
          if request.POST.get('cri13') <> '':
             scr_save13 = Screening_Criteria(criterion = request.POST.get('cri13'), fieldname = 'cri13', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save13.save()
       try:
          old_scr14 = Screening_Criteria.objects.get(fieldname = 'cri14', dec_id=dec_id)
          old_scr14.criterion = request.POST.get('cri14')
          old_scr14.updated_by = request.session['user']
          old_scr14.updated_date = datetime.datetime.now()
          old_scr14.save(update_fields=['criterion','updated_by','updated_date'])
          if old_scr14.criterion == '' or old_scr14.criterion is None:
             request.session['map_list'].append(old_scr14.id)
             old_scr14.delete()  
       except ObjectDoesNotExist:
          if request.POST.get('cri14') <> '':
             scr_save14 = Screening_Criteria(criterion = request.POST.get('cri14'), fieldname = 'cri14', dec_id = request.session['dec_id'], created_by =request.session['user'],created_date = datetime.datetime.now())         
             scr_save14.save()
       try:
          old_scr15 = Screening_Criteria.objects.get(fieldname = 'cri15', dec_id=dec_id)
          old_scr15.criterion = request.POST.get('cri15')
          old_scr15.updated_by = request.session['user']
          old_scr15.updated_date = datetime.datetime.now()
          old_scr15.save(update_fields=['criterion','updated_by','updated_date'])
          if old_scr15.criterion == '' or old_scr15.criterion is None:
             request.session['map_list'].append(old_scr15.id)
             old_scr15.delete()   
       except ObjectDoesNotExist:
          if request.POST.get('cri15') <> '':                                                                                                                                                                    
             scr_save15 = Screening_Criteria(criterion = request.POST.get('cri15'), fieldname = 'cri15', dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())         
             scr_save15.save()

       return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
    return render(request,'decisions/solution_options/add_scr_criteria.html',{'qset':qset,'qset2':qset2,'qset3':qset3,'dec_id':dec_id, 'dec_title':dec_title,'mapping_allowed': mapping_allowed, 'solopt_count':solopt_count,'created_by':created_by, 'loggedinuser':loggedinuser, 'stakeholdersNow':stakeholdersNow, 'shared':shared})


def add_eva_criteria(request):
    context = RequestContext(request)
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'created_by' in request.session:
       created_by = request.session['created_by']                                                                                                         
    else:
       created_by = 'not found'
    #if solopt_allowed == 'Y' and scrcr_allowed == 'Y': 
    if loggedinuser == created_by: 
       mapping_allowed = 'Y'
    else:
       mapping_allowed = 'N'
    if 'user_email' in request.session: 
       user_email = request.session['user_email']
    else:                                                                                                                                                
       user_email = 'not found'

    try: 
       SharedDec.objects.get(dec_id = dec_id, shared_user = loggedinuser)     
       shared = 'Y'
    except ObjectDoesNotExist:
       shared = 'N'

    try: 
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id)
       std_count = std.exclude(email = user_email).count()                                                                                               
       if std_count > 0:                                                                                                                                 
          stakeholdersNow = 'Y'
       else:
          stakeholdersNow = 'N'
    except ObjectDoesNotExist:                                                                                                                            
       stakeholdersNow = 'N'
    except MultipleObjectsReturned:
       stakeholdersNow = 'Y'
    solopt = Solution_Options.objects.filter(dec_id=dec_id, archived='N',deleted='N')
    dec = Decisions.objects.get(pk=dec_id) 
    something_saved = 'no'

    # Check if the records have been copied from the CBCSE Evaluation Criteria List to the Master List. If not, copy the records
    try: 
       firstrec = Master_Evaluation_Criteria.objects.get(dec_id=dec_id)
    except ObjectDoesNotExist:
       orig_qset = CBCSE_Evaluation_Criteria.objects.all()                                                                                                 
       for orig in orig_qset:
           orig_eva_save = Master_Evaluation_Criteria(overreaching_ec = orig.overreaching_ec, granular_ec = orig.granular_ec, suggested_evam = orig.suggested_evam, data = orig.data, dec_id = request.session['dec_id'], created_by = request.session['user'], created_date = datetime.datetime.now())        
           orig_eva_save.save()
    except MultipleObjectsReturned:                                                                                                                    
       print 'multiple rows in master evaluation criteria'
    
    # just to test
    obj= Evaluation_Criteria.objects.filter(dec_id=dec_id).order_by('-id')[:2]
    for o in obj:
        print o.id

    
    # stop test
    if request.method == 'POST':
       print request.POST
       # DELETION
       # Check if the hiddenfield list has any IDs. 
       # If there are IDs, check if there is an eva record with that id.  
       # Set the delete flag in evacr to Y for IDs existing
       # Delete the records in imp scores with those IDs
       # Set the delete flag in eva meas to Y for IDs existing
       for val in request.POST.getlist('hiddenfield'):                                                                                                                                                           
           if val.endswith('U'):
              print val
              print val[:-1]
              try: 
                 evadel = Evaluation_Criteria.objects.get(orig_eva_id = val[:-1], dec_id=dec_id)
                 evadel.deleted = 'Y'
                 evadel.updated_by = request.session['user']                                                                                                                                                    
                 evadel.updated_date = datetime.datetime.now()                                                                                                                                                  
                 evadel.save(update_fields=['deleted','updated_by','updated_date'])
                 something_saved ='yes'
                 for iss in Importance_Scores.objects.filter(eva_id = evadel.id, dec_id=dec_id):
                     iss.deleted = 'Y'
                     iss.updated_by = request.session['user']                                                                                                                                                    
                     iss.updated_date = datetime.datetime.now()                                                                                                                                                  
                     iss.save(update_fields=['deleted','updated_by','updated_date'])                
                 for evm in Evaluation_Measures.objects.filter(eva_id = evadel.id, dec_id=dec_id):
                     evm.deleted = 'Y'
                     evm.updated_by = request.session['user']
                     evm.updated_date = datetime.datetime.now()                                                                                                                                               
                     evm.save(update_fields=['deleted','updated_by','updated_date'])
              except ObjectDoesNotExist:
                 print 'cannot delete something that does not exist'  
       # INSERTION  
       # THIS IS FOR FIRST GROUP WHICH COMES FROM EXCEL
       for value in request.POST.getlist('evacr_q'):
           #print value
           # get the values of required fields from the Master List table 
           try:
              master_eva = Master_Evaluation_Criteria.objects.get(id = value, dec_id=dec_id)
              or_crit = master_eva.overreaching_ec
              crit =  master_eva.granular_ec
              sugg_evam = master_eva.suggested_evam
              data = master_eva.data
           except ObjectDoesNotExist:
              or_crit = ''
              crit = ''
              sugg_evam = ''
              data = ''

           # check if eva record already exists 
           try: 
              old_eva = Evaluation_Criteria.objects.get(criterion = crit, dec_id=dec_id)
              if old_eva.deleted == 'Y':
                 old_eva_exists = 'D'
              else:   
                 old_eva_exists = 'Y'   
           except ObjectDoesNotExist:
              old_eva_exists = 'N'

           # if exists, update record. Else insert record.
           # Insert into imp scores as well.
           fieldname = value + 'two'
           if old_eva_exists == 'N':
              if request.POST.get(fieldname) is not None and request.POST.get(fieldname) <> '':
                 crit2 = request.POST.get(fieldname).replace(",",";")
                 crit2 = crit2.replace("&#39;","'") 
                 combined = crit + ': ' + crit2
              else:   
                 combined = crit
              criterion2 = request.POST.get(fieldname).replace(",", ";")
              criterion2 = criterion2.replace("&#39;","'") 
              eva_save = Evaluation_Criteria(or_criterion = or_crit, criterion = crit, suggested_evam = sugg_evam, data = data, criterion2 = criterion2, combined = combined, orig_eva_id = value, dec_id = request.session['dec_id'], created_by = request.session['user'],created_date = datetime.datetime.now())  
              eva_save.save() 
              something_saved ='yes'
           else:
               if request.POST.get(fieldname) is not None and request.POST.get(fieldname) <> '':
                  crit2 = request.POST.get(fieldname).replace(",",";")
                  crit2 = crit2.replace("&#39;","'") 
                  old_eva.combined = crit + ': ' + crit2
               else:   
                  old_eva.combined = crit
               if old_eva_exists == 'D':
                  old_eva.deleted = 'N'
                  old_eva.criterion2 = request.POST.get(fieldname).replace(",",";")
                  old_eva.criterion2 = old_eva.criterion2.replace("&#39;","'") 
                  old_eva.orig_eva_id = value
                  old_eva.updated_by = request.session['user']
                  old_eva.updated_date = datetime.datetime.now()                                           
                  old_eva.save(update_fields=['criterion2','combined','orig_eva_id','deleted','updated_by','updated_date'])
                  something_saved ='yes'
               else:   
                  old_eva.criterion2 = request.POST.get(fieldname).replace(",",";")  
                  old_eva.criterion2 = old_eva.criterion2.replace("&#39;","'") 
                  old_eva.orig_eva_id = value
                  old_eva.updated_by = request.session['user']
                  old_eva.updated_date = datetime.datetime.now()                                                                                                                                                    
                  old_eva.save(update_fields=['criterion2','combined','orig_eva_id','updated_by','updated_date'])
                  something_saved ='yes' 
               try:
                  for isw in Importance_Scores.objects.filter(eva_id = old_eva.id, dec_id=dec_id, created_by = loggedinuser):
                      if old_eva.criterion2 is None or old_eva.criterion2 == '':
                          isw.criterion = old_eva.criterion
                      else:     
                          isw.criterion = old_eva.criterion + ': ' + old_eva.criterion2.replace(",",";")  
                      if old_eva_exists == 'D':  
                         isw.deleted = 'N'
                      else:
                         isw.deleted = isw.deleted 
                      isw.updated_by = request.session['user']
                      isw.updated_date = datetime.datetime.now()
                      isw.save(update_fields=['criterion','deleted','updated_by','updated_date'])   
               except ObjectDoesNotExist: 
                  print 'we know it does not exist'  
                  
               try: 
                  for evm in Evaluation_Measures.objects.filter(eva_id = old_eva.id, dec_id=dec_id):
                      if old_eva.criterion2 is None or old_eva.criterion2 == '':
                         evm.criterion = old_eva.criterion
                      else:   
                         evm.criterion = old_eva.criterion + ': ' + old_eva.criterion2.replace(",",";")  
                      if old_eva_exists == 'D':  
                         evm.deleted = 'N'
                      else:
                         evm.deleted = evm.deleted      
                      evm.updated_by = loggedinuser
                      evm.updated_date = datetime.datetime.now()    
                      evm.save(update_fields=['criterion','deleted','updated_by','updated_date']) 
               except ObjectDoesNotExist:
                  print 'evam does not exist' 

       # check if there is an existing text field record for cri1
       # if criterion for cri1 is '', then deleted is set to Y
       # all the fields are updated
       # if existing record does not exist, a new record is inserted for cri1 as long as it is not ''
       update_text_criteria(request, request.POST.get('cri1'),'cri1',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri2'), 'cri2',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri3'), 'cri3',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri4'), 'cri4',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri5'), 'cri5',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri6'), 'cri6',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri7'), 'cri7',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri8'), 'cri8',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri9'), 'cri9',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri10'), 'cri10',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri11'), 'cri11',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri12'), 'cri12',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri13'), 'cri13',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri14'), 'cri14',dec_id, loggedinuser)
       update_text_criteria(request, request.POST.get('cri15'), 'cri15',dec_id, loggedinuser)
       
       if 'submit' in request.POST:    
           if something_saved == 'yes':
              dec.updated_by = request.session['user']
              dec.updated_date = datetime.datetime.now()
              dec.save(update_fields=['updated_by','updated_date'])
           return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
       elif 'st' in request.POST:
           if stakeholdersNow == 'Y': 
              return HttpResponseRedirect('/utility_tool/decisions/solution_options/assign_tasks.html') 
           else:
              return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_st_all.html')    
    qset = Master_Evaluation_Criteria.objects.filter(dec_id=dec_id)
    qset_or = qset.values('overreaching_ec').distinct()
    #qset_or_count = qset.values('overreaching_ec').distinct().count() 
    for q in qset_or:
        if q['overreaching_ec'] == 'Addresses the identified need':
           qset_add = qset.filter(overreaching_ec=q['overreaching_ec'])
        elif q['overreaching_ec'] == 'Equity':
           qset_eq = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'External recommendations':                                                                                                                                            
           qset_ext = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Feasibility of implementation':                                                                                                                                            
           qset_feas = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Fit with local context':                                                                                                                                            
           qset_loc = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Impact on parent engagement':                                                                                                                                          
           qset_pe = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Impact on student academic performance':                                                                                                                                            
           qset_aced = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Impact on student socio-emotional development':                                                                                                                                            
           qset_emot = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Impact on student/staff engagement':                                                                                                                                            
           qset_staf = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Improves teacher performance':                                                                                                                                            
           qset_teac = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Meets required standards and regulations':                                                                                                                                            
           qset_stand = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Quality of implementation (for programs/strategies/tools already in place)':                                                                                                                                            
           qset_qual = qset.filter(overreaching_ec=q['overreaching_ec']) 
        elif q['overreaching_ec'] == 'Support from stakeholders':                                                                                                                                            
           qset_stx = qset.filter(overreaching_ec=q['overreaching_ec']) 
    for qq in qset_add:
        print qq.granular_ec
    qset_g = qset.filter(overreaching_ec="Fit with local context")
    qset2 = Evaluation_Criteria.objects.filter(dec_id=dec_id).exclude(deleted = 'Y') 
    qset3 = qset2.exclude(orig_eva_id__isnull=False)                                                                                                      
    return render(request,'decisions/solution_options/add_eva_criteria.html',{'qset':qset_or,'qset2':qset2,'qset3':qset3,'qset_add':qset_add,'qset_eq':qset_eq,'qset_ext':qset_ext,'qset_feas':qset_feas,'qset_loc':qset_loc,'qset_pe':qset_pe,'qset_aced':qset_aced,'qset_emot':qset_emot,'qset_staf':qset_staf,'qset_teac':qset_teac,'qset_stand':qset_stand,'qset_qual':qset_qual,'qset_stx':qset_stx,'dec_id':dec_id, 'dec_title':dec_title, 'solopt':solopt,'created_by':created_by, 'loggedinuser':loggedinuser, 'stakeholdersNow':stakeholdersNow, 'shared':shared}) 
 
def add_st_evacr(request):                                                                                                                                                                                      
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                  
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    stakeholders_decisions = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and evacr_type = 'Y' and deleted is null) order by firstName", [loggedinuser, dec_id])
    stakeholders = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id not in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and (evacr_type = 'Y' or deleted = 'Y')) order by firstName", [loggedinuser, dec_id])
   
    if request.method == 'POST':
        if 'id' in request.POST:
            if 'submit' in request.POST: 
               for value in request.POST.getlist('id'):
                  try:
                     old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value) 
                     old_stdec.evacr_type = 'Y' 
                     old_stdec.updated_by = request.session['user']
                     old_stdec.updated_date = datetime.datetime.now()
                     old_stdec.save(update_fields=['evacr_type','updated_by','updated_date'])
                  except ObjectDoesNotExist:
                     st = Stakeholders.objects.get(pk=value) 
                     st_dec = Stakeholders_Decisions(st_id = value, email=st.email, dec_id = request.session['dec_id'],evacr_type = 'Y',created_by = request.session['user'],created_date = datetime.datetime.now())
                     st_dec.save()
               #return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)      
            else:    
                for value in request.POST.getlist('id'):
                    try: 
                       old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value) 
                       old_stdec.evacr_type = 'N' 
                       old_stdec.updated_by = request.session['user']
                       old_stdec.updated_date = datetime.datetime.now()
                       old_stdec.save(update_fields=['evacr_type','updated_by','updated_date'])
                    except ObjectDoesNotExist:
                       print 'id does not exist'
                       return HttpResponse('Selected Id does not exist in database. Please contact your Administrator.')
        else:
            return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
    return render(request,'decisions/solution_options/add_st_evacr.html',{'stakeholders':stakeholders,'st_dec': stakeholders_decisions, 'dec_id':dec_id, 'dec_title':dec_title})

def add_scores(request):
    context = RequestContext(request)
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']          
    else:
       dec_title = 'not found'

    if 'user_email' in request.session:                                                                                                                                                                          
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'

    if 'user_email' in request.session: 
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    try:
       SharedDec.objects.get(dec_id = dec_id, shared_user = loggedinuser)
       shared = 'Y'
    except ObjectDoesNotExist:
       shared = 'N'

    try: 
       eva = Evaluation_Criteria.objects.get(dec_id=dec_id) 
    except ObjectDoesNotExist:
       print 'eva' 
       return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered evaluation criteria so you cannot view this screen.'})
    except MultipleObjectsReturned:                                                                                                                
       print 'eva multiple objects returned' 

    try: 
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id, iw_type = 'Y')                                                                  
       std_count = std.exclude(email = user_email).count()                                                                                          
       if std_count > 0: 
          stakeholdersNow = 'Y'   
       else:
          stakeholdersNow = 'N'  
    except ObjectDoesNotExist:
       stakeholdersNow = 'N'
    except MultipleObjectsReturned:                                                                                                                
       stakeholdersNow = 'Y'  

    try: 
       s = Scores_Setup.objects.get(dec_id = dec_id) 
       thinking = s.thinking
    except ObjectDoesNotExist:
       s = Scores_Setup(dec_id = dec_id, created_by=request.session['user'],created_date = datetime.datetime.now())
       s.save()
       thinking = ''

    dec = Decisions.objects.get(pk=dec_id) 
    something_saved = 'no'
    qset = Importance_Scores.objects.filter(dec_id=dec_id, created_by=request.session['user']).exclude(deleted = 'Y')
    eva1 = Evaluation_Criteria.objects.filter(dec_id=dec_id).exclude(deleted = 'Y') 
    eva = eva1.exclude(criterion = '')
    ids = set(e.id for e in eva)
    print ids
    ids2 = set(q.eva_id for q in qset)
    print ids2
    mylist = ids - ids2
    for l in mylist:
        print l
        e = Evaluation_Criteria.objects.get(id=l)
        if e.criterion2 is not None:
            criterion = e.combined
        else:
            criterion = e.criterion 
        sc = Importance_Scores(eva_id = e.id, criterion = criterion, score = 0, dec_id = dec_id, created_by=request.session['user'], email=request.session['user_email'],created_date = datetime.datetime.now())
        sc.save()    
    qset = Importance_Scores.objects.filter(dec_id=dec_id, created_by=request.session['user']).exclude(deleted = 'Y').order_by('eva_id') 

    MFormSet = modelformset_factory(Importance_Scores, form=ScoresForm, extra=0)
    if request.method == 'POST':
        scoresform = MFormSet(request.POST,request.FILES,prefix="scoresform" )
        if scoresform.is_valid():
           id = scoresform.save(commit=False)
           
           for recs in id:
               recs.updated_by = request.session['user'] 
               recs.updated_date = datetime.datetime.now()
               something_saved = 'yes'
               recs.save(update_fields=['score', 'updated_by','updated_date']) 
           if something_saved == 'yes':  
              dec.updated_by = request.session['user'] 
              dec.updated_date = datetime.datetime.now()
              dec.save(update_fields=['updated_by','updated_date'])  
           s.thinking = request.POST.get('thinking') 
           s.updated_by = request.session['user'] 
           s.updated_date = datetime.datetime.now()
           s.save(update_fields=['thinking','updated_by','updated_date'])   
           if 'submit' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
           elif 'votes' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_iw_votes.html')
           elif 'setup' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/pa_setup.html')  
           elif 'summary' in request.POST:
               return HttpResponseRedirect('/utility_tool/decisions/solution_options/summary.html') 
           elif 'st' in request.POST:                                                                                                                                                                            
              if stakeholdersNow == 'Y':
                 return HttpResponseRedirect('/utility_tool/decisions/solution_options/assign_tasks.html') 
              else:
                 return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_st_all.html')  
        else:
            print scoresform.errors
    else:
        scoresform = MFormSet(queryset = qset,prefix="scoresform")
        for form in scoresform:
           form.fields['criterion'].widget.attrs['readonly'] = True 
           #form.fields['created_by'].widget.attrs['readonly'] = True
    return render(request,'decisions/solution_options/add_scores.html',{'scoresform':scoresform,'dec_id':dec_id, 'dec_title':dec_title, 'created_by':created_by, 'loggedinuser':loggedinuser, 'stakeholdersNow':stakeholdersNow, 'thinking':thinking, 'shared':shared}) 

def add_st_iw(request):                                                                                                                                                                                       
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                  
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    stakeholders_decisions = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and iw_type = 'Y' and deleted is null) order by firstName", [loggedinuser, dec_id])
    stakeholders = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id not in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and (iw_type = 'Y' or deleted = 'Y')) order by firstName", [loggedinuser, dec_id]) 
   
    if request.method == 'POST':
        if 'id' in request.POST:
            if 'submit' in request.POST:
               for value in request.POST.getlist('id'):
                  try:
                     old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value) 
                     old_stdec.iw_type = 'Y' 
                     st = Stakeholders.objects.get(pk=value)
                     name = st.firstName + ' ' + st.lastName
                     old_stdec.name = name
                     old_stdec.updated_by = request.session['user']
                     old_stdec.updated_date = datetime.datetime.now()
                     old_stdec.save(update_fields=['iw_type','name','updated_by','updated_date'])
                  except ObjectDoesNotExist:
                     st = Stakeholders.objects.get(pk=value) 
                     name = st.firstName + ' ' + st.lastName
                     st_dec = Stakeholders_Decisions(st_id = value, name = name, email=st.email, dec_id = request.session['dec_id'],iw_type = 'Y',created_by = request.session['user'],created_date = datetime.datetime.now())
                     st_dec.save()
               #return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)      
            else:
                for value in request.POST.getlist('id'):
                    try:
                       old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value) 
                       old_stdec.iw_type = 'N' 
                       old_stdec.updated_by = request.session['user']
                       old_stdec.updated_date = datetime.datetime.now()
                       old_stdec.save(update_fields=['iw_type','updated_by','updated_date']) 
                    except ObjectDoesNotExist:
                       print 'id does not exist'
                       return HttpResponse('Selected Id does not exist in database. Please contact your Administrator.')
        else:
            return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)                                                                                                                         
    return render(request,'decisions/solution_options/add_st_iw.html',{'stakeholders':stakeholders,'st_dec': stakeholders_decisions, 'dec_id':dec_id, 'dec_title':dec_title})

'''
def add_iw_votes(request):                                                                                                                                                                                    
    context = RequestContext(request)
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']          
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'user_email' in request.session: 
       user_email = request.session['user_email']
    else:
       user_email = 'not found'

    try: 
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id)
       std_count = std.exclude(email = user_email).count()                                                                                                                                                       
       if std_count > 0: 
          stakeholdersNow = 'Y'   
       else:
          stakeholdersNow = 'N'  
    except ObjectDoesNotExist:
       stakeholdersNow = 'N'
    except MultipleObjectsReturned:                                                                                                                
       stakeholdersNow = 'Y'  

    qset = Stakeholders_Decisions.objects.filter(dec_id=dec_id, created_by=loggedinuser, iw_type = 'Y')
    qset_count = Stakeholders_Decisions.objects.filter(dec_id=dec_id, created_by=loggedinuser, iw_type = 'Y').count()
    total_votes = 10 * qset_count
    allowed_votes = 0

    MFormSet = modelformset_factory(Stakeholders_Decisions, form=VotesForm, extra=0)
    dec = Decisions.objects.get(pk=dec_id)
    something_saved = 'no'
 
    if request.method == 'POST':
        votesform = MFormSet(request.POST,request.FILES,prefix="votesform" )
        if votesform.is_valid():
           id = votesform.save(commit=False)
           for recs in id:
               if recs.votes is None:
                  errtext = 'Please enter the number of votes'
                  return render(request,'decisions/solution_options/add_iw_votes.html',{'votesform':votesform,'dec_id':dec_id, 'dec_title':dec_title, 'errtext':errtext, 'total_votes':total_votes, 'total_voters':qset_count})
               allowed_votes = recs.votes + allowed_votes
           print allowed_votes

           if allowed_votes > total_votes:
              errtext = 'The total number of votes cannot exceed '  + str(total_votes) + ', i.e., ten times the number of Stakeholders.' 
              return render(request,'decisions/solution_options/add_iw_votes.html',{'votesform':votesform,'dec_id':dec_id, 'dec_title':dec_title, 'errtext':errtext, 'total_votes':total_votes, 'total_voters':qset_count})
           elif allowed_votes <> total_votes:
              errtext = 'The total number of votes must be equal to ' + str(total_votes) + ', i.e., ten times the number of Stakeholders.'
              return render(request,'decisions/solution_options/add_iw_votes.html',{'votesform':votesform,'dec_id':dec_id, 'dec_title':dec_title, 'errtext':errtext, 'total_votes':total_votes, 'total_voters':qset_count}) 
           else:
               for recs in id:                                                                                                                                                                                   
                  if recs.votes <> '':
                     recs.updated_by = request.session['user'] 
                     recs.updated_date = datetime.datetime.now()
                     recs.save(update_fields=['votes','updated_by', 'updated_date'])
                     something_saved ='yes'
               if something_saved == 'yes':  
                  dec.updated_by = request.session['user'] 
                  dec.updated_date = datetime.datetime.now()
                  dec.save(update_fields=['updated_by','updated_date'])    
               if 'submit' in request.POST:
                  return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
               elif 'scores' in request.POST:
                  return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_scores.html')
               elif 'setup' in request.POST:
                  return HttpResponseRedirect('/utility_tool/decisions/pa_setup.html')                         
               elif 'st' in request.POST:     
                  if stakeholdersNow == 'Y': 
                     return HttpResponseRedirect('/utility_tool/decisions/solution_options/assign_tasks.html') 
                  else:
                     return HttpResponseRedirect('/utility_tool/decisions/solution_options/add_st_all.html')  
        else:
            print votesform.errors
    else:
        votesform = MFormSet(queryset = qset,prefix="votesform")
        for form in votesform:                                                                                                                                                                                   
           form.fields['name'].widget.attrs['readonly'] = True 
           form.fields['updated_by'].widget.attrs['readonly'] = True
           instance = getattr(form, 'instance', None)
           if not instance.votes:
              if instance.votes <> 0:
                 form.initial['votes'] = 10
    return render(request,'decisions/solution_options/add_iw_votes.html',{'votesform':votesform,'dec_id':dec_id, 'dec_title':dec_title, 'total_votes':total_votes, 'total_voters':qset_count})
'''   
def add_st_scrcr(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']                                                                                                                                                                  
    else:
       dec_title = 'not found'

    stakeholders_decisions = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and scrcr_type = 'Y'  and deleted is null) order by firstName", [loggedinuser, dec_id])

    stakeholders = Stakeholders.objects.raw("SELECT id, firstName, lastName, email, title, organisation from utility_tool_stakeholders where created_by=%s and id not in (SELECT st_id from utility_tool_stakeholders_decisions where dec_id = %s and (scrcr_type = 'Y' or deleted = 'Y')) order by firstName", [loggedinuser, dec_id])

    if request.method == 'POST':
        if 'id' in request.POST:
            if 'submit' in request.POST: 
                for value in request.POST.getlist('id'):
                   try:
                      old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value) 
                      old_stdec.scrcr_type = 'Y' 
                      old_stdec.updated_by = request.session['user']
                      old_stdec.updated_date = datetime.datetime.now()
                      old_stdec.save(update_fields=['scrcr_type','updated_by','updated_date'])  
                   except ObjectDoesNotExist:
                      st = Stakeholders.objects.get(pk=value) 
                      st_dec = Stakeholders_Decisions(st_id = value, email=st.email, dec_id = request.session['dec_id'],scrcr_type = 'Y',created_by = request.session['user'],created_date = datetime.datetime.now())
                      st_dec.save()
               #return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
            else:    
                for value in request.POST.getlist('id'):
                   try: 
                       old_stdec = Stakeholders_Decisions.objects.get(dec_id=dec_id,st_id=value) 
                       old_stdec.scrcr_type = 'N' 
                       old_stdec.updated_by = request.session['user']
                       old_stdec.updated_date = datetime.datetime.now()
                       old_stdec.save(update_fields=['iw_type','updated_by','updated_date'])  
                   except ObjectDoesNotExist:
                       print 'id does not exist'
                       return HttpResponse('Selected Id does not exist in database. Please contact your Administrator.')
        else:
            return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)

    return render(request,'decisions/solution_options/add_st_scrcr.html',{'stakeholders':stakeholders,'st_dec': stakeholders_decisions, 'dec_id':dec_id, 'dec_title':dec_title})

def handsontable(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']          
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'map_list' in request.session:
       map_list = request.session['map_list']
    else:
       map_list = 'not found'
    solopt_exists = ''
    scrcr_exists = ''
    try:
       SharedDec.objects.get(dec_id = dec_id, shared_user = loggedinuser)
       shared = 'Y'
    except ObjectDoesNotExist:
       shared = 'N'
    try: 
       ss = Solution_Options.objects.get(dec_id=dec_id)
       solopt_exists = 'Y'
    except ObjectDoesNotExist:
       print 'solopt'
       solopt_exists = 'N'
    except MultipleObjectsReturned:     
       solopt_exists = 'Y' 
       print 'solopt multiple objects returned' 

    try: 
       ss2 = Screening_Criteria.objects.get(dec_id=dec_id)
       if ss2.criterion == '':
          scrcr_exists = 'N'
       else:   
          scrcr_exists = 'Y'
    except ObjectDoesNotExist:
       print 'scrcr'
       scrcr_exists = 'N'
    except MultipleObjectsReturned:     
       print 'scrcr multiple objects returned'
       scrcr_exists = 'N'  
       for ss2 in Screening_Criteria.objects.filter(dec_id=dec_id):
           if ss2.criterion <> '':
              scrcr_exists = 'Y'   

    mess = ''
    if scrcr_exists == 'N' and solopt_exists == 'N':
       mess = 'You need to enter Solution Options and Screening Criteria before proceeding to this page.'    
    elif scrcr_exists == 'N' and solopt_exists == 'Y':
        mess = 'You need to enter Screening Criteria before proceeding to this page.'   
    elif scrcr_exists == 'Y' and solopt_exists == 'N':
        mess = 'You need to enter Solution Options before proceeding to this page.'   
    if mess <> '':    
       return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':mess})    
    dec = Decisions.objects.get(pk=dec_id) 
    something_saved = 'no'
    
    if request.method == 'POST':
       for array in request.POST.getlist('getdata'):
           print (array)
       # insert the handsontable into the mapping table. If it already exists, update it else, create it from Solution Options and Screening Criteria    
       try:
          h = MappingTable.objects.get(dec_id=dec_id)
          h.updated_by = request.session['user']
          h.updated_date = datetime.datetime.now()
          h.table = array
          h.save(update_fields=['table','updated_by', 'updated_date'])
       except ObjectDoesNotExist:
          h = MappingTable(table = array, dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
          h.save()     

       # remove the first [ from the array we got from ajax  
       arr = array[1:]                                                                                                                                                                                           
       # get the first and last postion of the solution options list
       firstpos = arr.find('[') + 1
       lastpos =  arr.find(']') 
       arr1 = arr[firstpos:lastpos]
       #print arr1
       #print 'arr1'
       solopt_list = []
       for l in arr1.split(','):
           ll = l.replace('"','')
           lj = ll.replace('Option: ', '')
           solopt_list.append(lj)
           #print 'in solopt list'
           #print lj    
       # remove the first and last array in the array of arrays - first one is the headings and last one is the empty row  
       a = arr.replace(arr1,'')
       b = a.replace('null,','')
       c = b.replace('[],','')
       y = c.replace(',[null]','')
       #print y
           
       firstpos = y.find('[') + 1
       lastpos =  y.find(']')
       w = y[firstpos:lastpos]
       #print 'w'
       #print w
       z = y.replace('[[','')
       v = z.replace(']]','')
       arr2 = v.replace('"],','",')
       #print 'arr2'
       #print arr2   
       temp_list = []
       # adding each array to a temporary list
       for l2 in arr2.split(','):
           l3 = l2.replace('"', '')
           #print 'l3'
           print l3
           if l3 == "Keep Option" or l3 =="Put it away for now":
              print l3 
              temp_list.append(l3)
       #for amm in temp_list:
           #print 'amm'
           #print amm   
       if 'Keep Option' in temp_list or 'Put it away for now' in temp_list:
          counter = 0
          for j in solopt_list: 
              try:
                 sol = Solution_Options.objects.get(dec_id=dec_id, sol_option = j)
                 #print 'temp list counter'
                 #print j
                 #print temp_list[counter]
                 if temp_list[counter] == "Keep Option":
                    sol.archived = "N"
                 elif  temp_list[counter] == "Put it away for now":
                    sol.archived = "Y"
                 sol.updated_by = request.session['user']
                 sol.updated_date = datetime.datetime.now() 
                 sol.save(update_fields=['archived','updated_by', 'updated_date'])
                 something_saved = 'yes'
                 counter = counter + 1
              except ObjectDoesNotExist:
                 print 'doesnotexist' 
   
       if something_saved == 'yes':  
          dec.updated_by = request.session['user'] 
          dec.updated_date = datetime.datetime.now()
          dec.save(update_fields=['updated_by','updated_date'])  

    solopt = Solution_Options.objects.filter(dec_id=dec_id)
    for s in solopt:
        print s.sol_option
    scrcr = Screening_Criteria.objects.filter(dec_id=dec_id) 
    scrcr_count = Screening_Criteria.objects.filter(dec_id=dec_id).count()
    test = []
    test2 = []
    try:
       mapping = MappingTable.objects.get(dec_id=dec_id)
       maptable =  mapping.table
       print "maptable"
       print maptable
       print 'in map list'                      
       print map_list
       for m in map_list:
           print 'am i in here'
           print m
           #print maptable.find('['+'"964"')
           #print maptable.find('["964"')
           firstposx = maptable.find('["'+ str(m) + '"') + 1
           print firstposx
           if firstposx <> 0:
              lastposx =  maptable.find(']', firstposx)
              print lastposx
              arr1 = maptable[firstposx:lastposx]
              print arr1
              test = maptable.replace(arr1, '')
              print test
              test2 = test.replace('[],', '')
              maptable = test2
              print "maptable and test2"
              print test2
              print maptable
    except ObjectDoesNotExist:
       maptable = 'doesnotexist'

    return render(request,'decisions/solution_options/handsontable.html', {'dec_id':dec_id, 'dec_title': dec_title, 'mapping':maptable, 'solopt':solopt, 'scrcr':scrcr, 'scrcr_count':scrcr_count, 'map_list':map_list, 'loggedinuser':loggedinuser, 'shared':shared})

def is_summary(request):                                                                                                                                                                                            
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'
    try: 
       eva = Evaluation_Criteria.objects.get(dec_id=dec_id)
    except ObjectDoesNotExist:
       print 'eva' 
       return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered evaluation criteria so you cannot view this screen.'})
    except MultipleObjectsReturned:          
       print 'eva multiple objects returned'
   
    if created_by == loggedinuser:
      try:
         Importance_Scores.objects.get(dec_id = dec_id, created_by = loggedinuser)
      except ObjectDoesNotExist:
         print 'is' 
         return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered importance scores so you cannot view this screen.'})
      except MultipleObjectsReturned:         
         print 'is multiple objects returned'
    eva_table = Evaluation_Criteria.objects.filter(dec_id = dec_id).exclude(deleted = 'Y').order_by('id')
    stdec = Stakeholders_Decisions.objects.filter(dec_id = dec_id)
    total_votes = 0
    stdec_list = []
    for st in stdec:
        if st.votes is None: 
           stvotes = 0
        else: 
           stvotes = st.votes   
        total_votes = stvotes + total_votes
    scores = Importance_Scores.objects.raw("SELECT i.id, i.score score FROM utility_tool_importance_scores i, utility_tool_stakeholders_decisions s WHERE i.dec_id = s.dec_id AND i.dec_id=%s AND s.iw_type = 'Y' AND i.email = s.email AND (i.deleted = 'N' OR i.deleted IS NULL) order by i.created_by, i.eva_id", [dec_id])
    #min_scores = Importance_Scores.objects.raw("SELECT id, MIN(score), eva_id FROM utility_tool_importance_scores WHERE dec_id=%s group by eva_id", [dec_id])
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")        
    cursor = database.cursor () 
    cursor2 = database.cursor () 
    cursor3 = database.cursor ()
    cursor.execute("""SELECT MAX(i.score) FROM utility_tool_importance_scores i, utility_tool_stakeholders_decisions s WHERE i.dec_id = s.dec_id AND i.dec_id=%s AND s.iw_type = 'Y' AND i.email = s.email AND (i.deleted = 'N' OR i.deleted IS NULL) group by i.eva_id""", [dec_id])
    cursor2.execute("""SELECT MIN(i.score) FROM utility_tool_importance_scores i, utility_tool_stakeholders_decisions s WHERE i.dec_id = s.dec_id AND i.dec_id=%s AND s.iw_type = 'Y' AND i.email = s.email AND (i.deleted = 'N' OR i.deleted IS NULL) group by i.eva_id""", [dec_id])
    cursor3.execute("""SELECT STDDEV(i.score) FROM utility_tool_importance_scores i, utility_tool_stakeholders_decisions s WHERE i.dec_id = s.dec_id AND i.dec_id=%s AND s.iw_type = 'Y' AND i.email = s.email AND (i.deleted = 'N' OR i.deleted IS NULL) group by i.eva_id""", [dec_id])
    max_scores = cursor.fetchall()
    min_scores = cursor2.fetchall()
    stdev_scores = cursor3.fetchall()
    return render(request,'decisions/solution_options/is_summary.html', {'dec_id':dec_id, 'dec_title': dec_title, 'scores':scores, 'eva_table':eva_table, 'stdec':stdec, 'total_votes':total_votes, 'loggedinuser':loggedinuser, 'min_scores':min_scores, 'max_scores':max_scores, 'stdev_scores':stdev_scores, 'created_by':created_by})


def summary(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']          
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'  
    try: 
       eva = Evaluation_Criteria.objects.get(dec_id=dec_id) 
    except ObjectDoesNotExist:
       print 'eva' 
       return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered evaluation criteria so you cannot view this screen.'}) 
    except MultipleObjectsReturned:                                                                                                                
       print 'eva multiple objects returned' 

    try:
      Importance_Scores.objects.get(dec_id = dec_id, created_by = loggedinuser)
    except ObjectDoesNotExist:
       print 'is' 
       return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered importance scores so you cannot view this screen.'})  
    except MultipleObjectsReturned:                                                                                                                
       print 'is multiple objects returned'
    eva_table = Evaluation_Criteria.objects.filter(dec_id = dec_id).exclude(deleted = 'Y') 
    scores = Importance_Scores.objects.raw("SELECT i.id, i.eva_id eva_id, i.criterion criterion, i.score score, i.created_by created_by, s.votes votes FROM utility_tool_importance_scores i, utility_tool_stakeholders_decisions s WHERE i.dec_id = s.dec_id AND i.dec_id=%s AND s.iw_type = 'Y' AND i.email = s.email AND (i.deleted = 'N' OR i.deleted IS NULL) order by i.criterion, i.created_by", [dec_id])
    return render(request,'decisions/solution_options/summary.html', {'dec_id':dec_id, 'dec_title': dec_title, 'scores':scores, 'eva_table':eva_table})  

def utility_results(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']               
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found' 

    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'

    # using a function here
    '''
    check_required(dec_id, loggedinuser, created_by)
    try:
       setup = PA_Setup.objects.get(dec_id = dec_id) 
       group_yn = setup.scores_group_yn
       votes_yn = setup.votes_yn
    except ObjectDoesNotExist:
       return render(request,'decisions/message.html', {'dec_id':dec_id, 'mess':'You have not entered the setup information for scores so you cannot view this screen. Please go back and complete the information'})
               
    if loggedinuser == created_by:
       if group_yn == 'N':
          group_cal(dec_id, loggedinuser, votes_yn)
       else:
          individual_cal(dec_id, loggedinuser)
       further_cal(dec_id, loggedinuser, request)   
    ''' 
    try: 
       evatable = EvaluationTable.objects.get(dec_id=dec_id)
       table =  evatable.table
    except ObjectDoesNotExist:
       table = 'doesnotexist'

    #TempTable.objects.all().delete()
    solopt = Solution_Options.objects.filter(dec_id=dec_id)
    solopt_count = Solution_Options.objects.filter(dec_id=dec_id).count()                                                                                                                           
    eva_count = Evaluation_Criteria.objects.filter(dec_id=dec_id).count() 
    #eva_count = eva_count1.exclude(deleted = 'Y').count()
    eva = Evaluation_Criteria.objects.filter(dec_id=dec_id)
    imp_scores = Importance_Scores.objects.filter(dec_id = dec_id, created_by = loggedinuser)
    util_res = Cost_Utility.objects.filter(dec_id = dec_id, archived='N').order_by('-weighted_utility')
    qset1 = Evaluation_Measures.objects.filter(dec_id=dec_id).order_by('opt_id', 'eva_id')

    return render(request,'decisions/solution_options/utility_results.html', {'dec_id':dec_id, 'dec_title': dec_title, 'table':table, 'evam':qset1, 'util_res':util_res, 'solopt':solopt, 'eva_count':eva_count, 'solopt_count':solopt_count, 'imp_scores':imp_scores, 'eva':eva, 'loggedinuser':loggedinuser})

def restore_idn(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    '''
    if 'old_idn_table' in request.session:
       h = IdentifyTable.objects.get(dec_id=dec_id)
       h.updated_by = request.session['user']
       h.updated_date = datetime.datetime.now()
       h.table = request.session['old_idn_table']
       #something_saved = 'yes'
       h.save(update_fields=['table','updated_by', 'updated_date'])  
    else:
       print 'nothing to do here as table isnt found'
    try:
       e_orig = Evaluation_Criteria.objects.filter(dec_id=dec_id)
       for e in e_orig.filter(orig_eva_id__isnull=False):
           print e.orig_eva_id 
           master_eva = Master_Evaluation_Criteria.objects.get(id = e.orig_eva_id, dec_id=dec_id)
           e.suggested_evam = master_eva.suggested_evam
           e.data = master_eva.data
           print e.data
           e.updated_by = request.session['user']
           e.updated_date = datetime.datetime.now()
           e.save(update_fields=['suggested_evam','data','updated_by','updated_date'])
       for e2 in e_orig.filter(orig_eva_id__isnull=True):
           e2.suggested_evam = ''
           e2.data = ''
           e2.updated_by = request.session['user']
           e2.updated_date = datetime.datetime.now()
           e2.save(update_fields=['suggested_evam','data','updated_by','updated_date'])
    except ObjectDoesNotExist:
       print 'probs' 
    '''
    IdentifyTable.objects.filter(dec_id=dec_id).last().delete()
    
    obj = IdentifyTable.objects.filter(dec_id=dec_id).order_by('-id')[:2]
    for o in obj:
        print o.id
    return HttpResponseRedirect('/utility_tool/decisions/solution_options/idn_measures.html') 

def export_results(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename=results.xls'      
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet("Results Table")
    row_num = 0  
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")        
    cursor = database.cursor ()
    sql = """SELECT sol_option, weighted_utility, cost, cost_utility_ratio  FROM utility_tool_cost_utility WHERE dec_id = %(dec_id)s AND (archived = 'N' or archived IS NULL)"""
    #Heading of tables
    a = xlwt.Alignment()
    a.wrap = True 
    a.vert = a.VERT_CENTER
    a.horz = a.HORZ_CENTER
    font_style = xlwt.XFStyle()
    font_style.font.bold = True 
    font_style.alignment = a
    pattern = xlwt.Pattern()
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = 22 
    pattern2 = xlwt.Pattern()
    pattern2.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern2.pattern_fore_colour = 22 
    font_style.pattern = pattern2                                                                                                                       
    pattern3 = xlwt.Pattern()
    pattern3.pattern_fore_colour = 1
    aL = xlwt.Alignment()                                                                                                                               
    aL.horz = a.HORZ_LEFT
    aL.wrap = True
    font_style4 = xlwt.XFStyle()
    font_style4.pattern = pattern3
    font_style4.alignment = aL
    num_style = xlwt.XFStyle()
    num_style.num_format_str = '0.00'
    row_num = 1
    columns = [
          (u"Solution Option", 5000),
          (u"Overall Utility Value or Stakeholder satisfaction 0-10", 5000),       
          (u"Total Cost", 5000),       
          (u"Cost per Unit of Utility", 5000),       
    ]
    try:
       cursor.execute(sql,{'dec_id' : dec_id})
       results = cursor.fetchall()
       if results != ():
          for col_num in xrange(len(columns)):
             ws.write(row_num, col_num, columns[col_num][0], font_style)
             # set column width
             ws.col(col_num).width = columns[col_num][1]                                                                                                
       for row in results:
          row_num += 1
          print row[0]
          sol_option = row[0]
          weighted_utility = round(row[1],2)
          cost = row[2]
          cost_utility_ratio = row[3]
          for col_num in xrange(len(row)):
              if col_num == 0: 
                 ws.write(row_num, col_num, row[col_num], font_style4)
              else:   
                 ws.write(row_num, col_num, row[col_num], num_style) 
    except:
       print "Error: unable to fetch data"
    # disconnect from server
    database.close()
    wb.save(response)                                                                                                                                   
    return response

'''def export_ut(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename=utility_results.xls'
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet("Detailed Utility Results")
    row_num = 0
    col_width = 256 * 50
    try:
       for i in itertools.count():
           ws.col(i).width = col_width
    except ValueError:
       pass
    if request.method == 'POST':
       print request.POST.getlist('hotvalue')
       for array in request.POST.getlist('hotvalue'):
           arr = array[1:]
           print 'arr'
           print arr
           for l in arr.split('\r\n'):
               ln = l.replace("'","")           
               print ln
               if not ln.startswith( ',,,,,,' ):
                  for m in ln.split(','):  
                      print m
                      h = TempTable(temptext=m)
                      h.save()
                  h = TempTable(temptext='NEXT ROW')
                  h.save()
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")        
    cursor = database.cursor ()
    sql = """SELECT temptext FROM utility_tool_temptable"""
    #Heading of tables
    a = xlwt.Alignment()
    a.wrap = True 
    a.vert = a.VERT_CENTER
    a.horz = a.HORZ_CENTER
    font_style = xlwt.XFStyle()
    font_style.font.bold = True 
    font_style.alignment = a
    pattern = xlwt.Pattern()
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = 22 
    pattern2 = xlwt.Pattern()
    pattern2.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern2.pattern_fore_colour = 22 
    font_style.pattern = pattern2                                                                                                                       
    pattern3 = xlwt.Pattern()
    pattern3.pattern_fore_colour = 1
    aL = xlwt.Alignment()                                                                                                                               
    aL.horz = a.HORZ_LEFT
    aL.wrap = True
    font_style4 = xlwt.XFStyle()
    font_style4.pattern = pattern3
    font_style4.alignment = aL
    num_style = xlwt.XFStyle()
    num_style.num_format_str = '0.00'
    row_num = 1
    col_num = 0
    try:                              
       cursor.execute(sql)
       results = cursor.fetchall()
       for row in results:
          if row[0] == 'NEXT ROW': 
             row_num += 1
             col_num = 0
          else:
             print 'row 0'
             print row[0]
             print row_num
             print col_num 
             ws.write(row_num, col_num, row[0], font_style4)
             col_num += 1 
    except:
       print "Error: unable to fetch data"
    # disconnect from server
    database.close()
    wb.save(response)                                                                                                                                   
    return response
'''
def export_mea(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename=evaluation_data.xls'      
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet("Evaluation Data Table")
    row_num = 0  
    col_width = 256 * 30     
    try: 
       for i in itertools.count():
           ws.col(i).width = col_width
    except ValueError:
       pass 
    #Heading of tables
    a = xlwt.Alignment()
    a.wrap = True 
    a.vert = a.VERT_CENTER
    a.horz = a.HORZ_CENTER
    font_style = xlwt.XFStyle()
    font_style.font.bold = True 
    font_style.alignment = a
    pattern = xlwt.Pattern()
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = 22 
    pattern2 = xlwt.Pattern()
    pattern2.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern2.pattern_fore_colour = 22 
    font_style.pattern = pattern2
    pattern3 = xlwt.Pattern()
    pattern3.pattern_fore_colour = 1
    aL = xlwt.Alignment()     
    aL.horz = a.HORZ_LEFT
    aL.wrap = True 
    font_style4 = xlwt.XFStyle()                                                                                                                                                                                 
    font_style4.pattern = pattern3
    font_style4.alignment = aL 
    #num_style = xlwt.XFStyle()
    #num_style.num_format_str = '0.00'
    row_num = 0
    col_num = 0
 
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    try:
       s = EvaluationTable.objects.get(dec_id=dec_id) 
       table =  s.table
    except ObjectDoesNotExist:
       table = 'doesnotexist'
    except MultipleObjectsReturned:
        s = EvaluationTable.objects.filter(dec_id=dec_id).last()
        table =  s.table
    print table
    arrx = table[1:]                                                                                                                                                                                           
    # get the first and last postion of the solution options list
    firstposx = arrx.find('[') + 1
    lastposx =  arrx.find(']') 
    arrx1 = arrx[firstposx:lastposx]
    #print 'arr1'
    pos = 0
    pos2 = 0
    archived_list = []
    for l in arrx1.split(','):    
        if l not in ("Evaluation Criterion", "How will you measure this?", "Data to collect", "Likely lowest score", "Likely highest score", "Higher scores are better? (Yes/No)"): 
           lm = l.replace(" (Average Rating or Score)","")
           ln = lm.replace('"','')    
           try: 
              solopt = Solution_Options.objects.get(sol_option = ln, dec_id=dec_id)
              if solopt.archived == 'Y' or solopt.deleted == 'Y':
                 archived_list.append(pos)
           except ObjectDoesNotExist:
              print 'nothing todo'  
           pos = pos + 1   
    new_list = []                                                                                                                                    
    '''print 'archived_list'
    for a in archived_list:
        print a'''
    # remove the first [ from the array we got from ajax  
    aa = table.replace('",null,"','",,"')
    #b = aa.replace('null,','')
    c = aa.replace('[],','')
    #w = c.replace(',[null]','')
    z = c.replace('[[','[')
    yy = z.replace('[""]','[')
    y = yy.replace('"],"','","')
    while len(y): 
       firstpos = y.find('[') + 1                                                                                                                                                                                
       lastpos =  y.find(']')
       arr2 = y[firstpos:lastpos]
       x = arr2.split(',')[0]
       print arr2
       for l2 in arr2.split(','):
           l3 = l2.replace('"', '')
           print l3 
           if ((pos2 not in archived_list) and (x <> '0' and x <> 0 and x <> 'None')):
              if row_num == 0:
                 ws.write(row_num, col_num, l3, font_style)
                 col_num = col_num + 1
              else:
                  try:
                     evacr = Evaluation_Criteria.objects.get(combined = l3, dec_id=dec_id)
                     if evacr.deleted <> 'Y':
                        ws.write(row_num, col_num, l3, font_style4) 
                        col_num = col_num + 1  
                     else:
                        break;   
                  except ObjectDoesNotExist:  
                     if (l3 == 'null'):
                          ws.write(row_num, col_num, ' ', font_style4)
                     else:     
                         ws.write(row_num, col_num, l3, font_style4) 
                     col_num = col_num + 1  
                  except MultipleObjectsReturned:      
                     if (l3 == 'null'):
                          ws.write(row_num, col_num, ' ', font_style4)
                     else:          
                         ws.write(row_num, col_num, l3, font_style4) 
                     col_num = col_num + 1  
           pos2 = pos2 + 1
       row_num = row_num + 1
       col_num = 0
       pos2 = 0
       z = y.replace(arr2, '')                                                                                                                       
       c = z.replace('[],','')
       # break out of the loop when only []] remains
       if (c == '[]]'):
          break;
       y = c
    wb.save(response)
    return response

def export_idn(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename=identify_measures.xls'                                                                        
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet("Identify Evaluation Measures")
    row_num = 0 

    col_width = 256 * 30     

    try:
       for i in itertools.count():
           ws.col(i).width = col_width
    except ValueError:
       pass

    #Heading of tables
    a = xlwt.Alignment()
    a.wrap = True
    a.vert = a.VERT_CENTER
    a.horz = a.HORZ_CENTER
    font_style = xlwt.XFStyle()
    font_style.font.bold = True
    font_style.alignment = a
    pattern = xlwt.Pattern()
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = 22
    pattern2 = xlwt.Pattern()
    pattern2.pattern = xlwt.Pattern.SOLID_PATTERN
    #pattern2.pattern_fore_colour = 22
    pattern2.pattern_fore_colour = xlwt.Style.colour_map['light_turquoise']   
    font_style.pattern = pattern2
    pattern3 = xlwt.Pattern()
    pattern3.pattern_fore_colour = 1
    aL = xlwt.Alignment()                                                                                                                               
    aL.horz = a.HORZ_LEFT
    aL.wrap = True
    font_style4 = xlwt.XFStyle()
    font_style4.pattern = pattern3
    font_style4.alignment = aL
    row_num = 0
    col_num = 0
 
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    try:
       s = IdentifyTable.objects.get(dec_id=dec_id) 
       table =  s.table
    except ObjectDoesNotExist:
       table = 'doesnotexist'
    except MultipleObjectsReturned:
        s = IdentifyTable.objects.filter(dec_id=dec_id).last()
        table =  s.table

    print table
    arrx = table[1:]                                                                                                                                                                                           
    # get the first and last postion of the solution options list
    firstposx = arrx.find('[') + 1
    lastposx =  arrx.find(']') 
    arrx1 = arrx[firstposx:lastposx]
    #print 'arr1'
    pos = 0
    pos2 = 0
    archived_list = []
    for l in arrx1.split(','):    

        if l not in ("Evaluation Criterion", "Common evaluation measure you can use across all options", "Data to collect"): 
           lm = l.replace("Describe the information you will use to evaluate  ","")
           lo = lm.replace('"','')    
           ln = lo.replace("  against this criterion and where you will get it from.","") 
           try: 
              solopt = Solution_Options.objects.get(sol_option = ln, dec_id=dec_id)
              if solopt.archived == 'Y' or solopt.deleted == 'Y':
                 archived_list.append(pos)
           except ObjectDoesNotExist:
              print 'nothing todo' 
           except MultipleObjectsReturned:  
              for s in Solution_Options.objects.filter(sol_option = ln, dec_id=dec_id):
                  if s.archived == 'Y' or s.deleted == 'Y':
                     archived_list.append(pos)
                     #pos = pos + 1
           pos = pos + 1   
    new_list = []                                                                                                                                    
    # remove the first [ from the array we got from ajax  
    aa = table.replace('",null,"','",,"')
    #b = aa.replace('null,','')
    c = aa.replace('[],','')
    #w = c.replace(',[null]','')
    z = c.replace('[[','[')
    yy = z.replace('[""]','[')
    y = yy.replace('"],"','","')

    while len(y): 
       firstpos = y.find('[') + 1
       lastpos =  y.find(']')
       arr2 = y[firstpos:lastpos]
       x = arr2.split(',')[0]
       print arr2
       for l2 in arr2.split(','):
           l3 = l2.replace('"', '')
           print l3 
           if ((pos2 not in archived_list) and (x <> '0' and x <> 0 and x <> 'None')):
              if row_num == 0:
                 ws.write(row_num, col_num, l3, font_style)
                 col_num = col_num + 1
              else:
                  try:
                     evacr = Evaluation_Criteria.objects.get(combined = l3, dec_id=dec_id)
                     if evacr.deleted <> 'Y':
                        ws.write(row_num, col_num, l3, font_style4) 
                        col_num = col_num + 1  
                     else:
                        break;   
                  except ObjectDoesNotExist:  
                     if (l3 == 'null'):
                          ws.write(row_num, col_num, ' ', font_style4)
                     else:     
                         ws.write(row_num, col_num, l3, font_style4) 
                     col_num = col_num + 1  
                  except MultipleObjectsReturned:      
                     if (l3 == 'null'):
                          ws.write(row_num, col_num, ' ', font_style4)
                     else:          
                         ws.write(row_num, col_num, l3, font_style4) 
                     col_num = col_num + 1  
           pos2 = pos2 + 1
       row_num = row_num + 1
       col_num = 0
       pos2 = 0
       z = y.replace(arr2, '')                                                                                                                       
       c = z.replace('[],','')
       # break out of the loop when only []] remains
       if (c == '[]]'):
          break;
       y = c
    wb.save(response)
    return response

def idn_measures(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']               
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'
    dec = Decisions.objects.get(pk=dec_id) 
    something_saved = 'no'
    try:
       s = IdentifyTable.objects.get(dec_id=dec_id) 
       table =  s.table
    except ObjectDoesNotExist:
       table = 'doesnotexist'
    except MultipleObjectsReturned:
        s = IdentifyTable.objects.filter(dec_id=dec_id).last()
        table =  s.table

    evacr = Evaluation_Criteria.objects.filter(dec_id=dec_id) 
    solopt = Solution_Options.objects.filter(dec_id=dec_id)

    if request.method == 'POST':
       for array in request.POST.getlist('getdata'):
            print array
            # insert the handsontable into identify table. If it already exists, update it else create it
            h = IdentifyTable(table = array, dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())                                                                       
            something_saved = 'yes'
            h.save()

       print 'Identify table'   
       new_list = [] 
       # remove the first [ from the array we got from ajax  
       arr = array[1:]
       # get the first and last postion of the solution options list
       firstpos = arr.find('[') + 1
       lastpos =  arr.find(']') 
       arr1 = arr[firstpos:lastpos]
       # remove the first and last array in the array of arrays - first one is the headings and last one is the empty row  
       a = arr.replace(arr1,'')
       aa = a.replace('",null,"','",,"')
       b = aa.replace('null,','')
       c = b.replace('[],','')
       w = c.replace(',[null]','')
       z = w.replace('[[','[')
       yy = z.replace('[""]','[')
       y = yy.replace('"],"','","')
       print 'y'
       print y

       # loop through the remaining array of arrays
       while len(y): 
          firstpos = y.find('[') + 1
          lastpos =  y.find(']')
          arr2 = y[firstpos:lastpos]
          temp_list = [] 

          # adding each array to a temporary list
          for l2 in arr2.split(','):
             l3 = l2.replace('"', '')
             temp_list.append(l3)
          # till here  
          print 'TEMP LIST 1'
          print temp_list[1]
          print temp_list[0]
          try:
             for evm in Evaluation_Measures.objects.filter(criterion = temp_list[0], dec_id=dec_id):
                 evm.measure = temp_list[1]
                 evm.unit = temp_list[2] 
                 evm.updated_by = loggedinuser
                 evm.updated_date = datetime.datetime.now()
                 evm.save(update_fields=['measure','unit','updated_by','updated_date'])
                 something_saved = 'yes'
          except ObjectDoesNotExist:
              print 'evam does not exist'  
          try:
             for e in Evaluation_Criteria.objects.filter(combined = temp_list[0], dec_id=dec_id):
                 e.suggested_evam = temp_list[1]
                 e.data = temp_list[2] 
                 e.updated_by = loggedinuser
                 e.updated_date = datetime.datetime.now()
                 e.save(update_fields=['suggested_evam','data','updated_by','updated_date'])
                 something_saved = 'yes'
          except ObjectDoesNotExist:
              print 'evam does not exist' 
          print 'out of that loop'
          z = y.replace(arr2, '')                                                                                                                                                                                
          c = z.replace('[],','')
          # break out of the loop when only []] remains
          if (c == '[]]'):
             break;
          y = c

    if something_saved == 'yes':
       dec.updated_by = request.session['user']
       dec.updated_date = datetime.datetime.now()
       dec.save(update_fields=['updated_by','updated_date'])

    return render(request,'decisions/solution_options/idn_measures.html', {'dec_id':dec_id, 'dec_title': dec_title, 'evacr':evacr, 'solopt':solopt, 'table':table, 'loggedinuser':loggedinuser, 'created_by':created_by})                                                                                                         


def add_measures(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']               
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'

    dec = Decisions.objects.get(pk=dec_id) 
    something_saved = 'no'

    try: 
       evatable = EvaluationTable.objects.get(dec_id=dec_id)
       table =  evatable.table
    except ObjectDoesNotExist:
       table = 'doesnotexist'
    print 'AAAAAAAAA'        
    evacr = Evaluation_Criteria.objects.filter(dec_id=dec_id) 
    solopt = Solution_Options.objects.filter(dec_id=dec_id)
    #evacr_count = Evaluation_Criteria.objects.filter(dec_id=dec_id).count()                                                                                                                                     
    solopt_count = Solution_Options.objects.filter(dec_id=dec_id,archived='N', deleted='N').count()
    if request.method == 'POST':
       print request.POST.getlist('getdata')
       for array in request.POST.getlist('getdata'):
            print array

       # insert the handsontable into evaluation table. If it already exists, update it else create it
       try: 
          h = EvaluationTable.objects.get(dec_id=dec_id)
          h.updated_by = request.session['user']
          h.updated_date = datetime.datetime.now()
          h.table = array
          something_saved = 'yes'
          h.save(update_fields=['table','updated_by', 'updated_date'])
       except ObjectDoesNotExist:
           h = EvaluationTable(table = array, dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
           something_saved = 'yes'
           h.save()
       # insert into evaluation_measures - individual records with only option / criterion combination
       for s in solopt:
          for e in evacr:
             try:
                 m = Evaluation_Measures.objects.get(opt_id=s.id, eva_id = e.id)
             except ObjectDoesNotExist:
                 print 'this is for testing'
                 print s.sol_option
                 if e.criterion2 is not None and e.criterion2 <> '':
                     criterion = e.criterion + ': ' + e.criterion2
                 else:   
                    criterion = e.criterion 
                 if s.deleted == 'Y' or s.archived == 'Y':
                    archived = 'Y'
                 else:
                    archived = 'N' 
                 print archived   
                 m = Evaluation_Measures(opt_id = s.id, sol_option = s.sol_option, archived = archived, deleted = e.deleted, eva_id = e.id, criterion = criterion,  dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
                 something_saved = 'yes' 
                 m.save() 
   
       new_list = []
       # remove the first [ from the array we got from ajax  
       arr = array[1:]
       # get the first and last postion of the solution options list
       firstpos = arr.find('[') + 1
       lastpos =  arr.find(']') 
       arr1 = arr[firstpos:lastpos]
       #print 'arr1'
       #print arr1
       #print 'FIRST STEP'  
       # remove the first and last array in the array of arrays - first one is the headings and last one is the empty row  
       a = arr.replace(arr1,'')
       aa = a.replace('",null,"','",0,"')
       b = aa.replace('null,','0,')
       c = b.replace('[],','')
       w = c.replace(',[null]','')
       #print 'w'
       #print w
       #print 'SECOND STEP'
       z = w.replace('[[','[')
       yy = z.replace('[""]','[0')
       #yy = z.replace('[""]','""')
       y = yy.replace('"],"','","')
       print y
       print 'AND AT THE Y'
       # loop through the remaining array of arrays
       while len(y): 
          firstpos = y.find('[') + 1
          lastpos =  y.find(']')
          #print firstpos
          #print lastpos
          arr2 = y[firstpos:lastpos]
          #print arr2
          temp_list = []
          # adding each array to a temporary list
          for l2 in arr2.split(','):
             l3 = l2.replace('"', '')
             temp_list.append(l3)
          #print temp_list  
          print len(temp_list)         
          if len(temp_list) >= 7:   
             entering_loop_first_time = 6 
             print 'inside len if'
             #print 'temp_list[0]'
             print temp_list[0]
             #print temp_list[1]
             #sol1 = solopt.exclude(archived = 'Y')    
             for s in solopt:
                #print 'do i enter this loop' 
                #print s.id
                #print temp_list[0]
                try:
                   print 'in try'
                   #print dec_id
                   print temp_list[0]
                   print s.id
                   em = Evaluation_Measures.objects.filter(dec_id = dec_id, criterion = temp_list[0], opt_id = s.id)
                   #em2 = em1.exclude(archived = 'Y')
                   #em = em2.exclude(deleted = 'Y')   
                   print 'before for'
                   for evam in em:
                       #print evam.id
                       # update each evaluation measure record with the values in the evaluation table
                       # lowest value, highest value etc. are float values
                       # higher_better is only one character
                       print 'inside evam'
                       m = Evaluation_Measures.objects.get(id = evam.id) 
                       #measure | unit | lowest_value | highest_value | higher_better | option_value
                       m.measure = temp_list[1]
                       '''print temp_list[1]
                       print temp_list[2]
                       print temp_list[3]
                       print temp_list[4]
                       print temp_list[5]
                       print temp_list[6]
                       print temp_list[7]'''
                       #print temp_list[8]
                       #print temp_list[9]
                       #print temp_list[10]
                       m.unit = temp_list[2]
                       m.lowest_value = float(temp_list[3])
                       m.highest_value = float(temp_list[4])
                       m.higher_better = temp_list[5][0][0]
                       print s.id
                       print evam.id
                       print 'entering_loop_first_time' 
                       print entering_loop_first_time
                       if entering_loop_first_time == 6:
                          m.option_value = float(temp_list[6]) 
                          #print 'if'
                          #print m.option_value
                          entering_loop_first_time = entering_loop_first_time + 1
                       else:
                          try: 
                             #print temp_list[entering_loop_first_time] 
                             m.option_value = float(temp_list[entering_loop_first_time])  
                          except:
                             #print 'ZERO' 
                             m.option_value = 0
                          entering_loop_first_time = entering_loop_first_time + 1
                          #print 'else'                                                                                                                                                                        
                       print 'option value'   
                       print m.option_value
                       m.updated_by = request.session['user']
                       m.updated_date = datetime.datetime.now() 
                       m.save(update_fields=['measure','unit','lowest_value','highest_value', 'higher_better', 'option_value', 'updated_by', 'updated_date'])
                       something_saved ='yes'
                except:
                   print 'cannot proceed - error'  
          # till here  
          print 'out of that loop'
          z = y.replace(arr2, '')
          c = z.replace('[],','')
          # break out of the loop when only []] remains
          if (c == '[]]'):
             break;
          y = c

       if something_saved == 'yes':  
          dec.updated_by = request.session['user'] 
          dec.updated_date = datetime.datetime.now()
          dec.save(update_fields=['updated_by','updated_date'])  
       #return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)       
    return render(request,'decisions/solution_options/add_measures.html', {'dec_id':dec_id, 'dec_title': dec_title, 'evacr':evacr, 'table':table, 'solopt':solopt,  'solopt_count':solopt_count, 'loggedinuser':loggedinuser, 'created_by':created_by})

'''
def add_eva_results(request):
    if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']               
    else:
       dec_title = 'not found'
    dec = Decisions.objects.get(pk=dec_id) 
    something_saved = 'no'
    if request.method == 'POST':
       print request.POST.getlist('getdata')
       for array in request.POST.getlist('getdata'):
          print (array)
          # insert the handsontable into evaluation_measures table. If it already exists, update it else create it
       try:
          h = Evaluation_Measures.objects.get(dec_id=dec_id)
          h.updated_by = request.session['user']
          h.updated_date = datetime.datetime.now()
          h.table = array
          something_saved = 'yes'
          h.save(update_fields=['table','updated_by', 'updated_date'])
       except ObjectDoesNotExist:
          h = Evaluation_Measures(table = array, dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
          something_saved = 'yes' 
          h.save()      
       if something_saved == 'yes':  
          dec.updated_by = request.session['user'] 
          dec.updated_date = datetime.datetime.now()
          dec.save(update_fields=['updated_by','updated_date'])  
    evacr = Evaluation_Criteria.objects.filter(dec_id=dec_id)      
    for e in evacr:
        print e.criterion
    try:
       evatable = Evaluation_Measures.objects.get(dec_id=dec_id)
       table =  evatable.table
    except ObjectDoesNotExist:
       table = 'doesnotexist'
    evacr = Evaluation_Criteria.objects.filter(dec_id=dec_id)   
    return render(request,'decisions/solution_options/add_eva_results.html', {'dec_id':dec_id, 'dec_title': dec_title, 'evacr':evacr, 'table':table})


def cost_setup(request):
    if 'dec_id' in request.session:                                                                                                                
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']     
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'
 
    context = RequestContext(request)

    try: 
       c = Cost_Setup.objects.get(dec_id = dec_id) 
    except ObjectDoesNotExist: 
       c = Cost_Setup(type_of_cost = 'Total', dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
       c.save()  
       c = Cost_Setup.objects.get(dec_id = dec_id) 

    dec = Decisions.objects.get(pk=dec_id) 
    if request.method == 'POST':
        print request.POST.get('no_of_part')
        setupform = CostSetupForm(request.POST)
        if setupform.is_valid():
           id = setupform.save(commit=False) 
           c.updated_by = request.session['user']
           c.updated_date = datetime.datetime.now()
           c.type_of_cost = id.type_of_cost
           c.save(update_fields=['type_of_cost', 'updated_by','updated_date',])
           dec.updated_by = request.session['user'] 
           dec.updated_date = datetime.datetime.now()
           dec.save(update_fields=['updated_by','updated_date'])
           return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)  
        else:
            print setupform.errors
    else:
        setupform = CostSetupForm(instance=c)

        if loggedinuser != created_by:
           setupform.fields['type_of_cost'].widget.attrs['disabled'] = True
    return render(request,'decisions/solution_options/costs1.html',{'dec_id':dec_id, 'dec_title':dec_title, 'setupform':setupform, 'loggedinuser':loggedinuser, 'created_by':created_by})
'''
def cost_table(request):                                                                                                                           
    if 'dec_id' in request.session:                                                                                                                
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']     
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'
 
    context = RequestContext(request)
    dec = Decisions.objects.get(pk=dec_id) 
    try: 
       c = Cost_Setup.objects.get(dec_id = dec_id) 
       type_of_cost = c.type_of_cost
       if type_of_cost == "Total":
          cost_text = "Total Cost"
       elif type_of_cost == "Avg":
          cost_text = "Average Cost"
       else:
          cost_text = "Marginal Cost"
    except ObjectDoesNotExist: 
       return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered costs so you cannot view this screen.'})   

    try:
       costs = Cost_Utility.objects.get(dec_id = dec_id)
    except ObjectDoesNotExist:
       for s in Solution_Options.objects.filter(dec_id=dec_id,archived='N', deleted='N'):
           costs  = Cost_Utility(opt_id = s.id, sol_option = s.sol_option,  dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
           costs.save()
       dec.updated_by = request.session['user'] 
       dec.updated_date = datetime.datetime.now()
       dec.save(update_fields=['updated_by','updated_date'])          
    except MultipleObjectsReturned:
       print 'multiple objects returned'

    costs1 = Cost_Utility.objects.filter(dec_id = dec_id)
    costs2 = costs1.exclude(archived = 'Y')
    costs = costs2.exclude(deleted = 'Y')
    if request.method == 'POST':
       #print request.POST.get('id') 
       if 'id' in request.POST:
          #changed = False
          c = Cost_Utility.objects.get(pk=request.POST.get('id'))
          '''
          if int(c.no_of_participants) != int(request.POST.get('no_of_participants')):
             c.no_of_participants = request.POST.get('no_of_participants')
             changed = True 
          if float(c.cost) != float(request.POST.get('cost')):
             c.cost = request.POST.get('cost')
             changed = True

          if changed == True:
          '''
          if request.POST.get('no_of_participants') == "None":
             c.no_of_participants = 0
          else:
             c.no_of_participants = request.POST.get('no_of_participants')
          c.cost = request.POST.get('cost')
          c.updated_by = request.session['user']
          c.updated_date = datetime.datetime.now()
          c.save(update_fields=['no_of_participants','cost','updated_by','updated_date'])
          dec.updated_by = request.session['user'] 
          dec.updated_date = datetime.datetime.now()
          dec.save(update_fields=['updated_by','updated_date'])
       #return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id) 

    return render(request,'decisions/solution_options/cost_table.html',{'dec_id':dec_id, 'dec_title':dec_title, 'costs_table':costs, 'type_of_cost':type_of_cost, 'cost_text':cost_text, 'loggedinuser':loggedinuser, 'created_by':created_by})

def cost_setup(request):
    if 'dec_id' in request.session:                                                                                                                     
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']          
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'
 
    context = RequestContext(request)
    source = ''
    print 'I AM IN COST SETUP'
    try: 
       c = Cost_Setup.objects.get(dec_id = dec_id) 
       type_of_cost = c.type_of_cost
       source = c.source
       if type_of_cost == "Total":
          cost_text = "Total Costs"
       elif type_of_cost == "Avg":
          cost_text = "Average Costs"
       else:                                                                                                                                                                                                     
          cost_text = "Marginal Costs"
    except ObjectDoesNotExist:
       c = Cost_Setup(type_of_cost = 'Total', dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
       c.save()  
       c = Cost_Setup.objects.get(dec_id = dec_id) 
       type_of_cost = c.type_of_cost
       if type_of_cost == "Total":
          cost_text = "Total Costs"
       elif type_of_cost == "Avg":
          cost_text = "Average Costs"
       else:
          cost_text = "Marginal Costs"
    dec = Decisions.objects.get(pk=dec_id) 
    for s in Solution_Options.objects.filter(dec_id=dec_id, deleted = 'N'):    
        try:
           costs = Cost_Utility.objects.get(opt_id = s.id)
           #if costs.sol_option <> s.sol_option:
           costs.sol_option = s.sol_option
           costs.archived = s.archived
           costs.updated_by = request.session['user']
           costs.updated_date = datetime.datetime.now()
           costs.save(update_fields=['sol_option','archived','updated_by','updated_date'])
        except ObjectDoesNotExist:
           costs  = Cost_Utility(opt_id = s.id, sol_option = s.sol_option, archived = s.archived,  dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
           costs.save()
          
        try: 
           detcosts = Detailed_Costs.objects.get(opt_id = s.id)
           if s.archived == 'Y':
              detcosts.archived = 'Y' 
              #Detailed_Costs.objects.get(opt_id = s.id).delete()
           else:
              detcosts.archived = 'N' 
           detcosts.save(update_fields=['archived'])       
           if detcosts.sol_option <> s.sol_option:   
              detcosts.sol_option = s.sol_option
              detcosts.updated_by = request.session['user']
              detcosts.updated_date = datetime.datetime.now()
              detcosts.save(update_fields=['sol_option','updated_by','updated_date'])
        except ObjectDoesNotExist:
           if s.archived == 'N':  
              detcosts  = Detailed_Costs(opt_id = s.id, sol_option = s.sol_option,archived ='N',dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
              detcosts.save()
    cost_table1 = Cost_Utility.objects.filter(dec_id = dec_id)
    cost_table = cost_table1.exclude(archived = 'Y')
    detcosts1 = Detailed_Costs.objects.filter(dec_id = dec_id)
    detcosts = detcosts1.exclude(archived = 'Y')
    print 'before post'
    if request.method == 'POST':
       print request.POST.get('id')
       print request.POST.get('part')
       print request.POST.get('cost')
       if request.POST.get('id'): 
          costs = Cost_Utility.objects.get(pk=request.POST.get('id'))
          if request.POST.get('part') == "None":
             costs.no_of_participants = 0
          else:
             costs.no_of_participants = request.POST.get('part')
          costs.cost = request.POST.get('cost')
          costs.updated_by = request.session['user']
          costs.updated_date = datetime.datetime.now()
          print request.POST.get('cost')
          costs.save(update_fields=['no_of_participants','cost','updated_by','updated_date'])
          c.type_of_cost = request.POST.get('radioValue')
          c.source = request.POST.get('source')
          c.updated_by = request.session['user']
          c.updated_date = datetime.datetime.now()
          c.save(update_fields=['type_of_cost','source','updated_by','updated_date'])
          type_of_cost = c.type_of_cost
          if type_of_cost == "Total":
             cost_text = "Total Cost"
          elif type_of_cost == "Avg":
             cost_text = "Average Cost"
          else:
             cost_text = "Marginal Cost"    
             
       if request.POST.get('d_id'):
          dcosts = Detailed_Costs.objects.get(pk=request.POST.get('d_id'))
          dcosts.personnel_cost = request.POST.get('personnel_cost')
          dcosts.facilities_cost = request.POST.get('facilities_cost')
          dcosts.materials_cost = request.POST.get('materials_cost')
          dcosts.training_cost = request.POST.get('training_cost')
          dcosts.other_cost = request.POST.get('other_cost')
          dcosts.total_cost = float(request.POST.get('personnel_cost')) + float(request.POST.get('facilities_cost'))  + float(request.POST.get('materials_cost')) + float(request.POST.get('training_cost'))  + float(request.POST.get('other_cost'))       
          dcosts.updated_by = request.session['user']
          dcosts.updated_date = datetime.datetime.now()
          dcosts.save(update_fields=['personnel_cost','facilities_cost','materials_cost','training_cost','other_cost','total_cost','updated_by','updated_date'])                                                 
    return render(request,'decisions/solution_options/costs1.html',{'dec_id':dec_id, 'dec_title':dec_title, 'loggedinuser':loggedinuser, 'created_by':created_by, 'type_of_cost':type_of_cost, 'source':source,'cost_text':cost_text, 'cost_table':cost_table, 'detcosts':detcosts})            
'''
def cost_setup(request):
    if 'dec_id' in request.session:                                                                                                                
       dec_id = request.session['dec_id']
    else:
       dec_id = 0
    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']     
    else:
       dec_title = 'not found'
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'
 
    context = RequestContext(request)
    print 'I AM IN COST SETUP'
    try: 
       c = Cost_Setup.objects.get(dec_id = dec_id) 
       type_of_cost = c.type_of_cost
       if type_of_cost == "Total":
          cost_text = "Total Costs"
       elif type_of_cost == "Avg":
          cost_text = "Average Costs"
       else:                                                                                                                                                                                                     
          cost_text = "Marginal Costs"
    except ObjectDoesNotExist:
       c = Cost_Setup(type_of_cost = 'Total', dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
       c.save()  
       c = Cost_Setup.objects.get(dec_id = dec_id) 
       type_of_cost = c.type_of_cost
       if type_of_cost == "Total":
          cost_text = "Total Costs"
       elif type_of_cost == "Avg":
          cost_text = "Average Costs"
       else:
          cost_text = "Marginal Costs"

    dec = Decisions.objects.get(pk=dec_id) 

    for s in Solution_Options.objects.filter(dec_id=dec_id):
        try:
           costs = Cost_Utility.objects.get(opt_id = s.id)
           #if costs.sol_option <> s.sol_option:
           costs.sol_option = s.sol_option
           costs.archived = s.archived
           costs.updated_by = request.session['user']
           costs.updated_date = datetime.datetime.now()
           costs.save(update_fields=['sol_option','archived','updated_by','updated_date'])
        except ObjectDoesNotExist:
           costs  = Cost_Utility(opt_id = s.id, sol_option = s.sol_option, archived = s.archived,  dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
           costs.save()
          
        try: 
           detcosts = Detailed_Costs.objects.get(opt_id = s.id)
           if s.archived == 'Y':
              Detailed_Costs.objects.get(opt_id = s.id).delete()
           elif detcosts.sol_option <> s.sol_option:
              detcosts.sol_option = s.sol_option
              detcosts.updated_by = request.session['user']
              detcosts.updated_date = datetime.datetime.now()
              detcosts.save(update_fields=['sol_option','updated_by','updated_date'])
        except ObjectDoesNotExist:
           if s.archived == 'N':  
              detcosts  = Detailed_Costs(opt_id = s.id, sol_option = s.sol_option,  dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
              detcosts.save()
    try:
       costs = Cost_Utility.objects.get(dec_id = dec_id)
    except ObjectDoesNotExist:
       for s in Solution_Options.objects.filter(dec_id=dec_id,archived='N'):
           costs  = Cost_Utility(opt_id = s.id, sol_option = s.sol_option,  dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
           costs.save()
       dec.updated_by = request.session['user'] 
       dec.updated_date = datetime.datetime.now()
       dec.save(update_fields=['updated_by','updated_date'])          
    except MultipleObjectsReturned:
       print 'multiple objects returned'
    cost_table1 = Cost_Utility.objects.filter(dec_id = dec_id)
    cost_table = cost_table1.exclude(archived = 'Y')
    try: 
       detcosts = Detailed_Costs.objects.get(dec_id = dec_id)
    except ObjectDoesNotExist:
       for s in Solution_Options.objects.filter(dec_id=dec_id,archived='N'):
           detcosts  = Detailed_Costs(opt_id = s.id, sol_option = s.sol_option,  dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
           detcosts.save()
    except MultipleObjectsReturned:
       print 'multiple objects returned'
    detcosts = Detailed_Costs.objects.filter(dec_id = dec_id)

    if request.method == 'POST':
       if request.POST.get('id'): 
          costs = Cost_Utility.objects.get(pk=request.POST.get('id'))
          if request.POST.get('no_of_participants') == "None":
             costs.no_of_participants = 0
          else:
             costs.no_of_participants = request.POST.get('no_of_participants')
          costs.cost = request.POST.get('cost')
          costs.updated_by = request.session['user']
          costs.updated_date = datetime.datetime.now()
          print request.POST.get('cost')
          costs.save(update_fields=['no_of_participants','cost','updated_by','updated_date'])
          c.type_of_cost = request.POST.get('radioValue')
          c.updated_by = request.session['user']
          c.updated_date = datetime.datetime.now()
          c.save(update_fields=['type_of_cost','updated_by','updated_date'])
          type_of_cost = c.type_of_cost
          if type_of_cost == "Total":
             cost_text = "Total Cost"
          elif type_of_cost == "Avg":
             cost_text = "Average Cost"
          else:
             cost_text = "Marginal Cost"    
       if request.POST.get('d_id'):
          dcosts = Detailed_Costs.objects.get(pk=request.POST.get('d_id'))
          dcosts.personnel_cost = request.POST.get('personnel_cost')
          dcosts.facilities_cost = request.POST.get('facilities_cost')
          dcosts.materials_cost = request.POST.get('materials_cost')
          dcosts.training_cost = request.POST.get('training_cost')
          dcosts.other_cost = request.POST.get('other_cost')
          dcosts.total_cost = float(request.POST.get('personnel_cost')) + float(request.POST.get('facilities_cost'))  + float(request.POST.get('materials_cost')) + float(request.POST.get('training_cost'))  + float(request.POST.get('other_cost')) 
          dcosts.updated_by = request.session['user']
          dcosts.updated_date = datetime.datetime.now()
          dcosts.save(update_fields=['personnel_cost','facilities_cost','materials_cost','training_cost','other_cost','total_cost','updated_by','updated_date'])
    return render(request,'decisions/solution_options/costs1.html',{'dec_id':dec_id, 'dec_title':dec_title, 'loggedinuser':loggedinuser, 'created_by':created_by, 'type_of_cost':type_of_cost, 'cost_text':cost_text, 'cost_table':cost_table, 'detcosts':detcosts})
'''
def decision_made(request):                                                                                                                                                                                         
    if 'dec_id' in request.session:                                                                                                                     
       dec_id = request.session['dec_id']
    else:
       dec_id = 0

    if 'dec_title' in request.session:
       dec_title = request.session['dec_title']          
    else:
       dec_title = 'not found'

    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'
 
    if 'user_email' in request.session: 
       user_email = request.session['user_email']                                                                                                     
    else:
       user_email = 'not found'

    context = RequestContext(request)
    
       # using a function here
    retval = check_required(request, dec_id)
    sol_mesg = ''
    eva_mesg = ''
    iw_mesg = ''
    evm_mesg = ''
    temp_mesg = 'You have not provided the following information to calculate the utility values. Please go back to the flowchart and complete the relevant step(s). '
    if 'solopt' in retval:
        sol_mesg = '- Identify Solution Options in the "Solution Options" step'
    if 'eva' in retval:    
        eva_mesg = '- Identify Evaluation Criteria in the "Evaluation Criteria" step' 
    if 'iw' in retval or 'listerr3' in retval:
        iw_mesg =  '- Add the Importance Scores for the Evaluation Criteria in the "Importance Scores" step' 
    if 'mea' in retval or 'listerr1' in retval or 'listerr2' in retval:
        print 'retval'
        print retval
        evm_mesg = '- Specify the evaluation measures and add the average rating/score for each Solution Option in the "Evidence-gathering to Evaluate Options" step' 
    if retval <> '':
        return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':temp_mesg, 'sol_mesg':sol_mesg,'eva_mesg':eva_mesg,'iw_mesg':iw_mesg,'evm_mesg':evm_mesg,})    
    ''' 
    if retval == 'listerr1':
        return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered average rating or score for one or more option values. Please go back to Evaluation Measures and complete the information.'})                   
    elif retval == 'listerr2':
        return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered evaluation measures for one or more evaluation criteria. Please go back to Evaluation Measures and complete the information.'})  
    elif retval == 'listerr3':
        return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered the score for a new evaluation criteria. Please go back and complete the information.'})
    
    
    try:
       setup = PA_Setup.objects.get(dec_id = dec_id) 
       group_yn = setup.scores_group_yn
       votes_yn = setup.votes_yn
    except ObjectDoesNotExist:
       return render(request,'decisions/message.html', {'dec_id':dec_id, 'mess':'You have not entered the setup information for scores so you cannot view this screen.'})
    ''' 
    try: 
       std = Stakeholders_Decisions.objects.filter(dec_id = dec_id, iw_type = 'Y')                                                                                   
       std_count = std.exclude(email = user_email).count()                                                                                            
       if std_count > 0: 
          stakeholdersNow = 'Y'   
       else:
          stakeholdersNow = 'N'  
    except ObjectDoesNotExist:
       stakeholdersNow = 'N'
    except MultipleObjectsReturned:                                                                                                                
       stakeholdersNow = 'Y'  
    print 'in dec made'
    cost_rec = ''
    # july 19
    #if loggedinuser == created_by:
    if stakeholdersNow == 'Y':
       individual_cal(dec_id, created_by, request)
    else:
       print 'before group cal' 
       group_cal(dec_id, created_by, request)
    retval = further_cal(dec_id, created_by, request)   
    if retval == 'em':
       evm_mesg = '- Specify the evaluation measures, indicate whether higher scores are better and add the average rating/score for each Solution Option in the "Evidence-gathering to Evaluate Options" step'
       return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':temp_mesg, 'evm_mesg':evm_mesg})    

    try: 
       c = Cost_Setup.objects.get(dec_id = dec_id) 
       type_of_cost = c.type_of_cost
       if type_of_cost == "Total":
          cost_text = "Total Cost"
       elif type_of_cost == "Avg":
          cost_text = "Average Cost"
       else:                                                                                                                                                                                                     
          cost_text = "Marginal Cost"
    except ObjectDoesNotExist:
       type_of_cost = '' 
       cost_rec = 'N'
       cost_text = 'Cost' 
       print 'cost does not exist'
       #return render(request,'decisions/message.html', {'dec_id':dec_id,'loggedinuser':loggedinuser,'mess':'You have not entered costs so you cannot view this screen.'})   
    
    try: 
       costs = Cost_Utility.objects.get(dec_id = dec_id)
       if costs.cost is None or costs.weighted_utility is None:
          print 'not checking costs'   
          #return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered costs so you cannot view this screen.'})   
    except ObjectDoesNotExist:
       cost_rec = 'N' 
       print 'not checking costs'   
       #return render(request,'decisions/message.html', {'dec_id':dec_id, 'loggedinuser':loggedinuser, 'mess':'You have not entered costs so you cannot view this screen.'})  
    except MultipleObjectsReturned:
       print 'multiple objects returned'

    wu = '' 
    so1 = '' 
    co = '' 
    so2 = '' 
    cur = ''                                                                                                                                                                                                     
    so3 = '' 
    cost_utility = ''
    dec = Decisions.objects.get(pk=dec_id)   
    cost_utility1 = Cost_Utility.objects.filter(dec_id = dec_id)
    cost_utility = cost_utility1.exclude(archived = 'Y')
    for cu in cost_utility:
        print 'sol option in dec made'
        print cu.sol_option
        cu.type_of_cost = type_of_cost
        if cu.weighted_utility is not None and cu.cost is not None:
           if cu.weighted_utility <> 0: 
              cu.cost_utility_ratio = "%.2f" %  round(float(cu.cost) / float(cu.weighted_utility),2)
           else:
              cu.cost_utility_ratio = None
           cu.updated_by = request.session['user']
           cu.updated_date = datetime.datetime.now()
           cu.save(update_fields=['type_of_cost','cost_utility_ratio','updated_by','updated_date'])    
           dec.updated_by = request.session['user']
           dec.updated_date = datetime.datetime.now()
           dec.save(update_fields=['updated_by','updated_date'])
    cc = cost_utility.order_by('-weighted_utility')
    
    onerec = cc.first()
    type_of_cost = onerec.type_of_cost
    try:
       wu = onerec.weighted_utility
    except:
       wu = "" 
    so1 = ""
    more_than_one = ""
    if onerec.cost == 0:
       costzero = 'Y'
    else:   
       costzero = 'N' 
    if wu == 0:
       wuzero = 'Y'
    else:   
       wuzero = 'N' 
    for c in cc:
        print 'INSIDE CC'
        print c.sol_option
        print wu
        print c.weighted_utility
        if c.cost <> 0:
           costzero = 'N'
        if c.weighted_utility == 0:                                                                                                                                                                                          
           wuzero = 'Y'
        if wu == c.weighted_utility:
           if so1 <> "":
              so1 = so1 + ", " + c.sol_option
              more_than_one = "Y"
           else:
              so1 = so1 + " " + c.sol_option 
           print so1   
        #else: 
           #so1 = onerec.sol_option         
               
    if costzero == 'N': 
       cd = cost_utility.exclude(cost = 0).order_by('cost')
       onerec2 = cd.first()                                                                                                                                                                                     
       try:
          co = onerec2.cost
       except:   
          co = "" 
       so2 = ""
       for c in cd:
          if co == c.cost:
             if so2 <> "":
                so2 = so2 + ", " + c.sol_option
             else:
                so2 = so2 + " " + c.sol_option

       ce1 = cost_utility.exclude(weighted_utility = 0) 
       ce = ce1.order_by('cost_utility_ratio') 
       onerec3 = ce.first()                                                                                                                                                                                     
       try:
          cur = onerec3.cost_utility_ratio
       except:
          cur = ""  
       so3 = ""
       for c in ce:
          if cur == c.cost_utility_ratio:
             if so3 <> "":
                so3 = so3 + ", " + c.sol_option
                more_than_one = "Y" 
             else:
                so3 = so3 + " " + c.sol_option 

    try: 
       dec_made = Decision_Made.objects.get(dec_id = dec_id)                                                                                                                                                         
       reason = dec_made.reason
       sol_opt = dec_made.sol_option
       none = dec_made.none
       primary_factor = dec_made.primary_factor
       other_cons = dec_made.other_cons
    except ObjectDoesNotExist:
       dec_made  = Decision_Made(dec_id = dec_id,created_by = request.session['user'],created_date = datetime.datetime.now())
       reason = ''
       sol_opt = ''
       none = ''
       primary_factor = ''
       other_cons = ''
       dec_made.save() 
       dec.updated_by = request.session['user']
       dec.updated_date = datetime.datetime.now()
       dec.save(update_fields=['updated_by','updated_date'])

    query = Solution_Options.objects.filter(dec_id = dec_id, archived = 'N', deleted='N')
    if request.method == 'POST':
       #print request.POST.getlist('id') 
       mystring = request.POST.get('reason')
       mystring = mystring.replace('\n', '##').replace('\r', '')
       dec_made.updated_by = request.session['user']
       dec_made.updated_date = datetime.datetime.now()
       dec_made.reason = mystring
       dec_made.primary_factor = request.POST.get('factor')
       dec_made.sol_option = request.POST.getlist('id') 
       dec_made.none = request.POST.get('none')
       dec_made.other_cons = request.POST.get('other_cons')
       dec_made.save(update_fields=['sol_option','reason', 'none','primary_factor','other_cons','updated_by','updated_date'])
       dec.updated_by = request.session['user']
       dec.updated_date = datetime.datetime.now()
       dec.save(update_fields=['updated_by','updated_date'])            
       if 'submit' in request.POST:                                                                                                            
          return HttpResponseRedirect('/utility_tool/decisions/%s/menu.html' % dec_id)
       elif 'ur' in request.POST:
          print 'UT' 
          return HttpResponseRedirect('/utility_tool/decisions/solution_options/utility_results.html') 
        
       #if loggedinuser != created_by:
           #decmadeform.fields['sol_option'].widget.attrs['disabled'] = True
           #decmadeform.fields['reason'].widget.attrs['disabled'] = True
    
    return render(request,'decisions/solution_options/decision_made.html',{'query':query,'dec_id':dec_id, 'dec_title':dec_title, 'cost_utility':cc, 'loggedinuser':loggedinuser, 'created_by':created_by, 'cost_text':cost_text, 'reason': reason, 'primary_factor' : primary_factor, 'other_cons':other_cons, 'sol_opt': sol_opt, 'noneX':none, 'wu':wu,'so1':so1, 'co':co, 'so2':so2, 'cur':cur, 'so3':so3, 'qcost':onerec.cost, 'more_than_one':more_than_one, 'wuzero':wuzero, 'costzero':costzero }) 

# Pa_setup, Add_iw_votes, Summary table, Utility results, Check add_measures, Cost_setup, Cost_table, Cost_utility, Decision_made
# when registration code is written, the new user created should be added as a stakeholder
def login(request):
    context = RequestContext(request)

    if request.method == 'POST':
        loginform = LoginForm(data=request.POST)
        if loginform.is_valid():
           login = loginform.save(commit=False)
           request.session['user'] = login.user
            #request.session['password'] = login.password
           try:
               login2 = Users.objects.filter(user=login.user).latest('startDate')
               #login2 = Users.objects.get(user = login.user)
               if login.user <> login2.user:
                  return render(request,'users/login.html',{'loginform': loginform, 'err': 'Invalid user name. Please enter the correct user name.'}) 
               if login.password <> login2.password:                                                                                                                                                                     
                  return render(request,'users/login.html',{'loginform': loginform, 'err': 'Invalid password. Please enter the correct password.'}) 
               if login2.endDate <= datetime.date.today():      
                  return render(request,'users/login.html',{'loginform': loginform, 'err': 'Your license agreement has expired.  Please re-register from the Home page. If you wish to continue using your existing account, re-register with the same User name. You may change your password and any other information that needs updating.'})                                                                                      
               request.session['user_email'] = login2.email 

               login2.lastLogin = datetime.datetime.now()
               if login2.timesLoggedin is None:
                  login2.timesLoggedin = 1
               else:
                  login2.timesLoggedin = login2.timesLoggedin + 1
               login2.save(update_fields=['lastLogin', 'timesLoggedin'])     
               log  = Login(user = request.session['user'],email = request.session['user_email'], loggedindate = datetime.datetime.now())
               log.save()
               return HttpResponseRedirect('/utility_tool/decisions/decisions_list.html')
           except ObjectDoesNotExist:
               #login2 = Users.objects.get(user = login.user)
               #request.session['user_email'] = login2.email 
               return render(request, 'users/login.html',{'loginform': loginform, 'err': 'Invalid user or password'})

        else:
           form_errors = 'Yes'
           print form_errors

    else:
        loginform = LoginForm()

    return render(request,'users/login.html', {'loginform':loginform})

def forgot(request):
    context = RequestContext(request)
    message = "" 
    if request.method == 'POST':
       forgotform = ForgotForm(request.POST)
       if forgotform.is_valid():
          forgot = forgotform.save(commit=False)
          try:
             r = Users.objects.get(email = forgot.email)
          except:
             print 'does not exist'
             return render_to_response('users/forgot.html',{'forgotform':forgotform,'err':'The email address you have entered does not match what we have in our records. Please enter again.'}, context)
          print 'jul31 test1'
          message = 'The User Name you used to log in to DecisionMaker is: "' + r.user + '". Your Password is "' + r.password + '". If you need to contact CBCSE, please email cbcse@tc.columbia.edu.'
          subject = 'Login details for DecisionMaker'
          from_email = 'decisionmaker.cbcse@gmail.com'
          to_email = forgot.email
          print to_email
          if subject and message and from_email and to_email:
             try: 
                send_mail(subject, message, from_email,[to_email])
             except BadHeaderError:
                return HttpResponse('Invalid header found.')
          else:
             return HttpResponse('Make sure all fields are entered and valid.') 
          return HttpResponseRedirect('/utility_tool/users/login.html')                                                                                                                                                       
       else:
          print forgotform.errors
          return render_to_response('users/forgot.html',{'forgotform':forgotform,'err':forgotform.errors}, context)
    else:
       forgotform = ForgotForm()
    return render(request,'users/forgot.html', {'forgotform':forgotform})

#replace render_to_response with render
def register(request):
   context = RequestContext(request)
   if request.method == 'POST':
      registerform = RegisterForm(data=request.POST)
      if registerform.is_valid():
         register = registerform.save(commit=False)
         try:
            login = Users.objects.filter(user=register.user).latest('startDate')
            if login.endDate > date.today():
               return render_to_response('users/register.html',{'registerform': registerform, 'err': 'The User Name you have entered already exists. Please select another one.'}, context)                                                                                      
         except ObjectDoesNotExist:
            print 'xyz' 
         try:
            r = Users.objects.filter(email = register.email)
            r_count = r.count()
            if r_count > 0:
               for records in r:
                   if records.user <> register.user: 
                      return render(request, 'users/register.html',{'registerform': registerform,'err':'Another user has the same email address entered. Please enter a different email address.'})
         except ObjectDoesNotExist:
             print 'something wrong in email unique check'
         '''try: 
            r = Users.objects.filter(user = register.user).count()                                                                                                                
            if r > 0: 
               return render(request, 'users/register.html',{'registerform': registerform,'err':'Another user has the same user name entered. Please enter a different user name.'})
         except ObjectDoesNotExist:
             print 'something wrong in email unique check' '''
         if register.password != register.passwordagain:                                                                                   
            return render(request, 'users/register.html',{'registerform': registerform, 'err': 'Password does not match Confirm Password.'})              
         if register.email != register.emailagain:
            return render(request, 'users/register.html',{'registerform': registerform, 'err': 'Email address does not match Confirm Email address.'})
         rand = random.randrange(1000,999999999)
         try:  
            unique = Users.objects.get(uniqueRandomId = rand)
            print 'RANDOM ID NOT CREATED'
         except ObjectDoesNotExist:
            register.uniqueRandomId = random.randrange(1000,999999999)
            register.endDate= datetime.datetime.now() + relativedelta(years=2)
         #register.save()
         #st = Stakeholders(firstName = register.firstName, lastName = register.lastName, email=register.email, created_by = register.user, created_date = datetime.datetime.now())
         #st.save()

         request.session['userR'] = register.user
         request.session['email'] = register.email
         request.session['passwordR'] = register.password
         request.session['firstName'] = register.firstName
         request.session['lastName'] = register.lastName
         request.session['state'] = register.state
         request.session['country'] = register.country
         request.session['organisation'] = register.organisation
         request.session['type_of_org'] = register.type_of_org
         request.session['other_org'] = register.other_org
         request.session['position'] = register.position
         request.session['other_pos'] = register.other_pos
         request.session['hearaboutus'] = register.hearaboutus
         request.session['other_hear'] = register.other_hear
         request.session['updates'] = register.updates
         request.session['education'] = register.education
         request.session['age'] = register.age
         request.session['gender'] = register.gender
         request.session['race'] = register.race
         request.session['other_race'] = register.other_race
         request.session['uniqueRandomId'] = register.uniqueRandomId
         request.session['publicOrPrivate'] = register.publicOrPrivate 
         return HttpResponseRedirect('/utility_tool/users/license.html') 
      else:
         print registerform.errors
          
   else:                                                                                                                            
      registerform = RegisterForm()
                                                             
   return render(request, 'users/register.html',{'registerform': registerform})

def license(request):
   context = RequestContext(request)                                                                                                                    
   if 'publicOrPrivate' in request.session:
      publicOrPrivate = request.session['publicOrPrivate']
   else:
      publicOrPrivate = 'Public' 
 
   if request.method == 'POST':
      licenseform = License(request.POST)
      if licenseform.is_valid():
         license = licenseform.save(commit=False)
         if license.licenseSigned == 'Yes':
            Users.objects.create(user=request.session['userR'], email=request.session['email'],password=request.session['passwordR'],firstName=request.session['firstName'],lastName=request.session['lastName'],state=request.session['state'],country=request.session['country'],organisation=request.session['organisation'],type_of_org=request.session['type_of_org'],other_org=request.session['other_org'],position=request.session['position'],other_pos=request.session['other_pos'],publicOrPrivate=request.session['publicOrPrivate'], licenseSigned ='Yes',endDate= datetime.datetime.now() + relativedelta(years=2), hearaboutus = request.session['hearaboutus'], other_hear =request.session['other_hear'],updates =request.session['updates'],education =request.session['education'], age =request.session['age'], gender =request.session['gender'], race = request.session['race'], other_race = request.session['other_race'],uniqueRandomId= request.session['uniqueRandomId'])
            st = Stakeholders(firstName = request.session['firstName'], lastName = request.session['lastName'], email=request.session['email'], created_by = request.session['userR'], created_date = datetime.datetime.now())
            st.save()
            return HttpResponseRedirect('/utility_tool/users/login.html')
         else:
            return HttpResponseRedirect('/Home.html')
      else:
         form_errors = 'Select Yes or No to proceed'
         print form_errors
         print licenseform.errors
         return render(request, 'users/license.html',{'licenseform': licenseform, 'form_errors':form_errors, 'publicOrPrivate':'publicOrPrivate'})
   else:
      licenseform = License()
   return render(request, 'users/license.html',{'licenseform': licenseform, 'publicOrPrivate':'publicOrPrivate'})
def return_pdf(request):
   publicOrPrivate=request.session['publicOrPrivate']
   if publicOrPrivate == 'Public':
      with open('/home/amritha/costutility/documents/DM Online Public Institution Tool Kit License.pdf', 'r') as pdf: 
         response = HttpResponse(pdf.read(), content_type='application/pdf')
         response['Content-Disposition'] = 'inline;filename=DM Online Public Institution Tool Kit License.pdf'
         return response
      pdf.closed
   else:
      with open('/home/amritha/costutility/documents/DM Online Private Institution Tool Kit License.pdf', 'r') as pdf: 
         response = HttpResponse(pdf.read(), content_type='application/pdf')
         response['Content-Disposition'] = 'inline;filename=DM Online Private Institution Tool Kit License.pdf'
         return response
      pdf.closed
def private_pdf(request):
      with open('/home/amritha/costutility/documents/DM Online Private Institution Tool Kit License.pdf', 'r') as pdf: 
         response = HttpResponse(pdf.read(), content_type='application/pdf')
         response['Content-Disposition'] = 'inline;filename=DM Online Private Institution Tool Kit License.pdf'
         return response
      pdf.closed

def public_pdf(request):
      with open('/home/amritha/costutility/documents/DM Online Public Institution Tool Kit License.pdf', 'r') as pdf: 
         response = HttpResponse(pdf.read(), content_type='application/pdf')
         response['Content-Disposition'] = 'inline;filename=DM Online Public Institution Tool Kit License.pdf'
         return response
      pdf.closed
     
def license2(request):                                                                                                                                                                                     
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'
    return render(request, 'license2.html', {'loggedinuser':loggedinuser})
 

def logout(request):
    if 'user' in request.session:
        del request.session['user']
    if 'user_email' in request.session:
       del request.session['user_email']
    if 'dec_id' in request.session:
       del request.session['dec_id']

    return render(request,'Home.html')

def import_excel(request):                                                                                                                                                                                       
    # Open the workbook and define the worksheet
    book = xlrd.open_workbook("/home/amritha/costutility/documents/Stakeholders.xlsx")
    sheet = book.sheet_by_name("Stakeholders")
    
    # Establish a MySQL connection
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")
    # Get the cursor, which is used to traverse the database, line by line
    cursor = database.cursor()
    
    # Create the INSERT INTO sql query
    query = """INSERT INTO utility_tool_stakeholders (firstName, lastName, title, organisation, email, phone, notes, created_by, created_date) VALUES ( %s, %s, %s, %s, %s, %s, %s, %s, %s)""" 
    # Create a For loop to iterate through each row in the XLS file, starting at row 2 to skip the headers
    
    for r in range(1, sheet.nrows):
        firstName = sheet.cell(r,0).value
        lastName = sheet.cell(r,1).value
        title = sheet.cell(r,2).value
        organisation = sheet.cell(r,3).value
        email = sheet.cell(r,4).value
        phone = sheet.cell(r,5).value
        notes = sheet.cell(r,6).value
        created_by = request.session['user']
        created_date = datetime.datetime.now()
        # Assign values from each row
        values = (firstName, lastName, title, organisation, email, phone, notes,created_by, created_date)
        # Execute sql Query
        cursor.execute(query, values)
    # Close the cursor
    cursor.close()
    # Commit the transaction
    database.commit()
    
    # Close the database connection
    database.close()
    columns = str(sheet.ncols)
    rows = str(sheet.nrows)
    return HttpResponseRedirect('/utility_tool/stakeholders/stakeholders.html')

def export_stakeholders(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename=mystakeholders.xls'                                                                                                                            
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet("Stakeholders")
    row_num = 0
    
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")
    user = request.session['user']
    cursor = database.cursor ()
    # Create the INSERT INTO sql query
    sql = """SELECT firstName, lastName, title, organisation, email, phone, notes FROM utility_tool_stakeholders WHERE created_by = %(user)s"""
    columns = [
        (u"First Name", 6000),
        (u"Last Name", 6000),
        (u"Title", 6000),
        (u"Organization", 12000),
        (u"Email", 8000),
        (u"Phone", 6000),
        (u"Notes", 20000)
    ]
    
    a = xlwt.Alignment()
    a.wrap = True 
    a.vert = a.VERT_CENTER
    a.horz = a.HORZ_CENTER
    font_style = xlwt.XFStyle()
    font_style.font.bold = True 
    font_style.alignment = a

    aL = xlwt.Alignment()
    aL.horz = a.HORZ_LEFT
    aL.wrap = True
    font_style2 = xlwt.XFStyle()
    font_style2.alignment = aL

    for col_num in xrange(len(columns)):
        ws.write(row_num, col_num, columns[col_num][0], font_style)
        # set column width
        ws.col(col_num).width = columns[col_num][1]
    try:
    # Execute the SQL command
       cursor.execute(sql,{'user' : user})
       # Fetch all the rows in a list of lists.
       results = cursor.fetchall()
       for row in results:
          row_num += 1
          firstName = row[0]
          lastName = row[1]
          title = row[2]
          organisation = row[3]
          email = row[4]
          phone = row[5] 
          notes = row[6]
          for col_num in xrange(len(row)):
             ws.write(row_num, col_num, row[col_num], font_style2)
    except:
       print "Error: unable to fetch data"
    # disconnect from server
    database.close()
    wb.save(response)
    return response

def export_users(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename=users.xls'                                                                                                                                 
    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet("Users")
    row_num = 0 
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")
    cursor = database.cursor ()
    sql = """SELECT id, user, email, firstName, lastName, addressline1, addressline2, city, state, zip, country, phone, organisation, type_of_org, other_org, position, other_pos, hearaboutus, other_hear, updates, education, age, gender, race, other_race,  publicOrPrivate, startDate, endDate, lastLogin, timesLoggedin FROM utility_tool_users"""

    #Heading of tables
    a = xlwt.Alignment()
    a.wrap = True
    a.vert = a.VERT_CENTER
    a.horz = a.HORZ_CENTER
    font_style = xlwt.XFStyle()
    font_style.font.bold = True
    font_style.alignment = a
    pattern = xlwt.Pattern()
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = 22
    pattern2 = xlwt.Pattern()
    pattern2.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern2.pattern_fore_colour = 22
    font_style.pattern = pattern2
    pattern3 = xlwt.Pattern()
    pattern3.pattern_fore_colour = 1
    font_style3 = xlwt.XFStyle()
    aL = xlwt.Alignment()
    aL.horz = a.HORZ_LEFT
    aL.wrap = True
    font_style3.alignment = aL
    font_style3.pattern = pattern2
    font_style4 = xlwt.XFStyle()
    font_style4.pattern = pattern3
    font_style4.alignment = aL
    date_style4 = xlwt.XFStyle()
    date_style4.pattern = pattern3
    date_style4.num_format_str = 'mm/dd/yyyy'
    date_style3 = xlwt.XFStyle()
    date_style3.pattern = pattern2
    date_style3.num_format_str = 'mm/dd/yyyy'
    font_style5 = xlwt.XFStyle()
    font_style5.font.bold = True
    font_style5.pattern = pattern

    users = Users.objects.all()                                                                                                                                                                                  
    uscount = users.count()
    
    ws.write(0, 0, "Number of Users:", font_style5)
    ws.write(0, 1, "", font_style5)
    ws.write(0, 2, "", font_style5)
    ws.write(0, 3, uscount, font_style4)
    ws.write(0, 4, "", font_style4)

    row_num = 3
    columns = [
          (u"Id", 2000),
          (u"User Name", 7000),
          (u"Email", 7000),
          (u"First Name", 7000),
          (u"Last Name", 7000),
          (u"Address Line1", 7000),
          (u"Address Line2", 7000),
          (u"City", 7000),
          (u"State", 7000),
          (u"Zip", 7000),
          (u"Country", 7000),
          (u"Phone", 7000),
          (u"Organization", 7000),
          (u"Type of Organization", 7000),
          (u"Other Organization", 7000),
          (u"Position", 7000),
          (u"Other Position", 7000),
          (u"How did you hear about DecisionMaker?", 7000),
          (u"Any other way?", 7000),
          (u"Would you like to receive occasional updates?", 7000),
          (u"Highest level of education completed", 7000),
          (u"Age", 7000),
          (u"Gender", 7000),
          (u"Race/Ethnicity", 7000),
          (u"Other Race/Ethnicity", 7000),
          (u"Public or Private", 7000),
          (u"Start Date of Licence", 7000),
          (u"Licence Expiry Date", 7000),
          (u"Date Last Logged in", 7000),
          (u"Number of times Logged in", 7000)
    ]

    try:
       cursor.execute(sql)
       results = cursor.fetchall()
       if results != ():
          for col_num in xrange(len(columns)):
             ws.write(row_num, col_num, columns[col_num][0], font_style)
             # set column width
             ws.col(col_num).width = columns[col_num][1]
       for row in results:
          row_num += 1
          id = row[0]
          user = row[1]
          email = row[2]
          firstName = row[3]
          lastName = row[4]
          addressline1 = row[5]
          addressline2 = row[6]
          city = row[7]
          state = row[8]
          zip = row[9]
          country = row[10]
          phone = row[11]
          organisation = row[12]
          type_of_org = row[13]
          other_org = row[14]
          position = row[15]
          other_pos = row[16]
          hearaboutus = row[17]
          other_hear = row[18] 
          updates = row[19] 
          education = row[20] 
          age = row[21] 
          gender = row[22] 
          race = row[23] 
          other_race = row[24]
          publicOrPrivate = row[25]
          startDate = row[26]
          endDate = row[27]
          lastLogin = row[28]
          timesLoggedin = row[29]
          for col_num in xrange(len(row)):
              if col_num == 26 or col_num == 27 or col_num == 28:
                ws.write(row_num, col_num, row[col_num],date_style4)
              else:
                ws.write(row_num, col_num, row[col_num],font_style4)
    except:
       print "Error: unable to fetch data"
    # disconnect from server
    database.close()
    wb.save(response)                                                                                                                                                                                            
    return response


def imports(request):                                                                                                                                    
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    return render(request,'admin/imports.html', {'loggedinuser':loggedinuser})

def import_cbcse_scrcr(request):
    # Open the workbook and define the worksheet                                                                                                         
    book = xlrd.open_workbook("/home/amritha/costutility/costutility/static/CBCSE_screening_criteria.xlsx")
    sheet = book.sheet_by_name("Sheet1")
    CBCSE_Screening_Criteria.objects.all().delete()
    # Establish a MySQL connection
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")
    # Get the cursor, which is used to traverse the database, line by line
    cursor = database.cursor()
    # Create the INSERT INTO sql query
    query = """INSERT INTO utility_tool_cbcse_screening_criteria (criterion) VALUES (%s)"""
    # Create a For loop to iterate through each row in the XLS file, starting at row 2 to skip the headers
    for r in range(1, sheet.nrows):
        print sheet.cell(r,0).value
        list_one      = str(sheet.cell(r,0).value)
        # Execute sql Query
        cursor.execute(query, (list_one,))
    # Close the cursor
    cursor.close()
    # Commit the transaction
    database.commit()
    # Close the database connection
    database.close()
    columns = str(sheet.ncols)                                                                                                                           
    rows = str(sheet.nrows)
    return HttpResponseRedirect('/utility_tool/admin/upload.html')

def import_cbcse_evacr(request):
    # Open the workbook and define the worksheet                                                                                                         
    book = xlrd.open_workbook("/home/amritha/costutility/costutility/static/CBCSE_evaluation_criteria.xlsx")
    sheet = book.sheet_by_name("Granular evaluation criteria")
    CBCSE_Evaluation_Criteria.objects.all().delete()
    # Establish a MySQL connection
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")    
    # Get the cursor, which is used to traverse the database, line by line
    cursor = database.cursor()
    # Create the INSERT INTO sql query
    query = """INSERT INTO utility_tool_cbcse_evaluation_criteria (overreaching_ec, granular_ec, suggested_evam, data) VALUES (%s, %s, %s, %s)"""
    # Create a For loop to iterate through each row in the XLS file, starting at row 2 to skip the headers
    for r in range(1, sheet.nrows):
        orec = sheet.cell(r,0).value
        orec = str(orec.replace(",", ";"))
        gran = sheet.cell(r,1).value
        gran = str(gran.replace(",", ";"))
        sugg = sheet.cell(r,2).value
        sugg = str(sugg.replace(",", ";"))
        data = sheet.cell(r,3).value
        data = str(data.replace(",", ";"))
        #orec = str(sheet.cell(r,0).value)
        #gran = str(sheet.cell(r,1).value)                                                                                
        #sugg = str(sheet.cell(r,2).value)
        #data = str(sheet.cell(r,3).value)        
        values  = (orec, gran, sugg, data)
        # Execute sql Query
        cursor.execute(query, values)
    # Close the cursor
    cursor.close()
    # Commit the transaction
    database.commit()
    # Close the database connection
    database.close()
    columns = str(sheet.ncols)                                                                                                                           
    rows = str(sheet.nrows)
    return HttpResponseRedirect('/utility_tool/admin/upload.html')

def appendices(request,dec_id):     
    '''if 'dec_id' in request.session:
       dec_id = request.session['dec_id']
    else:
       dec_id = 0'''   
    dec_id = int(dec_id)
    if 'user' in request.session:                                                                                                                                                                                
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    if 'created_by' in request.session:
       created_by = request.session['created_by']
    else:
       created_by = 'not found'

    if 'whereamI2' in request.session:
       from_where = request.session['whereamI2'] 
    else:
       from_where = 'here' 

    d = Decisions.objects.get(pk=dec_id)  
    if d.real_dec_yn == 'R':
       filename = 'appendices-' + str(dec_id) + '-R.xls'
    elif d.real_dec_yn == 'T':
       filename = 'appendices-' + str(dec_id) + '-T.xls'         
    else:
       filename = 'appendices-' + str(dec_id) + '-X.xls'         

    response = HttpResponse(content_type='application/ms-excel')                                                                                                                                                 
    response['Content-Disposition'] = 'attachment; filename="%s"' % filename  
    wb = xlwt.Workbook(encoding='utf-8')
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")
    print 'am i in here'
    #Heading of tables
    borders = xlwt.Borders()
    borders.left = 1
    borders.right = 1
    borders.top = 1
    borders.bottom = 1
    a = xlwt.Alignment()
    a.wrap = True 
    a.vert = a.VERT_TOP
    a.horz = a.HORZ_CENTER
    font_style = xlwt.XFStyle()
    font_style.font.bold = True 
    font_style.alignment = a
    font_style.borders = borders
    font_style100 = xlwt.XFStyle()
    font_style100.alignment = a
    font_style100.borders = borders
    pattern = xlwt.Pattern()
    pattern.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern.pattern_fore_colour = xlwt.Style.colour_map['pale_blue']
    pattern2 = xlwt.Pattern()
    pattern2.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern2.pattern_fore_colour = xlwt.Style.colour_map['light_turquoise']  
    pattern3 = xlwt.Pattern()
    pattern3.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern3.pattern_fore_colour = xlwt.Style.colour_map['tan']
    pattern4 = xlwt.Pattern()
    pattern4.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern4.pattern_fore_colour = xlwt.Style.colour_map['silver_ega'] 
    pattern_basic = xlwt.Pattern()
    pattern_basic.pattern_fore_colour = 1
    pattern6 = xlwt.Pattern()
    pattern6.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern6.pattern_fore_colour = 1
    pattern7 = xlwt.Pattern()
    pattern7.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern7.pattern_fore_colour = xlwt.Style.colour_map['light_green']
    font_style.pattern = pattern2
    font_style3 = xlwt.XFStyle()
    aL = xlwt.Alignment()
    aL.vert = a.VERT_TOP   
    aL.horz = a.HORZ_LEFT
    aL.wrap = True 
    aR = xlwt.Alignment()
    aR.vert = a.VERT_TOP   
    aR.horz = a.HORZ_RIGHT
    aR.wrap = True
    font_style3.alignment = aR 
    font_style3.pattern = pattern3
    font_style3.borders = borders
    num_style5 = xlwt.XFStyle()
    num_style5.num_format_str = '0.00'
    num_style5.alignment = aR 
    num_style5.pattern = pattern4 
    num_style5.borders = borders
    num_style6 = xlwt.XFStyle()
    num_style6.num_format_str = '0.00'
    num_style6.alignment = aR 
    num_style6.pattern = pattern7
    num_style6.borders = borders
    font_style4 = xlwt.XFStyle()
    font_style4.pattern = pattern_basic
    font_style4.alignment = aL 
    font_style4.borders = borders 
    date_style4 = xlwt.XFStyle()
    date_style4.borders = borders
    date_style4.pattern = pattern3
    date_style4.num_format_str = 'mm/dd/yyyy'
    date_style3 = xlwt.XFStyle()
    date_style3.pattern = pattern2
    date_style3.borders = borders
    date_style3.num_format_str = 'mm/dd/yyyy'
    font_style5 = xlwt.XFStyle()
    font_style5.font.bold = True 
    font_style5.pattern = pattern
    font_style5.borders = borders
    num_style = xlwt.XFStyle()
    num_style.num_format_str = '0.00'
    num_style.alignment = aR
    num_style.borders = borders
    font_style6 = xlwt.XFStyle()
    font_style6.pattern = pattern6
    font_style6.borders = borders
    font_style12 = xlwt.XFStyle()
    font_style12.alignment.wrap = 1
    font_style12.pattern = pattern6
    font_style12.borders = borders
    font_style7 = xlwt.XFStyle()
    font_style7.alignment = a
    font_style7.pattern = pattern7
    font_style7.borders = borders
    font_style8 = xlwt.XFStyle()
    font_style8.alignment = a
    font_style8.pattern = pattern3
    font_style8.borders = borders
    pattern8 = xlwt.Pattern()
    pattern8.pattern = xlwt.Pattern.SOLID_PATTERN
    pattern8.pattern_fore_colour = xlwt.Style.colour_map['ice_blue']
    font_style9 = xlwt.XFStyle()
    font_style9.pattern = pattern8
    font_style9.font.bold = True 
    font_style9.borders = borders
    font_style10 = xlwt.XFStyle()
    font_style10.alignment = aR
    font_style10.pattern = pattern_basic
    font_style10.borders = borders 
    font_style11 = xlwt.XFStyle()
    font_style11.pattern = pattern6
    try:
       s = Solution_Options.objects.filter(dec_id=dec_id)                                                                                                                                                       
       s1 = s.exclude(deleted = 'Y') 
       s2 = s1.exclude(archived = 'Y')
       s1_count = s1.count()
       s_count = s2.count()
       print 's_count'
       print s_count
    except ObjectDoesNotExist:
       s_count = 3
       s1_count = 0

    #Stakeholders
    ws = wb.add_sheet("Appendix A Stakeholders")
    row_num = 0
    col_num = 0
    #Heading
    ws.write(0, 0, "Appendix A: Stakeholders Participating in Decision ID " + str(dec_id), font_style5)
    counter = 1
    while counter < 8:
        ws.write(0, counter, "", font_style5)
        counter = counter + 1

    ws.write(1, 0, "The table below shows what tasks stakeholders were invited to participate in, and whether they participated.", font_style6)
    counter = 1
    while counter < 8:
        ws.write(1, counter, "", font_style6)
        counter = counter + 1
    
    row_num = 2
    col_num = 0

    cursor5 = database.cursor ()
    sql = """SELECT email, created_by, '', '', solopt_type, scrcr_type, evacr_type, iw_type  from utility_tool_stakeholders_decisions where dec_id = %(dec_id)s"""

    if from_where == 'here':
       columns = [
             (u"First Name", 5000),
             (u"Last Name", 5000),
             (u"Title", 5000),
             (u"Organization", 5000),
             (u"Suggest Solution Options", 5000),
             (u"Suggest Screening Criteria", 5000),
             (u"Suggest Evaluation Criteria", 5000),
             (u"Contribute Importance Scores", 5000),   
       ]
    else:
       columns = [
             (u"ID", 5000),
             (u"Name", 5000),
             (u"Title", 5000),
             (u"Organization", 5000),
             (u"Suggest Solution Options", 5000),
             (u"Suggest Screening Criteria", 5000),
             (u"Suggest Evaluation Criteria", 5000),
             (u"Contribute Importance Scores", 5000),   
       ]

    try:
       cursor5.execute(sql, {'dec_id' : dec_id})
       results = cursor5.fetchall()
       if results != ():
          for col_num in xrange(len(columns)):
             ws.write(row_num, col_num, columns[col_num][0], font_style)
             # set column width
             ws.col(col_num).width = columns[col_num][1]
       for row in results:
          row_num += 1
          email = row[0]
          created_by = row[1]  
          temp1 = row[2]
          temp3 = row[3]   
          solopt_type = row[4]
          scrcr_type = row[5]
          evacr_type = row[6]
          iw_type = row[7]
          for col_num in xrange(len(row)):
              if col_num == 0 or col_num == 1 or col_num == 2 or col_num == 3:
                 ws.write(row_num, col_num, StValue(row[0], row[1], col_num, from_where),font_style6)
              else:
                 res = PartValue(dec_id, email, created_by, row[col_num], col_num)
                 if res == 'P':
                    ws.write(row_num, col_num, res,font_style7)
                 elif res == 'N':
                    ws.write(row_num, col_num, res,font_style8) 
                 else:
                    ws.write(row_num, col_num, res,font_style100)
                 '''if row[col_num] == '':
                    ws.write(row_num, col_num, 'X',font_style4)
                 elif row[col_num] == 'Y':
                    ws.write(row_num, col_num, row[col_num],font_style4)'''
    except:
       print "Error: unable to fetch data"

    row_num += 2
    ws.write(row_num, 0, "Key", font_style5)
    counter = 1
    while counter < 3:
        ws.write(row_num, counter, "", font_style5)
        counter = counter + 1
    row_num += 1
    ws.write(row_num, 0, "Invited but did not participate", font_style8)
    ws.write(row_num, 1, "", font_style8)
    ws.write(row_num, 2, "N", font_style8) 
    row_num += 1
    ws.write(row_num, 0, "Invited and participated", font_style7)
    ws.write(row_num, 1, "", font_style7)
    ws.write(row_num, 2, "P", font_style7)  
    row_num += 1
    ws.write(row_num, 0, "Not invited to participate", font_style100)
    ws.write(row_num, 1, "", font_style100)
    ws.write(row_num, 2, "X", font_style100)  
    #Mapping Table
    try:                                                                                                                                              
       mapp = MappingTable.objects.get(dec_id=dec_id)
       table =  mapp.table
    except ObjectDoesNotExist:
       table = 'doesnotexist'
    # hiding the unwanted columns; first add them to a list and set width to 1
    if table <> 'doesnotexist':
       ws = wb.add_sheet("Appendix B Screening SOs")                                                                                                  
       row_num = 0 
       col_width = 256 * 30     
       try:
          for i in itertools.count():
             ws.col(i).width = col_width
       except ValueError:
          pass
       col_num = 0
       #Heading
       ws.write(0, 1, "Appendix B: How Solution Options Performed Against the Screening Criteria", font_style5)
       counter = 2
       s1_count = s1_count + 1
       while counter < s1_count:
          ws.write(0, counter, "", font_style5)
          counter = counter + 1
       ws.write(1, 1, "The table below indicates whether each Solution Option met or “passed” each Screening Criterion:", font_style6)
       counter = 2
       while counter < s1_count:
          ws.write(1, counter, "", font_style6)
          counter = counter + 1
    
       row_num = 2
       col_num = 0
       scrcr_list = []
       scrcr_list.append("Screening Criteria")
       scrcr_list.append("2186")
       scrcr_list.append("2187")
       try:
          for scrcr in Screening_Criteria.objects.filter(dec_id=dec_id):
              scrcr_list.append(str(scrcr.id))
              print scrcr.id 
       except ObjectDoesNotExist:
          print 'nothing todo scrcr'
    
       print 'SCRCR LIST'
       for i in scrcr_list:
          print i
       print 'APPENDIX B'
       print table
       arrx = table[1:]
       # get the first and last postion of the solution options list
       firstposx = arrx.find('[') + 1
       lastposx =  arrx.find(']')
       arrx1 = arrx[firstposx:lastposx]
       #print 'arr1'
       pos = 0
       pos2 = 0
       new_list = []              
       col_not_disp = []
       col_count = 0                                                                                                                      
       # remove the first [ from the array we got from ajax  
       aa = table.replace('",null,"','",,"')
       c = aa.replace('[],','')
       z = c.replace('[[','[')
       yy = z.replace('[""]','[')
       y = yy.replace('"],"','","')                                                                                                                   
       print 'this is the Y june27'
       print y
       copy_of_y = y
       while len(y): 
          firstpos = y.find('[') + 1
          lastpos =  y.find(']')
          arr2 = y[firstpos:lastpos]
          x = arr2.split(',')[0]
          print arr2
          for l2 in arr2.split(','):
             l3 = l2.replace('"', '')
             print 'looking for l3'
             print l3
             col_count = col_count + 1
             if (l3.find('has been deleted') != -1):
                print 'in find'
                print l3
                print col_count
                col_not_disp.append(col_count - 1)
          z = y.replace(arr2, '')                                                                                                                     
          c = z.replace('[],','')
          # break out of the loop when only []] remains
          if (c == '[]]'):
            break;
          y = c  
       print 'col not disp LIST'
       for i in col_not_disp:
          print i  
       y = copy_of_y
       col_not_disp.append(0)
       while len(y): 
          firstpos = y.find('[') + 1
          lastpos =  y.find(']')
          arr2 = y[firstpos:lastpos]
          x = arr2.split(',')[0]
          print arr2                                                                                                                                  
          for l2 in arr2.split(','):
             l3 = l2.replace('"', '')
             print l3
             if col_num in col_not_disp:
                print 'skip this column'
                ws.col(col_num).width = 1 
                col_num = col_num + 1 
             else:
                if (l3 == 'null'):
                   if row_num == 2:
                      ws.write(row_num, col_num, ' ', font_style)
                   else: 
                      ws.write(row_num, col_num, ' ', font_style100)
                else:
                   if row_num == 2:   
                      if col_num == 1:
                         ws.write(row_num, col_num, 'Screening Criteria', font_style)
                      else:
                         ws.write(row_num, col_num, l3, font_style)                                                                                   
                   elif row_num == 3: 
                      if l3 == 'Keep Option or Put it away for now':
                         ws.write(row_num, col_num, l3, font_style4) 
                      elif l3 == 'Keep Option':
                         ws.write(row_num, col_num, l3, font_style7)  
                      else: 
                         ws.write(row_num, col_num, l3, font_style8)  
                   else: 
                      ws.write(row_num, col_num, l3, font_style100)
                col_num = col_num + 1
                pos2 = pos2 + 1
          row_num = row_num + 1
          col_num = 0
          pos2 = 0
          z = y.replace(arr2, '')                                                                                                                     
          c = z.replace('[],','')                                                                                                                     
          # break out of the loop when only []] remains
          if (c == '[]]'):
            break;
          y = c
    else:
       ws = wb.add_sheet("Appendix B Screening SOs")                                                                                                  
       row_num = 0 
       col_width = 256 * 30     
       try:
          for i in itertools.count():
             ws.col(i).width = col_width
       except ValueError:
          pass
       col_num = 0
       #Heading
       ws.write(0, 0, "Appendix B: How Solution Options Performed Against the Screening Criteria", font_style5)
       counter = 1
       s1_count = 6
       while counter < s1_count:
          ws.write(0, counter, "", font_style5)
          counter = counter + 1
       ws.write(1, 0, "The user has not created a table to indicates whether each Solution Option met or “passed” each Screening Criterion.", font_style6)     
       counter = 1
       while counter < s1_count:
          ws.write(1, counter, "", font_style6)
          counter = counter + 1     

    #Importance Scores
    eva_table = Evaluation_Criteria.objects.filter(dec_id = dec_id).exclude(deleted = 'Y').order_by('id')                                             
                                                           
    stdec = Stakeholders_Decisions.objects.filter(dec_id = dec_id)
    stdec_count = stdec.count()
    if stdec_count > 1:
       one_stakeholder = 'N'
    else:
       one_stakeholder = 'Y' 
    total_votes = 0
    stdec_list = []
    vote_changed = 'N'
    for st in stdec:
        if st.votes is not None:
           total_votes = st.votes + total_votes
        if st.votes <> 10:
           vote_changed = 'Y' 
    scores = Importance_Scores.objects.raw("SELECT i.id, i.score score FROM utility_tool_importance_scores i, utility_tool_stakeholders_decisions s WHERE i.dec_id = s.dec_id AND i.dec_id=%s AND s.iw_type = 'Y' AND i.email = s.email AND (i.deleted = 'N' OR i.deleted IS NULL) order by i.created_by, i.eva_id", [dec_id])
    #min_scores = Importance_Scores.objects.raw("SELECT id, MIN(score), eva_id FROM utility_tool_importance_scores WHERE dec_id=%s group by eva_id", [dec_id])
    cursor4 = database.cursor () 
    cursor2 = database.cursor () 
    cursor3 = database.cursor ()
    cursor4.execute("""SELECT MAX(i.score) FROM utility_tool_importance_scores i, utility_tool_stakeholders_decisions s WHERE i.dec_id = s.dec_id AND
i.dec_id=%s AND s.iw_type = 'Y' AND i.email = s.email AND (i.deleted = 'N' OR i.deleted IS NULL) group by i.eva_id""", [dec_id])                      
    
    cursor2.execute("""SELECT MIN(i.score) FROM utility_tool_importance_scores i, utility_tool_stakeholders_decisions s WHERE i.dec_id = s.dec_id AND i.dec_id=%s AND s.iw_type = 'Y' AND i.email = s.email AND (i.deleted = 'N' OR i.deleted IS NULL) group by i.eva_id""", [dec_id])
    cursor3.execute("""SELECT STDDEV(i.score) FROM utility_tool_importance_scores i, utility_tool_stakeholders_decisions s WHERE i.dec_id = s.dec_id AND i.dec_id=%s AND s.iw_type = 'Y' AND i.email = s.email AND (i.deleted = 'N' OR i.deleted IS NULL) group by i.eva_id""", [dec_id])
    #cursor4.execute("""SELECT MAX(score) score FROM utility_tool_importance_scores WHERE dec_id=%s group by eva_id""", [dec_id])
    #cursor2.execute("""SELECT MIN(score) FROM utility_tool_importance_scores WHERE dec_id=%s group by eva_id""", [dec_id])
    #cursor3.execute("""SELECT STDDEV(score) FROM utility_tool_importance_scores WHERE dec_id=%s group by eva_id""", [dec_id])
    max_scores = cursor4.fetchall()
    min_scores = cursor2.fetchall()
    stdev_scores = cursor3.fetchall()
    ws = wb.add_sheet("Appendix C Importance Scores")
    row_num = 0
    col_width = 256 * 30
    try:
       for i in itertools.count():
           ws.col(i).width = col_width
    except ValueError:
       pass
    col_num = 0
    #Heading
    ws.write(0, 0, "Appendix C: Importance Scores", font_style5)
    counter = 1
    eva_count = eva_table.count()
    print 'EVA COUNT'
    print eva_count
    eva_count = eva_count + 3
    while counter < eva_count:
        ws.write(0, counter, "", font_style5)
        counter = counter + 1
    if one_stakeholder == 'Y':
       ws.write(1, 0, "The table below shows a summary of Importance Scores for Evaluation Criteria", font_style6)
    else:
        if vote_changed == 'Y':
           ws.write(1, 0, "The table below shows a summary of the Importance Scores for Evaluation Criteria provided by each stakeholder. Stakeholders were given a different number of votes in determining the relative importance of each Evaluation Criterion. The more votes a stakeholder is assigned, the more this stakeholder influences the final importance weights.", font_style12)                                                                      
        else:
           ws.write(1, 0, "The table below shows a summary of the Importance Scores for Evaluation Criteria provided by each stakeholder. Stakeholderswere each given an equal number of votes in determining the overall importance weights.", font_style12)
    counter = 1
    while counter < eva_count:
        ws.write(1, counter, "", font_style6)
        counter = counter + 1
    ws.write(3, 0, '', font_style)
    ws.write(3, 1, 'Number of Votes assigned', font_style)
    ws.write(3, 2, '% of Votes assigned', font_style)
    counter = 3
    for e in eva_table:
       ws.write(3, counter, e.criterion, font_style)               
       counter = counter + 1
    ws.write(4, 0, 'Summary', font_style9)
    counter = 1
    while counter < eva_count:
        ws.write(4, counter, "", font_style9)
        counter = counter + 1
    ws.write(5, 0, 'Importance score (adjusted by the votes)', font_style4)
    ws.write(5, 1, '', font_style4)
    ws.write(5, 2, '', font_style4)
    counter = 3
    for e in eva_table:
       ws.write(5, counter, e.weight, num_style)               
       counter = counter + 1
    ws.write(6, 0, 'Importance weight (adjusted by the votes)', font_style4)
    ws.write(6, 1, '', font_style4)
    ws.write(6, 2, '', font_style4)
    counter = 3
    for e in eva_table:
       ws.write(6, counter, e.adjusted_weight, num_style)
       counter = counter + 1
    ws.write(7, 0, 'Individual importance scores', font_style9)
    counter = 1                                                                                                                                       
    while counter < eva_count:
        ws.write(7, counter, "", font_style9)
        counter = counter + 1

    counter = 0
    row_num = 8
    percvotes = ''
    for st in stdec:
       if from_where == 'here':
          ws.write(row_num, counter, st.name, font_style4)    
       else:
          ws.write(row_num, counter, st.st_id, font_style4) 
       counter = counter + 1 
       ws.write(row_num, counter, st.votes, font_style10)     
       counter = counter + 1
       if st.votes is not None:
          percvotes = (float(st.votes) / total_votes) * 100 
          percvotes = str(percvotes) + '%'
       ws.write(row_num, counter, percvotes, num_style) 
       counter = counter + 1
       for s in scores:
           if s.email == st.email:
              ws.write(row_num, counter, s.score, num_style)  
              counter = counter + 1
       row_num = row_num + 1
       counter = 0
    ws.write(row_num, 0, 'How much consensus is there among stakeholders on the relative importance of the criterion?', font_style9)
    counter = 1
    while counter < eva_count:
        ws.write(row_num, counter, "", font_style9)
        counter = counter + 1
    row_num = row_num + 1
    ws.write(row_num, 0, 'Lowest individual Importance Score', font_style4)
    ws.write(row_num, 1, '', font_style4)
    ws.write(row_num, 2, '', font_style4)
    counter = 3
    for value in min_scores:
       ws.write(row_num, counter, value[0], num_style)
       counter = counter + 1
    row_num = row_num + 1                                                                                                                             
    ws.write(row_num, 0, 'Highest individual Importance Score', font_style4)
    ws.write(row_num, 1, '', font_style4)
    ws.write(row_num, 2, '', font_style4)
    counter = 3
    for value in max_scores:
       ws.write(row_num, counter, value[0], num_style)               
       counter = counter + 1
    row_num = row_num + 1
    ws.write(row_num, 0, 'Standard deviation of the importance scores assigned', font_style4)
    ws.write(row_num, 1, '', font_style4)
    ws.write(row_num, 2, '', font_style4)
    counter = 3
    for value in stdev_scores:
       ws.write(row_num, counter, value[0], num_style)               
       counter = counter + 1

    #Identify Measures
    ws = wb.add_sheet("Appendix D Info to Assess SOs")
    row_num = 0 
    col_width =400 * 30     
    try:
       for i in itertools.count():
           ws.col(i).width = col_width
    except ValueError:
       pass
    col_num = 0
    #Heading
    ws.write(0, 0, "Appendix D: Information Used to Assess Each Solution Option against Each Evaluation Criterion", font_style5)
    counter = 1
    s_count = s_count + 3 
    while counter < s_count:
        ws.write(0, counter, "", font_style5)
        counter = counter + 1
    row_num = 1
    col_num = 0
 
    try:
       s = IdentifyTable.objects.get(dec_id=dec_id) 
       table =  s.table
    except ObjectDoesNotExist:
       table = 'doesnotexist'
    except MultipleObjectsReturned:
        s = IdentifyTable.objects.filter(dec_id=dec_id).last()
        table =  s.table

    print table
    if table <> 'doesnotexist':
       arrx = table[1:]                                                                                                                                    
       # get the first and last postion of the solution options list
       firstposx = arrx.find('[') + 1
       lastposx =  arrx.find(']') 
       arrx1 = arrx[firstposx:lastposx]
       #print 'arr1'
       pos = 0
       pos2 = 0
       archived_list = []
       for l in arrx1.split(','):    
           if l not in ("Evaluation Criterion", "Common evaluation measure you can use across all options", "Data to collect"): 
              lm = l.replace("Describe the information you will use to evaluate  ","")
              lo = lm.replace('"','')    
              ln = lo.replace("  against this criterion and where you will get it from.","") 
              try: 
                 solopt = Solution_Options.objects.get(sol_option = ln, dec_id=dec_id)                                                                     
                 if solopt.archived == 'Y' or solopt.deleted == 'Y':
                    archived_list.append(pos)
              except ObjectDoesNotExist:
                 print 'nothing todo' 
              except MultipleObjectsReturned:
                 for s in Solution_Options.objects.filter(sol_option = ln, dec_id=dec_id):
                     if s.archived == 'Y' or s.deleted == 'Y':
                        archived_list.append(pos) 
              pos = pos + 1   
       new_list = []                                                                                                                                    
       # remove the first [ from the array we got from ajax  
       aa = table.replace('",null,"','",,"')
       #b = aa.replace('null,','')
       c = aa.replace('[],','')
       #w = c.replace(',[null]','')
       z = c.replace('[[','[')
       yy = z.replace('[""]','[')
       y = yy.replace('"],"','","')
       while len(y): 
          firstpos = y.find('[') + 1
          lastpos =  y.find(']')
          arr2 = y[firstpos:lastpos]
          x = arr2.split(',')[0]
          print arr2
          for l2 in arr2.split(','):
              l3 = l2.replace('"', '')
              print l3 
              if ((pos2 not in archived_list) and (x <> '0' and x <> 0 and x <> 'None')):
                if row_num == 1:
                   ws.write(row_num, col_num, l3, font_style)
                   col_num = col_num + 1
                else:
                    try:
                       evacr = Evaluation_Criteria.objects.get(combined = l3, dec_id=dec_id)
                       if evacr.deleted <> 'Y':
                          ws.write(row_num, col_num, l3, font_style4) 
                          col_num = col_num + 1  
                       else:
                          break;   
                    except ObjectDoesNotExist:                                                                                                            
                       if (l3 == 'null'):
                           ws.write(row_num, col_num, ' ', font_style4)
                       else:
                           ws.write(row_num, col_num, l3, font_style4) 
                       col_num = col_num + 1  
                    except MultipleObjectsReturned:      
                        if (l3 == 'null'):
                            ws.write(row_num, col_num, ' ', font_style4)
                        else:          
                            ws.write(row_num, col_num, l3, font_style4) 
                        col_num = col_num + 1  
              pos2 = pos2 + 1
          row_num = row_num + 1
          col_num = 0
          pos2 = 0
          z = y.replace(arr2, '')                                                                                                                       
          c = z.replace('[],','')
          # break out of the loop when only []] remains
          if (c == '[]]'):
             break;
          y = c

    #Detailed Costs
    ws = wb.add_sheet("Appendix E Detailed Costs")
    row_num = 1  
    cursor = database.cursor ()
    sql = """SELECT sol_option, personnel_cost, facilities_cost, materials_cost, training_cost, other_cost, total_cost from utility_tool_detailed_costs where dec_id = %(dec_id)s AND (archived = 'N' or archived IS NULL)"""

    columns = [
          (u"Solution Option", 9000),
          (u"Personnel Cost", 5000),
          (u"Facilities Cost", 5000),
          (u"Materials Cost", 5000),
          (u"Training Cost", 5000),
          (u"Other Cost", 5000),
          (u"Total Cost", 5000),
    ]
    #Heading
    ws.write(0, 0, "Appendix E: Detailed Cost Information", font_style5)
    ws.write(0, 1, "", font_style5)
    ws.write(0, 2, "", font_style5)
    ws.write(0, 3, "", font_style5)
    ws.write(0, 4, "", font_style5)  
    ws.write(0, 5, "", font_style5)
    ws.write(0, 6, "", font_style5)  

    try:
       cursor.execute(sql, {'dec_id' : dec_id})
       results = cursor.fetchall()
       if results != ():
          for col_num in xrange(len(columns)):
             ws.write(row_num, col_num, columns[col_num][0], font_style)
             # set column width
             ws.col(col_num).width = columns[col_num][1]
       for row in results:
          row_num += 1
          sol_option = row[0]
          personnel_cost = row[1]
          facilities_cost = row[2]
          materials_cost = row[3]
          training_cost = row[4]
          other_cost = row[5]
          total_cost = row[6]
          for col_num in xrange(len(row)):                                                                                                                                                                       
              if col_num == 0:
                 ws.write(row_num, col_num, row[col_num],font_style4)
              else:
                 ws.write(row_num, col_num, row[col_num],num_style)
    except:
       print "Error: unable to fetch data"

    database.close()
    wb.save('/home/amritha/' + filename)
    with open('/home/amritha/' + filename, 'r') as xlsx:
       response = HttpResponse(xlsx.read(), content_type='application/xlsx')
       response['Content-Disposition'] = 'inline;filename="%s"' % filename
       return response 

def export_sr(request):
    dec = Decisions.objects.filter(real_dec_yn = 'R')
    real_dec = dec.exclude(demoDec = 'Y')
    f = io.BytesIO()
    zipObj = ZipFile(f, 'a', compression=zipfile.ZIP_DEFLATED)
    for r in real_dec:
        try:
          summary_report(request, str(r.id))
          filename = 'SummaryReport-' + str(r.id) + '-R.docx'  
          zipObj.write(filename)
        except:
           print 'could not write error SR'
           print r.id 
    zipObj.close()
    response = HttpResponse(f.getvalue(),content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=SummaryReportsReal.zip'
    return response    

def export_sr_t(request):
    dec = Decisions.objects.filter(real_dec_yn = 'T')                                                                                     
    real_dec = dec.exclude(demoDec = 'Y')
    f = io.BytesIO()
    zipObj = ZipFile(f, 'a', compression=zipfile.ZIP_DEFLATED)
    for r in real_dec:
        print 'Nov the 6th'
        print r.id 
        try:
           summary_report(request, str(r.id))
           filename = 'SummaryReport-' + str(r.id) + '-T.docx'  
           zipObj.write(filename)
        except:
           print 'could not write error SR T'
           print r.id 
    zipObj.close()
    response = HttpResponse(f.getvalue(),content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=SummaryReportsTraining.zip'
    return response    

def export_app(request):
    dec = Decisions.objects.filter(real_dec_yn = 'R')
    real_dec = dec.exclude(demoDec = 'Y')
    f = io.BytesIO()
    zipObj = ZipFile(f, 'a', compression=zipfile.ZIP_DEFLATED)
    request.session['whereamI2'] = 'export_app'
    for r in real_dec:
        try:
          appendices(request, str(r.id))
          filename = 'appendices-' + str(r.id) + '-R.xls'  
          zipObj.write(filename)
        except:
           print 'could not write error APP'
           print r.id 
    zipObj.close()
    response = HttpResponse(f.getvalue(),content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=AppendicesReal.zip'
    return response

def export_app_t(request):
    dec = Decisions.objects.filter(real_dec_yn = 'T')
    real_dec = dec.exclude(demoDec = 'Y')
    f = io.BytesIO()
    zipObj = ZipFile(f, 'a', compression=zipfile.ZIP_DEFLATED)
    request.session['whereamI2'] = 'export_app'
    for r in real_dec:
        try:
          appendices(request, str(r.id))
          filename = 'appendices-' + str(r.id) + '-T.xls'  
          zipObj.write(filename)
        except:
           print 'could not write error APP T'
           print r.id 
    zipObj.close()
    response = HttpResponse(f.getvalue(),content_type='application/zip')
    response['Content-Disposition'] = 'attachment; filename=AppendicesTraining.zip'
    return response

def export_test(request):     
    if 'user' in request.session:
       loggedinuser = request.session['user']
    else:
       loggedinuser = 'not found'

    row_num = 0 
    database = MySQLdb.connect (host="amritha.mysql.pythonanywhere-services.com", user = "amritha", passwd = "lilies19", charset="utf8", db = "amritha$costutility")
    cursor2 = database.cursor ()
    #sql2 = """SELECT * FROM utility_tool_mappingtable where dec_id not in %(demo_dec_list)s"""
    sql2 = """SELECT dec_id, sol_id, sol_option, sc_id, criterion, result, created_date, created_by, updated_date, updated_by FROM utility_tool_mapping_data"""                

    try:
       cursor2.execute(sql2,{'demo_dec_list' :demo_dec_list})
       results = cursor2.fetchall()
       for row in results:
          row_num += 1
          for col_num in xrange(len(row)):
              print row[col_num]
    except BaseException as error:
       print('An exception occurred: {}'.format(error))
    return HttpResponseRedirect('/utility_tool/admin/options.html')

def extract_idn_data(request):
    Identify_Data.objects.all().delete()
    i = 1
    qset = IdentifyTable.objects.all().values('dec_id').distinct()
    for q in qset:
    #for idtable in IdentifyTable.objects.all().order_by('-id').distinct('dec_id'):
    #while i < 2:
        Temp_Mapping.objects.all().delete()
        Cri_Temp_Mapping.objects.all().delete()
        idtable = IdentifyTable.objects.filter(dec_id=q['dec_id']).last()  
        #idtable = IdentifyTable.objects.filter(dec_id=692).last()  
        #dec_id = 761 767 805
        #jul16 692 556 782 608
        table =  idtable.table
        dec_id = idtable.dec_id
        #print 'dec id and tanle id'
        #print dec_id
        #print idtable.id 
        #print 'first row'
        # remove the first [ from the array we got from ajax  
        # get the first and last postion of the solution options list
        # add all the solution options into a temporary table called Temp_Mapping
        firstpos = table.find('[') + 2                                                                                                    
        lastpos =  table.find(']')
        arr1 = table[firstpos:lastpos]
        pos = 3
        archived = 'N'
        deleted = 'N' 
        for lx in arr1.split(','):
           l4 = lx.replace('"', '')
           if l4 not in ("Evaluation Criterion", "Common evaluation measure you can use across all options", "Data to collect"):    
              l5 = l4.replace('Describe the information you will use to evaluate  ','')
              l3 = l5.replace('  against this criterion and where you will get it from.','')
              print 'l3'
              print l3
              try:
                 sol = Solution_Options.objects.get(dec_id=dec_id, sol_option = l3)  
                 #print sol.archived
                 print sol.dec_id
                 print sol.id 
                 if sol.archived == 'Y' or sol.deleted == 'Y':
                    archived = 'Y'  
                 else:
                    archived = 'N' 
                 tm = Temp_Mapping(dec_id = dec_id, sol_id = sol.id, sol_position = pos, sol_option = l3, archived = archived) 
                 tm.save()
                 pos = pos + 1
              except:
                 print 'not inserting' 
    
        # Loop through the rest of the rows. Ignore the Keep Option row.
        a = table.replace(arr1,'')
        #aa = a.replace('",null,"','",0,"')
        #print aa
        #b = aa.replace('null,','0,')
        #print b
        c = a.replace('[],','')
        '''w0 = c.replace(',[null]','')
        print w0
        w1 = w0.replace('[0','')
        print w1
        w2 = w1.replace('0','')
        print w2
        w3 = w2.replace('null]','')
        print w3'''
        w4 = c.replace('[[["', '["')
        w = w4.replace('"],','",')  
        firstpos = w.find('[') + 2
        lastpos =  w.find(']')
        arr2 = w[firstpos:lastpos]
        '''aa = w.replace(arr2,'')
        cc = aa.replace('[],','')
        c = cc.replace('["],','')'''
        #print 'c'
        #print c
        c = w
        while len(c):
           '''firstpos = c.find('[') + 2
           lastpos =  c.find(']') '''
           #print 'c jun8'
           #print c
           '''arr3 = c[firstpos:lastpos]
           print arr3
           print 'arr3' '''
           # how many rows are there in the solution options for that decision
           max_pos = Temp_Mapping.objects.filter(dec_id = dec_id).count()
           max_pos = max_pos + 2
           #print 'july 9'
           #print max_pos
           pos2 = 0
           ec_id = 0
           ec_measure = ''                                                                                                              
           ec_data = ''
           criteria = ''
           for ly in w.split(','):
             l5 = ly.replace('[','')
             l4 = l5.replace('"', '')                            
             #print 'forget everything'
             #print l4
             if l4 <> '' and l4 <> 'null' and l4 is not None and l4 <> 'null]' and l4 <> 'null]]':        
                #print 'jun 30 in l4'
                #print pos2
                if pos2 == 0:
                  #print 'crtierion'
                  #print l4
                  try:  
                     ec = Evaluation_Criteria.objects.get(dec_id=dec_id, combined = l4)   
                     deleted = ec.deleted
                     ec_id = ec.id
                     #print ec_id
                  except:
                     print 'ec does not exist' 
                  criteria = l4   
                elif pos2 == 1:                                                                                                              
                  ec_measure = l4
                elif pos2 == 2:
                  ec_data = l4  
                # add all mapping postions for the scr - solopt combination in this temporary table for one option   
                if pos2 > 2 and pos2 <= max_pos:
                  print 'jul 9 222'
                  print pos2
                  print max_pos
                  print l4
                  if ec_id <>  0 and ec_id is not None and ec_id <> ' ' and ec_id <> '':
                     ecm = Cri_Temp_Mapping(cri_id = ec_id, position = pos2, value = l4)
                     ecm.save()
             if pos2 < max_pos:
                #print 'if' 
                #print pos2          
                pos2 = pos2 + 1
             else:
                #print 'else' 
                #print pos2  
                pos2 = 0 
                # create mapping rows for all the options and criteria without the result    
                for temp in Temp_Mapping.objects.all():
                  #print 'jun 30'
                  #print temp.sol_id
                  if ec_id <>  0 and ec_id is not None and ec_id <> ' ' and ec_id <> '':
                     try:
                       m = Identify_Data.objects.get(dec_id = dec_id, sol_id = temp.sol_id, ec_id = ec_id, sol_position = temp.sol_position)
                     except ObjectDoesNotExist:
                       try:
                           cr = Cri_Temp_Mapping.objects.get(cri_id = ec_id, position = temp.sol_position)
                           m = Identify_Data(dec_id = dec_id, sol_id = temp.sol_id, sol_position = temp.sol_position, sol_option = temp.sol_option, ec_id = ec_id, measure = ec_measure, data = ec_data, criterion = criteria, created_date = idtable.created_date,created_by = idtable.created_by,updated_date = idtable.updated_date,updated_by = idtable.updated_by, archived = temp.archived, deleted = deleted, result = cr.value) 
                           m.save()
                       except ObjectDoesNotExist:
                          print 'jun 30 part 2'
                          print dec_id
                          print ec_id
                          print temp.sol_position
                          print criteria 
                          print pos2                           
                # update the result for this row 
                '''for temp in Cri_Temp_Mapping.objects.all():
                  if temp.value <> '' and temp.value is not None and temp.value <> 'null':  
                     try:  
                       m = Identify_Data.objects.get(sol_position = temp.position, ec_id = temp.cri_id, dec_id = dec_id)                        
                       m.result = temp.value                                                                                                 
                       print 'jun 9'
                       print temp.value
                       print dec_id
                       print temp.cri_id
                       print temp.position
                       print temp.value
                       m.save(update_fields=['result'])
                     except ObjectDoesNotExist:
                       print 'PROBLEM' '''
           print 'out of that loop'
           z = c.replace(w, '')
           y = z.replace('[],','')
           print 'this is y'
           print y
           # break out of the loop when only []] remains
           if (y == '[,]' or y == '[[]'):
             break;
           c = y
           if 1 == 1:
              break; 
        #i = 2   
    return HttpResponseRedirect('/utility_tool/admin/options.html')   

def extract_data(request):
    Mapping_Data.objects.all().delete() 
    #i = 1
    for mapptable in MappingTable.objects.all():
    #while i < 2:
        Temp_Mapping.objects.all().delete() 
        Cri_Temp_Mapping.objects.all().delete() 
        #mapptable = MappingTable.objects.get(dec_id=767)  
        #dec_id = 761 
        table =  mapptable.table
        dec_id = mapptable.dec_id
        #print 'first row'
        # remove the first [ from the array we got from ajax  
        # get the first and last postion of the solution options list
        # add all the solution options into a temporary table called Temp_Mapping
        firstpos = table.find('[') + 2
        lastpos =  table.find(']') 
        arr1 = table[firstpos:lastpos]
        pos = 2
        for lx in arr1.split(','):
           l3 = lx.replace('"', '')
           #print 'l3'
           #print l3
           try:
              sol = Solution_Options.objects.get(dec_id=dec_id, sol_option = l3)   
              tm = Temp_Mapping(dec_id = dec_id, sol_id = sol.id, sol_position = pos, sol_option = l3) 
              tm.save()
              pos = pos + 1
           except:
              print 'not inserting' 
    
        # Loop through the rest of the rows. Ignore the Keep Option row.
        a = table.replace(arr1,'')
        aa = a.replace('",null,"','",0,"')
        b = aa.replace('null,','0,')
        c = b.replace('[],','')
        w0 = c.replace(',[null]','')
        w1 = w0.replace('[0','')
        w2 = w1.replace('0','')     
        w = w2.replace('null]','')     
        #print 'w'
        #print w
        firstpos = w.find('[') + 2
        lastpos =  w.find(']') 
        arr2 = w[firstpos:lastpos]
        #print arr2
        #print 'arr2'
        aa = w.replace(arr2,'')
        c = aa.replace('[],','')
        while len(c):
           firstpos = c.find('[') + 2
           lastpos =  c.find(']') 
           print 'c'
           print c

           arr3 = c[firstpos:lastpos]
           print arr3
           print 'arr3'
           pos2 = 0
           scr_id = 0
           criteria = ''
           res_f = ''
           res = ''
           # how many rows are there in the solution options for that decision
           max_pos = Temp_Mapping.objects.filter(dec_id = dec_id).count()
           max_pos = max_pos + 2
           for ly in arr3.split(','):
             l4 = ly.replace('"', '')                             
             if pos2 == 0:
               scr_id = l4
             elif pos2 == 1:
               criteria = l4
             # add all mapping postions for the scr - solopt combination in this temporary table for one option   
             if pos2 > 1 and pos2 < max_pos:
               #print 'am i in here'
               #print scr_id
               if scr_id <>  0 and scr_id is not None and scr_id <> ' ' and scr_id <> '':
                  stm = Cri_Temp_Mapping(cri_id = scr_id, position = pos2, value = l4)
                  stm.save()
             pos2 = pos2 + 1
           # create mapping rows for all the options and criteria without the result    
           for temp in Temp_Mapping.objects.all():
             if scr_id <>  0 and scr_id is not None and scr_id <> ' ' and scr_id <> '':  
                try:
                   m = Mapping_Data.objects.get(dec_id = dec_id, sol_id = temp.sol_id, sc_id = scr_id)
                except ObjectDoesNotExist:
                   m = Mapping_Data(dec_id = dec_id, sol_id = temp.sol_id, sol_position = temp.sol_position, sol_option = temp.sol_option, sc_id = scr_id, criterion = criteria, created_date = mapptable.created_date,created_by = mapptable.created_by,updated_date = mapptable.updated_date,updated_by = mapptable.updated_by)   
                   m.save()

           # update the result for this row 
           for temp in Cri_Temp_Mapping.objects.all():    
              try:  
                 m = Mapping_Data.objects.get(sol_position = temp.position, sc_id = temp.cri_id, dec_id = dec_id)                                                          
                 m.result = temp.value
                 m.save(update_fields=['result'])
              except ObjectDoesNotExist:
                 print 'PROBLEM'  
           print 'out of that loop'
           z = c.replace(arr3, '')
           y = z.replace('[],','')
           print 'this is y'
           print y
           # break out of the loop when only []] remains
           if (y == '[,]' or y == '[[]'):
             break;
           c = y
        #i = 2   
    return HttpResponseRedirect('/utility_tool/admin/options.html')
